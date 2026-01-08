import os
import json
import gc
import shutil
import re
import secrets
from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify, send_file, make_response
from werkzeug.utils import secure_filename
import requests
import uuid
from uuid import uuid4
import tempfile
from datetime import datetime
from io import BytesIO
from password_generator import PasswordGenerator

# Lazy loading - sadece gerektiğinde yükle
pandas_loaded = False
matplotlib_loaded = False
plt = None
pd = None
np = None
mdates = None
Rectangle = None

def load_pandas():
    global pd, pandas_loaded
    if not pandas_loaded:
        import pandas as pd
        pandas_loaded = True
    return pd

def load_matplotlib():
    global plt, np, mdates, Rectangle, matplotlib_loaded
    if not matplotlib_loaded:
        import matplotlib
        matplotlib.use('Agg')  # GUI olmadan çalışması için
        import matplotlib.pyplot as plt
        import matplotlib.dates as mdates
        import numpy as np
        from matplotlib.patches import Rectangle
        matplotlib_loaded = True
    return plt, np, mdates, Rectangle
# Conditional imports for optional dependencies - Lazy loading
DOCX_AVAILABLE = False
WEASYPRINT_AVAILABLE = False

def load_docx():
    global DOCX_AVAILABLE
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        DOCX_AVAILABLE = True
        print(f"DEBUG: load_docx başarılı, DOCX_AVAILABLE = {DOCX_AVAILABLE}")
        return Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    except ImportError as e:
        DOCX_AVAILABLE = False
        print(f"DEBUG: load_docx hatası: {e}")
        return None, None, None, None, None, None, None

app = Flask(__name__, template_folder='templates', static_folder='static')
app.secret_key = os.environ.get('SECRET_KEY', 'your_secret_key_development')  # Production'da ortam değişkeninden al

# Production-safe Flask ayarları
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0
is_dev = os.environ.get('FLASK_ENV') == 'development'
app.config['TEMPLATES_AUTO_RELOAD'] = True if is_dev else False
app.config['DEBUG'] = True if is_dev else (os.environ.get('FLASK_DEBUG', 'false').lower() == 'true')
app.jinja_env.auto_reload = True if is_dev else False

# HESAP Projesi Route'ları
# @app.route('/olcum') - Gereksiz sayfa silindi
# @app.route('/ayarlar') - Gereksiz sayfa silindi
# @app.route('/birimler') - Gereksiz sayfa silindi
# @app.route('/formuller') - Gereksiz sayfa silindi
# @app.route('/hesaplama_tablosu') - Gereksiz sayfa silindi

# @app.route('/olc_istenilen') - Gereksiz sayfa silindi
# def olc_istenilen(): - Gereksiz sayfa silindi

def allowed_file(filename, allowed_extensions):
    """Dosya uzantısının izin verilen uzantılar arasında olup olmadığını kontrol eder."""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in allowed_extensions

DATA_DIR = os.environ.get('DATA_DIR')
if DATA_DIR:
    try:
        os.makedirs(DATA_DIR, exist_ok=True)
    except Exception as _e:
        print(f"DATA_DIR oluşturulamadı: {DATA_DIR}, hata: {_e}")

def data_path(filename: str) -> str:
    return os.path.join(DATA_DIR, filename) if DATA_DIR else filename

USERS_FILE = data_path('users.json')
# EMISSIONS_FILE = data_path('emissions.json') - Emission sistemi silindi
PARAMETERS_FILE = data_path('parameters.json')
# MEASUREMENTS_FILE = data_path('measurements.json') - Ölçüm oluşturma sistemi silindi
FIRMA_OLCUM_FILE = data_path('firma_olcum.json')
FIRMA_KAYIT_FILE = data_path('firma_kayit.json')
SAHA_OLC_FILE = data_path('saha_olc.json')
BACA_BILGILERI_FILE = data_path('baca_bilgileri.json')
PARAMETRE_OLCUM_FILE = data_path('parametre_olcum.json')
# TEKLIF_FILE ve USED_TEKLIF_NUMBERS_FILE silindi
BACA_PARALAR_FILE = data_path('baca_paralar.json')
PARAMETRE_SAHABIL_FILE = data_path('parametre_sahabil.json')
PARAMETRE_FIELDS_FILE = data_path('parametre_fields.json')
ASGARI_FIYATLAR_FILE = data_path('asgari_fiyatlar.json')
PAR_SAHA_HEADERS_FILE = data_path('par_saha_header_groups.json')
FORMS_FILE = data_path('forms.json')
IL_ILCE_FILE = data_path('il-ilce.json')
RAPOR2_FILE = data_path('rapor2_sablonu.json')
ACCESS_CODE_FILE = data_path('access_code.json')

def ensure_data_files():
    # İlk çalıştırmada mevcut repo kökünden DATA_DIR'e veri kopyala
    if not DATA_DIR:
        return
    try:
        # İç içe import ile global import tekrarlarını önle
        import shutil
        base_files = [
            'users.json', 'parameters.json', 'excel_data.json',
            'firma_olcum.json', 'firma_kayit.json', 'saha_olc.json', 'baca_bilgileri.json',
            'parametre_olcum.json', # 'teklif.json', 'used_teklif_numbers.json' silindi
            'parametre_sahabil.json', 'forms.json', 'parametre_fields.json', 'baca_paralar.json',
            'par_saha_header_groups.json'
        ]
        for fname in base_files:
            src = fname  # repo kökü
            dst = data_path(fname)
            # DATA_DIR'de yoksa ve repo kökünde varsa kopyala
            if not os.path.exists(dst) and os.path.exists(src):
                try:
                    shutil.copy2(src, dst)
                    print(f"DATA_DIR bootstrap: {src} -> {dst}")
                except Exception as e:
                    print(f"DATA_DIR kopyalama hatası {src} -> {dst}: {e}")
    except Exception as e:
        print(f"DATA_DIR başlangıç kopyalama hatası: {e}")

# Başlangıçta tek seferlik kontrol
ensure_data_files()

# RAPOR2 not defteri için yardımcı fonksiyonlar
RAPOR2_FILE = data_path('rapor2.json')

def load_rapor2_content():
    if not os.path.exists(RAPOR2_FILE):
        return ''
    try:
        with open(RAPOR2_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if isinstance(data, dict):
                return data.get('content', '')
            if isinstance(data, str):
                return data
    except Exception as exc:
        print(f"rapor2.json okunamadı: {exc}")
    return ''

def save_rapor2_content(content):
    try:
        with open(RAPOR2_FILE, 'w', encoding='utf-8') as f:
            json.dump({'content': content}, f, ensure_ascii=False, indent=2)
        return True
    except Exception as exc:
        print(f"rapor2.json kaydedilemedi: {exc}")
        return False

def load_users():
    """Kullanıcıları JSON dosyasından yükler."""
    if not os.path.exists(USERS_FILE):
        # Dosya yoksa, varsayılan admin kullanıcısı ile oluştur
        default_users = {'admin': {'password': '1111', 'role': 'admin'}}
        save_users(default_users)
        return default_users
    with open(USERS_FILE, 'r', encoding='utf-8') as f:
        return json.load(f)

def save_users(users_data):
    """Kullanıcıları JSON dosyasına kaydeder."""
    with open(USERS_FILE, 'w', encoding='utf-8') as f:
        json.dump(users_data, f, indent=4, ensure_ascii=False)


def load_access_code():
    """Aktif giriş şifresini JSON dosyasından yükler."""
    if not os.path.exists(ACCESS_CODE_FILE):
        return None
    try:
        with open(ACCESS_CODE_FILE, 'r', encoding='utf-8') as f:
            content = f.read().strip()
            if not content:
                return None
            return json.loads(content)
    except (FileNotFoundError, json.JSONDecodeError):
        return None


def save_access_code(code_data: dict) -> None:
    """Aktif giriş şifresini JSON dosyasına kaydeder."""
    with open(ACCESS_CODE_FILE, 'w', encoding='utf-8') as f:
        json.dump(code_data, f, indent=2, ensure_ascii=False)

# def load_emissions(): - Emission sistemi silindi
# def save_emissions(emissions_data): - Emission sistemi silindi



def load_parameters():
    """Parametreleri excel_data.json dosyasından yükler."""
    excel_data_file = data_path('excel_data.json')
    if not os.path.exists(excel_data_file):
        return []
    
    try:
        with open(excel_data_file, 'r', encoding='utf-8') as f:
            excel_data = json.load(f)
        
        # Excel verisini parametre formatına dönüştür
        parameters = []
        if excel_data and len(excel_data) > 1:  # İlk satır başlık
            for i, row in enumerate(excel_data[1:], 1):  # İlk satırı atla
                if len(row) >= 2 and row[1]:  # A sütunu (index 1) boş değilse
                    param_name = row[1]  # B sütunu (index 1) - Parametre Adı
                    method = row[2] if len(row) > 2 else ""  # C sütunu - Metot
                    
                    # Excel verisinden parametre objesi oluştur
                    param_obj = {
                        "Parametre Adı": param_name,
                        "Metot": method,
                        "İzo Oran": row[3] if len(row) > 3 else "",
                        "Nozzle": row[4] if len(row) > 4 else "",
                        "1. İmp": row[5] if len(row) > 5 else "",
                        "2. İmp": row[6] if len(row) > 6 else "",
                        "3. İmp": row[7] if len(row) > 7 else "",
                        "4. İmp": row[8] if len(row) > 8 else "",
                        "L/DAK": row[9] if len(row) > 9 else "",
                        "T.HAC": row[10] if len(row) > 10 else "",
                        "LOQ": row[11] if len(row) > 11 else "",
                        "id": f"excel_{i}_{hash(param_name) % 10000}",
                        "KK": "",
                        "-3S": "",
                        "-2S": "",
                        "+2S": "",
                        "+3S": ""
                    }
                    parameters.append(param_obj)
        
        return parameters
    except Exception as e:
        print(f"Excel verisi yüklenirken hata: {e}")
        return []

def save_parameters(parameters_data):
    """Parametreleri JSON dosyasına kaydeder."""
    with open(PARAMETERS_FILE, 'w', encoding='utf-8') as f:
        json.dump(parameters_data, f, indent=4, ensure_ascii=False)

# def load_measurements(): - Ölçüm oluşturma sistemi silindi
# def save_measurements(measurements_data): - Ölçüm oluşturma sistemi silindi

def load_firma_olcum():
    """Firma ölçüm verilerini JSON dosyasından yükler."""
    try:
        if not os.path.exists(FIRMA_OLCUM_FILE):
            return []
        with open(FIRMA_OLCUM_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if not isinstance(data, list):
                return []
            return data
    except Exception as e:
        print(f"Firma ölçüm verileri yüklenirken hata: {e}")
        return []

def save_firma_olcum(firma_olcum_data):
    """Firma ölçüm verilerini JSON dosyasına kaydeder."""
    try:
        if not isinstance(firma_olcum_data, list):
            print("Hata: firma_olcum_data liste olmalıdır")
            return False
        with open(FIRMA_OLCUM_FILE, 'w', encoding='utf-8') as f:
            json.dump(firma_olcum_data, f, indent=4, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Firma ölçüm verileri kaydedilirken hata: {e}")
        return False

def load_firma_kayit():
    """Firma kayıt verilerini JSON dosyasından yükler."""
    try:
        if not os.path.exists(FIRMA_KAYIT_FILE):
            return []
        with open(FIRMA_KAYIT_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if not isinstance(data, list):
                return []
            return data
    except Exception as e:
        print(f"Firma kayıt verileri yüklenirken hata: {e}")
        return []

def save_firma_kayit(firma_kayit_data):
    """Firma kayıt verilerini JSON dosyasına kaydeder."""
    try:
        if not isinstance(firma_kayit_data, list):
            print("Hata: firma_kayit_data liste olmalıdır")
            return False
        with open(FIRMA_KAYIT_FILE, 'w', encoding='utf-8') as f:
            json.dump(firma_kayit_data, f, indent=4, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Firma kayıt verileri kaydedilirken hata: {e}")
        return False

def load_saha_olc():
    """Saha ölçüm verilerini JSON dosyasından yükler."""
    try:
        if not os.path.exists(SAHA_OLC_FILE):
            return []
        with open(SAHA_OLC_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if not isinstance(data, list):
                return []
            return data
    except Exception as e:
        print(f"Saha ölçüm verileri yüklenirken hata: {e}")
        return []

def save_saha_olc(saha_olc_data):
    """Saha ölçüm verilerini JSON dosyasına kaydeder."""
    try:
        if not isinstance(saha_olc_data, list):
            print("Hata: saha_olc_data liste olmalıdır")
            return False
        with open(SAHA_OLC_FILE, 'w', encoding='utf-8') as f:
            json.dump(saha_olc_data, f, indent=4, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Saha ölçüm verileri kaydedilirken hata: {e}")
        return False

def load_baca_bilgileri():
    """Baca bilgilerini JSON dosyasından yükler."""
    try:
        if not os.path.exists(BACA_BILGILERI_FILE):
            return []
        with open(BACA_BILGILERI_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if not isinstance(data, list):
                return []

            # Eski kayıtlarda id yoksa otomatik üret ve dosyaya geri yaz
            changed = False
            for rec in data:
                if not rec.get('id'):
                    rec['id'] = str(uuid4())
                    changed = True

            if changed:
                try:
                    with open(BACA_BILGILERI_FILE, 'w', encoding='utf-8') as wf:
                        json.dump(data, wf, indent=4, ensure_ascii=False)
                except Exception as e:
                    print(f"Baca bilgileri ID güncellemesi kaydedilemedi: {e}")

            return data
    except Exception as e:
        print(f"Baca bilgileri yüklenirken hata: {e}")
        return []

def save_baca_bilgileri(baca_bilgileri_data):
    """Baca bilgilerini JSON dosyasına kaydeder."""
    try:
        if not isinstance(baca_bilgileri_data, list):
            print("Hata: baca_bilgileri_data liste olmalıdır")
            return False
        with open(BACA_BILGILERI_FILE, 'w', encoding='utf-8') as f:
            json.dump(baca_bilgileri_data, f, indent=4, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Baca bilgileri kaydedilirken hata: {e}")
        return False


def get_saha_olcum_summary():
    """Kayıtlı saha ölçümleri tablosu için gruplanmış özet liste döner.

    Öncelik sırası:
    - Eğer baca_bilgileri.json içinde saha kayıtları varsa, bunlar kullanılır.
    - Eğer baca_bilgileri.json boşsa, firma_olcum.json içindeki ölçüm
      tanımlarından özet kayıtlar üretilir (her firma/ölçüm/baca için).
    """
    # 1) Önce gerçek saha kayıtlarını (baca_bilgileri.json) dene
    raw_records = load_baca_bilgileri()
    groups = {}
    if raw_records:
        for rec in raw_records:
            firma = (rec.get('firma_adi') or rec.get('firma') or '').strip()
            olcum_kodu = str(rec.get('olcum_kodu') or '').strip()
            baca_val = (rec.get('baca_adi') or rec.get('baca_no') or '').strip()
            full_tarih = (rec.get('kayit_tarihi') or
                          rec.get('created_at') or
                          rec.get('updated_at') or '')

            # Aynı saniye içindeki ölçümleri tek grup saymak için
            tarih_key = full_tarih[:19] if full_tarih else ''

            key = (firma, olcum_kodu, baca_val, tarih_key)

            if key not in groups:
                groups[key] = {
                    'id': None,          # İlk kaydın ID'si (görsel amaçlı)
                    'ids': [],           # Bu kombinasyona ait tüm ID'ler
                    'firma_adi': firma,
                    'firma': firma,
                    'olcum_kodu': olcum_kodu,
                    'baca_adi': rec.get('baca_adi') or baca_val,
                    'baca_no': rec.get('baca_no') or baca_val,
                    'parametre_adi': [],
                    'parametre': [],
                    'kayit_tarihi': full_tarih,
                    'personel_adi': (rec.get('personel_adi') or '').strip(),
                }

            group = groups[key]

            rec_id = str(rec.get('id') or '').strip()
            if rec_id:
                group['ids'].append(rec_id)
                if group['id'] is None:
                    group['id'] = rec_id

            param_val = (rec.get('parametre_adi') or rec.get('parametre') or '').strip()
            if param_val and param_val not in group['parametre_adi']:
                group['parametre_adi'].append(param_val)
    else:
        # 2) Baca verisi yoksa, firma_olcum.json üzerinden özet üret
        try:
            firma_olcumler = load_firma_olcum()
        except Exception:
            firma_olcumler = []

        for olcum in firma_olcumler:
            firma = (olcum.get('firma_adi') or '').strip()
            olcum_kodu = str(olcum.get('olcum_kodu') or '').strip()
            bas_tarih = (olcum.get('baslangic_tarihi') or '').strip()
            created_at = (olcum.get('olusturma_tarihi') or '').strip()
            full_tarih = created_at or bas_tarih
            tarih_key = full_tarih[:19] if full_tarih else ''

            baca_parametreleri = olcum.get('baca_parametreleri') or {}
            if not isinstance(baca_parametreleri, dict):
                continue

            for baca_no, param_list in baca_parametreleri.items():
                baca_val = str(baca_no or '').strip()
                key = (firma, olcum_kodu, baca_val, tarih_key)

                if key not in groups:
                    # Firma ölçüm kayıtlarının kendi ID'sini grup ID'si olarak kullan
                    group_id = f"{olcum.get('id','')}::{baca_val or '-'}::{tarih_key}"
                    groups[key] = {
                        'id': group_id,
                        'ids': [],  # Henüz gerçek baca_bilgileri kaydı yok
                        'firma_adi': firma,
                        'firma': firma,
                        'olcum_kodu': olcum_kodu,
                        'baca_adi': baca_val,
                        'baca_no': baca_val,
                        'parametre_adi': [],
                        'parametre': [],
                        'kayit_tarihi': full_tarih,
                        'personel_adi': ', '.join(olcum.get('personel') or []),
                    }

                group = groups[key]
                # Parametre isimlerini topla
                for p in (param_list or []):
                    p_val = str(p or '').strip()
                    if p_val and p_val not in group['parametre_adi']:
                        group['parametre_adi'].append(p_val)

    result = []
    for group in groups.values():
        raw_params = group.get('parametre_adi') or []

        def _format_param(p):
            text = str(p or '')
            text = text.replace('_', ' ')
            # Harf ve rakamın bitişik olduğu son ekleri ayır ("EPA5" -> "EPA 5")
            text = re.sub(r"([A-Za-zÇĞİÖŞÜçğıöşü])([0-9])", r"\1 \2", text)
            return text.upper().strip()

        formatted_params = [_format_param(p) for p in raw_params if str(p or '').strip()]
        joined = ', '.join(formatted_params)

        group['parametre_adi'] = joined
        group['parametre'] = joined
        result.append(group)

    # Firma / ölçüm / baca sırasına göre sırala
    result.sort(key=lambda r: (
        (r.get('firma_adi') or ''),
        (r.get('olcum_kodu') or ''),
        (r.get('baca_adi') or r.get('baca_no') or ''),
    ))

    return result

def load_parametre_olcum():
    """Parametre ölçüm verilerini JSON dosyasından yükler."""
    try:
        if not os.path.exists(PARAMETRE_OLCUM_FILE):
            return []
        with open(PARAMETRE_OLCUM_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
            if not isinstance(data, list):
                return []
            return data
    except Exception as e:
        print(f"Parametre ölçüm verileri yüklenirken hata: {e}")
        return []

def save_parametre_olcum(parametre_olcum_data):
    """Parametre ölçüm verilerini JSON dosyasına kaydeder."""
    try:
        if not isinstance(parametre_olcum_data, list):
            print("Hata: parametre_olcum_data liste olmalıdır")
            return False
        with open(PARAMETRE_OLCUM_FILE, 'w', encoding='utf-8') as f:
            json.dump(parametre_olcum_data, f, indent=4, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Parametre ölçüm verileri kaydedilirken hata: {e}")
        return False

# load_teklif ve save_teklif fonksiyonları silindi

# load_used_teklif_numbers fonksiyonu silindi
# save_used_teklif_numbers fonksiyonu silindi

# generate_teklif_no fonksiyonu silindi
# generate_teklif_no fonksiyonu içeriği silindi

# reserve_teklif_no fonksiyonu silindi
# reserve_teklif_no fonksiyonu içeriği silindi

# release_teklif_no fonksiyonu silindi
# release_teklif_no fonksiyonu içeriği silindi

# migrate_existing_teklif_numbers fonksiyonu silindi
# migrate_existing_teklif_numbers fonksiyonu içeriği silindi

# convert_teklif_numbers_to_3_digit fonksiyonu silindi
# convert_teklif_numbers_to_3_digit fonksiyonu içeriği silindi

# resequence_teklif_numbers fonksiyonu silindi
# resequence_teklif_numbers fonksiyonu içeriği silindi

def format_tarih_gg_aa_yyyy(tarih_str):
    """YYYY-MM-DD formatındaki tarihi GG.AA.YYYY formatına çevirir"""
    if not tarih_str:
        return ''
    try:
        # YYYY-MM-DD formatını parse et
        tarih_obj = datetime.strptime(tarih_str, '%Y-%m-%d')
        # GG.AA.YYYY formatına çevir
        return tarih_obj.strftime('%d.%m.%Y')
    except:
        return tarih_str

# Template'de kullanılacak yardımcı fonksiyonları Flask'a ekle
app.jinja_env.globals.update(format_tarih_gg_aa_yyyy=format_tarih_gg_aa_yyyy)

@app.route('/api/formlar/genel-hukumler', methods=['GET'])
def get_genel_hukumler():
    """Formlar sayfasındaki genel hükümler metnini döndürür"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        # Varsayılan genel hükümler metni
        default_text = """<strong>GENEL HÜKÜMLER</strong>
<ul>
    <li><strong>Fiyatlara KDV dahil değildir.</strong></li>
    <li><strong>Teklifimizin geçerlilik süresi 30 gündür.</strong></li>
    <li><strong>(*) işaretli parametrelerin ölçümü ve analizleri yetkili taşeron işbirliği laboratuvarı tarafından yapılacaktır.</strong></li>
    <li><strong>Fiyat tarifemiz; Çevre, Şehircilik ve İklim Değişikliği Bakanlığı'nın 27.12.2024 tarihinde yayınlanan duyurusu itibarıyla, 01.01.2025 tarihi itibariyle geçerli olacak şekil tarifi öz sınırları kaynaklamıştır.</strong></li>
    <li><strong>Yetkili Ölçüm ve Analiz Laboratuvarlarından 02.12.2024 tarihinden sonra Fiyat Tarifesi 13. Maddesi gereği, ödeme rapor tesliminden önce para yapılacaktır.</strong></li>
    <li><strong>Ölçüm esnasında firmanızdan kaynaklanan bir hangi bir sebepten dolayı ölçüm yapılamaması durumunda ücret iadeesi yapılmayacaktır.</strong></li>
    <li><strong>Tüm anlaşmazlıklarda İstanbul (Çağlayan) Adliyeleri yetkilidir.</strong></li>
    <li><strong>İş Güvenliği Firmanız sorumluluğundadır.</strong></li>
    <li><strong>Numuneler 30 gün saklanacaktır.</strong></li>
    <li><strong>Raporlara itiraz süresi teslim tarihinden itibaren 15 gündür.</strong></li>
    <li><strong>Raporlar 15 iş günü içinde teslim edilecektır.</strong></li>
    <li><strong>Laboratuvarımız TS EN ISO/IEC 17025 standardına uygun çalışmaktadır.</strong></li>
    <li><strong>Karar kuralları: Ölçüm sonuçları belirsizlik hesaplamaları ile birlikte değerlendirilir.</strong></li>
</ul>"""
        
        return jsonify({'success': True, 'content': default_text})
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

def format_date_with_day(dt_str):
    try:
        if not dt_str or dt_str is None:
            return ''
        dt_str = str(dt_str).strip()
        if not dt_str:
            return ''
        # Eğer zaten GG.AA.YY (GUN) formatındaysa, olduğu gibi döndür
        if '(' in dt_str and ')' in dt_str:
            return dt_str
        # Eğer zaten GG.AA.YY formatındaysa, gün ekle
        if '.' in dt_str and len(dt_str) == 8:
            dt = datetime.strptime(dt_str, '%d.%m.%y')
        # Sadece tarih formatı (YYYY-MM-DD)
        elif len(dt_str) == 10 and dt_str.count('-') == 2:
            dt = datetime.strptime(dt_str, '%Y-%m-%d')
        # Tarih-saat formatı (YYYY-MM-DDTHH:MM veya YYYY-MM-DD HH:MM)
        elif 'T' in dt_str:
            dt = datetime.strptime(dt_str, '%Y-%m-%dT%H:%M')
        else:
            dt = datetime.strptime(dt_str, '%Y-%m-%d %H:%M')
        # Türkçe kısa günler
        gunler = ['PZT', 'SAL', 'ÇAR', 'PER', 'CUM', 'CTS', 'PAZ']
        gun_index = dt.weekday()  # Pazartesi=0, Pazar=6
        gun_kisa = gunler[gun_index]
        return dt.strftime('%d.%m.%y') + f' ({gun_kisa})'
    except Exception as e:
        print(f"Tarih formatı hatası: {dt_str}, Hata: {e}")
        return str(dt_str) if dt_str else ''

# Yetki kontrolü fonksiyonları
def can_read(role):
    """Rol 1, 2, 3 veya admin ise okuma yetkisi var"""
    return role in ['1', '2', '3', 'admin']

def can_write(role):
    """Rol 2, 3 veya admin ise yazma yetkisi var"""
    return role in ['2', '3', 'admin']

def can_edit(role):
    """Rol 3 veya admin ise düzenleme yetkisi var"""
    return role in ['3', 'admin']

def can_delete(role):
    """Rol 3 veya admin ise silme yetkisi var"""
    return role in ['3', 'admin']

# Uygulama başlangıcında verileri yükle
users = load_users()

# İl-ilçe verilerini yükle
CITIES_DATA = []
try:
    # Verileri yerel dosyadan okumayı dene
    with open('il-ilce.json', 'r', encoding='utf-8') as f:
        CITIES_DATA = json.load(f)
except FileNotFoundError:
    # Yerel dosya yoksa URL'den çek
    try:
        response = requests.get('https://raw.githubusercontent.com/furkan-dogu/Turkiye-Sehir-ve-Ilceleri/main/il-ilce.json')
        response.raise_for_status() # HTTP hatalarını kontrol et
        CITIES_DATA = response.json()
        # Gelen veriyi kalıcı veri klasörüne kaydet
        with open(IL_ILCE_FILE, 'w', encoding='utf-8') as f:
            json.dump(CITIES_DATA, f, ensure_ascii=False, indent=2)
    except requests.exceptions.RequestException as e:
        print(f"Uyarı: İl-ilçe verileri yüklenemedi. Hata: {e}")

@app.route('/login', methods=['GET', 'POST'])
def login():
    error = None
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        # Kullanıcı adını küçük harfe çevir ve büyük/küçük harf duyarlılığını kaldır
        username_lower = username.lower()
        
        # Kullanıcıları yükle ve kullanıcı adını küçük harfe çevirerek kontrol et
        current_users = load_users()
        user = None
        original_username = None
        
        # Kullanıcı adını büyük/küçük harf duyarlılığı olmadan ara
        for stored_username, user_data in current_users.items():
            if stored_username.lower() == username_lower:
                user = user_data
                original_username = stored_username
                break

        login_ok = False

        if user:
            # Admin kendi şifresiyle giriş yapar
            if original_username == 'admin':
                if user.get('password') == password:
                    login_ok = True
            else:
                # Diğer kullanıcılar için aktif 6 haneli giriş şifresi zorunlu
                access_info = load_access_code() or {}
                active_code = str(access_info.get('code') or '').strip()

                if not active_code:
                    error = 'Giriş şifresi henüz oluşturulmadı. Lütfen admin ile iletişime geçin.'
                elif password == active_code:
                    login_ok = True

        if login_ok and user and original_username:
            session['logged_in'] = True
            session['username'] = original_username  # Orijinal kullanıcı adını kullan
            session['role'] = user.get('role')
            if original_username != 'admin':
                access_info = load_access_code() or {}
                session['access_code_code'] = str(access_info.get('code') or '').strip()
            else:
                session.pop('access_code_code', None)
            return redirect(url_for('index'))

        if not error:
            error = 'Geçersiz kullanıcı adı veya şifre'

    return render_template('login.html', error=error)

@app.before_request
def enforce_access_code_session():
    if not session.get('logged_in'):
        return
    username = session.get('username')
    role = session.get('role')
    # Admin bu kontrolden muaf
    if username == 'admin' or role == 'admin':
        return
    access_info = load_access_code() or {}
    current_code = str(access_info.get('code') or '').strip()
    previous_code = str(session.get('access_code_code') or '').strip()
    if current_code and previous_code and current_code != previous_code:
        flash('Giriş şifresi değiştirildi. Lütfen yeni şifre ile tekrar giriş yapınız.', 'error')
        session.pop('logged_in', None)
        session.pop('username', None)
        session.pop('role', None)
        session.pop('access_code_code', None)
        return redirect(url_for('login'))

@app.route('/')
def index():
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    access_info = load_access_code() or {}
    return render_template(
        'index.html',
        username=session.get('username'),
        role=session.get('role'),
        access_code=access_info.get('code')
    )
@app.route('/api/generate_access_code', methods=['POST'])
def api_generate_access_code():
    """Admin için yeni 6 haneli giriş şifresi üretir ve kaydeder."""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum bulunamadı'}), 401

    # Hem role hem kullanıcı adına bakarak admin kontrolü yap
    if session.get('role') != 'admin' and session.get('username') != 'admin':
        return jsonify({'success': False, 'message': 'Yetkiniz yok'}), 403

    try:
        # Her çağrıda rastgele ve bağımsız 6 haneli kod üret
        code = f"{secrets.randbelow(1_000_000):06d}"

        info = {
            'code': code,
            'generated_at': datetime.now().isoformat(),
            'generated_by': session.get('username') or 'admin'
        }
        save_access_code(info)

        return jsonify({'success': True, 'code': code, 'generated_at': info['generated_at']})
    except Exception as e:
        print(f"Erişim şifresi üretilirken hata: {e}")
        return jsonify({'success': False, 'message': 'Şifre oluşturulurken hata oluştu'}), 500

@app.route('/logout')
def logout():
    session.pop('logged_in', None)
    session.pop('username', None)
    session.pop('role', None)
    return redirect(url_for('login'))


# @app.route('/api/pivot/summary') - Pivot sistemi silindi
# def pivot_summary(): - Pivot sistemi silindi
    if not session.get('logged_in'):
        return jsonify({'error': 'auth'}), 401
    if session.get('role') != 'admin':
        return jsonify({'error': 'forbidden'}), 403
    try:
        start = request.args.get('start')
        end = request.args.get('end')
        if not start or not end:
            return jsonify({'error': 'start/end required'}), 400
        from datetime import datetime
        def in_range(dt):
            try:
                d = datetime.fromisoformat(dt.replace('Z', '+00:00')).date()
            except Exception:
                try:
                    d = datetime.strptime(dt, '%Y-%m-%d').date()
                except Exception:
                    return False
            return datetime.strptime(start, '%Y-%m-%d').date() <= d <= datetime.strptime(end, '%Y-%m-%d').date()
        def _norm(sval: str) -> str:
            m = (sval or '').strip().lower()
            tr_map = str.maketrans({'ı':'i','İ':'i','ş':'s','Ş':'s','ç':'c','Ç':'c','ğ':'g','Ğ':'g','ü':'u','Ü':'u','ö':'o','Ö':'o'})
            return m.translate(tr_map)
        # teklifler = load_teklif() - teklif sistemi silindi
        # filt_teklif = [t for t in teklifler if t.get('teklif_tarihi') and in_range(t.get('teklif_tarihi'))]
        toplam_teklif_adedi = 0  # teklif sistemi silindi
        toplam_teklif_tutari = 0  # teklif sistemi silindi
        accepted_list, rejected_list = [], []  # teklif sistemi silindi
        kabul_adet = 0  # teklif sistemi silindi
        red_adet = 0  # teklif sistemi silindi
        kapsam_ici = []
        kapsam_disi = []
        for t in accepted_list:
            tip = _norm(t.get('teklif_tipi',''))
            if 'kapsam' in tip and 'ici' in tip:
                kapsam_ici.append(t)
            elif 'kapsam' in tip and ('disi' in tip or 'dış' in tip):
                kapsam_disi.append(t)
        kapsam_ici_adet = len(kapsam_ici)
        kapsam_disi_adet = len(kapsam_disi)
        kapsam_ici_tutar = sum(float(t.get('netToplam', 0) or 0) for t in kapsam_ici)
        kapsam_disi_tutar = sum(float(t.get('netToplam', 0) or 0) for t in kapsam_disi)
        from collections import defaultdict
        parametre_ozet = defaultdict(lambda: {'adet': 0, 'toplam': 0.0})
        # for t in filt_teklif: - teklif sistemi silindi
            # for pr in t.get('parametreler', []) or []: - teklif sistemi silindi
        parametre_list = [{'parametre': k, 'adet': v['adet'], 'toplam': round(v['toplam'], 2)} for k, v in sorted(parametre_ozet.items(), key=lambda x: x[0].lower())]
        try:
            baca_list = load_baca_bilgileri()
        except Exception:
            baca_list = []
        baca_index = {}
        for b in baca_list:
            if not (b.get('firma_adi') and b.get('olcum_kodu') and b.get('baca_adi')):
                continue
            key = (b.get('firma_adi'), b.get('olcum_kodu'), b.get('baca_adi'))
            baca_index[key] = (b.get('personel_adi') or '').strip() or 'Bilinmiyor'
        personel_set = set()
        matrix = defaultdict(lambda: defaultdict(int))
        for o in [o for o in load_parametre_olcum() if o.get('created_at') and in_range(o.get('created_at')[:10])]:
            param = (o.get('parametre_adi') or 'Bilinmiyor').strip() or 'Bilinmiyor'
            key = (o.get('firma_adi'), o.get('olcum_kodu'), o.get('baca_adi'))
            personel = (o.get('personel_adi') or '').strip()
            if not personel:
                personel = baca_index.get(key, 'Bilinmiyor')
            if not personel:
                personel = 'Bilinmiyor'
            personel_set.add(personel)
            matrix[param][personel] += 1
        personeller = sorted(list(personel_set), key=lambda x: x.lower())
        matrix_rows = []
        for param, counts in sorted(matrix.items(), key=lambda x: x[0].lower()):
            row = {'parametre': param, 'counts': {p: counts.get(p, 0) for p in personeller}, 'toplam': sum(counts.values())}
            matrix_rows.append(row)
        toplam_by_person = defaultdict(int)
        for counts in matrix.values():
            for p, v in counts.items():
                toplam_by_person[p] += v
        personel_list = [{'personel': p, 'adet': toplam_by_person[p]} for p in sorted(toplam_by_person.keys(), key=lambda x: (-toplam_by_person[x], x.lower()))]
        return jsonify({'summary': {'toplam_teklif_adedi': toplam_teklif_adedi, 'toplam_teklif_tutari': round(toplam_teklif_tutari, 2), 'kapsam_ici_adet': kapsam_ici_adet, 'kapsam_ici_tutar': round(kapsam_ici_tutar, 2), 'kapsam_disi_adet': kapsam_disi_adet, 'kapsam_disi_tutar': round(kapsam_disi_tutar, 2), 'kabul_adet': kabul_adet, 'red_adet': red_adet, 'toplam_olcum_adedi': len([o for o in load_parametre_olcum() if o.get('created_at') and in_range(o.get('created_at')[:10])])}, 'parametreler': parametre_list, 'personeller': personel_list, 'personel_parametre': {'personel_headers': personeller, 'rows': matrix_rows}})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# @app.route('/admin', methods=['GET', 'POST']) - Admin sistemi silindi
# def admin(): - Admin sistemi silindi
    # Protect this route
    if session.get('role') != 'admin':
        return redirect(url_for('index'))

    current_users = load_users()

    if request.method == 'POST':
        action = request.form.get('action')

        if action == 'add_user':
            username = request.form['username']
            password = request.form['password']
            surname = request.form.get('surname', '')
            gorev = request.form.get('gorev', '')
            role = request.form['role']
            
            # İmza dosyasını işle
            imza_filename = None
            if 'imza' in request.files and request.files['imza'].filename:
                imza_file = request.files['imza']
                if imza_file and allowed_file(imza_file.filename, {'png', 'jpg', 'jpeg', 'gif'}):
                    # Dosya adını güvenli hale getir
                    filename = secure_filename(imza_file.filename)
                    imza_filename = f"{username}_{filename}"
                    
                    # Static/images/signatures klasörünü oluştur
                    import os
                    signatures_dir = os.path.join('static', 'images', 'signatures')
                    os.makedirs(signatures_dir, exist_ok=True)
                    
                    # Dosyayı kaydet
                    imza_file.save(os.path.join(signatures_dir, imza_filename))
            
            if username and password and username not in current_users:
                current_users[username] = {
                    'password': password, 
                    'role': role,
                    'surname': surname,
                    'gorev': gorev,
                    'imza': imza_filename
                }
                save_users(current_users)
                flash(f'Kullanıcı "{username}" başarıyla eklendi!', 'success')
            else:
                flash('Kullanıcı adı zaten mevcut veya eksik bilgi!', 'danger')

        elif action == 'update_user':
            username = request.form.get('edit_username') or request.form['username']
            if username != 'admin' and username in current_users:
                # Şifre sadece girilmişse güncelle
                if request.form.get('password'):
                    current_users[username]['password'] = request.form['password']
                current_users[username]['role'] = request.form['role']
                current_users[username]['surname'] = request.form.get('surname', '')
                current_users[username]['gorev'] = request.form.get('gorev', '')
                
                # İmza dosyasını işle
                if 'imza' in request.files and request.files['imza'].filename:
                    imza_file = request.files['imza']
                    if imza_file and allowed_file(imza_file.filename, {'png', 'jpg', 'jpeg', 'gif'}):
                        # Dosya adını güvenli hale getir
                        filename = secure_filename(imza_file.filename)
                        imza_filename = f"{username}_{filename}"
                        
                        # Static/images/signatures klasörünü oluştur
                        import os
                        signatures_dir = os.path.join('static', 'images', 'signatures')
                        os.makedirs(signatures_dir, exist_ok=True)
                        
                        # Eski imzayı sil
                        if current_users[username].get('imza'):
                            old_imza_path = os.path.join(signatures_dir, current_users[username]['imza'])
                            if os.path.exists(old_imza_path):
                                os.remove(old_imza_path)
                        
                        # Yeni dosyayı kaydet
                        imza_file.save(os.path.join(signatures_dir, imza_filename))
                        current_users[username]['imza'] = imza_filename
                
                save_users(current_users)
                flash(f'Kullanıcı "{username}" başarıyla güncellendi!', 'success')
            else:
                flash('Kullanıcı güncellenemedi!', 'danger')

        elif action == 'delete_user':
            username = request.form['username']
            if username != 'admin' and username in current_users:
                del current_users[username]
                save_users(current_users)
                flash(f'Kullanıcı "{username}" başarıyla silindi!', 'success')
            else:
                flash('Kullanıcı silinemedi!', 'danger')
        
        return redirect(url_for('admin'))

    # For GET request, pass the list of users to the template
    return render_template('admin.html', username=session.get('username'), users=current_users)

# @app.route('/add_emission', methods=['GET', 'POST']) - Emission sistemi silindi
# def add_emission(): - Emission sistemi silindi

# @app.route('/edit_emission/<int:emission_id>', methods=['GET', 'POST']) - Emission sistemi silindi
# def edit_emission(emission_id): - Emission sistemi silindi

# @app.route('/delete_emission/<int:emission_id>') - Emission sistemi silindi
# def delete_emission(emission_id): - Emission sistemi silindi



@app.route('/parametre')
def parametre():
    if not session.get('logged_in') or not can_read(session.get('role')):
        return redirect(url_for('login'))
    current_parameters = load_parameters()
    return render_template('parametre.html', parameters=current_parameters, username=session.get('username'), role=session.get('role'))
@app.route('/import_parameters', methods=['POST'])
def import_parameters():
    if 'username' not in session:
        return redirect(url_for('login'))

    if 'excel_file' not in request.files:
        flash('Dosya bölümü bulunamadı!', 'danger')
        return redirect(url_for('parametre'))

    file = request.files['excel_file']

    if not file or not file.filename or file.filename == '':
        flash('Dosya seçilmedi!', 'danger')
        return redirect(url_for('parametre'))

    if (file.filename and (file.filename.endswith('.xlsx') or file.filename.endswith('.xls'))):
        try:
            pd = load_pandas()
            df = pd.read_excel(file)
            
            # İstenen sütun sırası
            expected_columns = [
                'Sıra',
                'Parametre Adı',
                'Metot',
                'İzo Oran',
                'Nozzle',
                '1. İmp',
                '2. İmp',
                '3. İmp',
                '4. İmp',
                'L/DAK',
                'T.HAC',
                'LOQ',
                'KK',
                '-3S',
                '-2S',
                '+2S',
                '+3S'
            ]
            
            # Sadece zorunlu sütunları kontrol et
            required_columns = ['Parametre Adı', 'Metot']
            if not all(col in df.columns for col in required_columns):
                flash(f'Excel dosyasında gerekli sütunlar bulunamadı. En az "Parametre Adı" ve "Metot" sütunları olmalıdır.', 'danger')
                return redirect(url_for('parametre'))

            # NaN değerleri boş string ile değiştir
            df.fillna('', inplace=True)

            new_parameters = []
            for index, row in df.iterrows():
                param_dict = {}
                
                # Sadece beklenen sütunları al, diğerlerini atla
                for col in expected_columns:
                    if col in df.columns:
                        cell_value = row[col]
                        if pd.isna(cell_value) or cell_value == '' or cell_value is None:
                            param_dict[col] = ''
                        else:
                            param_dict[col] = str(cell_value)
                    else:
                        param_dict[col] = ''
                
                # Sıra sütununu parametre verilerinden çıkar
                if 'Sıra' in param_dict:
                    del param_dict['Sıra']
                
                param_dict['id'] = str(uuid4())
                new_parameters.append(param_dict)
            
            save_parameters(new_parameters)

            flash('Parametreler başarıyla Excel dosyasından yüklendi.', 'success')
        except Exception as e:
            flash(f'Dosya okunurken bir hata oluştu: {e}', 'danger')
        
        return redirect(url_for('parametre'))
    else:
        flash('Geçersiz dosya formatı. Lütfen .xlsx veya .xls uzantılı bir dosya yükleyin.', 'danger')
        return redirect(url_for('parametre'))


@app.route('/export_parameters')
def export_parameters():
    if 'username' not in session:
        return redirect(url_for('login'))

    parameters = load_parameters()
    if not parameters:
        flash('Dışarı aktarılacak veri bulunmuyor.', 'warning')
        return redirect(url_for('parametre'))

    pd = load_pandas()
    df = pd.DataFrame(parameters)
    # 'Lab LOQ' sütunu varsa kaldır
    if 'Lab LOQ' in df.columns:
        df = df.drop(columns=['Lab LOQ'])
    if 'id' in df.columns:
        df = df.drop(columns=['id'])

    # İstenen sıralamaya göre sütunları yeniden düzenle
    desired_order = [
        'Parametre Adı',
        'Metot', 
        'İzo Oran',
        'Nozzle',
        '1. İmp',
        '2. İmp',
        '3. İmp',
        '4. İmp',
        'L/DAK',
        'T.HAC',
        'LOQ',
        'KK',
        '-3S',
        '-2S',
        '+2S',
        '+3S'
    ]
    for col in desired_order:
        if col not in df.columns:
            df[col] = ''
    df = df[desired_order]

    with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
        df.to_excel(tmp.name, index=False, sheet_name='Parametreler', engine='openpyxl')
        tmp.flush()
        
        return send_file(
            tmp.name,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='parametreler.xlsx'
        )





@app.route('/parametre/add', methods=['GET', 'POST'])
def add_parameter():
    if not session.get('logged_in') or not can_write(session.get('role')):
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        new_param = {
            'id': str(uuid4()),
            'Parametre Adı': request.form.get('Parametre Adı'),
            'Metot': request.form.get('Metot'),
            'İzo Oran': request.form.get('İzo Oran'),
            'LOQ': request.form.get('LOQ'),
            'KK': request.form.get('KK'),
            '1. İmp': request.form.get('1. İmp'),
            '2. İmp': request.form.get('2. İmp'),
            '3. İmp': request.form.get('3. İmp'),
            '4. İmp': request.form.get('4. İmp'),
            'L/DAK': request.form.get('L/DAK'),
            'Nozzle': request.form.get('Nozzle'),
            'T.HAC': request.form.get('T.HAC'),
            '-3S': request.form.get('-3S'),
            '-2S': request.form.get('-2S'),
            '+2S': request.form.get('+2S'),
            '+3S': request.form.get('+3S')
        }
        
        parameters = load_parameters()
        parameters.append(new_param)
        save_parameters(parameters)
        flash('Yeni parametre başarıyla eklendi.', 'success')
        return redirect(url_for('parametre'))

    return render_template('add_edit_parameter.html', username=session.get('username'))

@app.route('/parametre/edit/<parameter_id>', methods=['GET', 'POST'])
def edit_parameter(parameter_id):
    if not session.get('logged_in') or not can_edit(session.get('role')):
        return redirect(url_for('login'))

    parameters = load_parameters()
    param_to_edit = next((p for p in parameters if p['id'] == parameter_id), None)

    if not param_to_edit:
        flash('Parametre bulunamadı.', 'danger')
        return redirect(url_for('parametre'))

    if request.method == 'POST':
        param_to_edit['Parametre Adı'] = request.form.get('Parametre Adı')
        param_to_edit['Metot'] = request.form.get('Metot')
        param_to_edit['İzo Oran'] = request.form.get('İzo Oran')
        param_to_edit['LOQ'] = request.form.get('LOQ')
        param_to_edit['KK'] = request.form.get('KK')
        param_to_edit['1. İmp'] = request.form.get('1. İmp')
        param_to_edit['2. İmp'] = request.form.get('2. İmp')
        param_to_edit['3. İmp'] = request.form.get('3. İmp')
        param_to_edit['4. İmp'] = request.form.get('4. İmp')
        param_to_edit['L/DAK'] = request.form.get('L/DAK')
        param_to_edit['Nozzle'] = request.form.get('Nozzle')
        param_to_edit['T.HAC'] = request.form.get('T.HAC')
        param_to_edit['-3S'] = request.form.get('-3S')
        param_to_edit['-2S'] = request.form.get('-2S')
        param_to_edit['+2S'] = request.form.get('+2S')
        param_to_edit['+3S'] = request.form.get('+3S')
        
        save_parameters(parameters)
        flash('Parametre başarıyla güncellendi.', 'success')
        return redirect(url_for('parametre'))

    return render_template('add_edit_parameter.html', parameter=param_to_edit, username=session.get('username'))

@app.route('/parametre/delete/<parameter_id>', methods=['POST'])
def delete_parameter(parameter_id):
    if not session.get('logged_in') or not can_delete(session.get('role')):
        return redirect(url_for('login'))

    parameters = load_parameters()
    parameters_to_keep = [p for p in parameters if p['id'] != parameter_id]

    if len(parameters) == len(parameters_to_keep):
        flash('Silinecek parametre bulunamadı.', 'danger')
    else:
        save_parameters(parameters_to_keep)
        flash('Parametre başarıyla silindi.', 'success')
    
    return redirect(url_for('parametre'))



@app.route('/api/ilceler/<il_adi>')
def api_ilceler(il_adi):
    # Gelen il adına göre ilçeleri bul ve sadece ilçe adlarını döndür
    for il in CITIES_DATA:
        if il['il_adi'].upper() == il_adi.upper():
            ilceler = [ilce['ilce_adi'] for ilce in il['ilceler']]
            return jsonify(sorted(ilceler))
    return jsonify([])



@app.route('/olcum_olustur')
def olcum_olustur():
    if not session.get('logged_in') or not can_read(session.get('role')):
        return redirect(url_for('login'))
    measurements = load_measurements()
    # Ölçüm başlangıç ve bitiş tarihlerini formatla
    for m in measurements:
        m['olcum_baslangic_fmt'] = format_date_with_day(m.get('olcum_baslangic', ''))
        m['olcum_bitis_fmt'] = format_date_with_day(m.get('olcum_bitis', ''))
    # Sadece görevi "Saha" olan kullanıcıları filtrele
    current_users = load_users()
    saha_users = []
    for username, user_data in current_users.items():
        if user_data.get('gorev') == 'Saha':
            saha_users.append(username)
    user_list = sorted(saha_users)
    # Parametreleri yükle ve isimleri benzersiz yap
    parameters = load_parameters()
    unique_param_names = []
    seen = set()
    for p in parameters:
        name = p.get('Parametre Adı') or p.get('parametre_adi') or p.get('ad') or p.get('isim')
        if name and name not in seen:
            unique_param_names.append(name)
            seen.add(name)
    return render_template('olcum_olustur_simple.html', username=session.get('username'), measurements=measurements, users=user_list, unique_param_names=unique_param_names, CITIES_DATA=CITIES_DATA)

@app.route('/add_measurement', methods=['POST'])
def add_measurement():
    if not session.get('logged_in') or not can_write(session.get('role')):
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        print("=== ÖLÇÜM EKLEME DEBUG ===")
        print("Form verileri:", dict(request.form))
        
        # Çoklu personel seçimini al ve virgülle birleştir
        selected_personnel = request.form.getlist('olcumPersoneli')
        print("Seçilen personeller (getlist):", selected_personnel)
        
        personnel_str = ', '.join(selected_personnel) if selected_personnel else ''
        print("Personel string:", personnel_str)
        
        new_measurement = {
            'id': str(uuid4()),
            'firma_adi': request.form.get('firmaAdi'),
            'il': request.form.get('il'),
            'ilce': request.form.get('ilce'),
            'il_ilce': f"{request.form.get('il')} - {request.form.get('ilce')}" if request.form.get('il') and request.form.get('ilce') else '',
            'olcum_kodu': request.form.get('olcumKodu'),
            'baca_sayisi': request.form.get('bacaSayisi', ''),
            'olcum_baslangic': request.form.get('olcumBaslangic'),
            'olcum_bitis': request.form.get('olcumBitis'),
            'olcum_baslangic_fmt': format_date_with_day(request.form.get('olcumBaslangic')),
            'olcum_bitis_fmt': format_date_with_day(request.form.get('olcumBitis')),
            'yetkili': request.form.get('yetkili', ''),
            'telefon_no': request.form.get('telefonNo', ''),
            'olcumPersoneli': personnel_str,
            'parametre': request.form.get('parametre', ''),
            'tarih': request.form.get('olcumBaslangic', '')[:10] if request.form.get('olcumBaslangic') else '',  # Sadece tarih kısmı
            'durum': request.form.get('durum', 'Aktif')
        }
        
        measurements = load_measurements()
        measurements.append(new_measurement)
        save_measurements(measurements)
        
        # AJAX/fetch ile gelirse JSON döndür
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest' or request.is_json:
            return jsonify({'success': True, 'measurement_id': new_measurement['id']})
        
        flash('Yeni ölçüm başarıyla eklendi!', 'success')
        return redirect(url_for('olcum_olustur'))
    
    return redirect(url_for('olcum_olustur'))

@app.route('/save_baca_data', methods=['POST'])
def save_baca_data():
    if not session.get('logged_in') or not can_write(session.get('role')):
        return jsonify({'error': 'Unauthorized'}), 401
    
    try:
        data = request.get_json()
        olcum_id = data.get('olcum_id')
        baca_data = data.get('baca_data', [])
        
        # Saha verilerini kaydet (basitleştirilmiş)
        return jsonify({'success': True, 'message': 'Baca bilgileri kaydedildi'})
        
        return jsonify({'success': True, 'message': 'Baca bilgileri kaydedildi'})
        
    except Exception as e:
        print(f"Baca veri kaydetme hatası: {e}")
        return jsonify({'error': f'Kaydetme hatası: {str(e)}'}), 500

@app.route('/edit_measurement/<measurement_id>', methods=['POST'])
def edit_measurement(measurement_id):
    if not session.get('logged_in') or not can_edit(session.get('role')):
        return redirect(url_for('login'))
    measurements = load_measurements()
    measurement = next((m for m in measurements if m['id'] == measurement_id), None)
    if not measurement:
        flash('Ölçüm kaydı bulunamadı.', 'danger')
        return redirect(url_for('olcum_olustur'))
    
    # Çoklu personel seçimini al ve virgülle birleştir
    selected_personnel = request.form.getlist('olcumPersoneli')
    personnel_str = ', '.join(selected_personnel) if selected_personnel else ''
    
    # Formdan gelen verilerle güncelle
    measurement['firma_adi'] = request.form.get('firmaAdi')
    measurement['olcum_kodu'] = request.form.get('olcumKodu')
    measurement['baca_sayisi'] = request.form.get('bacaSayisi')
    measurement['olcum_baslangic'] = request.form.get('olcumBas')
    measurement['olcum_bitis'] = request.form.get('olcumBit')
    measurement['olcumPersoneli'] = personnel_str
    measurement['tarih'] = request.form.get('olcumBas', '')[:10] if request.form.get('olcumBas') else ''
    save_measurements(measurements)
    flash('Ölçüm başarıyla güncellendi!', 'success')
    return redirect(url_for('olcum_olustur'))

@app.route('/delete_measurement/<measurement_id>', methods=['POST'])
def delete_measurement(measurement_id):
    if not session.get('logged_in') or not can_delete(session.get('role')):
        return jsonify({'error': 'Unauthorized'}), 401
    
    try:
        measurements = load_measurements()
        new_measurements = [m for m in measurements if m['id'] != measurement_id]
        
        if len(measurements) == len(new_measurements):
            return jsonify({'error': 'Silinecek ölçüm kaydı bulunamadı'}), 404
        else:
            save_measurements(new_measurements)
            return jsonify({'success': True, 'message': 'Ölçüm başarıyla silindi'}), 200
            
    except Exception as e:
        return jsonify({'error': f'Silme işlemi sırasında hata: {str(e)}'}), 500

@app.route('/delete_selected_measurements', methods=['POST'])
def delete_selected_measurements():
    if not session.get('logged_in') or not can_delete(session.get('role')):
        return jsonify({'error': 'Unauthorized'}), 401
    
    try:
        data = request.get_json()
        selected_ids = data.get('ids', [])
        
        if not selected_ids:
            return jsonify({'error': 'Seçilen ölçüm bulunamadı'}), 400
        
        measurements = load_measurements()
        original_count = len(measurements)
        
        # Seçilen ID'leri sil
        new_measurements = [m for m in measurements if m['id'] not in selected_ids]
        
        if len(new_measurements) < original_count:
            save_measurements(new_measurements)
            deleted_count = original_count - len(new_measurements)
            return jsonify({'success': True, 'message': f'{deleted_count} ölçüm başarıyla silindi'}), 200
        else:
            return jsonify({'error': 'Silinecek ölçüm bulunamadı'}), 404
            
    except Exception as e:
        return jsonify({'error': f'Silme işlemi sırasında hata: {str(e)}'}), 500

@app.route('/get_measurement_details/<measurement_id>')
def get_measurement_details(measurement_id):
    try:
        if not session.get('logged_in') or not can_read(session.get('role')):
            return jsonify({'error': 'Unauthorized'}), 401
        
        measurements = load_measurements()
        measurement = next((m for m in measurements if m['id'] == measurement_id), None)
        
        if not measurement:
            return jsonify({'error': 'Ölçüm bulunamadı'}), 404
        
        # Tarihleri formatla
        measurement['olcum_baslangic_fmt'] = format_date_with_day(measurement.get('olcum_baslangic', ''))
        measurement['olcum_bitis_fmt'] = format_date_with_day(measurement.get('olcum_bitis', ''))
        

        
        # Baca bilgilerini al (basitleştirilmiş)
        baca_bilgileri = []
        
        # HTML içeriği oluştur
        html_content = f"""
        <div class="container-fluid">
            <div class="row">
                <div class="col-md-6">
                    <div class="card mb-3">
                        <div class="card-header bg-primary text-white">
                            <h6 class="mb-0"><i class="fas fa-info-circle me-2"></i>Ölçüm Bilgileri</h6>
                        </div>
                        <div class="card-body">
                            <table class="table table-sm">
                                <tr>
                                    <td><strong>Firma Adı:</strong></td>
                                    <td>{measurement.get('firma_adi', 'Belirtilmemiş')}</td>
                                </tr>
                                <tr>
                                    <td><strong>Ölçüm Kodu:</strong></td>
                                    <td>{measurement.get('olcum_kodu', 'Belirtilmemiş')}</td>
                                </tr>
                                <tr>
                                    <td><strong>Baca Sayısı:</strong></td>
                                    <td>{measurement.get('baca_sayisi', 'Belirtilmemiş')}</td>
                                </tr>
                                <tr>
                                    <td><strong>Ölçüm Başlangıç:</strong></td>
                                    <td>{measurement.get('olcum_baslangic_fmt', 'Belirtilmemiş')}</td>
                                </tr>
                                <tr>
                                    <td><strong>Ölçüm Bitiş:</strong></td>
                                    <td>{measurement.get('olcum_bitis_fmt', 'Belirtilmemiş')}</td>
                                </tr>
                                <tr>
                                    <td><strong>Personel:</strong></td>
                                    <td>{measurement.get('olcumPersoneli', 'Belirtilmemiş')}</td>
                                </tr>
                                <tr>
                                    <td><strong>Durum:</strong></td>
                                    <td><span class="badge bg-success">{measurement.get('durum', 'Aktif')}</span></td>
                                </tr>
                            </table>
                        </div>
                    </div>
                </div>
                <div class="col-md-6">
                    <div class="card mb-3">
                        <div class="card-header bg-info text-white">
                            <h6 class="mb-0"><i class="fas fa-building me-2"></i>Firma Bilgileri</h6>
                        </div>
                        <div class="card-body">
        """
        html_content += """
                            <div class="text-center text-muted">
                                <i class="fas fa-info-circle fa-2x mb-2"></i>
                                <p>Firma bilgileri artık kullanılmıyor</p>
                            </div>
        """
        html_content += """
                        </div>
                    </div>
                </div>
            </div>
        """
        
        # Baca bilgileri varsa ekle
        if baca_bilgileri:
            html_content += """
            <div class="row">
                <div class="col-12">
                    <div class="card">
                        <div class="card-header bg-success text-white">
                            <h6 class="mb-0"><i class="fas fa-chimney me-2"></i>Baca ve Parametre Bilgileri</h6>
                        </div>
                        <div class="card-body">
                            <div class="table-responsive">
                                <table class="table table-sm table-bordered">
                                    <thead class="table-light">
                                        <tr>
                                            <th style="width: 50px;">Sıra</th>
                                            <th>Baca Adı</th>
                                            <th>Parametreler</th>
                                        </tr>
                                    </thead>
                                    <tbody>
            """
            
            for index, baca in enumerate(baca_bilgileri, 1):
                parametreler_html = ""
                if baca.get('parametreler'):
                    # Parametre sayılarını hesapla
                    param_counts = {}
                    for param in baca['parametreler']:
                        param_counts[param] = param_counts.get(param, 0) + 1
                    
                    # Her parametreyi sayısıyla birlikte göster
                    for param, count in param_counts.items():
                        parametreler_html += f'<span class="badge bg-primary me-1" style="position: relative;">{param}<span style="position: absolute; top: -8px; right: -8px; background: #dc3545; color: white; border-radius: 50%; width: 18px; height: 18px; font-size: 0.7rem; display: flex; align-items: center; justify-content: center; font-weight: bold;">{count}</span></span>'
                else:
                    parametreler_html = '<span class="text-muted">Parametre seçilmemiş</span>'
                
                html_content += f"""
                                        <tr>
                                            <td class="text-center">
                                                <span class="badge bg-secondary" style="min-width: 25px; font-size: 0.85rem; font-weight: bold; padding: 2px 6px;">{index}</span>
                                            </td>
                                            <td><strong>{baca.get('baca_adi', 'Belirtilmemiş')}</strong></td>
                                            <td>{parametreler_html}</td>
                                        </tr>
                """
            
            html_content += """
                                    </tbody>
                                </table>
                            </div>
                        </div>
                    </div>
                </div>
            </div>
            """
        else:
            html_content += """
            <div class="row">
                <div class="col-12">
                    <div class="card">
                        <div class="card-header bg-warning text-dark">
                            <h6 class="mb-0"><i class="fas fa-exclamation-triangle me-2"></i>Baca Bilgileri</h6>
                        </div>
                        <div class="card-body">
                            <div class="text-center text-muted">
                                <i class="fas fa-info-circle fa-2x mb-2"></i>
                                <p>Bu ölçüm için baca ve parametre kaydı bulunamadı.</p>
                            </div>
                        </div>
                    </div>
                </div>
            </div>
            """
        
        html_content += """
        </div>
        """
        
        return jsonify({
            'success': True,
            'html': html_content
        })
    except Exception as e:
        print(f"Hata: {e}")
        return jsonify({'error': f'Server hatası: {str(e)}'}), 500

@app.route('/export_measurement/<measurement_id>')
def export_measurement(measurement_id):
    if 'username' not in session:
        return redirect(url_for('login'))
    
    measurements = load_measurements()
    measurement = next((m for m in measurements if m['id'] == measurement_id), None)
    
    if not measurement:
        flash('Ölçüm bulunamadı.', 'danger')
        return redirect(url_for('olcum_olustur'))
    
    # Ölçüm verilerini DataFrame'e çevir
    data = {
        'Firma Adı': [measurement['firma_adi']],
        'Ölçüm Kodu': [measurement['olcum_kodu']],
        'Baca Sayısı': [measurement['baca_sayisi']],
        'Ölçüm Başlangıç': [measurement['olcum_baslangic']],
        'Ölçüm Bitiş': [measurement['olcum_bitis']],
        'Personel': [measurement['olcumPersoneli']],
        'Durum': [measurement['durum']],
        'Tarih': [measurement['tarih']]
    }
    
    df = pd.DataFrame(data)
    
    import tempfile
    import os
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
        df.to_excel(tmp.name, index=False, sheet_name='Ölçüm Verileri', engine='openpyxl')
        tmp.flush()
        
        filename = f"olcum_{measurement['olcum_kodu']}_{measurement['firma_adi']}.xlsx"
        return send_file(
            tmp.name,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=filename
        )

@app.route('/export_all_measurements')
def export_all_measurements():
    if 'username' not in session:
        return redirect(url_for('login'))
    
    measurements = load_measurements()
    if not measurements:
        flash('Dışarı aktarılacak ölçüm bulunmuyor.', 'warning')
        return redirect(url_for('olcum_olustur'))
    
    # Ölçüm verilerini DataFrame'e çevir
    data_list = []
    for m in measurements:
        data_list.append({
            'Firma Adı': m['firma_adi'],
            'Ölçüm Kodu': m['olcum_kodu'],
            'Baca Sayısı': m['baca_sayisi'],
            'Ölçüm Başlangıç': m['olcum_baslangic'],
            'Ölçüm Bitiş': m['olcum_bitis'],
            'Personel': m['olcumPersoneli'],
            'Durum': m['durum'],
            'Tarih': m['tarih']
        })
    
    df = pd.DataFrame(data_list)
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
        df.to_excel(tmp.name, index=False, sheet_name='Tüm Ölçümler', engine='openpyxl')
        tmp.flush()
        
        return send_file(
            tmp.name,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='tum_olcumler.xlsx'
        )

@app.route('/import_measurement_data', methods=['POST'])
def import_measurement_data():
    if 'username' not in session:
        return redirect(url_for('login'))
    
    if 'excel_file' not in request.files:
        flash('Dosya seçilmedi.', 'danger')
        return redirect(url_for('olcum_olustur'))
    
    file = request.files['excel_file']
    if file.filename == '':
        flash('Dosya seçilmedi.', 'danger')
        return redirect(url_for('olcum_olustur'))
    
    if file and file.filename and file.filename.endswith(('.xlsx', '.xls')):
        try:
            df = pd.read_excel(file)
            measurements = load_measurements()
            
            # Excel'den gelen verileri ölçümlere uygula
            for index, row in df.iterrows():
                # Ölçüm kodu ile eşleşen kaydı bul
                matching_measurement = next((m for m in measurements if m['olcum_kodu'] == row.get('Ölçüm Kodu', '')), None)
                
                if matching_measurement:
                    # Verileri güncelle
                    matching_measurement['firma_adi'] = row.get('Firma Adı', matching_measurement['firma_adi'])
                    matching_measurement['baca_sayisi'] = str(row.get('Baca Sayısı', matching_measurement['baca_sayisi']))
                    matching_measurement['olcum_baslangic'] = str(row.get('Ölçüm Başlangıç', matching_measurement['olcum_baslangic']))
                    matching_measurement['olcum_bitis'] = str(row.get('Ölçüm Bitiş', matching_measurement['olcum_bitis']))
                    matching_measurement['olcumPersoneli'] = str(row.get('Personel', matching_measurement['olcumPersoneli']))
                    matching_measurement['durum'] = row.get('Durum', matching_measurement['durum'])
                    matching_measurement['tarih'] = str(row.get('Tarih', matching_measurement['tarih']))
            
            save_measurements(measurements)
            flash('Ölçüm verileri başarıyla güncellendi.', 'success')
            
        except Exception as e:
            flash(f'Dosya okunurken hata oluştu: {e}', 'danger')
    else:
        flash('Geçersiz dosya formatı. Lütfen .xlsx veya .xls uzantılı dosya yükleyin.', 'danger')
    
    return redirect(url_for('olcum_olustur'))

@app.route('/export_selected_measurements')
def export_selected_measurements():
    if 'username' not in session:
        return redirect(url_for('login'))
    
    ids_param = request.args.get('ids', '')
    if not ids_param:
        flash('Seçilen ölçüm bulunamadı.', 'warning')
        return redirect(url_for('olcum_olustur'))
    
    selected_ids = ids_param.split(',')
    measurements = load_measurements()
    selected_measurements = [m for m in measurements if m['id'] in selected_ids]
    
    if not selected_measurements:
        flash('Seçilen ölçüm bulunamadı.', 'warning')
        return redirect(url_for('olcum_olustur'))
    
    # Seçilen ölçümleri DataFrame'e çevir
    data_list = []
    for m in selected_measurements:
        data_list.append({
            'Firma Adı': m['firma_adi'],
            'Ölçüm Kodu': m['olcum_kodu'],
            'Baca Sayısı': m['baca_sayisi'],
            'Ölçüm Başlangıç': m['olcum_baslangic'],
            'Ölçüm Bitiş': m['olcum_bitis'],
            'Personel': m['olcumPersoneli'],
            'Durum': m['durum'],
            'Tarih': m['tarih']
        })
    
    df = pd.DataFrame(data_list)
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
        df.to_excel(tmp.name, index=False, sheet_name='Seçilen Ölçümler', engine='openpyxl')
        tmp.flush()
        
        return send_file(
            tmp.name,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'secilen_olcumler_{len(selected_measurements)}_kayit.xlsx'
        )



@app.route('/export_selected_parameters')
def export_selected_parameters():
    if 'username' not in session:
        return redirect(url_for('login'))
    
    ids_param = request.args.get('ids', '')
    selected_ids = ids_param.split(',') if ids_param else []
    parameters = load_parameters()
    selected_parameters = [p for p in parameters if p['id'] in selected_ids]
    
    columns = [
        'Parametre Adı', 'Metot', 'İzo Oran', 'Nozzle', '1. İmp', '2. İmp', '3. İmp', '4. İmp', 'L/DAK', 'T.HAC', 'LOQ', 'KK', '-3S', '-2S', '+2S', '+3S'
    ]
    # Her parametrede başlıkları normalize et ve Lab LOQ'yu LOQ'ya kopyala
    for p in selected_parameters:
        for col in columns:
            if col not in p:
                p[col] = ''
        # Eski veri desteği: 'Lab LOQ' varsa 'LOQ'ya kopyala
        if not p['LOQ'] and ('Lab LOQ' in p and p['Lab LOQ']):
            p['LOQ'] = p['Lab LOQ']
        # Yanlış anahtarları düzelt (ör: nozzle, NOZZLE, Nozzle vs.)
        for key in list(p.keys()):
            if key.lower().replace(' ', '') == 'nozzle':
                p['Nozzle'] = p[key]
            if key.lower().replace(' ', '') == 't.hac':
                p['T.HAC'] = p[key]
    if not selected_parameters:
        df = pd.DataFrame()
    else:
        df = pd.DataFrame(selected_parameters)
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
        df.to_excel(tmp.name, index=False, sheet_name='Seçilen Parametreler', engine='openpyxl')
        tmp.flush()
        
        return send_file(
            tmp.name,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'secilen_parametreler_{len(selected_parameters)}_kayit.xlsx'
        )
# def delete_selected_users(): - Admin sistemi silindi
    if 'username' not in session:
        return jsonify({'success': False, 'error': 'Unauthorized'}), 401
    
    data = request.get_json()
    usernames = data.get('usernames', []) if data else []
    
    if not usernames:
        return jsonify({'success': False, 'error': 'No usernames provided'}), 400
    
    current_users = load_users()
    deleted_count = 0
    
    for username in usernames:
        if username != 'admin' and username in current_users:
            del current_users[username]
            deleted_count += 1
    
    save_users(current_users)
    
    if deleted_count > 0:
        flash(f'{deleted_count} kullanıcı başarıyla silindi!', 'success')
    
    return jsonify({'success': True, 'deleted_count': deleted_count})

# @app.route('/export_selected_users') - Admin sistemi silindi
# def export_selected_users(): - Admin sistemi silindi
    if 'username' not in session:
        return redirect(url_for('login'))
    
    usernames_param = request.args.get('usernames', '')
    if not usernames_param:
        flash('Seçilen kullanıcı bulunamadı.', 'warning')
        return redirect(url_for('admin'))
    
    selected_usernames = usernames_param.split(',')
    users_data = load_users()
    selected_users = []
    
    for username in selected_usernames:
        if username in users_data:
            selected_users.append({
                'Kullanıcı Adı': username,
                'Rol': users_data[username]['role']
            })
    
    if not selected_users:
        flash('Seçilen kullanıcı bulunamadı.', 'warning')
        return redirect(url_for('admin'))
    
    # Seçilen kullanıcıları DataFrame'e çevir
    df = pd.DataFrame(selected_users)
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
        df.to_excel(tmp.name, index=False, sheet_name='Seçilen Kullanıcılar', engine='openpyxl')
        tmp.flush()
        
        return send_file(
            tmp.name,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'secilen_kullanicilar_{len(selected_users)}_kayit.xlsx'
        )

# @app.route('/export_selected_emissions') - Emission sistemi silindi
# def export_selected_emissions(): - Emission sistemi silindi
    if 'username' not in session:
        return redirect(url_for('login'))
    
    ids_param = request.args.get('ids', '')
    if not ids_param:
        flash('Seçilen emisyon bulunamadı.', 'warning')
        return redirect(url_for('index'))
    
    selected_ids = ids_param.split(',')
    emissions = load_emissions()
    selected_emissions = [e for e in emissions if e['id'] in selected_ids]
    
    if not selected_emissions:
        flash('Seçilen emisyon bulunamadı.', 'warning')
        return redirect(url_for('index'))
    
    # Seçilen emisyonları DataFrame'e çevir
    data_list = []
    for e in selected_emissions:
        data_list.append({
            'Tesis Adı': e['tesis_adi'],
            'Tarih': e['tarih'],
            'Parametre': e['parametre'],
            'Sonuç': e['sonuc'],
            'Birim': e['birim']
        })
    
    df = pd.DataFrame(data_list)
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
        df.to_excel(tmp.name, index=False, sheet_name='Seçilen Emisyonlar', engine='openpyxl')
        tmp.flush()
        
        return send_file(
            tmp.name,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'secilen_emisyonlar_{len(selected_emissions)}_kayit.xlsx'
        )

@app.route('/delete_selected_parameters', methods=['POST'])
def delete_selected_parameters():
    if 'username' not in session:
        return jsonify({'success': False, 'error': 'Unauthorized'}), 401
    data = request.get_json()
    ids = data.get('ids', []) if data else []
    if not ids:
        return jsonify({'success': False, 'error': 'No IDs provided'}), 400
    parameters = load_parameters()
    print('Silinecek parametre idleri:', ids)
    new_parameters = [p for p in parameters if str(p.get('id')) not in ids and p.get('id') not in ids]
    save_parameters(new_parameters)
    return jsonify({'success': True})

@app.route('/add_selected_parameters', methods=['POST'])
def add_selected_parameters():
    if 'username' not in session:
        return jsonify({'success': False, 'error': 'Unauthorized'}), 401
    
    data = request.get_json()
    ids = data.get('ids', []) if data else []
    
    if not ids:
        return jsonify({'success': False, 'error': 'No IDs provided'}), 400
    
    try:
        parameters = load_parameters()
        selected_parameters = [p for p in parameters if str(p.get('id')) in ids or p.get('id') in ids]
        
        if not selected_parameters:
            return jsonify({'success': False, 'error': 'Seçilen parametreler bulunamadı'}), 404
        
        # Seçilen parametreleri kopyala ve yeni ID'ler ver
        new_parameters = []
        for param in selected_parameters:
            new_param = param.copy()
            new_param['id'] = str(uuid4())  # Yeni benzersiz ID
            new_param['created_at'] = datetime.now().isoformat()
            new_parameters.append(new_param)
        
        # Yeni parametreleri mevcut listeye ekle
        parameters.extend(new_parameters)
        save_parameters(parameters)
        
        print(f'{len(new_parameters)} adet parametre başarıyla eklendi')
        return jsonify({'success': True, 'message': f'{len(new_parameters)} adet parametre başarıyla eklendi'})
        
    except Exception as e:
        print(f'Parametre ekleme hatası: {e}')
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/update_parameter_field', methods=['POST'])
def update_parameter_field():
    if 'username' not in session:
        return jsonify({'success': False, 'error': 'Unauthorized'}), 401
    
    data = request.get_json()
    parameter_id = data.get('parameter_id')
    field = data.get('field')
    value = data.get('value')
    
    if not parameter_id or not field:
        return jsonify({'success': False, 'error': 'Missing required fields'}), 400
    
    parameters = load_parameters()
    parameter = next((p for p in parameters if str(p.get('id')) == str(parameter_id)), None)
    
    if not parameter:
        return jsonify({'success': False, 'error': 'Parameter not found'}), 404
    
    # Yeni sütunları güncelle
    parameter[field] = value
    save_parameters(parameters)
    
    return jsonify({'success': True})

@app.route('/firma_kayit')
def firma_kayit():
    if not session.get('logged_in') or not can_read(session.get('role')):
        return redirect(url_for('login'))
    
    # Firma kayıt verilerini yükle
    firma_kayitlar = load_firma_kayit()
    
    # İstediğiniz illeri en üste taşı
    oncelikli_iller = ['KOCAELİ', 'SAKARYA', 'DÜZCE', 'BOLU', 'İSTANBUL', 'BURSA', 'BİLECİK', 'KÜTAHYA']
    
    # İlleri öncelik sırasına göre düzenle
    sirali_cities = []
    
    # Önce öncelikli illeri ekle
    for il_adi in oncelikli_iller:
        for il in CITIES_DATA:
            if il['il_adi'] == il_adi:
                sirali_cities.append(il)
                break
    
    # Sonra diğer illeri ekle
    for il in CITIES_DATA:
        if il['il_adi'] not in oncelikli_iller:
            sirali_cities.append(il)
    
    return render_template('firma_kayit.html', 
                         username=session.get('username'), 
                         role=session.get('role'),
                         cities_data=sirali_cities,
                         firma_kayitlar=firma_kayitlar)




@app.route('/api/asgari_fiyatlar')
def get_asgari_fiyatlar():
    """Asgari fiyat verilerini döndürür"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        current_year = datetime.now().year
        raw = load_asgari_fiyatlar()
        # UI uyumluluğu: {parametre, metot, fiyat} listesine dönüştür
        asgari_fiyatlar = []
        for item in raw:
            yillik = item.get('yillik', {}) or {}
            fiyat = yillik.get(str(current_year))
            if fiyat is None:
                # en yeni yıla fallback
                if yillik:
                    latest_year = sorted(yillik.keys())[-1]
                    fiyat = yillik.get(latest_year)
            asgari_fiyatlar.append({
                'parametre': item.get('parametre', ''),
                'metot': item.get('metot', ''),
                'fiyat': float(fiyat or 0)
            })
        return jsonify({'success': True, 'asgari_fiyatlar': asgari_fiyatlar, 'current_year': current_year})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/asgari_fiyatlar/add', methods=['POST'])
def add_asgari_fiyat():
    """Yeni asgari fiyat parametresi ekler"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        data = request.get_json()
        
        # Yeni parametre
        yeni_parametre = {
            'parametre': data.get('parametre', ''),
            'metot': data.get('metot', ''),
            'fiyat': float(data.get('fiyat', 0))
        }
        
        # Validasyon
        if not yeni_parametre['parametre']:
            return jsonify({'success': False, 'message': 'Parametre adı gerekli'})
        
        # Dosyadan oku, ekle ve kaydet
        data_file = load_asgari_fiyatlar()
        # Aynı parametre varsa güncelle
        updated = False
        for item in data_file:
            if item.get('parametre', '').upper() == yeni_parametre['parametre'].upper():
                item['metot'] = yeni_parametre.get('metot', item.get('metot', ''))
                yillik = item.get('yillik', {}) or {}
                yillik[str(datetime.now().year)] = yeni_parametre['fiyat']
                item['yillik'] = yillik
                updated = True
                break
        if not updated:
            data_file.append({
                'parametre': yeni_parametre['parametre'],
                'metot': yeni_parametre['metot'],
                'yillik': { str(datetime.now().year): yeni_parametre['fiyat'] }
            })
        save_asgari_fiyatlar(data_file)
        return jsonify({'success': True})
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/asgari_fiyatlar/delete', methods=['POST'])
def delete_asgari_fiyat():
    """Asgari fiyat parametresini siler"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        data = request.get_json()
        parametre_adi = data.get('parametre', '')
        
        if not parametre_adi:
            return jsonify({'success': False, 'message': 'Parametre adı gerekli'})
        
        # Sabit parametreler (silinemez)
        sabit_parametreler = ['TABAN FİYAT', 'YOL']
        
        if parametre_adi.upper() in [p.upper() for p in sabit_parametreler]:
            return jsonify({'success': False, 'message': 'Bu parametre silinemez'})
        
        # Gerçek uygulamada veritabanından silinir
        return jsonify({
            'success': True, 
            'message': 'Parametre başarıyla silindi'
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})



# Asgari fiyat tarifesi - yıllara göre
def load_asgari_fiyatlar():
    """Asgari fiyatları yıllara göre yükler"""
    try:
        with open('asgari_fiyatlar.json', 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"Asgari fiyatlar yüklenirken hata: {e}")
        return []

def get_parametre_fiyati(parametre_adi, yil):
    """Belirli bir yıl için parametre fiyatını döndürür"""
    asgari_fiyatlar = load_asgari_fiyatlar()
    
    # Parametre adı eşleştirmesi
    parametre_eslestirme = {
        'TOZ': 'TOZ',
        'YG': 'YANMA GAZI', 
        'YANMA GAZI': 'YANMA GAZI',  # Eksik olan eşleştirme eklendi
        'VOC': 'VOC',
        'AĞIR METAL': 'AĞIR METAL',
        'TOC': 'TOC',
        'SO2': 'YANMA GAZI',  # SO2 YANMA GAZI içinde
        'NO': 'YANMA GAZI',   # NO YANMA GAZI içinde  
        'CO': 'YANMA GAZI',   # CO YANMA GAZI içinde
        'O2': 'YANMA GAZI',   # O2 YANMA GAZI içinde
        'PM10': 'PM10',
        'ÇT': 'ÇÖKEN TOZ',
        'Çt': 'ÇÖKEN TOZ',    # Küçük t ile yazım
        'NEM': 'NEM',
        'HF': 'HF',
        'HCL': 'HCL',
        'AMON.': 'AMONYAK',
        'AMONYAK': 'AMONYAK',  # Tam yazım
        'FORM.': 'FORMALDEHİT',
        'FORMALDEHİT': 'FORMALDEHİT',  # Tam yazım
        'CR6': 'CR+6',
        'Cr6': 'CR+6',        # Küçük c ile yazım
        'FOSF.A.': 'FOSFORİK ASİT',
        'FOSFORİK ASİT': 'FOSFORİK ASİT',  # Tam yazım
        'SÜLFÜRİK ASİT': 'SÜLFÜRİK ASİT',  # Eksik olan eşleştirme eklendi
        'SÜLF.A': 'SÜLFÜRİK ASİT',  # Alternatif yazım
        'HCN': 'HCN',
        'PAH': 'PAH',         # Tam yazım
        'DİOKSİN FURAN': 'DİOKSİN FURAN'  # Tam yazım
    }
    
    eslestirilen_parametre = parametre_eslestirme.get(parametre_adi, parametre_adi)
    print(f"DEBUG: Parametre '{parametre_adi}' -> '{eslestirilen_parametre}' için {yil} yılı fiyatı aranıyor")
    
    for fiyat_kaydi in asgari_fiyatlar:
        if fiyat_kaydi['parametre'] == eslestirilen_parametre:
            yillik_fiyatlar = fiyat_kaydi.get('yillik', {})
            fiyat = yillik_fiyatlar.get(str(yil), 0)
            print(f"DEBUG: Bulundu! {eslestirilen_parametre} -> {yil}: {fiyat} TL")
            return fiyat
    
    print(f"DEBUG: Parametre '{eslestirilen_parametre}' bulunamadı!")
    return 0

@app.route('/api/parametre-fiyatlari', methods=['GET'])
def get_parametre_fiyatlari():
    """Parametre fiyat tarifesini döndürür"""
    # 2025 yılı için varsayılan fiyatlar
    varsayilan_fiyatlar = {
        'TOZ': 6655,
        'YG': 3855,
        'VOC': 13171,
        'AĞIR METAL': 19000,
        'TOC': 8500,
        'SO2': 4500,
        'NO': 4200,
        'CO': 3800,
        'O2': 3500,
        'PM10': 7200,
        'ÇT': 5500,
        'NEM': 1430,
        'HF': 12000,
        'HCL': 9500,
        'AMON.': 6800,
        'FORM.': 7500,
        'CR6': 15000,
        'FOSF.A.': 11000,
        'HCN': 13000
    }
    
    return jsonify({
        'success': True,
        'fiyatlar': varsayilan_fiyatlar
    })

@app.route('/api/asgari-fiyatlar/<yil>', methods=['GET'])
def get_asgari_fiyatlar_yil(yil):
    """Belirli bir yıl için asgari fiyatları döndürür"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        yil = int(yil)
        asgari_fiyatlar = load_asgari_fiyatlar()
        
        # Yıllık fiyatları filtrele
        yillik_fiyatlar = {}
        for fiyat_kaydi in asgari_fiyatlar:
            parametre = fiyat_kaydi['parametre']
            yillik_fiyat = fiyat_kaydi.get('yillik', {}).get(str(yil), 0)
            yillik_fiyatlar[parametre] = yillik_fiyat
        
        return jsonify({
            'success': True,
            'yil': yil,
            'fiyatlar': yillik_fiyatlar
        })
        
    except ValueError:
        return jsonify({'success': False, 'message': 'Geçersiz yıl formatı'})
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/otomatik-fiyat-hesapla', methods=['POST'])
def otomatik_fiyat_hesapla():
    """Seçilen parametreler için otomatik fiyat hesaplar"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        data = request.get_json()
        secilen_parametreler = data.get('parametreler', [])
        yil = data.get('yil', 2025)  # Varsayılan 2025
        
        if not secilen_parametreler:
            return jsonify({'success': False, 'message': 'Seçilen parametre bulunamadı'})
        
        # Her parametre için fiyat hesapla
        hesaplanan_fiyatlar = []
        toplam_fiyat = 0
        
        for parametre in secilen_parametreler:
            parametre_adi = parametre.get('parametre', '')
            adet = parametre.get('adet', 1)
            
            # Yıllık fiyat tarifesinden fiyatı al
            birim_fiyat = get_parametre_fiyati(parametre_adi, yil)
            top_fiyat = birim_fiyat * adet
            
            hesaplanan_fiyatlar.append({
                'parametre': parametre_adi,
                'metot': parametre.get('metot', ''),
                'birimFiyat': birim_fiyat,
                'adet': adet,
                'topFiyat': top_fiyat
            })
            
            toplam_fiyat += top_fiyat
        
        return jsonify({
            'success': True,
            'parametreler': hesaplanan_fiyatlar,
            'toplam': toplam_fiyat,
            'yil': yil,
            'message': f'{len(secilen_parametreler)} parametre için {yil} yılı toplam {toplam_fiyat:,.2f} TL hesaplandı'
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/baca-bilgileri/<firma_adi>', methods=['GET'])
def get_baca_bilgileri(firma_adi):
    """Firma için baca bilgilerini ve parametrelerini döndürür"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        # Baca bilgilerini yükle
        baca_bilgileri = load_baca_bilgileri()
        
        # Firma adına göre filtrele
        firma_bacalar = [baca for baca in baca_bilgileri if baca.get('firma_adi', '').upper() == firma_adi.upper()]
        
        if not firma_bacalar:
            return jsonify({
                'success': False, 
                'message': f'{firma_adi} firması için baca bilgisi bulunamadı',
                'bacalar': []
            })
        
        # Her baca için parametreleri topla
        tum_parametreler = set()
        baca_detaylari = []
        
        for baca in firma_bacalar:
            baca_adi = baca.get('baca_adi', '')
            parametreler = []
            
            # Baca bilgilerinden parametreleri çıkar
            for key, value in baca.items():
                if key.startswith('parametre_') and value and value.strip():
                    parametre_adi = key.replace('parametre_', '').replace('_', ' ').upper()
                    parametreler.append(parametre_adi)
                    tum_parametreler.add(parametre_adi)
            
            baca_detaylari.append({
                'baca_adi': baca_adi,
                'parametreler': parametreler
            })
        
        return jsonify({
            'success': True,
            'firma_adi': firma_adi,
            'bacalar': baca_detaylari,
            'tum_parametreler': list(tum_parametreler),
            'message': f'{len(firma_bacalar)} baca bulundu, {len(tum_parametreler)} farklı parametre tespit edildi'
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})



@app.route('/api/firma_kayit/add', methods=['POST'])
def add_firma_kayit():
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        data = request.get_json()
        
        # Yeni firma kaydı oluştur
        yeni_firma = {
            'id': str(uuid.uuid4()),
            'firmaAdi': data.get('firmaAdi', ''),
            'adres': data.get('adres', ''),
            'il': data.get('il', ''),
            'ilce': data.get('ilce', ''),
            'vergiDairesi': data.get('vergiDairesi', ''),
            'vergiNo': data.get('vergiNo', ''),
            'yetkiliAdi': data.get('yetkiliAdi', ''),
            'yetkiliTel': data.get('yetkiliTel', ''),
            'yetkiliMail': data.get('yetkiliMail', ''),
            'danismanAdi': data.get('danismanAdi', ''),
            'danismanMail': data.get('danismanMail', ''),
            'danismanTel': data.get('danismanTel', ''),
            'kayitTarihi': datetime.now().strftime('%d.%m.%Y')
        }
        
        # Mevcut verileri yükle
        firma_kayitlar = load_firma_kayit()
        firma_kayitlar.append(yeni_firma)
        
        # Verileri kaydet
        if save_firma_kayit(firma_kayitlar):
            return jsonify({'success': True, 'message': 'Firma başarıyla kaydedildi', 'firma': yeni_firma})
        else:
            return jsonify({'success': False, 'message': 'Firma kaydedilirken hata oluştu'})
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/firma_kayit/update', methods=['POST'])
def update_firma_kayit():
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        data = request.get_json()
        firma_id = data.get('id')
        
        # Mevcut verileri yükle
        firma_kayitlar = load_firma_kayit()
        
        # Firma bul ve güncelle
        for firma in firma_kayitlar:
            if firma.get('id') == firma_id:
                firma.update({
                    'firmaAdi': data.get('firmaAdi', ''),
                    'adres': data.get('adres', ''),
                    'il': data.get('il', ''),
                    'ilce': data.get('ilce', ''),
                    'vergiDairesi': data.get('vergiDairesi', ''),
                    'vergiNo': data.get('vergiNo', ''),
                    'yetkiliAdi': data.get('yetkiliAdi', ''),
                    'yetkiliTel': data.get('yetkiliTel', ''),
                    'yetkiliMail': data.get('yetkiliMail', ''),
                    'danismanAdi': data.get('danismanAdi', ''),
                    'danismanMail': data.get('danismanMail', ''),
                    'danismanTel': data.get('danismanTel', '')
                })
                break
        
        # Verileri kaydet
        if save_firma_kayit(firma_kayitlar):
            return jsonify({'success': True, 'message': 'Firma başarıyla güncellendi'})
        else:
            return jsonify({'success': False, 'message': 'Firma güncellenirken hata oluştu'})
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/firma_kayit/delete', methods=['POST'])
def delete_firma_kayit():
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        data = request.get_json()
        firma_id = data.get('id')
        
        # Mevcut verileri yükle
        firma_kayitlar = load_firma_kayit()
        
        # Firma bul ve sil
        firma_kayitlar = [firma for firma in firma_kayitlar if firma.get('id') != firma_id]
        
        # Verileri kaydet
        if save_firma_kayit(firma_kayitlar):
            return jsonify({'success': True, 'message': 'Firma başarıyla silindi'})
        else:
            return jsonify({'success': False, 'message': 'Firma silinirken hata oluştu'})
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/firma_kayit/import', methods=['POST'])
def import_firma_kayit():
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'})
    
    try:
        data = request.get_json()
        imported_firmalar = data.get('firmalar', [])
        
        if not imported_firmalar:
            return jsonify({'success': False, 'message': 'İçe aktarılacak firma verisi bulunamadı'})
        
        # Mevcut verileri yükle
        firma_kayitlar = load_firma_kayit()
        
        # Yeni firmaları ekle
        for firma in imported_firmalar:
            # ID kontrolü - eğer aynı ID varsa yeni ID oluştur
            if any(existing_firma.get('id') == firma.get('id') for existing_firma in firma_kayitlar):
                firma['id'] = str(uuid.uuid4())
            firma_kayitlar.append(firma)
        
        # Verileri kaydet
        if save_firma_kayit(firma_kayitlar):
            return jsonify({
                'success': True, 
                'message': f'{len(imported_firmalar)} adet firma başarıyla içe aktarıldı',
                'imported_count': len(imported_firmalar)
            })
        else:
            return jsonify({'success': False, 'message': 'Firmalar içe aktarılırken hata oluştu'})
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/firma_olcum')
def firma_olcum():
    if not session.get('logged_in') or not can_read(session.get('role')):
        return redirect(url_for('login'))
    
    firma_olcumler = load_firma_olcum()
    
    # Tarihleri GG.AA.YY formatında formatla
    for olcum in firma_olcumler:
        # Tarihleri GG.AA.YY formatında formatla
        if 'baslangic_tarihi' in olcum and olcum['baslangic_tarihi']:
            olcum['baslangic_tarihi'] = format_date_with_day(olcum['baslangic_tarihi'])
        if 'bitis_tarihi' in olcum and olcum['bitis_tarihi']:
            olcum['bitis_tarihi'] = format_date_with_day(olcum['bitis_tarihi'])
    
    # Parametreleri yükle
    parameters = load_parameters()
    
    return render_template('firma_olcum.html', 
                         firma_olcumler=firma_olcumler, 
                         username=session.get('username'), 
                         role=session.get('role'),
                         parameters=parameters)

@app.route('/firma_olcum/add_step1', methods=['GET', 'POST'])
def add_firma_olcum_step1():
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        # Form verilerini al
        firma_adi = request.form.get('firma_adi', '').strip()
        olcum_kodu = request.form.get('olcum_kodu', '').strip()
        baslangic_tarihi = request.form.get('baslangic_tarihi', '')
        bitis_tarihi = request.form.get('bitis_tarihi', '')
        durum = request.form.get('durum', 'Aktif')
        secilen_personel = request.form.getlist('personel')
        
        # Zorunlu alanları kontrol et
        if not firma_adi or not olcum_kodu:
            flash('Firma adı ve ölçüm kodu zorunludur!', 'error')
            return render_template('add_firma_olcum_step1.html', username=session.get('username'), role=session.get('role'))
        
        # Aynı ölçüm kodu herhangi bir firmada daha önce kaydedilmiş mi kontrol et (global benzersizlik)
        firma_olcumler = load_firma_olcum()
        for olcum in firma_olcumler:
            if olcum.get('olcum_kodu') == olcum_kodu:
                flash('Bu ölçüm kodu başka bir firmada zaten kayıtlı. Lütfen farklı bir ölçüm kodu girin.', 'error')
                return render_template('add_firma_olcum_step1.html', username=session.get('username'), role=session.get('role'))
        
        # Session'a geçici veri kaydet
        session['temp_firma_olcum'] = {
            'firma_adi': firma_adi,
            'olcum_kodu': olcum_kodu,
            'baslangic_tarihi': baslangic_tarihi,
            'bitis_tarihi': bitis_tarihi,
            'durum': durum,
            'secilen_personel': secilen_personel
        }
        
        return redirect(url_for('add_firma_olcum_step2'))
    
    # Saha personeli listesini al (görevi "Saha" olan kullanıcılar)
    users = load_users()
    saha_personeli = []
    for username, user_data in users.items():
        if user_data.get('gorev') == 'Saha':
            saha_personeli.append({
                'username': username,
                'surname': user_data.get('surname', ''),
                'gorev': user_data.get('gorev', '')
            })
    
    # İstediğiniz illeri en üste taşı
    oncelikli_iller = ['KOCAELİ', 'SAKARYA', 'DÜZCE', 'BOLU', 'İSTANBUL', 'BURSA', 'BİLECİK', 'KÜTAHYA']
    
    # İlleri öncelik sırasına göre düzenle
    sirali_cities = []
    
    # Önce öncelikli illeri ekle
    for il_adi in oncelikli_iller:
        for il in CITIES_DATA:
            if il['il_adi'] == il_adi:
                sirali_cities.append(il)
                break
    
    # Sonra diğer illeri ekle
    for il in CITIES_DATA:
        if il['il_adi'] not in oncelikli_iller:
            sirali_cities.append(il)
    
    # Firma listesini yükle ve formatını düzenle
    firma_kayitlar = load_firma_kayit()
    
    # Firma verilerini template için uygun formata çevir
    formatted_firmalar = []
    for firma in firma_kayitlar:
        formatted_firma = {
            'firma_adi': firma.get('firmaAdi', ''),
            'yetkili_adi': firma.get('yetkiliAdi', ''),
            'yetkili_tel': firma.get('yetkiliTel', ''),
            'il': firma.get('il', ''),
            'ilce': firma.get('ilce', ''),
            'adres': firma.get('adres', ''),
            'vergi_dairesi': firma.get('vergiDairesi', ''),
            'vergi_no': firma.get('vergiNo', ''),
            'yetkili_mail': firma.get('yetkiliMail', ''),
            'danisman_adi': firma.get('danismanAdi', ''),
            'danisman_mail': firma.get('danismanMail', ''),
            'danisman_tel': firma.get('danismanTel', ''),
            'kayit_tarihi': firma.get('kayitTarihi', '')
        }
        formatted_firmalar.append(formatted_firma)
    
    return render_template('add_firma_olcum_step1.html', 
                         username=session.get('username'), 
                         role=session.get('role'),
                         saha_personeli=saha_personeli,
                         cities=sirali_cities,
                         firmalar=formatted_firmalar)
@app.route('/firma_olcum/add_step2', methods=['GET', 'POST'])
def add_firma_olcum_step2():
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    
    # Session'dan geçici veriyi al
    temp_data = session.get('temp_firma_olcum')
    if not temp_data:
        return redirect(url_for('add_firma_olcum_step1'))
    
    if request.method == 'POST':
        try:
            # 2. aşama verilerini al
            baca_sayisi = request.form.get('baca_sayisi', '')
            notlar = request.form.get('notlar', '')
            baca_parametreleri_json = request.form.get('baca_parametreleri', '{}')
            
            # Baca sayısı kontrolü
            if not baca_sayisi or int(baca_sayisi) < 1:
                raise ValueError('Baca sayısı en az 1 olmalıdır')
            
            # JSON'dan baca parametrelerini parse et
            try:
                baca_parametreleri = json.loads(baca_parametreleri_json)
            except json.JSONDecodeError:
                baca_parametreleri = {}
            
            # Yeni kayıt oluştur
            yeni_kayit = {
                'id': str(uuid4()),
                'firma_adi': temp_data['firma_adi'],
                'olcum_kodu': temp_data['olcum_kodu'],
                'baslangic_tarihi': temp_data['baslangic_tarihi'],
                'bitis_tarihi': temp_data['bitis_tarihi'],
                'durum': temp_data['durum'],
                'personel': temp_data['secilen_personel'],
                'baca_sayisi': baca_sayisi,
                'baca_parametreleri': baca_parametreleri,
                'notlar': notlar,
                'olusturma_tarihi': datetime.now().isoformat()
            }
            
            # Veriyi kaydet
            firma_olcumler = load_firma_olcum()
            firma_olcumler.append(yeni_kayit)
            if not save_firma_olcum(firma_olcumler):
                raise Exception('Veriler kaydedilemedi')
            
            # Baca bilgilerini baca_bilgileri.json dosyasına kaydet
            try:
                # Mevcut baca bilgilerini yükle
                try:
                    with open('baca_bilgileri.json', 'r', encoding='utf-8') as f:
                        baca_bilgileri = json.load(f)
                except (FileNotFoundError, json.JSONDecodeError):
                    baca_bilgileri = []
                
                # Her baca için parametreleri ekle
                # baca_parametreleri yapısı: {"baca_adi": ["param1", "param2"]} şeklinde
                for baca_adi, parametreler in baca_parametreleri.items():
                    # parametreler bir liste olmalı
                    if not isinstance(parametreler, list):
                        parametreler = [parametreler]
                    
                    # Her parametre için kayıt oluştur
                    for parametre in parametreler:
                        baca_kayit = {
                            'firma_adi': temp_data['firma_adi'],
                            'olcum_kodu': temp_data['olcum_kodu'],
                            'baca_no': baca_adi,  # baca_adi'yi baca_no olarak kullan
                            'baca_adi': baca_adi,
                            'parametre': parametre,
                            'personel_adi': ', '.join(temp_data.get('secilen_personel', [])),
                            'kayit_tarihi': datetime.now().isoformat()
                        }
                        baca_bilgileri.append(baca_kayit)
                
                # Dosyaya kaydet
                with open('baca_bilgileri.json', 'w', encoding='utf-8') as f:
                    json.dump(baca_bilgileri, f, ensure_ascii=False, indent=2)
                    
            except Exception as e:
                print(f"Baca bilgileri kaydedilirken hata: {e}")
                # Hata olsa da devam et
            
            # Session'dan geçici veriyi temizle
            session.pop('temp_firma_olcum', None)
            
            # AJAX isteği ise JSON döndür
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': True, 'message': 'Firma ölçüm kaydı başarıyla oluşturuldu!'})
            
            flash('Firma ölçüm kaydı başarıyla oluşturuldu!', 'success')
            return redirect(url_for('firma_olcum'))
            
        except Exception as e:
            print(f"Kayıt hatası: {e}")
            import traceback
            traceback.print_exc()
            
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': False, 'error': f'Kayıt hatası: {str(e)}'})
            
            flash(f'Kayıt sırasında hata oluştu: {str(e)}', 'error')
            return redirect(url_for('add_firma_olcum_step2'))
    
    # Parametreleri yükle
    parameters = load_parameters()
    
    return render_template('add_firma_olcum_step2.html', 
                         username=session.get('username'), 
                         role=session.get('role'),
                         temp_data=temp_data,
                         parameters=parameters)

@app.route('/firma_olcum/detail/<olcum_id>')
def firma_olcum_detail(olcum_id):
    if not session.get('logged_in') or not can_read(session.get('role')):
        return redirect(url_for('login'))
    
    firma_olcumler = load_firma_olcum()
    olcum = next((o for o in firma_olcumler if o['id'] == olcum_id), None)
    
    if not olcum:
        flash('Ölçüm kaydı bulunamadı!', 'error')
        return redirect(url_for('firma_olcum'))
    
    # Parametreleri yükle
    parameters = load_parameters()
    
    # Sadece seçili parametrelerin benzersiz isimlerini al
    selected_param_names = set()
    if olcum.get('baca_parametreleri'):
        for baca_adi, parametreler in olcum['baca_parametreleri'].items():
            for param_name in parametreler:
                # Eğer parametre tam format ise (Parametre Adı (Metot)), sadece parametre adını al
                if ' (' in param_name and param_name.endswith(')'):
                    param_adi = param_name.split(' (')[0]
                    selected_param_names.add(param_adi)
                else:
                    # Kısa isim ise, olduğu gibi kullan
                    selected_param_names.add(param_name)
    
    # Benzersiz parametre listesi oluştur (sadece seçili parametreler)
    unique_parameters = []
    for param_name in sorted(selected_param_names):
        # Her parametre için kaç bacada seçildiğini hesapla
        param_count = 0
        if olcum.get('baca_parametreleri'):
            for baca_adi, parametreler in olcum['baca_parametreleri'].items():
                if param_name in parametreler:
                    param_count += 1
        
        unique_parameters.append({
            'name': param_name,
            'full_name': param_name,
            'count': param_count
        })
    

    
    return render_template('firma_olcum_detail.html', 
                         username=session.get('username'), 
                         role=session.get('role'),
                         olcum=olcum,
                         parameters=unique_parameters)

@app.route('/firma_olcum/delete/<olcum_id>', methods=['POST'])
def delete_firma_olcum(olcum_id):
    if not session.get('logged_in') or not can_delete(session.get('role')):
        return jsonify({'success': False, 'error': 'Yetkisiz erişim'}), 401
    
    try:
        firma_olcumler = load_firma_olcum()
        olcum = next((o for o in firma_olcumler if o['id'] == olcum_id), None)
        
        if not olcum:
            return jsonify({'success': False, 'error': 'Ölçüm kaydı bulunamadı'}), 404
        
        # Kaydı sil
        firma_olcumler = [o for o in firma_olcumler if o['id'] != olcum_id]
        
        if save_firma_olcum(firma_olcumler):
            return jsonify({'success': True, 'message': 'Ölçüm kaydı başarıyla silindi'})
        else:
            return jsonify({'success': False, 'error': 'Kayıt silinirken hata oluştu'}), 500
            
    except Exception as e:
        print(f"Silme hatası: {e}")
        return jsonify({'success': False, 'error': f'Silme hatası: {str(e)}'}), 500

@app.route('/firma_olcum/delete_selected', methods=['POST'])
def delete_selected_firma_olcum():
    if not session.get('logged_in') or not can_delete(session.get('role')):
        return jsonify({'success': False, 'error': 'Yetkisiz erişim'}), 401
    
    try:
        data = request.get_json()
        print(f"Gelen veri: {data}")  # Debug için log
        ids = data.get('ids', []) if data else []
        print(f"Silinecek ID'ler: {ids}")  # Debug için log
        
        if not ids:
            return jsonify({'success': False, 'error': 'Silinecek kayıt seçilmedi'}), 400
        
        firma_olcumler = load_firma_olcum()
        original_count = len(firma_olcumler)
        print(f"Orijinal kayıt sayısı: {original_count}")  # Debug için log
        
        # Seçilen kayıtları sil
        firma_olcumler = [o for o in firma_olcumler if o['id'] not in ids]
        print(f"Silme sonrası kayıt sayısı: {len(firma_olcumler)}")  # Debug için log
        
        if save_firma_olcum(firma_olcumler):
            deleted_count = original_count - len(firma_olcumler)
            print(f"Silinen kayıt sayısı: {deleted_count}")  # Debug için log
            return jsonify({'success': True, 'message': f'{deleted_count} kayıt başarıyla silindi'})
        else:
            return jsonify({'success': False, 'error': 'Kayıtlar silinirken hata oluştu'}), 500
            
    except Exception as e:
        print(f"Toplu silme hatası: {e}")
        return jsonify({'success': False, 'error': f'Silme hatası: {str(e)}'}), 500

@app.route('/firma_olcum/edit/<olcum_id>', methods=['GET', 'POST'])
def edit_firma_olcum(olcum_id):
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    
    # Firma ölçüm verilerini yükle
    firma_olcumler = load_firma_olcum()
    olcum = next((o for o in firma_olcumler if o['id'] == olcum_id), None)
    
    if not olcum:
        flash('Ölçüm kaydı bulunamadı!', 'error')
        return redirect(url_for('firma_olcum'))
    
    if request.method == 'POST':
        try:
            # Form verilerini al
            firma_adi = request.form.get('firma_adi', '').strip()
            olcum_kodu = request.form.get('olcum_kodu', '').strip()
            baslangic_tarihi = request.form.get('baslangic_tarihi', '')
            bitis_tarihi = request.form.get('bitis_tarihi', '')
            durum = request.form.get('durum', 'Aktif')
            secilen_personel = request.form.getlist('personel')
            baca_sayisi = request.form.get('baca_sayisi', '')
            notlar = request.form.get('notlar', '')
            baca_parametreleri_json = request.form.get('baca_parametreleri', '{}')
            
            # Zorunlu alanları kontrol et
            if not firma_adi or not olcum_kodu:
                flash('Firma adı ve ölçüm kodu zorunludur!', 'error')
                return render_template('edit_firma_olcum.html', 
                                     username=session.get('username'), 
                                     role=session.get('role'),
                                     olcum=olcum,
                                     cities=CITIES_DATA)
            
            # Aynı ölçüm kodu herhangi bir firmada başka bir kayıt için kullanılıyor mu kontrol et (global benzersizlik)
            for diger in firma_olcumler:
                if diger.get('id') != olcum_id and diger.get('olcum_kodu') == olcum_kodu:
                    flash('Bu ölçüm kodu başka bir firmada zaten kayıtlı. Lütfen farklı bir ölçüm kodu girin.', 'error')
                    return render_template('edit_firma_olcum.html', 
                                         username=session.get('username'), 
                                         role=session.get('role'),
                                         olcum=olcum,
                                         cities=CITIES_DATA)
            
            # JSON'dan baca parametrelerini parse et
            try:
                baca_parametreleri = json.loads(baca_parametreleri_json)
            except json.JSONDecodeError:
                baca_parametreleri = {}
            
            # Kaydı güncelle
            olcum['firma_adi'] = firma_adi
            olcum['olcum_kodu'] = olcum_kodu
            olcum['baslangic_tarihi'] = baslangic_tarihi
            olcum['bitis_tarihi'] = bitis_tarihi
            olcum['durum'] = durum
            olcum['personel'] = secilen_personel
            olcum['baca_sayisi'] = baca_sayisi
            olcum['baca_parametreleri'] = baca_parametreleri
            olcum['notlar'] = notlar
            
            # Veriyi kaydet
            if save_firma_olcum(firma_olcumler):
                flash('Firma ölçüm kaydı başarıyla güncellendi!', 'success')
                return redirect(url_for('firma_olcum'))
            else:
                flash('Kayıt güncellenirken hata oluştu!', 'error')
                
        except Exception as e:
            print(f"Güncelleme hatası: {e}")
            flash(f'Güncelleme sırasında hata oluştu: {str(e)}', 'error')
    
    # Saha personeli listesini al
    users = load_users()
    saha_personeli = []
    for username, user_data in users.items():
        if user_data.get('gorev') == 'Saha':
            saha_personeli.append({
                'username': username,
                'surname': user_data.get('surname', ''),
                'gorev': user_data.get('gorev', '')
            })
    
    # İstediğiniz illeri en üste taşı
    oncelikli_iller = ['KOCAELİ', 'SAKARYA', 'DÜZCE', 'BOLU', 'İSTANBUL', 'BURSA', 'BİLECİK', 'KÜTAHYA']
    
    # İlleri öncelik sırasına göre düzenle
    sirali_cities = []
    
    # Önce öncelikli illeri ekle
    for il_adi in oncelikli_iller:
        for il in CITIES_DATA:
            if il['il_adi'] == il_adi:
                sirali_cities.append(il)
                break
    
    # Sonra diğer illeri ekle
    for il in CITIES_DATA:
        if il['il_adi'] not in oncelikli_iller:
            sirali_cities.append(il)
    
    # Parametreleri yükle
    parameters = load_parameters()
    
    return render_template('edit_firma_olcum.html', 
                         username=session.get('username'), 
                         role=session.get('role'),
                         olcum=olcum,
                         saha_personeli=saha_personeli,
                         cities=sirali_cities,
                         parameters=parameters)

@app.route('/saha_olc')
def saha_olc():
    """Saha ölçüm sayfası"""
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    
    # Özet tablo için firma/ölçüm/baca bazında gruplanmış kayıtlar
    saha_olcumler = get_saha_olcum_summary()
    response = make_response(render_template('saha_olc.html', 
                         username=session.get('username'), 
                         role=session.get('role'),
                         saha_olcumler=saha_olcumler))
    
    # Cache'i tamamen devre dışı bırak - Her zaman yeni dosya yüklesin
    response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate, max-age=0'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    
    return response




# @app.route('/api/kk_grafik_olustur', methods=['POST']) - KK Eğri sistemi silindi
# def api_kk_grafik_olustur(): - KK Eğri sistemi silindi
    try:
        data = request.get_json()
        parametre = data.get('parametre')
        baslangic_tarih = data.get('baslangic_tarih')
        bitis_tarih = data.get('bitis_tarih')
        localStorage_verileri = data.get('localStorage_verileri', [])  # localStorage verilerini al
        
        if not parametre:
            return jsonify({'error': 'Parametre seçilmedi'}), 400
        
        # Parametre verilerini yükle
        parameters = load_parameters()
        
        # Parametre yönetiminden KK değerlerini dinamik olarak al
        parametre_eslesme = {
            'SO2': 'So2',
            'NO': 'No', 
            'CO': 'Co',
            'O2': 'O2',
            'TOC': 'TOC'
        }
        
        # Parametre adını eşleştir
        metot_adi = parametre_eslesme.get(parametre, parametre)
        
        # Parametre yönetiminden KK değerlerini bul
        kk_degerleri = None
        for param in parameters:
            if param.get('Metot') == metot_adi and param.get('KK'):
                try:
                    kk_degerleri = {
                        'kk': float(param.get('KK', 0)),
                        'minus_3s': float(param.get('-3S', 0)),
                        'minus_2s': float(param.get('-2S', 0)),
                        'plus_2s': float(param.get('+2S', 0)),
                        'plus_3s': float(param.get('+3S', 0))
                    }
                    break
                except (ValueError, TypeError):
                    continue
        
        # Eğer parametre yönetiminde bulunamazsa, varsayılan değerleri kullan
        if not kk_degerleri:
            varsayilan_degerler = {
                'O2': {'kk': 7.0, 'minus_3s': 6.6, 'minus_2s': 6.8, 'plus_2s': 7.2, 'plus_3s': 7.4},
                'CO': {'kk': 500, 'minus_3s': 400, 'minus_2s': 450, 'plus_2s': 550, 'plus_3s': 600},
                'NO': {'kk': 500, 'minus_3s': 400, 'minus_2s': 450, 'plus_2s': 550, 'plus_3s': 600},
                'SO2': {'kk': 500, 'minus_3s': 400, 'minus_2s': 450, 'plus_2s': 550, 'plus_3s': 600},
                'TOC': {'kk': 100, 'minus_3s': 80, 'minus_2s': 90, 'plus_2s': 110, 'plus_3s': 120}
            }
            kk_degerleri = varsayilan_degerler.get(parametre, {})
        
        if not kk_degerleri:
            return jsonify({'error': f'{parametre} parametresi için KK değerleri bulunamadı'}), 400
        
        # KK değerlerini al
        kk = kk_degerleri['kk']
        minus_3s = kk_degerleri['minus_3s']
        minus_2s = kk_degerleri['minus_2s']
        plus_2s = kk_degerleri['plus_2s']
        plus_3s = kk_degerleri['plus_3s']
        
        # Tarih aralığında veriler oluştur
        from datetime import datetime, timedelta
        
        baslangic = datetime.strptime(baslangic_tarih, '%Y-%m-%d')
        bitis = datetime.strptime(bitis_tarih, '%Y-%m-%d')
        
        # Saha ölçümlerinden KK verilerini al
        parametre_olcumleri = load_parametre_olcum()
        baca_bilgileri = load_baca_bilgileri()
        gercek_tarihler = []
        gercek_degerler = []
        gercek_firmalar = []
        gercek_kodlar = []
        gercek_bacalar = []  # Baca listesi ekle
        gercek_cihazlar = []
        gercek_personeller = []
        
        # localStorage verilerini işle
        for veri in localStorage_verileri:
            try:
                # Tarih formatını kontrol et
                tarih_str = veri.get('tarih', '')
                if tarih_str:
                    olcum_tarih = datetime.strptime(tarih_str, '%Y-%m-%d')
                    
                    if baslangic <= olcum_tarih <= bitis:
                        # Parametre tipine göre değeri al
                        kk_deger = None
                        
                        if parametre == 'O2':
                            kk_deger = veri.get('o2')
                        elif parametre == 'CO':
                            kk_deger = veri.get('co')
                        elif parametre == 'NO':
                            kk_deger = veri.get('no')
                        elif parametre == 'SO2':
                            kk_deger = veri.get('so2')
                        elif parametre == 'TOC':
                            kk_deger = veri.get('toc')
                        
                        if kk_deger is not None and kk_deger != '' and str(kk_deger).strip() != '':
                            try:
                                # Virgülü noktaya çevir (float için gerekli)
                                deger_str = str(kk_deger).replace(',', '.')
                                deger = float(deger_str)
                                # Değeri virgül ile formatla (2 ondalık basamak)
                                deger = round(deger, 2)
                                
                                gercek_tarihler.append(olcum_tarih)
                                gercek_degerler.append(deger)
                                gercek_firmalar.append(veri.get('firma', ''))
                                gercek_kodlar.append(veri.get('kod', ''))
                                gercek_bacalar.append(veri.get('baca', ''))  # Baca bilgisini ekle
                                gercek_cihazlar.append(veri.get('cihaz', ''))
                                gercek_personeller.append(veri.get('personel', 'Admin'))
                            except ValueError:
                                continue
            except:
                continue
        
        # Seçilen parametre için saha ölçümlerini filtrele
        for olcum in parametre_olcumleri:
            # Tarih aralığında mı kontrol et
            try:
                # Tarih formatını kontrol et
                tarih_str = olcum.get('tarih', '')
                if not tarih_str:
                    # parametre_verileri içindeki TARİH'i kontrol et
                    parametre_verileri = olcum.get('parametre_verileri', {})
                    tarih_str = parametre_verileri.get('TARİH', '')
                
                if tarih_str:
                    # GG.AA.YY formatını YYYY-MM-DD'ye çevir
                    if '.' in tarih_str and len(tarih_str) == 8:
                        olcum_tarih = datetime.strptime(tarih_str, '%d.%m.%y')
                    else:
                        olcum_tarih = datetime.strptime(tarih_str, '%Y-%m-%d')
                    
                    if baslangic <= olcum_tarih <= bitis:
                        # Parametre tipine göre KK değerini al
                        kk_deger = None
                        parametre_verileri = olcum.get('parametre_verileri', {})
                    
                        if parametre == 'O2':
                            kk_deger = parametre_verileri.get('KK1-O2')
                        elif parametre == 'CO':
                            kk_deger = parametre_verileri.get('KK1-CO')
                        elif parametre == 'NO':
                            kk_deger = parametre_verileri.get('KK1-NO')
                        elif parametre == 'SO2':
                            kk_deger = parametre_verileri.get('KK1-SO2')
                        elif parametre == 'TOC':
                            kk_deger = parametre_verileri.get('KK1-SPAN')
                        
                        if kk_deger and str(kk_deger).strip() and str(kk_deger) != '':
                            try:
                                # Tire ile ayrılmış değerler varsa (örn: 7,02-7,06-7,04) ilkini al
                                deger_str = str(kk_deger)
                                if '-' in deger_str:
                                    deger_str = deger_str.split('-')[0]
                                
                                # Virgülü noktaya çevir (float için gerekli)
                                deger_str = deger_str.replace(',', '.')
                                deger = float(deger_str)
                                
                                # Baca bilgilerinden personel adını bul
                                personel_adi = 'Admin'  # Varsayılan
                                for baca_bilgi in baca_bilgileri:
                                    if (baca_bilgi.get('firma_adi') == olcum.get('firma_adi') and 
                                        baca_bilgi.get('olcum_kodu') == olcum.get('olcum_kodu') and
                                        baca_bilgi.get('baca_adi') == olcum.get('baca_adi')):
                                        personel_adi = baca_bilgi.get('personel_adi', 'Admin')
                                        break
                                
                                gercek_tarihler.append(olcum_tarih)
                                gercek_degerler.append(deger)
                                gercek_firmalar.append(olcum.get('firma_adi', ''))
                                gercek_kodlar.append(olcum.get('olcum_kodu', ''))
                                gercek_bacalar.append(olcum.get('baca_adi', ''))  # Baca adını ekle
                                gercek_cihazlar.append(olcum.get('cihaz_adi', '-'))  # Cihaz adını doğru alanından al
                                gercek_personeller.append(personel_adi)
                            except ValueError:
                                continue
            except:
                continue
        
        # Standart sapma hesapla (KK değerinden)
        standart_sapma = (plus_3s - kk) / 3  # +3S = KK + 3σ
        
        # Gerçek veriler varsa onları kullan, yoksa boş liste
        if gercek_tarihler and gercek_degerler:
            tarihler = gercek_tarihler
            degerler = gercek_degerler
        else:
            # Veri yoksa sadece başlangıç ve bitiş tarihlerini göster
            tarihler = [baslangic, bitis]
            degerler = [None, None]
        
        # Matplotlib'i lazy loading ile yükle
        plt, np, mdates, Rectangle = load_matplotlib()
        
        # Grafik oluştur (çok daha geniş ve yüksek)
        plt.figure(figsize=(18, 8))
        plt.rcParams['font.size'] = 12  # Font boyutu
        
        # Gerçek verileri mavi nokta olarak çiz
        if gercek_tarihler and gercek_degerler:
            # Tarihleri sırala
            sorted_data = sorted(zip(gercek_tarihler, gercek_degerler))
            sorted_tarihler = [x[0] for x in sorted_data]
            sorted_degerler = [x[1] for x in sorted_data]
            
            plt.plot(sorted_tarihler, sorted_degerler, 'bo', markersize=16, linewidth=4, label='Ölçüm Değerleri')
            # Noktaları birleştiren çizgi
            plt.plot(sorted_tarihler, sorted_degerler, 'b-', linewidth=3, alpha=0.9)
        
        # Ana veri çizgisi (gerçek veri olmadığında çizilmez)
        # plt.plot(tarihler, degerler, 'b-o', linewidth=2, markersize=6)
        
        # Outlier'ı kırmızı ile işaretle (gerçek veri olmadığında çizilmez)
        # if len(degerler) >= 10:
        #     plt.plot(tarihler[9], degerler[9], 'ro', markersize=8)
        
        # Kontrol limitleri (çok daha kalın çizgiler)
        plt.axhline(y=plus_3s, color='red', linestyle='--', linewidth=4)
        plt.axhline(y=kk, color='green', linewidth=4)
        plt.axhline(y=minus_3s, color='red', linestyle='--', linewidth=4)
        
        # 2-sigma limitleri (çok daha kalın mavi tireli çizgiler)
        plt.axhline(y=plus_2s, color='blue', linestyle='--', linewidth=4)
        plt.axhline(y=minus_2s, color='blue', linestyle='--', linewidth=4)
        
        # Kontrol limitlerini Y eksenine yaz (çok daha büyük font)
        plt.text(-0.05, float(plus_3s), '+3S', fontsize=16, ha='right', va='center', color='red', fontweight='bold', transform=plt.gca().get_yaxis_transform())
        plt.text(-0.05, float(plus_2s), '+2S', fontsize=16, ha='right', va='center', color='blue', fontweight='bold', transform=plt.gca().get_yaxis_transform())
        plt.text(-0.05, float(kk), 'KK', fontsize=16, ha='right', va='center', color='green', fontweight='bold', transform=plt.gca().get_yaxis_transform())
        plt.text(-0.05, float(minus_2s), '-2S', fontsize=16, ha='right', va='center', color='blue', fontweight='bold', transform=plt.gca().get_yaxis_transform())
        plt.text(-0.05, float(minus_3s), '-3S', fontsize=16, ha='right', va='center', color='red', fontweight='bold', transform=plt.gca().get_yaxis_transform())
        
        # 1-sigma limitleri (hesaplanmış)
        plus_1s = kk + standart_sapma
        minus_1s = kk - standart_sapma
        plt.axhline(y=plus_1s, color='gray', linestyle=':', linewidth=1, alpha=0.5)
        plt.axhline(y=minus_1s, color='gray', linestyle=':', linewidth=1, alpha=0.5)
        
        # Bölgeleri işaretle (lejant olmadan) - sabit değerlerle
        plt.fill_between(range(len(tarihler)), plus_3s, plus_2s, alpha=0.1, color='red')
        plt.fill_between(range(len(tarihler)), plus_2s, plus_1s, alpha=0.1, color='orange')
        plt.fill_between(range(len(tarihler)), plus_1s, minus_1s, alpha=0.1, color='green')
        plt.fill_between(range(len(tarihler)), minus_1s, minus_2s, alpha=0.1, color='orange')
        plt.fill_between(range(len(tarihler)), minus_2s, minus_3s, alpha=0.1, color='red')
        
        # Grafik ayarları
        # plt.xlabel('Tarih', fontsize=12, fontweight='bold')  # X ekseni etiketi kaldırıldı
        # plt.ylabel('Değer', fontsize=12, fontweight='bold')  # Y ekseni etiketi kaldırıldı
        
        # Başlık ve kontrol limitleri (çok daha büyük font)
        title_text = f'{parametre} KK GRAF.    KK: {kk}    -2S: {minus_2s}    +2S: {plus_2s}    -3S: {minus_3s}    +3S: {plus_3s}'
        plt.title(title_text, fontsize=18, fontweight='bold', loc='left')
        plt.grid(True, alpha=0.3)
        # Lejantı kaldır
        # plt.legend(loc='upper right')
        
        # X ekseni formatı - sadece veri olan tarihleri göster
        plt.gca().xaxis.set_major_formatter(mdates.DateFormatter('%d/%m'))
        # Tüm tarihleri tick olarak göster (çok daha büyük font)
        plt.xticks(tarihler, rotation=90, fontsize=14)
        
        # X ekseni limitlerini ayarla - veri olan tarihlerin öncesi ve sonrası için çok daha fazla boşluk
        if tarihler:
            min_date = min(tarihler)
            max_date = max(tarihler)
            # 10 gün öncesi ve sonrası ekle (çok daha geniş görünüm)
            plt.xlim(min_date - timedelta(days=10), max_date + timedelta(days=10))
        
        # Y ekseni limitleri (parametreye göre özelleştirilmiş)
        if parametre == 'TOC':
            # TOC için sabit aralık: 91-109
            y_min = 91
            y_max = 109
        else:
            # Diğer parametreler için dinamik aralık
            # -3S ve +3S değerlerinin üstünde ve altında %30 boşluk
            margin = (plus_3s - minus_3s) * 0.3
            y_min = minus_3s - margin
            y_max = plus_3s + margin
        
        plt.ylim(y_min, y_max)
        
        # Y ekseni tick'lerini dinamik olarak ayarla (çok daha büyük font)
        if parametre == 'O2':
            plt.yticks([6.6, 6.8, 7.0, 7.2, 7.4], fontsize=14)
        elif parametre == 'TOC':
            # TOC için daha dar aralık: 91-109
            plt.yticks([91, 94, 97, 100, 103, 106, 109], fontsize=14)
        else:
            # CO, NO, SO2 için dinamik tick'ler
            step = (y_max - y_min) / 8  # 8 eşit aralık
            ticks = []
            for i in range(9):
                tick_value = y_min + (i * step)
                ticks.append(round(tick_value, 1))
            plt.yticks(ticks, fontsize=14)
        
        # Grafik alanını temizle ve kaydet
        plt.tight_layout(pad=3.0)  # Çok daha fazla boşluk
        
        # Grafiği base64 formatında döndür
        from io import BytesIO
        import base64
        
        img_buffer = BytesIO()
        plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
        img_buffer.seek(0)
        img_base64 = base64.b64encode(img_buffer.getvalue()).decode()
        plt.close()
        
        # RAM temizleme
        gc.collect()
        plt.clf()
        plt.cla()
        
        # İstatistiksel veriler (gerçek veriler varsa onları kullan)
        if gercek_degerler:
            min_deger = min(gercek_degerler)
            max_deger = max(gercek_degerler)
            ortalama = sum(gercek_degerler) / len(gercek_degerler)
        else:
            min_deger = '-'
            max_deger = '-'
            ortalama = round(float(kk), 1)
        
        # Gerçek verileri tablo için hazırla
        gercek_veriler = []
        if gercek_tarihler and gercek_degerler:
            for i, tarih in enumerate(gercek_tarihler):
                gercek_veriler.append({
                    'tarih': tarih.strftime('%Y-%m-%d'),
                    'firma': gercek_firmalar[i] if i < len(gercek_firmalar) else '',
                    'kod': gercek_kodlar[i] if i < len(gercek_kodlar) else '',
                    'baca': gercek_bacalar[i] if i < len(gercek_bacalar) else '',
                    'deger': round(gercek_degerler[i], 1),
                    'cihaz': gercek_cihazlar[i] if i < len(gercek_cihazlar) else '',
                    'personel': gercek_personeller[i] if i < len(gercek_personeller) else 'Admin'
                })
        
        istatistikler = {
            'ortalama': round(float(ortalama), 1) if isinstance(ortalama, (int, float)) else ortalama,
            'standart_sapma': round(float(standart_sapma), 1),
            'min_deger': round(float(min_deger), 1) if isinstance(min_deger, (int, float)) else min_deger,
            'max_deger': round(float(max_deger), 1) if isinstance(max_deger, (int, float)) else max_deger,
            'ucl': round(float(plus_3s), 1),
            'cl': round(float(kk), 1),
            'lcl': round(float(minus_3s), 1),
            'gercek_veriler': gercek_veriler
        }
        
        return jsonify({
            'grafik': img_base64,
            'istatistikler': istatistikler,
            'veriler': {
                'tarihler': [t.strftime('%Y-%m-%d') for t in tarihler],
                'degerler': [round(d, 1) if d is not None else None for d in degerler]
            }
        })
        
    except Exception as e:
        print(f"KK grafik oluşturma hatası: {e}")
        return jsonify({'error': f'Grafik oluşturulurken hata oluştu: {str(e)}'}), 500
# @app.route('/api/kk_rapor_olustur', methods=['POST']) - KK Eğri sistemi silindi
# def api_kk_rapor_olustur(): - KK Eğri sistemi silindi
    try:
        data = request.get_json()
        parametre = data.get('parametre')
        tarih_baslangic = data.get('tarih_baslangic')
        tarih_bitis = data.get('tarih_bitis')
        format_type = data.get('format', 'excel')
        
        if not parametre:
            return jsonify({'error': 'Parametre seçilmedi'}), 400
        
        # Parametre verilerini yükle
        parameters = load_parameters()
        
        # Seçilen parametre için KK değerlerini parametre yönetiminden al
        kk_degerleri = {}
        
        # Parametre yönetiminden KK değerlerini al
        parametre_eslesme = {
            'SO2': 'So2',
            'NO': 'No', 
            'CO': 'Co',
            'O2': 'O2',
            'TOC': 'TOC'
        }
        
        # Parametre adını eşleştir
        metot_adi = parametre_eslesme.get(parametre, parametre)
        
        # Parametre yönetiminden KK değerlerini bul
        for param in parameters:
            if param.get('Metot') == metot_adi and param.get('KK'):
                try:
                    kk_degerleri[parametre] = {
                        'kk': float(param.get('KK', 0)),
                        'minus_3s': float(param.get('-3S', 0)),
                        'minus_2s': float(param.get('-2S', 0)),
                        'plus_2s': float(param.get('+2S', 0)),
                        'plus_3s': float(param.get('+3S', 0))
                    }
                    break
                except (ValueError, TypeError):
                    continue
        
        # Eğer parametre yönetiminde bulunamazsa, varsayılan değerleri kullan
        if parametre not in kk_degerleri:
            varsayilan_degerler = {
                'O2': {'kk': 7.0, 'minus_3s': 6.6, 'minus_2s': 6.8, 'plus_2s': 7.2, 'plus_3s': 7.4},
                'CO': {'kk': 500, 'minus_3s': 400, 'minus_2s': 450, 'plus_2s': 550, 'plus_3s': 600},
                'NO': {'kk': 500, 'minus_3s': 400, 'minus_2s': 450, 'plus_2s': 550, 'plus_3s': 600},
                'SO2': {'kk': 500, 'minus_3s': 400, 'minus_2s': 450, 'plus_2s': 550, 'plus_3s': 600},
                'TOC': {'kk': 100, 'minus_3s': 80, 'minus_2s': 90, 'plus_2s': 110, 'plus_3s': 120}
            }
            kk_degerleri[parametre] = varsayilan_degerler.get(parametre, {})
        
        if parametre not in kk_degerleri:
            return jsonify({'error': f'{parametre} parametresi için KK değerleri tanımlanmamış'}), 400
        
        # KK değerlerini al
        degerler = kk_degerleri[parametre]
        kk = degerler['kk']
        minus_3s = degerler['minus_3s']
        minus_2s = degerler['minus_2s']
        plus_2s = degerler['plus_2s']
        plus_3s = degerler['plus_3s']
        
        # Tarih aralığında veriler oluştur
        from datetime import datetime, timedelta
        
        baslangic = datetime.strptime(tarih_baslangic, '%Y-%m-%d')
        bitis = datetime.strptime(tarih_bitis, '%Y-%m-%d')
        
        # Sağdaki tabloda görünen gerçek verileri al
        gercek_veriler = []
        
        # Saha ölçümlerinden KK verilerini al (api_kk_grafik_olustur ile aynı mantık)
        parametre_olcumleri = load_parametre_olcum()
        baca_bilgileri = load_baca_bilgileri()
        
        # localStorage verilerini de dahil et (manuel elle girilen veriler)
        # Bu veriler frontend'den gönderilmeli, şimdilik boş bırakıyoruz
        # localStorage_verileri = data.get('localStorage_verileri', [])
        
        # Seçilen parametre için saha ölçümlerini filtrele
        for olcum in parametre_olcumleri:
            # Tarih aralığında mı kontrol et
            try:
                # Tarih formatını kontrol et
                tarih_str = olcum.get('tarih', '')
                if not tarih_str:
                    # parametre_verileri içindeki TARİH'i kontrol et
                    parametre_verileri = olcum.get('parametre_verileri', {})
                    tarih_str = parametre_verileri.get('TARİH', '')
                
                if tarih_str:
                    # GG.AA.YY formatını YYYY-MM-DD'ye çevir
                    if '.' in tarih_str and len(tarih_str) == 8:
                        olcum_tarih = datetime.strptime(tarih_str, '%d.%m.%y')
                    else:
                        olcum_tarih = datetime.strptime(tarih_str, '%Y-%m-%d')
                    
                    if baslangic <= olcum_tarih <= bitis:
                        # Parametre tipine göre KK değerini al
                        kk_deger = None
                        parametre_verileri = olcum.get('parametre_verileri', {})
                    
                        if parametre == 'O2':
                            kk_deger = parametre_verileri.get('KK1-O2')
                        elif parametre == 'CO':
                            kk_deger = parametre_verileri.get('KK1-CO')
                        elif parametre == 'NO':
                            kk_deger = parametre_verileri.get('KK1-NO')
                        elif parametre == 'SO2':
                            kk_deger = parametre_verileri.get('KK1-SO2')
                        elif parametre == 'TOC':
                            kk_deger = parametre_verileri.get('KK1-SPAN')
                        
                        if kk_deger and str(kk_deger).strip() and str(kk_deger) != '0':
                            try:
                                # Tire ile ayrılmış değerler varsa (örn: 7,02-7,06-7,04) ilkini al
                                deger_str = str(kk_deger)
                                if '-' in deger_str:
                                    deger_str = deger_str.split('-')[0]
                                
                                # Virgülü noktaya çevir (float için gerekli)
                                deger_str = deger_str.replace(',', '.')
                                deger = float(deger_str)
                                
                                # Baca bilgilerinden personel adını bul
                                personel_adi = 'Admin'  # Varsayılan
                                for baca_bilgi in baca_bilgileri:
                                    if (baca_bilgi.get('firma_adi') == olcum.get('firma_adi') and 
                                        baca_bilgi.get('olcum_kodu') == olcum.get('olcum_kodu') and
                                        baca_bilgi.get('baca_adi') == olcum.get('baca_adi')):
                                        personel_adi = baca_bilgi.get('personel_adi', 'Admin')
                                        break
                                
                                gercek_veriler.append({
                                    'tarih': olcum_tarih.strftime('%Y-%m-%d'),
                                    'firma': olcum.get('firma_adi', ''),
                                    'kod': olcum.get('olcum_kodu', ''),
                                    'deger': round(deger, 1),
                                    'cihaz': olcum.get('baca_adi', ''),
                                    'personel': personel_adi
                                })
                            except ValueError:
                                continue
            except:
                continue
        
        # localStorage verilerini de dahil et (manuel elle girilen veriler)
        localStorage_verileri = data.get('localStorage_verileri', [])
        for veri in localStorage_verileri:
            try:
                # Tarih formatını kontrol et
                tarih_str = veri.get('tarih', '')
                if tarih_str:
                    # YYYY-MM-DD formatını kontrol et
                    if '-' in tarih_str and len(tarih_str) == 10:
                        olcum_tarih = datetime.strptime(tarih_str, '%Y-%m-%d')
                    else:
                        continue
                    
                    if baslangic <= olcum_tarih <= bitis:
                        # localStorage'da parametre adları farklı: CO, NO, SO2, O2, TOC
                        # Saha ölçümlerinde: KK1-CO, KK1-NO, KK1-SO2, KK1-O2, KK1-SPAN
                        kk_deger = None
                        
                        # Parametre adını localStorage formatına çevir (hem büyük hem küçük harf kontrol et)
                        if parametre == 'O2':
                            kk_deger = veri.get('O2') or veri.get('o2')
                        elif parametre == 'CO':
                            kk_deger = veri.get('CO') or veri.get('co')
                        elif parametre == 'NO':
                            kk_deger = veri.get('NO') or veri.get('no')
                        elif parametre == 'SO2':
                            kk_deger = veri.get('SO2') or veri.get('so2')
                        elif parametre == 'TOC':
                            kk_deger = veri.get('TOC') or veri.get('toc')
                        
                        if kk_deger and str(kk_deger).strip() and str(kk_deger) != '':
                            try:
                                # Virgülü noktaya çevir (float için gerekli)
                                deger_str = str(kk_deger).replace(',', '.')
                                deger = float(deger_str)
                                
                                gercek_veriler.append({
                                    'tarih': olcum_tarih.strftime('%Y-%m-%d'),
                                    'firma': veri.get('firma', ''),
                                    'kod': veri.get('kod', ''),
                                    'deger': round(deger, 1),
                                    'baca': veri.get('baca', ''),
                                    'personel': veri.get('personel', '')
                                })
                            except ValueError:
                                continue
            except:
                continue
        
        # İstatistiksel veriler hesapla
        if gercek_veriler:
            degerler_list = [v['deger'] for v in gercek_veriler]
            min_deger = min(degerler_list)
            max_deger = max(degerler_list)
            ortalama = sum(degerler_list) / len(degerler_list)
            standart_sapma = (plus_3s - kk) / 3  # +3S = KK + 3σ
        else:
            min_deger = '-'
            max_deger = '-'
            ortalama = round(float(kk), 1)
            standart_sapma = (plus_3s - kk) / 3
        
        # Format'a göre rapor oluştur
        if format_type == 'excel':
            return create_kk_excel_report(parametre, gercek_veriler, kk, minus_3s, minus_2s, plus_2s, plus_3s, ortalama, standart_sapma, min_deger, max_deger, tarih_baslangic, tarih_bitis)
        elif format_type == 'word':
            return create_kk_word_report(parametre, gercek_veriler, kk, minus_3s, minus_2s, plus_2s, plus_3s, ortalama, standart_sapma, min_deger, max_deger, tarih_baslangic, tarih_bitis)
        elif format_type == 'pdf':
            return create_kk_pdf_report(parametre, gercek_veriler, kk, minus_3s, minus_2s, plus_2s, plus_3s, ortalama, standart_sapma, min_deger, max_deger, tarih_baslangic, tarih_bitis)
        else:
            return jsonify({'error': 'Geçersiz format türü'}), 400
            
    except Exception as e:
        print(f"KK rapor oluşturma hatası: {e}")
        return jsonify({'error': f'Rapor oluşturulurken hata oluştu: {str(e)}'}), 500

def create_kk_excel_report(parametre, gercek_veriler, kk, minus_3s, minus_2s, plus_2s, plus_3s, ortalama, standart_sapma, min_deger, max_deger, tarih_baslangic, tarih_bitis):
    """KK Excel raporu oluştur - sadece veri"""
    try:
        import pandas as pd
        from io import BytesIO
        from datetime import datetime
        
        # Veri oluştur - tarih formatını düzelt
        def format_tarih_for_excel(tarih_str):
            if not tarih_str:
                return ''
            try:
                # YYYY-MM-DD formatından GG.AA.YY formatına çevir
                if '-' in tarih_str and len(tarih_str) == 10:
                    tarih_obj = datetime.strptime(tarih_str, '%Y-%m-%d')
                    return tarih_obj.strftime('%d.%m.%y')
                return tarih_str
            except:
                return tarih_str
        
        data = {
            "Sıra": list(range(1, len(gercek_veriler) + 1)),
            "Tarih": [format_tarih_for_excel(veri.get('tarih', '')) for veri in gercek_veriler],
            "Firma": [veri.get('firma', '') for veri in gercek_veriler],
            "Kod": [veri.get('kod', '') for veri in gercek_veriler],
            "Baca": [veri.get('baca', '') for veri in gercek_veriler],
            "Personel": [veri.get('personel', '') for veri in gercek_veriler],
            "Cihaz": [veri.get('cihaz', '') for veri in gercek_veriler],
            "Değer": [float(veri.get('deger', 0)) if veri.get('deger') else 0 for veri in gercek_veriler]
        }
        
        df = pd.DataFrame(data)
        
        # Excel dosyasını oluştur (sadece veri)
        output = BytesIO()
        df.to_excel(output, index=False, engine='openpyxl')
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'KK_Rapor_{parametre}_{tarih_baslangic}_{tarih_bitis}.xlsx'
        )
        
    except Exception as e:
        print(f"KK Excel rapor oluşturma hatası: {e}")
        return jsonify({'error': f'Excel rapor oluşturulurken hata oluştu: {str(e)}'}), 500

def create_kk_word_report(parametre, gercek_veriler, kk, minus_3s, minus_2s, plus_2s, plus_3s, ortalama, standart_sapma, min_deger, max_deger, tarih_baslangic, tarih_bitis):
    """KK Word raporu oluştur"""
    try:
        # Docx kütüphanesini yükle
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        
        if not Document:
            return jsonify({'error': 'Word rapor oluşturma için python-docx kütüphanesi gerekli'}), 500
        
        # Word dosyası oluştur
        doc = Document()
        
        # Sayfa kenar boşluklarını ayarla (1 cm = 28.35 points)
        from docx.shared import Inches
        sections = doc.sections
        for section in sections:
            section.left_margin = Inches(1)    # Sol kenar 1 cm
            section.right_margin = Inches(1)   # Sağ kenar 1 cm
            section.top_margin = Inches(1)     # Üst kenar 1 cm
            section.bottom_margin = Inches(1)  # Alt kenar 1 cm
        
        # Başlık
        title = doc.add_heading(f'{parametre} KALİTE KONTROL RAPORU', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Başlık font boyutunu artır
        for run in title.runs:
            run.font.size = Pt(24)
            run.font.bold = True
        
        # Grafik ekle
        doc.add_paragraph()
        # Tarih formatını GG.AA.YY yap
        from datetime import datetime
        baslangic_tarih = datetime.strptime(tarih_baslangic, '%Y-%m-%d').strftime('%d.%m.%y')
        bitis_tarih = datetime.strptime(tarih_bitis, '%Y-%m-%d').strftime('%d.%m.%y')
        
        doc.add_heading(f'{parametre} Kalite Kontrol Grafiği ({baslangic_tarih} - {bitis_tarih})', level=1)
        
        # Grafik oluştur ve Word'e ekle
        try:
            # Grafik oluştur
            from datetime import datetime, timedelta
            
            baslangic = datetime.strptime(tarih_baslangic, '%Y-%m-%d')
            bitis = datetime.strptime(tarih_bitis, '%Y-%m-%d')
            
            # Gerçek verileri hazırla
            gercek_tarihler = []
            gercek_degerler = []
            
            if gercek_veriler:
                for veri in gercek_veriler:
                    try:
                        tarih = datetime.strptime(veri['tarih'], '%Y-%m-%d')
                        gercek_tarihler.append(tarih)
                        gercek_degerler.append(veri['deger'])
                    except:
                        continue
            
            # Grafik oluştur
            plt.figure(figsize=(10, 6))
            plt.rcParams['font.size'] = 12
            
            # Gerçek verileri mavi nokta olarak çiz
            if gercek_tarihler and gercek_degerler:
                # Tarihleri sırala
                sorted_data = sorted(zip(gercek_tarihler, gercek_degerler))
                sorted_tarihler = [x[0] for x in sorted_data]
                sorted_degerler = [x[1] for x in sorted_data]
                
                plt.plot(sorted_tarihler, sorted_degerler, 'bo', markersize=8, linewidth=2, label='Ölçüm Değerleri')
                # Noktaları birleştiren çizgi
                plt.plot(sorted_tarihler, sorted_degerler, 'b-', linewidth=1, alpha=0.7)
            
            # Kontrol limitleri
            plt.axhline(y=plus_3s, color='red', linestyle='--', linewidth=2)
            plt.axhline(y=kk, color='green', linewidth=2)
            plt.axhline(y=minus_3s, color='red', linestyle='--', linewidth=2)
            
            # 2-sigma limitleri
            plt.axhline(y=plus_2s, color='blue', linestyle='--', linewidth=2)
            plt.axhline(y=minus_2s, color='blue', linestyle='--', linewidth=2)
            
            # Kontrol limitlerini grafik üzerine yaz
            plt.text(-0.05, float(plus_3s), '+3S', fontsize=12, ha='right', va='center', color='red', fontweight='bold', transform=plt.gca().get_yaxis_transform())
            plt.text(-0.05, float(plus_2s), '+2S', fontsize=12, ha='right', va='center', color='blue', fontweight='bold', transform=plt.gca().get_yaxis_transform())
            plt.text(-0.05, float(kk), 'KK', fontsize=12, ha='right', va='center', color='green', fontweight='bold', transform=plt.gca().get_yaxis_transform())
            plt.text(-0.05, float(minus_2s), '-2S', fontsize=12, ha='right', va='center', color='blue', fontweight='bold', transform=plt.gca().get_yaxis_transform())
            plt.text(-0.05, float(minus_3s), '-3S', fontsize=12, ha='right', va='center', color='red', fontweight='bold', transform=plt.gca().get_yaxis_transform())
            
            # 1-sigma limitleri
            plus_1s = kk + standart_sapma
            minus_1s = kk - standart_sapma
            plt.axhline(y=plus_1s, color='gray', linestyle=':', linewidth=1, alpha=0.5)
            plt.axhline(y=minus_1s, color='gray', linestyle=':', linewidth=1, alpha=0.5)
            
            # Bölgeleri işaretle
            if gercek_tarihler:
                tarihler = gercek_tarihler
            else:
                tarihler = [baslangic, bitis]
            
            plt.fill_between(range(len(tarihler)), plus_3s, plus_2s, alpha=0.1, color='red')
            plt.fill_between(range(len(tarihler)), plus_2s, plus_1s, alpha=0.1, color='orange')
            plt.fill_between(range(len(tarihler)), plus_1s, minus_1s, alpha=0.1, color='green')
            plt.fill_between(range(len(tarihler)), minus_1s, minus_2s, alpha=0.1, color='orange')
            plt.fill_between(range(len(tarihler)), minus_2s, minus_3s, alpha=0.1, color='red')
            
            # Başlık
            title_text = f'{parametre} KK GRAF.    KK: {kk}    -2S: {minus_2s}    +2S: {plus_2s}    -3S: {minus_3s}    +3S: {plus_3s}'
            plt.title(title_text, fontsize=14, fontweight='bold', loc='left')
            plt.grid(True, alpha=0.3)
            
            # X ekseni formatı
            plt.gca().xaxis.set_major_formatter(mdates.DateFormatter('%d/%m'))
            plt.xticks(tarihler, rotation=45, fontsize=10)
            
            # X ekseni limitlerini ayarla
            if tarihler:
                min_date = min(tarihler)
                max_date = max(tarihler)
                plt.xlim(min_date - timedelta(days=2), max_date + timedelta(days=2))
            
            # Y ekseni limitleri
            margin = (plus_3s - minus_3s) * 0.1
            y_min = minus_3s - margin
            y_max = plus_3s + margin
            plt.ylim(y_min, y_max)
            
            # Grafik alanını temizle ve kaydet
            plt.tight_layout()
            
            # Grafiği sabit dosya adıyla kaydet (örnek kod yaklaşımı)
            grafik_path = 'kk_grafik_temp.png'
            plt.savefig(grafik_path, format='png', dpi=150, bbox_inches='tight')
            plt.close()
            
            # Grafiği Word'e ekle
            doc.add_picture(grafik_path, width=Inches(6))
            
            # Geçici dosyayı temizle
            try:
                if os.path.exists(grafik_path):
                    os.remove(grafik_path)
            except Exception as cleanup_error:
                print(f"Grafik dosyası temizleme hatası: {cleanup_error}")
                # Dosya silinemezse devam et
                
        except Exception as e:
            print(f"Grafik ekleme hatası: {e}")
            doc.add_paragraph("Grafik oluşturulamadı.")
        
        # Ölçüm verileri tablosu
        doc.add_heading('Ölçüm Verileri', level=1)
        if gercek_veriler:
            data_table = doc.add_table(rows=1, cols=7)
            data_table.style = 'Table Grid'
            
            # Tabloya sol kenar boşluğu ekle (1 cm)
            data_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            data_table.left_margin = Inches(1)   # Sol kenar 1 cm
            
            # Tablo genişliklerini ayarla - Firma adı 3.1 cm
            data_table.columns[0].width = Inches(0.5)   # Sıra No - dar
            data_table.columns[1].width = Inches(1.0)   # Tarih
            data_table.columns[2].width = Inches(3.1)   # Firma - 3.1 cm
            data_table.columns[3].width = Inches(1.5)   # Kod
            data_table.columns[4].width = Inches(1.5)   # Baca
            data_table.columns[5].width = Inches(1.2)   # Personel
            data_table.columns[6].width = Inches(1.0)   # Değer
            
            data_hdr_cells = data_table.rows[0].cells
            headers = ['Sıra', 'Tarih', 'Firma', 'Kod', 'Baca', 'Personel', 'Değer']
            
            for i, header in enumerate(headers):
                data_hdr_cells[i].text = header
            
            for i, veri in enumerate(gercek_veriler, 1):
                row_cells = data_table.add_row().cells
                row_cells[0].text = str(i)
                
                # Tarih formatını GG.AA.YY yap
                try:
                    from datetime import datetime
                    tarih_str = veri.get('tarih', '')
                    if tarih_str:
                        tarih_obj = datetime.strptime(tarih_str, '%Y-%m-%d')
                        formatted_tarih = tarih_obj.strftime('%d.%m.%y')
                        row_cells[1].text = formatted_tarih
                    else:
                        row_cells[1].text = ''
                except:
                    row_cells[1].text = veri.get('tarih', '')
                
                # Firma adının sadece 2 kelimesini al
                firma_adi = veri.get('firma', '')
                firma_kelimeler = firma_adi.split()
                firma_kisaltilmis = ' '.join(firma_kelimeler[:2]) if len(firma_kelimeler) >= 2 else firma_adi
                row_cells[2].text = firma_kisaltilmis
                row_cells[3].text = veri.get('kod', '')
                row_cells[4].text = veri.get('baca', '')
                row_cells[5].text = veri.get('personel', '')
                row_cells[6].text = str(veri.get('deger', ''))
        else:
            doc.add_paragraph('Bu tarih aralığında ölçüm verisi bulunamadı.')
        
        # Word dosyasını geçici olarak kaydet
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_word:
            doc.save(tmp_word.name)
            tmp_word_path = tmp_word.name
        
        # Word dosyasını oku
        with open(tmp_word_path, 'rb') as f:
            word_content = f.read()
        
        # Geçici dosyayı sil
        os.unlink(tmp_word_path)
        
        # Response oluştur
        response = make_response(word_content)
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename="KK_Rapor_{parametre}_{tarih_baslangic}_{tarih_bitis}.docx"'
        
        return response
        
    except Exception as e:
        print(f"KK Word rapor oluşturma hatası: {e}")
        return jsonify({'error': f'Word rapor oluşturulurken hata oluştu: {str(e)}'}), 500

def create_kk_pdf_report(parametre, gercek_veriler, kk, minus_3s, minus_2s, plus_2s, plus_3s, ortalama, standart_sapma, min_deger, max_deger, tarih_baslangic, tarih_bitis):
    """KK PDF raporu oluştur - Word'den PDF'e çevir"""
    try:
        from docx2pdf import convert
        import tempfile
        import os
        
        # Önce Word dosyası oluştur
        word_content = create_kk_word_report(parametre, gercek_veriler, kk, minus_3s, minus_2s, plus_2s, plus_3s, ortalama, standart_sapma, min_deger, max_deger, tarih_baslangic, tarih_bitis)
        
        # Word dosyasını geçici olarak kaydet
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_word:
            tmp_word.write(word_content.data)
            tmp_word_path = tmp_word.name
        
        # PDF dosyası için geçici dosya oluştur
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_pdf:
            tmp_pdf_path = tmp_pdf.name
        
        # Word'den PDF'e çevir
        convert(tmp_word_path, tmp_pdf_path)
        
        # PDF dosyasını oku
        with open(tmp_pdf_path, 'rb') as f:
            pdf_content = f.read()
        
        # Geçici dosyaları sil
        try:
            os.unlink(tmp_word_path)
            os.unlink(tmp_pdf_path)
        except:
            pass
        
        # Response oluştur
        response = make_response(pdf_content)
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = f'attachment; filename="KK_Rapor_{parametre}_{tarih_baslangic}_{tarih_bitis}.pdf"'
        
        return response
        
    except Exception as e:
        print(f"KK PDF rapor oluşturma hatası: {e}")
        return jsonify({'error': f'PDF rapor oluşturulurken hata oluştu: {str(e)}'}), 500

@app.route('/api/firmalar')
def api_firmalar():
    """Firma listesini döndür"""
    try:
        firma_olcumler = load_firma_olcum()
        firmalar = []
        firma_adi_set = set()
        
        for olcum in firma_olcumler:
            if olcum.get('firma_adi') and olcum.get('firma_adi') not in firma_adi_set:
                firmalar.append({'firma_adi': olcum['firma_adi']})
                firma_adi_set.add(olcum['firma_adi'])
        
        return jsonify(firmalar)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/olcum_kodlari/<firma_adi>')
def api_olcum_kodlari(firma_adi):
    """Belirli firma için ölçüm kodlarını döndür"""
    try:
        firma_olcumler = load_firma_olcum()
        olcum_kodlari = []
        olcum_kodu_set = set()
        
        for olcum in firma_olcumler:
            if (olcum.get('firma_adi') == firma_adi and 
                olcum.get('olcum_kodu') and 
                olcum.get('olcum_kodu') not in olcum_kodu_set):
                olcum_kodlari.append({'olcum_kodu': olcum['olcum_kodu']})
                olcum_kodu_set.add(olcum['olcum_kodu'])
        
        return jsonify(olcum_kodlari)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/baca_listesi/<firma_adi>/<olcum_kodu>')
def api_baca_listesi(firma_adi, olcum_kodu):
    """Belirli firma ve ölçüm kodu için baca listesini döndür"""
    try:
        firma_olcumler = load_firma_olcum()
        baca_listesi = []
        
        for olcum in firma_olcumler:
            if (olcum.get('firma_adi') == firma_adi and 
                olcum.get('olcum_kodu') == olcum_kodu and
                olcum.get('baca_parametreleri')):
                
                baca_parametreleri = olcum['baca_parametreleri']
                if isinstance(baca_parametreleri, dict):
                    for baca_adi in baca_parametreleri.keys():
                        baca_listesi.append({'baca_adi': baca_adi})
        
        return jsonify(baca_listesi)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/parametre_listesi/<firma_adi>/<olcum_kodu>/<baca_adi>')
def api_parametre_listesi(firma_adi, olcum_kodu, baca_adi):
    """Belirli firma, ölçüm kodu ve baca için parametre listesini döndür"""
    try:
        firma_olcumler = load_firma_olcum()
        parametre_listesi = []
        
        for olcum in firma_olcumler:
            if (olcum.get('firma_adi') == firma_adi and 
                olcum.get('olcum_kodu') == olcum_kodu and
                olcum.get('baca_parametreleri')):
                
                baca_parametreleri = olcum['baca_parametreleri']
                if isinstance(baca_parametreleri, dict) and baca_adi in baca_parametreleri:
                    baca_parametreleri_list = baca_parametreleri[baca_adi]
                    if isinstance(baca_parametreleri_list, list):
                        for parametre in baca_parametreleri_list:
                            # Parametre string olarak tutuluyor
                            parametre_listesi.append({'parametre_adi': parametre})
        
        return jsonify(parametre_listesi)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/baca_bilgileri')
def api_baca_bilgileri():
    """Kaydedilen baca bilgileri listesini döndürür."""
    try:
        saved_baca_bilgileri = load_baca_bilgileri()
        return jsonify(saved_baca_bilgileri)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/baca_bilgileri/export_excel', methods=['POST'])
def export_baca_bilgileri_excel():
    try:
        data = request.get_json()
        selected_ids = data.get('selected_ids', [])
        
        # Baca bilgilerini yükle
        baca_bilgileri = load_baca_bilgileri()
        
        # Seçili kayıtları filtrele
        if selected_ids:
            filtered_data = [item for item in baca_bilgileri if item.get('id') in selected_ids]
        else:
            filtered_data = baca_bilgileri
        
        if not filtered_data:
            return jsonify({'error': 'Dışa aktarılacak veri bulunamadı'}), 404
        
        # Excel için DataFrame oluştur - baca_bilgileri objesi içindeki verileri doğru şekilde map et
        df_data = []
        for item in filtered_data:
            # Baca bilgileri objesini al
            baca_bilgileri = item.get('baca_bilgileri', {})
            
            # UUID'leri alan adlarına map et - gerçek UUID'leri kullan
            field_mapping = {
                '597fad80-d28f-40ea-bd28-a76c61c5203d': 'BACA NO',
                'ddca398d-0e55-4662-b661-3731e0975bd2': 'YAKIT TÜRÜ',
                '22867c9a-ca3c-4d80-b017-b73dafdd7fef': 'ISIL GÜÇ (MW)',
                '6b3546e0-184c-49de-82e4-e2835e81923b': 'ÇATI ŞEKLİ',
                '98399625-5bbc-465e-8e09-de454f231ae4': 'KAYNAK TÜRÜ',
                'd9958774-43f7-4bc3-8e12-436614a6193a': 'BACA ŞEKLİ',
                'b1b6fc38-98c0-4048-8b8e-795cf7d44c48': 'BACA ÖLÇÜSÜ',
                '6a301d72-f21b-485b-b8fb-116ad5cb223f': 'YERDEN YÜK.',
                '8ec9ecc9-ecda-4bf2-9802-02fa2e3fda4c': 'ÇATI YÜK',
                'ab2b67dd-16a5-4bee-9b6e-b60b8cfc2d0c': 'RÜZGAR HIZ (M/S)',
                'eca60e54-ec39-4412-8884-caa17faed0be': 'ORT. SIC.',
                '64238e0a-6387-4c31-9bf4-d7f800ef17e1': 'ORT. NEM',
                'b09ad69a-e4d4-4219-b055-2cf923ffd499': 'ORT. BAS.',
                '9c8c8bcf-c98e-4109-8b10-63b08b26460e': 'A BACA',
                'af55c55f-f83b-4b90-a655-ee76bf6bb2ac': 'B BACA',
                '20881447-f7c8-4a6b-8583-76c7246082ef': 'C DELİK'
            }
            
            # Temel alanlar
            row = {
                'Firma': item.get('firma_adi', ''),
                'Ölçüm Kodu': item.get('olcum_kodu', ''),
                'Baca': item.get('baca_adi', ''),
            }
            
            # Baca bilgileri objesindeki her UUID için değeri al
            for uuid, field_name in field_mapping.items():
                value = baca_bilgileri.get(uuid, '')
                if value and value.strip() != '':
                    row[field_name] = value
                else:
                    row[field_name] = '*'
            
            # Ek alanlar
            created_at = item.get('created_at', '')
            updated_at = item.get('updated_at', '')
            
            # Tarih formatını düzelt
            if created_at:
                try:
                    created_date = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
                    created_at = created_date.strftime('%d.%m.%Y %H:%M')
                except:
                    created_at = created_at
            
            if updated_at:
                try:
                    updated_date = datetime.fromisoformat(updated_at.replace('Z', '+00:00'))
                    updated_at = updated_date.strftime('%d.%m.%Y %H:%M')
                except:
                    updated_at = updated_at
            
            row.update({
                'Fotoğraf': item.get('photo_path', '') or '*',
                'Kayıt Tarihi': created_at or '*',
                'Güncelleme Tarihi': updated_at or '*'
            })
            
            df_data.append(row)
        
        df = pd.DataFrame(df_data)
        
        # Geçici dosya oluştur
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
            # Excel dosyasını oluştur
            with pd.ExcelWriter(tmp_file.name, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='Baca Bilgileri')
                
                # Sütun genişliklerini ayarla
                worksheet = writer.sheets['Baca Bilgileri']
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)  # Maksimum 50 karakter
                    worksheet.column_dimensions[column_letter].width = adjusted_width
            
            tmp_file_path = tmp_file.name
        
        # Dosyayı gönder
        return send_file(
            tmp_file_path,
            as_attachment=True,
            download_name=f'baca_bilgileri_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx',
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
    
    except Exception as e:
        return jsonify({'error': f'Excel dışa aktarma hatası: {str(e)}'}), 500
@app.route('/api/baca_bilgileri/export_pdf', methods=['POST'])
def export_baca_bilgileri_pdf():
    try:
        data = request.get_json()
        selected_ids = data.get('selected_ids', [])
        
        # Baca bilgilerini yükle
        baca_bilgileri = load_baca_bilgileri()
        
        # Seçili kayıtları filtrele
        if selected_ids:
            filtered_data = [item for item in baca_bilgileri if item.get('id') in selected_ids]
        else:
            filtered_data = baca_bilgileri
        
        if not filtered_data:
            return jsonify({'error': 'Dışa aktarılacak veri bulunamadı'}), 404
        
        # İlk seçili baca için Word şablonunu kullanarak Word dosyası oluştur
        item = filtered_data[0]  # İlk seçili kayıt
        
        # Bu bacaya ait parametreleri bul
        parametre_olcumleri = load_parametre_olcum()
        baca_parametreleri = []
        
        for parametre in parametre_olcumleri:
            if (parametre.get('firma_adi') == item.get('firma_adi') and 
                parametre.get('olcum_kodu') == item.get('olcum_kodu') and 
                parametre.get('baca_adi') == item.get('baca_adi')):
                baca_parametreleri.append(parametre)
        
        # Form bilgilerini al
        forms = load_forms()
        form_bilgisi = None
        if forms:
            for form in forms:
                if 'EMİSYON' in form.get('form_adi', '').upper():
                    form_bilgisi = form
                    break
            if not form_bilgisi:
                form_bilgisi = forms[0]
        
        # Word dosyasını oluştur
        word_doc = create_baca_word_document_from_template(item, baca_parametreleri, form_bilgisi)
        
        if word_doc is None:
            return jsonify({'error': 'Word dosyası oluşturulamadı'}), 500
        
        # Word dosyasını geçici olarak kaydet
        import tempfile
        import os
        from io import BytesIO
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_word:
            # Document objesini BytesIO'ya kaydet
            doc_buffer = BytesIO()
            word_doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            # Geçici dosyaya yaz
            tmp_word.write(doc_buffer.getvalue())
            tmp_word_path = tmp_word.name
        
        # Word dosyasını PDF'e çevir
        try:
            from docx2pdf import convert
            import pythoncom
            
            # COM bileşenlerini başlat
            pythoncom.CoInitialize()
            
            # PDF dosya adını oluştur
            firma_adi = item.get('firma_adi', 'Bilinmeyen')
            olcum_kodu = item.get('olcum_kodu', 'Bilinmeyen')
            baca_adi = item.get('baca_adi', 'Bilinmeyen')
            
            # Dosya adını temizle
            def clean_filename(text):
                tr_to_en = {
                    'ç': 'c', 'ğ': 'g', 'ı': 'i', 'ö': 'o', 'ş': 's', 'ü': 'u',
                    'Ç': 'C', 'Ğ': 'G', 'I': 'I', 'İ': 'I', 'Ö': 'O', 'Ş': 'S', 'Ü': 'U'
                }
                for tr_char, en_char in tr_to_en.items():
                    text = text.replace(tr_char, en_char)
                cleaned = "".join(c for c in text if c.isalnum() or c in (' ', '-', '_')).strip()
                cleaned = cleaned.replace(' ', '-')
                while '--' in cleaned:
                    cleaned = cleaned.replace('--', '-')
                return cleaned
            
            firma_adi_clean = clean_filename(firma_adi)
            olcum_kodu_clean = clean_filename(olcum_kodu)
            baca_adi_clean = clean_filename(baca_adi)
            
            pdf_filename = f"{firma_adi_clean}-{baca_adi_clean}-{olcum_kodu_clean}.pdf"
            pdf_path = os.path.join(tempfile.gettempdir(), pdf_filename)
            
            # Word dosyasını PDF'e çevir
            convert(tmp_word_path, pdf_path)
            
            # PDF dosyasını oku
            with open(pdf_path, 'rb') as pdf_file:
                pdf_content = pdf_file.read()
            
            # Geçici dosyaları temizle
            os.unlink(tmp_word_path)
            os.unlink(pdf_path)
            
            # PDF dosyasını döndür
            response = make_response(pdf_content)
            response.headers['Content-Type'] = 'application/pdf'
            response.headers['Content-Disposition'] = f'attachment; filename="{pdf_filename}"'
            return response
            
        except Exception as pdf_error:
            print(f"PDF dönüştürme hatası: {pdf_error}")
            
            # COM bileşenlerini temizle (hata durumunda da)
            try:
                pythoncom.CoUninitialize()
            except:
                pass
                
            # PDF dönüştürme çalışmazsa, Word dosyasını döndür
            with open(tmp_word_path, 'rb') as word_file:
                word_content = word_file.read()
            
            # Geçici dosyayı temizle
            os.unlink(tmp_word_path)
            
            # Word dosyasını döndür
            response = make_response(word_content)
            response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            response.headers['Content-Disposition'] = f'attachment; filename="{firma_adi_clean}-{baca_adi_clean}-{olcum_kodu_clean}.docx"'
            return response
        
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"PDF Export Error: {str(e)}")
        print(f"Error Details: {error_details}")
        return jsonify({'error': f'PDF dışa aktarma hatası: {str(e)}'}), 500

@app.route('/api/baca_bilgileri/import_excel', methods=['POST'])
def import_baca_bilgileri_excel():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'Dosya seçilmedi'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Dosya seçilmedi'}), 400
        
        if not file.filename.endswith('.xlsx'):
            return jsonify({'error': 'Sadece Excel (.xlsx) dosyaları desteklenir'}), 400
        
        # Excel dosyasını oku
        df = pd.read_excel(file, engine='openpyxl')
        
        # Gerekli sütunları kontrol et
        required_columns = ['Firma', 'Ölçüm Kodu', 'Baca']
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            return jsonify({'error': f'Eksik sütunlar: {", ".join(missing_columns)}'}), 400
        
        # Mevcut baca bilgilerini yükle
        existing_data = load_baca_bilgileri()
        
        # Yeni kayıtları ekle - baca_bilgileri objesi formatında
        new_records = []
        for index, row in df.iterrows():
            try:
                # UUID mapping'i (export ile aynı)
                field_mapping = {
                    '7425ab86-4cfb-4796-bd2c-a84c3f1c61a4': 'BACA NO',
                    '2517064c-c286-4210-8f60-fa0a6b9e22a9': 'YAKIT TÜRÜ',
                    '2360aa46-81dd-41db-a6b6-5c23e62900fe': 'ISIL GÜÇ (MW)',
                    'f1e7b875-ddf4-4628-87a3-eda6afa775b7': 'ÇATI ŞEKLİ',
                    '7647078a-4f69-4908-99f7-3fff5651cc9b': 'KAYNAK TÜRÜ',
                    '394272bc-a8dc-46c4-a1a8-15ed4c0dbf29': 'BACA ŞEKLİ',
                    '6774fc7e-c124-4272-b826-482d064f3215': 'BACA ÖLÇÜSÜ',
                    '943f1111-1be7-4dd3-b62d-b734963c768e': 'YERDEN YÜK.',
                    '3f622008-8178-460c-a6a1-c5f4eb357231': 'ÇATI YÜK',
                    'eba873d9-1398-4fde-9ef0-851359037990': 'RÜZGAR HIZ (M/S)',
                    '5ab155d9-a52d-4359-abbf-bd3083e216da': 'ORT. SIC.',
                    'dea8eb12-71c2-4af4-a280-acd5eac4660b': 'ORT. NEM',
                    '922ab008-e6ec-442c-a2ee-4159b6ce948e': 'ORT. BAS.',
                    '0c47bc1d-afdf-477a-94b8-abb2c1d6f92c': 'A BACA',
                    'f59ec1f6-f93e-49a7-ac11-c33ce327987a': 'B BACA',
                    '0e9de9b8-d7a4-4125-b299-747e47c71b4b': 'C DELİK'
                }
                
                # Reverse mapping oluştur
                reverse_mapping = {v: k for k, v in field_mapping.items()}
                
                # Baca bilgileri objesi oluştur
                baca_bilgileri = {}
                for field_name, uuid in reverse_mapping.items():
                    value = row.get(field_name, '')
                    if value and str(value).strip() and str(value).lower() != 'nan':  # Boş değilse ve NaN değilse ekle
                        baca_bilgileri[uuid] = str(value)
                
                # Temel alanları kontrol et
                firma_adi = str(row.get('Firma', '')).strip()
                olcum_kodu = str(row.get('Ölçüm Kodu', '')).strip()
                baca_adi = str(row.get('Baca', '')).strip()
                
                if not firma_adi or not olcum_kodu or not baca_adi:
                    return jsonify({'error': f'Satır {index + 1}: Firma, Ölçüm Kodu ve Baca alanları zorunludur'}), 400
                
                # Fotoğraf alanını kontrol et
                photo_path = row.get('Fotoğraf', '')
                if photo_path and str(photo_path).strip() and str(photo_path).lower() != 'nan':
                    photo_path = str(photo_path)
                else:
                    photo_path = None
                
                new_record = {
                    'id': str(uuid4()),
                    'firma_adi': firma_adi,
                    'olcum_kodu': olcum_kodu,
                    'baca_adi': baca_adi,
                    'baca_bilgileri': baca_bilgileri,
                    'photo_path': photo_path,
                    'created_at': str(row.get('Kayıt Tarihi', datetime.now().isoformat())),
                    'updated_at': str(row.get('Güncelleme Tarihi', datetime.now().isoformat()))
                }
                new_records.append(new_record)
                
            except Exception as row_error:
                return jsonify({'error': f'Satır {index + 1} işlenirken hata: {str(row_error)}'}), 400
        
        # Yeni kayıtları mevcut veriye ekle
        existing_data.extend(new_records)
        
        # Veriyi kaydet
        if save_baca_bilgileri(existing_data):
            return jsonify({
                'success': True,
                'message': f'{len(new_records)} kayıt başarıyla içe aktarıldı',
                'imported_count': len(new_records)
            })
        else:
            return jsonify({'error': 'Veriler kaydedilemedi'}), 500
    
    except Exception as e:
        return jsonify({'error': f'Excel içe aktarma hatası: {str(e)}'}), 500

@app.route('/api/baca_bilgileri/<record_id>', methods=['DELETE'])
def api_delete_baca_bilgileri(record_id):
    """Belirtilen ID'ye sahip baca bilgilerini siler."""
    try:
        saved_baca_bilgileri = load_baca_bilgileri()
        
        # Kaydı bul
        record_index = None
        for i, record in enumerate(saved_baca_bilgileri):
            if record.get('id') == record_id:
                record_index = i
                break
        
        if record_index is None:
            return jsonify({'success': False, 'error': 'Kayıt bulunamadı'}), 404
        
        # Kaydı sil
        deleted_record = saved_baca_bilgileri.pop(record_index)
        
        # Dosyaya kaydet
        if save_baca_bilgileri(saved_baca_bilgileri):
            print(f"Baca bilgileri silindi: {deleted_record.get('firma_adi')} - {deleted_record.get('olcum_kodu')} - {deleted_record.get('baca_adi')}")
            return jsonify({'success': True, 'message': 'Baca bilgileri başarıyla silindi'})
        else:
            return jsonify({'success': False, 'error': 'Veriler kaydedilemedi'}), 500
            
    except Exception as e:
        print(f"Baca bilgileri silinirken hata: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/baca_parametreleri')
def api_baca_parametreleri():
    """Baca parametrelerini döndürür."""
    try:
        baca_paralar = load_baca_paralar()
        return jsonify(baca_paralar)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/parametre_bilgileri')
def api_parametre_bilgileri():
    """Tüm parametre bilgilerini döndürür."""
    try:
        parametreler = load_parameters()
        return jsonify(parametreler)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/save_parametre_olcum_old', methods=['POST'])
def save_parametre_olcum_old():
    """Parametre ölçümlerini kaydeder."""
    try:
        firma_adi = request.form.get('firma_adi')
        olcum_kodu = request.form.get('olcum_kodu')
        baca_adi = request.form.get('baca_adi')
        parametre_adi = request.form.get('parametre_adi')
        olcum_sonucu = request.form.get('olcum_sonucu')
        olcum_birimi = request.form.get('olcum_birimi')
        olcum_tarihi = request.form.get('olcum_tarihi')
        olcum_saati = request.form.get('olcum_saati')
        olcum_notlari = request.form.get('olcum_notlari')
        
        if not all([firma_adi, olcum_kodu, baca_adi, parametre_adi, olcum_sonucu, olcum_birimi]):
            return jsonify({'success': False, 'error': 'Eksik bilgi'})
        
        # Saha ölçüm verilerini yükle
        saha_olc_data = load_saha_olc()
        
        # Yeni ölçüm kaydı
        new_olcum = {
            'id': str(uuid4()),
            'firma_adi': firma_adi,
            'olcum_kodu': olcum_kodu,
            'baca_adi': baca_adi,
            'parametre_adi': parametre_adi,
            'olcum_sonucu': olcum_sonucu,
            'olcum_birimi': olcum_birimi,
            'olcum_tarihi': olcum_tarihi,
            'olcum_saati': olcum_saati,
            'olcum_notlari': olcum_notlari,
            'created_at': datetime.now().isoformat()
        }
        
        # Veriyi ekle
        saha_olc_data.append(new_olcum)
        
        # Dosyaya kaydet
        if save_saha_olc(saha_olc_data):
            print(f"Parametre ölçümü kaydedildi: {firma_adi} - {olcum_kodu} - {baca_adi} - {parametre_adi}")
            return jsonify({
                'success': True, 
                'message': 'Parametre ölçümü başarıyla kaydedildi'
            })
        else:
            return jsonify({'success': False, 'error': 'Veriler kaydedilemedi'}), 500
            
    except Exception as e:
        print(f"Parametre ölçümü kaydedilirken hata: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/save_parametre_olcum_saha', methods=['POST'])
def save_parametre_olcum_saha():
    """Saha ölçümü için parametre ölçümlerini kaydeder."""
    try:
        firma_adi = request.form.get('firma_adi')
        olcum_kodu = request.form.get('olcum_kodu')
        baca_adi = request.form.get('baca_adi')
        parametre_adi = request.form.get('parametre_adi')
        parametre_verileri_json = request.form.get('parametre_verileri')
        
        if not all([firma_adi, olcum_kodu, baca_adi, parametre_adi, parametre_verileri_json]):
            return jsonify({'success': False, 'error': 'Eksik bilgi'})
        
        parametre_verileri = json.loads(parametre_verileri_json)
        
        # Mevcut parametre ölçümlerini yükle
        saved_parametre_olcum = load_parametre_olcum()
        
        # İlgili baca bilgisinden personel adını bul (parametre ölçümü için varsayılan)
        personel_adi_default = ''
        try:
            for b in load_baca_bilgileri():
                if (
                    b.get('firma_adi') == firma_adi and
                    b.get('olcum_kodu') == olcum_kodu and
                    b.get('baca_adi') == baca_adi
                ):
                    personel_adi_default = b.get('personel_adi', '')
                    break
        except Exception:
            personel_adi_default = ''
        
        # Kaydedilecek yeni veri
        new_record = {
            'id': str(uuid4()),
            'firma_adi': firma_adi,
            'olcum_kodu': olcum_kodu,
            'baca_adi': baca_adi,
            'parametre_adi': parametre_adi,
            'parametre_verileri': parametre_verileri,
            # Personel adı kayıt içinde bulunmuyorsa baca bilgilerindeki personele düş
            'personel_adi': personel_adi_default,
            'created_at': datetime.now().isoformat(),
            'updated_at': datetime.now().isoformat()
        }
        
        # Aynı firma-ölçüm-baca-parametre kombinasyonu varsa güncelle, yoksa ekle
        record_found = False
        for i, record in enumerate(saved_parametre_olcum):
            if (record.get('firma_adi') == firma_adi and 
                record.get('olcum_kodu') == olcum_kodu and 
                record.get('baca_adi') == baca_adi and
                record.get('parametre_adi') == parametre_adi):
                # Mevcut kaydı güncelle
                new_record['id'] = record.get('id', str(uuid4()))
                new_record['created_at'] = record.get('created_at', datetime.now().isoformat())
                saved_parametre_olcum[i] = new_record
                record_found = True
                break
        
        if not record_found:
            # Yeni kayıt ekle
            saved_parametre_olcum.append(new_record)
        
        # Dosyaya kaydet
        if save_parametre_olcum(saved_parametre_olcum):
            print(f"Parametre ölçümü kaydedildi: {firma_adi} - {olcum_kodu} - {baca_adi} - {parametre_adi}")
            return jsonify({
                'success': True, 
                'message': 'Parametre ölçümü başarıyla kaydedildi'
            })
        else:
            return jsonify({'success': False, 'error': 'Veriler kaydedilemedi'}), 500
            
    except Exception as e:
        print(f"Parametre ölçümü kaydedilirken hata: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/parametre_olcumleri')
def api_parametre_olcumleri():
    """Tüm parametre ölçümlerini döndürür."""
    try:
        parametre_olcumleri = load_parametre_olcum()
        return jsonify(parametre_olcumleri)
    except Exception as e:
        print(f"Parametre ölçümleri yüklenirken hata: {e}")
        return jsonify([])

@app.route('/api/parametre_olcumleri/bulk_delete', methods=['POST'])
def api_bulk_delete_parametre_olcumleri():
    """Seçilen parametre ölçümlerini toplu olarak siler."""
    try:
        data = request.get_json()
        ids_to_delete = data.get('ids', [])
        
        if not ids_to_delete:
            return jsonify({'success': False, 'error': 'Silinecek kayıt seçilmedi'})
        
        # Mevcut parametre ölçümlerini yükle
        parametre_olcumleri = load_parametre_olcum()
        
        # Silinecek kayıtları filtrele
        original_count = len(parametre_olcumleri)
        parametre_olcumleri = [record for record in parametre_olcumleri if record.get('id') not in ids_to_delete]
        deleted_count = original_count - len(parametre_olcumleri)
        
        # Dosyaya kaydet
        if save_parametre_olcum(parametre_olcumleri):
            print(f"Parametre ölçümleri toplu silindi: {deleted_count} kayıt")
            return jsonify({
                'success': True, 
                'message': f'{deleted_count} adet parametre ölçümü başarıyla silindi',
                'deleted_count': deleted_count
            })
        else:
            return jsonify({'success': False, 'error': 'Veriler kaydedilemedi'})
            
    except Exception as e:
        print(f"Parametre ölçümleri toplu silme hatası: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/saha_olcumleri/bulk_delete', methods=['POST'])
def api_bulk_delete_saha_olcumleri():
    """Saha özet ekranından seçilen ölçümler için 1-4. aşama verilerini toplu olarak siler."""
    try:
        data = request.get_json() or {}
        ids = data.get('ids', [])

        if not ids:
            return jsonify({'success': False, 'message': 'Silinecek kayıt seçilmedi'})

        id_set = {str(i) for i in ids}

        # 1) Baca bilgileri (özet kayıtlar)
        baca_bilgileri = load_baca_bilgileri()
        targets = []  # Sonraki aşamalar için firma/ölçüm/baca/parametre kombinasyonları
        remaining_baca = []

        for rec in baca_bilgileri:
            rec_id = rec.get('id')
            rec_id_str = str(rec_id) if rec_id is not None else None

            if rec_id_str and rec_id_str in id_set:
                firma = (rec.get('firma_adi') or rec.get('firma') or '').strip()
                olcum_kodu = (rec.get('olcum_kodu') or '').strip()
                baca = (rec.get('baca_adi') or rec.get('baca_no') or '').strip()
                parametre = (rec.get('parametre_adi') or rec.get('parametre') or '').strip()

                targets.append({
                    'firma': firma,
                    'olcum_kodu': olcum_kodu,
                    'baca': baca,
                    'parametre': parametre,
                })
            else:
                remaining_baca.append(rec)

        deleted_baca = len(baca_bilgileri) - len(remaining_baca)
        save_baca_bilgileri(remaining_baca)

        if not targets:
            return jsonify({'success': False, 'message': 'Seçilen ID\'lere ait kayıt bulunamadı'}), 404

        # Ortak anahtar kümeleri
        combo_set = {
            (t['firma'], t['olcum_kodu'], t['baca'])
            for t in targets
        }
        param_combo_set = {
            (t['firma'], t['olcum_kodu'], t['baca'], t['parametre'])
            for t in targets if t['parametre']
        }

        def load_json_list(path):
            if not os.path.exists(path):
                return []
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    content = f.read().strip()
                    if not content:
                        return []
                    return json.loads(content)
            except (FileNotFoundError, json.JSONDecodeError):
                return []

        def save_json_list(path, data):
            with open(path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)

        # 2) 1. Aşama verileri (asama_verileri.json, asama == 1)
        asama1_path = 'asama_verileri.json'
        asama1_veriler = load_json_list(asama1_path)
        before_asama1 = len(asama1_veriler)
        asama1_veriler = [
            r for r in asama1_veriler
            if not (
                r.get('asama') == 1 and
                ( (r.get('firma') or '').strip(),
                  (r.get('olcum_kodu') or '').strip(),
                  (r.get('baca') or r.get('bacaNo') or '').strip() ) in combo_set
            )
        ]
        deleted_asama1 = before_asama1 - len(asama1_veriler)
        if deleted_asama1:
            save_json_list(asama1_path, asama1_veriler)

        # 3) 2. Aşama verileri (asama2_verileri.json, asama == 2)
        asama2_path = 'asama2_verileri.json'
        asama2_veriler = load_json_list(asama2_path)
        before_asama2 = len(asama2_veriler)
        asama2_veriler = [
            r for r in asama2_veriler
            if not (
                r.get('asama') == 2 and
                ( (r.get('firma') or '').strip(),
                  (r.get('olcum_kodu') or '').strip(),
                  (r.get('baca') or '').strip() ) in combo_set
            )
        ]
        deleted_asama2 = before_asama2 - len(asama2_veriler)
        if deleted_asama2:
            save_json_list(asama2_path, asama2_veriler)

        # 4) 3. Aşama verileri (asama3_verileri.json, asama == 3)
        asama3_path = 'asama3_verileri.json'
        asama3_veriler = load_json_list(asama3_path)
        before_asama3 = len(asama3_veriler)
        asama3_veriler = [
            r for r in asama3_veriler
            if not (
                r.get('asama') == 3 and
                ( (r.get('firma') or '').strip(),
                  (r.get('olcum_kodu') or '').strip(),
                  (r.get('baca') or '').strip() ) in combo_set
            )
        ]
        deleted_asama3 = before_asama3 - len(asama3_veriler)
        if deleted_asama3:
            save_json_list(asama3_path, asama3_veriler)

        # 5) 4. Aşama / parametre ölçümleri (parametre_olcum.json)
        parametre_olcumleri = load_parametre_olcum()
        before_param = len(parametre_olcumleri)
        parametre_olcumleri = [
            r for r in parametre_olcumleri
            if (
                ( (r.get('firma_adi') or r.get('firma') or '').strip(),
                  (r.get('olcum_kodu') or '').strip(),
                  (r.get('baca_adi') or r.get('baca_no') or r.get('baca') or '').strip(),
                  (r.get('parametre_adi') or r.get('parametre') or '').strip() )
                not in param_combo_set
            )
        ]
        deleted_param = before_param - len(parametre_olcumleri)
        if deleted_param:
            save_parametre_olcum(parametre_olcumleri)

        return jsonify({
            'success': True,
            'message': 'Seçili ölçüm ve ilgili aşama verileri silindi',
            'deleted': {
                'baca_bilgileri': deleted_baca,
                'asama1': deleted_asama1,
                'asama2': deleted_asama2,
                'asama3': deleted_asama3,
                'parametre_olcum': deleted_param,
            }
        })

    except Exception as e:
        print(f"Saha ölçümleri toplu silme hatası: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/save_baca_bilgileri_saha', methods=['POST'])
def save_baca_bilgileri_saha():
    """Saha ölçümü için baca bilgilerini kaydeder."""
    try:
        firma_adi = request.form.get('firma_adi')
        olcum_kodu = request.form.get('olcum_kodu')
        baca_adi = request.form.get('baca_adi')
        baca_bilgileri_json = request.form.get('baca_bilgileri')
        personel_adi = request.form.get('personel_adi', '')
        is_edit = request.form.get('is_edit', 'false') == 'true'
        
        if not all([firma_adi, olcum_kodu, baca_adi, baca_bilgileri_json]):
            return jsonify({'success': False, 'error': 'Eksik bilgi'})
        
        print(f"Gelen baca bilgileri JSON: {baca_bilgileri_json}")
        baca_bilgileri = json.loads(baca_bilgileri_json)
        print(f"Parse edilen baca bilgileri: {baca_bilgileri}")
        
        # Fotoğraf işleme
        photo_path = None
        if 'photo' in request.files:
            photo = request.files['photo']
            if photo and photo.filename:
                # Fotoğraf klasörü oluştur
                photo_dir = os.path.join(app.root_path, 'static', 'uploads', 'photos')
                os.makedirs(photo_dir, exist_ok=True)
                
                # Benzersiz dosya adı oluştur
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f"{firma_adi}_{olcum_kodu}_{baca_adi}_{timestamp}.jpg"
                photo_path = os.path.join(photo_dir, filename)
                
                # Fotoğrafı kaydet
                photo.save(photo_path)
                photo_path = f"uploads/photos/{filename}"  # Web erişimi için relative path
        
        # Mevcut baca bilgilerini yükle
        saved_baca_bilgileri = load_baca_bilgileri()
        
        # Kaydedilecek yeni veri
        new_record = {
            'id': str(uuid4()),
            'firma_adi': firma_adi,
            'olcum_kodu': olcum_kodu,
            'baca_adi': baca_adi,
            'baca_bilgileri': baca_bilgileri,
            'personel_adi': personel_adi,
            'photo_path': photo_path,
            'created_at': datetime.now().isoformat(),
            'updated_at': datetime.now().isoformat()
        }
        
        # Eğer güncelleme ise, mevcut kaydı bul ve güncelle
        if is_edit:
            record_found = False
            for i, record in enumerate(saved_baca_bilgileri):
                if (record.get('firma_adi') == firma_adi and 
                    record.get('olcum_kodu') == olcum_kodu and 
                    record.get('baca_adi') == baca_adi):
                    # Mevcut kaydı güncelle
                    new_record['id'] = record.get('id', str(uuid4()))
                    new_record['created_at'] = record.get('created_at', datetime.now().isoformat())
                    # Fotoğraf değişmemişse eski fotoğrafı koru
                    if not photo_path:
                        new_record['photo_path'] = record.get('photo_path')
                    saved_baca_bilgileri[i] = new_record
                    record_found = True
                    break
            
            if not record_found:
                # Kayıt bulunamadıysa yeni kayıt olarak ekle
                saved_baca_bilgileri.append(new_record)
        else:
            # Yeni kayıt ekle
            saved_baca_bilgileri.append(new_record)
        
        # Baca bilgilerini dosyaya kaydet
        if save_baca_bilgileri(saved_baca_bilgileri):
            print(f"Baca bilgileri kaydedildi: {firma_adi} - {olcum_kodu} - {baca_adi}")
            
            # Personel adı değiştiyse, o bacaya ait parametre ölçüm kayıtlarını güncelle
            if personel_adi:
                update_parametre_olcum_personel(firma_adi, olcum_kodu, baca_adi, personel_adi)
            
            return jsonify({
                'success': True, 
                'message': 'Baca bilgileri başarıyla kaydedildi',
                'photo_path': photo_path
            })
        else:
            return jsonify({'success': False, 'error': 'Veriler kaydedilemedi'})
        
    except Exception as e:
        print(f"Baca bilgileri kaydetme hatası: {str(e)}")
        return jsonify({'success': False, 'error': str(e)})

def update_parametre_olcum_personel(firma_adi, olcum_kodu, baca_adi, personel_adi):
    """Baca bilgilerinde personel adı değiştiğinde, o bacaya ait parametre ölçüm kayıtlarını günceller."""
    try:
        # Mevcut parametre ölçüm kayıtlarını yükle
        parametre_olcumleri = load_parametre_olcum()
        
        # Güncellenen kayıt sayısı
        updated_count = 0
        
        # O bacaya ait tüm parametre ölçüm kayıtlarını bul ve güncelle
        for record in parametre_olcumleri:
            if (record.get('firma_adi') == firma_adi and 
                record.get('olcum_kodu') == olcum_kodu and 
                record.get('baca_adi') == baca_adi):
                # Personel adını güncelle
                record['personel_adi'] = personel_adi
                updated_count += 1
        
        # Güncellenmiş verileri kaydet
        if updated_count > 0:
            save_parametre_olcum(parametre_olcumleri)
            print(f"Parametre ölçüm kayıtları güncellendi: {updated_count} kayıt - {firma_adi} - {olcum_kodu} - {baca_adi} - Personel: {personel_adi}")
        else:
            print(f"Güncellenecek parametre ölçüm kaydı bulunamadı: {firma_adi} - {olcum_kodu} - {baca_adi}")
            
    except Exception as e:
        print(f"Parametre ölçüm kayıtları güncellenirken hata: {str(e)}")

def cleanup_orphaned_parametre_personel():
    """Baca bilgilerinde olmayan personel isimlerini parametre ölçümlerinden temizler."""
    try:
        # Baca bilgilerini yükle
        baca_bilgileri = load_baca_bilgileri()
        
        # Geçerli personel isimlerini topla
        valid_personel = set()
        for baca in baca_bilgileri:
            personel = baca.get('personel_adi', '').strip()
            if personel:
                valid_personel.add(personel)
        
        print(f"Geçerli personel isimleri: {valid_personel}")
        
        # Parametre ölçümlerini yükle
        parametre_olcumleri = load_parametre_olcum()
        
        # Temizlenen kayıt sayısı
        cleaned_count = 0
        
        # Geçersiz personel isimlerini temizle
        for record in parametre_olcumleri:
            current_personel = record.get('personel_adi', '').strip()
            if current_personel and current_personel not in valid_personel:
                print(f"Geçersiz personel ismi temizleniyor: {current_personel} -> (boş)")
                record['personel_adi'] = ''
                cleaned_count += 1
        
        # Temizlenmiş verileri kaydet
        if cleaned_count > 0:
            save_parametre_olcum(parametre_olcumleri)
            print(f"Geçersiz personel isimleri temizlendi: {cleaned_count} kayıt")
        else:
            print("Temizlenecek geçersiz personel ismi bulunamadı")
            
    except Exception as e:
        print(f"Geçersiz personel isimleri temizlenirken hata: {str(e)}")

@app.route('/ayarlar', methods=['GET', 'POST'])
def ayarlar():
    """Ayarlar sayfası."""
    if not session.get('logged_in'):
        return redirect(url_for('login'))

    users = load_users()
    parameters = load_parameters()
    baca_paralar = load_baca_paralar()
    access_info = load_access_code() or {}

    unique_parameters = []
    seen_names = set()

    for param in parameters:
        param_name = param.get('Parametre Adı', '')
        if param_name and param_name not in seen_names:
            seen_names.add(param_name)
            unique_parameters.append(param)

    # Kullanıcı yönetimi formundan gelen POST isteklerini işle
    if request.method == 'POST':
        # Sadece admin kullanıcı kullanıcı yönetimi yapabilsin
        if session.get('username') != 'admin' and session.get('role') != 'admin':
            flash('Bu işlem için yetkiniz yok.', 'danger')
            return redirect(url_for('ayarlar'))

        current_users = load_users()
        action = request.form.get('action')

        if action == 'add_user':
            username = (request.form.get('username') or '').strip()
            password = (request.form.get('password') or '').strip()
            surname = (request.form.get('surname') or '').strip()
            gorev = (request.form.get('gorev') or '').strip()
            role = (request.form.get('role') or '').strip()
            asg_fiyat = (request.form.get('asgFiyat') or '').strip()

            imza_filename = None
            if 'imza' in request.files and request.files['imza'].filename:
                imza_file = request.files['imza']
                if imza_file and allowed_file(imza_file.filename, {'png', 'jpg', 'jpeg', 'gif'}):
                    filename = secure_filename(imza_file.filename)
                    imza_filename = f"{username}_{filename}"
                    signatures_dir = os.path.join('static', 'images', 'signatures')
                    os.makedirs(signatures_dir, exist_ok=True)
                    imza_file.save(os.path.join(signatures_dir, imza_filename))

            if username and password and username not in current_users:
                current_users[username] = {
                    'password': password,
                    'role': role or '1',
                    'surname': surname,
                    'gorev': gorev,
                    'imza': imza_filename,
                    'asgFiyat': asg_fiyat
                }
                save_users(current_users)
                flash(f'Kullanıcı "{username}" başarıyla eklendi!', 'success')
            else:
                flash('Kullanıcı adı zaten mevcut veya eksik bilgi!', 'danger')

        elif action == 'update_user':
            username = (request.form.get('edit_username') or request.form.get('username') or '').strip()
            if username in current_users:
                if request.form.get('password'):
                    current_users[username]['password'] = request.form.get('password')
                current_users[username]['role'] = (request.form.get('role') or current_users[username].get('role'))
                current_users[username]['surname'] = request.form.get('surname', '')
                current_users[username]['gorev'] = request.form.get('gorev', '')
                current_users[username]['asgFiyat'] = request.form.get('asgFiyat', '')

                if 'imza' in request.files and request.files['imza'].filename:
                    imza_file = request.files['imza']
                    if imza_file and allowed_file(imza_file.filename, {'png', 'jpg', 'jpeg', 'gif'}):
                        filename = secure_filename(imza_file.filename)
                        imza_filename = f"{username}_{filename}"
                        signatures_dir = os.path.join('static', 'images', 'signatures')
                        os.makedirs(signatures_dir, exist_ok=True)
                        old_imza = current_users[username].get('imza')
                        if old_imza:
                            old_imza_path = os.path.join(signatures_dir, old_imza)
                            if os.path.exists(old_imza_path):
                                os.remove(old_imza_path)
                        imza_file.save(os.path.join(signatures_dir, imza_filename))
                        current_users[username]['imza'] = imza_filename

                save_users(current_users)
                flash(f'Kullanıcı "{username}" başarıyla güncellendi!', 'success')
            else:
                flash('Kullanıcı güncellenemedi!', 'danger')

        return redirect(url_for('ayarlar'))

    return render_template(
        'ayarlar.html',
        username=session.get('username'),
        role=session.get('role'),
        users=users,
        parameters=parameters,
        unique_parameters=unique_parameters,
        baca_paralar=baca_paralar,
        access_code=access_info.get('code')
    )


@app.route('/delete_user', methods=['POST'])
def delete_user():
    """Tek bir kullanıcıyı siler (admin hariç)."""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'error': 'Oturum bulunamadı'}), 401

    # Sadece admin veya rolü admin olan kullanıcı silebilsin
    if session.get('role') != 'admin' and session.get('username') != 'admin':
        return jsonify({'success': False, 'error': 'Yetkiniz yok'}), 403

    data = request.get_json(silent=True) or {}
    username = (data.get('username') or '').strip()

    if not username:
        return jsonify({'success': False, 'error': 'Kullanıcı adı belirtilmedi'}), 400
    if username == 'admin':
        return jsonify({'success': False, 'error': 'admin kullanıcısı silinemez'}), 400

    current_users = load_users()
    if username not in current_users:
        return jsonify({'success': False, 'error': 'Kullanıcı bulunamadı'}), 404

    imza_filename = None
    if isinstance(current_users.get(username), dict):
        imza_filename = current_users[username].get('imza')
    if imza_filename:
        signatures_dir = os.path.join('static', 'images', 'signatures')
        imza_path = os.path.join(signatures_dir, imza_filename)
        if os.path.exists(imza_path):
            os.remove(imza_path)

    del current_users[username]
    save_users(current_users)

    flash(f'Kullanıcı "{username}" başarıyla silindi!', 'success')
    return jsonify({'success': True})


@app.route('/delete_selected_users', methods=['POST'])
def delete_selected_users():
    """Birden fazla kullanıcıyı toplu olarak siler (admin hariç)."""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'error': 'Oturum bulunamadı'}), 401

    if session.get('role') != 'admin' and session.get('username') != 'admin':
        return jsonify({'success': False, 'error': 'Yetkiniz yok'}), 403

    data = request.get_json(silent=True) or {}
    usernames = data.get('usernames') or []
    if not usernames:
        return jsonify({'success': False, 'error': 'Kullanıcı listesi boş'}), 400

    current_users = load_users()
    deleted_count = 0
    signatures_dir = os.path.join('static', 'images', 'signatures')

    for raw_username in usernames:
        username = (raw_username or '').strip()
        if not username or username == 'admin':
            continue
        if username in current_users:
            imza_filename = None
            if isinstance(current_users.get(username), dict):
                imza_filename = current_users[username].get('imza')
            if imza_filename:
                imza_path = os.path.join(signatures_dir, imza_filename)
                if os.path.exists(imza_path):
                    os.remove(imza_path)
            del current_users[username]
            deleted_count += 1

    save_users(current_users)
    if deleted_count > 0:
        flash(f'{deleted_count} kullanıcı başarıyla silindi!', 'success')
    else:
        flash('Silinecek kullanıcı bulunamadı.', 'warning')

    return jsonify({'success': True, 'deleted_count': deleted_count})


@app.route('/tum_sonuclar')
def tum_sonuclar():
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    try:
        with open('asama_verileri.json', 'r', encoding='utf-8') as f:
            content = f.read().strip()
            if content:
                asama_verileri = json.loads(content)
            else:
                asama_verileri = []
    except (FileNotFoundError, json.JSONDecodeError):
        asama_verileri = []

    asama1_index = {}
    for r in asama_verileri:
        if r.get('asama') != 1:
            continue
        key = (
            (r.get('firma') or '').strip(),
            (r.get('olcum_kodu') or '').strip(),
            (r.get('baca') or r.get('bacaNo') or '').strip()
        )
        asama1_index[key] = r

    # 2. Aşama verileri (parametre bazlı süre, travers, nozzle, sayaç hacmi vb.)
    try:
        with open('asama2_verileri.json', 'r', encoding='utf-8') as f:
            content = f.read().strip()
            if content:
                asama2_verileri = json.loads(content)
            else:
                asama2_verileri = []
    except (FileNotFoundError, json.JSONDecodeError):
        asama2_verileri = []

    # 3. Aşama verileri
    try:
        with open('asama3_verileri.json', 'r', encoding='utf-8') as f:
            content = f.read().strip()
            if content:
                asama3_verileri = json.loads(content)
            else:
                asama3_verileri = []
    except (FileNotFoundError, json.JSONDecodeError):
        asama3_verileri = []

    # 2. Aşama parametre verileri için index (firma, ölçüm kodu, baca, temiz_parametre_adi)
    def _clean_param_name(name: str) -> str:
        if not name:
            return ''
        s = str(name).split('\n')[0].split('|')[0].strip()
        if '-' in s:
            s = s.split('-')[0].strip()
        return s

    asama2_index = {}
    for a2 in asama2_verileri:
        if a2.get('asama') != 2:
            continue
        firma2 = (a2.get('firma') or '').strip()
        olcum2 = (a2.get('olcum_kodu') or '').strip()
        baca2 = (a2.get('baca') or '').strip()
        if not firma2 or not olcum2 or not baca2:
            continue
        params2 = a2.get('parametreler') or []
        for p2 in params2:
            ad_full = (p2.get('parametre') or p2.get('parametre_adi') or '').strip()
            ad_clean = _clean_param_name(ad_full)
            if not ad_clean:
                continue
            asama2_index[(firma2, olcum2, baca2, ad_clean)] = p2

    sonuc_kayitlari = []
    sira = 1
    for as3 in asama3_verileri:
        if as3.get('asama') != 3:
            continue

        firma = (as3.get('firma') or '').strip()
        olcum_kodu = (as3.get('olcum_kodu') or '').strip()
        baca_adi = (as3.get('baca') or '').strip()
        if not firma or not olcum_kodu or not baca_adi:
            continue

        key = (firma, olcum_kodu, baca_adi)
        asama1 = asama1_index.get(key, {})
        baca_no = (asama1.get('bacaNo') or asama1.get('baca') or '').strip()
        cihaz_no = (asama1.get('cihazSeri') or '').strip()

        raw_tarih = (as3.get('kayit_tarihi') or '').strip()
        if raw_tarih:
            try:
                formatted = format_date_for_eurometric(raw_tarih)
                olcum_tarihi = (formatted or '').replace('/', '-')
            except Exception:
                olcum_tarihi = raw_tarih
        else:
            olcum_tarihi = ''

        parametreler = as3.get('parametreler') or []
        for idx, p in enumerate(parametreler):
            parametre_adi = (p.get('parametre') or p.get('parametre_adi') or '').strip()
            if not parametre_adi:
                continue

            parametre_clean = _clean_param_name(parametre_adi)
            a2_param = asama2_index.get((firma, olcum_kodu, baca_adi, parametre_clean), {})

            hiz = (asama1.get('ortalamaBacaHizi') or '').strip()
            trav = (a2_param.get('traversSayisi') or '').strip()
            nozzle = (a2_param.get('nozzleCapi') or '').strip()
            top_hac = ''
            if a2_param.get('sayacHacmi') is not None:
                top_hac = str(a2_param.get('sayacHacmi')).strip()

            sonuc_kayitlari.append({
                'sira': sira,
                'firma': firma,
                'olcum_kodu': olcum_kodu,
                'baca_adi': baca_adi,
                'baca_no': baca_no,
                'parametre': parametre_clean,
                'baca': baca_adi,
                'param_index': idx,
                'hiz': hiz,
                'trav': trav,
                'nozzle': nozzle,
                'top_hac': top_hac,
                'cihaz_no': cihaz_no,
                'olcum_tarihi': olcum_tarihi,
            })
            sira += 1

    return render_template(
        'tum_sonuclar.html',
        username=session.get('username'),
        role=session.get('role'),
        sonuc_kayitlari=sonuc_kayitlari
    )


@app.route('/parametre_sahabil')
def parametre_sahabil():
    """Parametre Sahabil sayfası"""
    if not session.get('logged_in'):
        return redirect(url_for('login'))
    
    # Parametreleri yükle
    parameters = load_parameters()
    
    # Tekrar eden parametre adlarını kaldır (sadece benzersiz olanları al)
    unique_parameters = []
    seen_names = set()
    
    for param in parameters:
        param_name = param.get('Parametre Adı', '')
        if param_name and param_name not in seen_names:
            seen_names.add(param_name)
            unique_parameters.append(param)
    
    return render_template('parametre_sahabil.html', 
                         username=session.get('username'), 
                         role=session.get('role'),
                         parameters=unique_parameters)
@app.route('/save_parametre_sahabil', methods=['POST'])
def save_parametre_sahabil():
    """Parametre sahabil ölçümlerini kaydeder."""
    try:
        parametre_id = request.form.get('parametre_id')
        parametre_adi = request.form.get('parametre_adi')
        
        if not parametre_id or not parametre_adi:
            return jsonify({'success': False, 'error': 'Eksik parametre bilgisi'})
        
        # Özel parametreler listesi
        specialParams = ['yg', 'voc', 'toc', 'pm10', 'ctoz'];
        isSpecialParam = any(param in parametre_adi.lower() for param in specialParams);
        
        # Ölçüm verilerini hazırla
        olcum_data = {
            'id': str(uuid4()),
            'parametre_id': parametre_id,
            'parametre_adi': parametre_adi,
            'is_special_param': isSpecialParam,
            'created_at': datetime.now().isoformat()
        }
        
        if isSpecialParam:
            # Özel parametreler için basit veri
            olcum_data.update({
                'olcum_tarihi': request.form.get('olcum_tarihi'),
                'olcum_saati': request.form.get('olcum_saati'),
                'olcum_degeri': request.form.get('olcum_degeri'),
                'birim': request.form.get('birim'),
                'notlar': request.form.get('notlar')
            })
        else:
            # Diğer parametreler için detaylı veri
            olcum_data.update({
                'olcum_tarihi': request.form.get('olcum_tarihi'),
                'notlar': request.form.get('notlar'),
                'measurements': {}
            })
            
            # 3 ölçüm için veri topla
            for i in range(1, 4):
                olcum_data['measurements'][f'olcum_{i}'] = {
                    'metot': request.form.get(f'metot_{i}'),
                    'nozzle_cap': request.form.get(f'nozzle_cap_{i}'),
                    'travers': request.form.get(f'travers_{i}'),
                    'b_hiz': request.form.get(f'b_hiz_{i}'),
                    'b_sic': request.form.get(f'b_sic_{i}'),
                    'b_bas': request.form.get(f'b_bas_{i}'),
                    'b_nem_gm3': request.form.get(f'b_nem_gm3_{i}'),
                    'b_nem_yuzde': request.form.get(f'b_nem_yuzde_{i}'),
                    'syc_hac': request.form.get(f'syc_hac_{i}'),
                    'syc_ilk': request.form.get(f'syc_ilk_{i}'),
                    'syc_son': request.form.get(f'syc_son_{i}'),
                    'syc_sic': request.form.get(f'syc_sic_{i}'),
                    'debi': request.form.get(f'debi_{i}'),
                    'isdl': request.form.get(f'isdl_{i}')
                }
        
        # Parametre sahabil verilerini yükle
        parametre_sahabil_data = load_parametre_sahabil()
        parametre_sahabil_data.append(olcum_data)
        
        # Dosyaya kaydet
        if save_parametre_sahabil(parametre_sahabil_data):
            print(f"Parametre sahabil ölçümü kaydedildi: {parametre_adi}")
            return jsonify({
                'success': True, 
                'message': 'Parametre ölçümü başarıyla kaydedildi'
            })
        else:
            return jsonify({'success': False, 'error': 'Veriler kaydedilemedi'}), 500
            
    except Exception as e:
        print(f"Parametre sahabil ölçümü kaydedilirken hata: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/parametre_metodlari')
def api_parametre_metodlari():
    """Parametre metodlarını döndürür."""
    try:
        # Parametreleri yükle
        parameters = load_parameters()
        
        # Parametre adı ve metodunu eşleştir
        parametre_metodlari = {}
        for param in parameters:
            parametre_adi = param.get('Parametre Adı', '')
            metot = param.get('Metot', '')
            if parametre_adi and metot:
                parametre_metodlari[parametre_adi.upper()] = metot
        
        return jsonify(parametre_metodlari)
    except Exception as e:
        print(f"Parametre metodları yüklenirken hata: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/personel_listesi')
def api_personel_listesi():
    """Personel listesini döndürür - sadece görevi 'Saha' olanlar."""
    try:
        # users.json dosyasından personel verilerini yükle
        users_data = load_users()
        
        # Sadece görevi 'Saha' olan personeli filtrele
        saha_personel = []
        for username, user_info in users_data.items():
            # Admin kullanıcısını atla
            if username == 'admin':
                continue
                
            # Görevi 'Saha' olanları al
            if user_info.get('gorev') == 'Saha':
                # Ad ve soyadı birleştir
                ad = username
                soyad = user_info.get('surname', '')
                tam_ad = f"{ad} {soyad}".strip()
                
                saha_personel.append({
                    'id': username,
                    'personel_adi': tam_ad,
                    'gorev': user_info.get('gorev', 'Saha')
                })
        
        return jsonify(saha_personel)
    except Exception as e:
        print(f"Personel listesi yüklenirken hata: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/parametre_sahabil_list')
def api_parametre_sahabil_list():
    """Parametre sahabil ölçümlerini döndürür."""
    try:
        parametre_sahabil_data = load_parametre_sahabil()
        return jsonify(parametre_sahabil_data)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/delete_parametre_sahabil/<record_id>', methods=['DELETE'])
def delete_parametre_sahabil(record_id):
    """Parametre sahabil ölçümünü siler."""
    try:
        parametre_sahabil_data = load_parametre_sahabil()
        
        # Kaydı bul ve sil
        parametre_sahabil_data = [record for record in parametre_sahabil_data if record.get('id') != record_id]
        
        if save_parametre_sahabil(parametre_sahabil_data):
            return jsonify({'success': True, 'message': 'Kayıt başarıyla silindi'})
        else:
            return jsonify({'success': False, 'error': 'Veriler kaydedilemedi'}), 500
            
    except Exception as e:
        print(f"Parametre sahabil silme hatası: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/delete_selected_parametre_sahabil', methods=['POST'])
def delete_selected_parametre_sahabil():
    """Seçilen parametre sahabil ölçümlerini siler."""
    try:
        data = request.get_json()
        selected_ids = data.get('ids', [])
        
        if not selected_ids:
            return jsonify({'success': False, 'error': 'Seçilen kayıt bulunamadı'})
        
        parametre_sahabil_data = load_parametre_sahabil()
        
        # Seçilen kayıtları sil
        parametre_sahabil_data = [record for record in parametre_sahabil_data if record.get('id') not in selected_ids]
        
        if save_parametre_sahabil(parametre_sahabil_data):
            return jsonify({'success': True, 'message': f'{len(selected_ids)} kayıt başarıyla silindi'})
        else:
            return jsonify({'success': False, 'error': 'Veriler kaydedilemedi'}), 500
            
    except Exception as e:
        print(f"Seçilen parametre sahabil silme hatası: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/import_parametre_sahabil', methods=['POST'])
def import_parametre_sahabil():
    """Parametre sahabil verilerini Excel dosyasından içe aktarır."""
    # Session kontrolünü geçici olarak kaldırıyoruz
    # if not session.get('logged_in') or not can_write(session.get('role')):
    #     return jsonify({'success': False, 'error': 'Yetkisiz erişim'}), 403
    
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'Dosya seçilmedi'})
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': 'Dosya seçilmedi'})
        
        if not file.filename.endswith(('.xlsx', '.xls')):
            return jsonify({'success': False, 'error': 'Sadece Excel dosyaları (.xlsx, .xls) kabul edilir'})
        
        # Excel dosyasını oku
        df = pd.read_excel(file)
        
        # Mevcut parametre sahabil verilerini yükle
        parametre_sahabil_data = load_parametre_sahabil()
        
        # Excel verilerini parametre sahabil formatına dönüştür
        imported_count = 0
        print(f"Excel dosyası okundu. Toplam satır: {len(df)}")
        
        for index, row in df.iterrows():
            # Debug: İlk 10 satırı yazdır
            if index < 10:
                print(f"Satır {index}: {[str(val) if pd.notna(val) else 'NaN' for val in row]}")
            
            # Boş satırları atla
            if pd.isna(row.iloc[1]) or str(row.iloc[1]).strip() == '':
                continue
            
            parametre_turu = str(row.iloc[1]).strip()  # Sütun 2 (Unnamed: 1)
            parametre_adi = str(row.iloc[2]).strip() if len(row) > 2 and pd.notna(row.iloc[2]) else ''  # Sütun 3 (Unnamed: 2)
            
            # Ölçüm verilerini topla (sütun 4, 5, 6)
            olcum_verileri = {}
            for i in range(3, min(6, len(row))):  # Sütun 4, 5, 6 (Unnamed: 3, 4, 5)
                if pd.notna(row.iloc[i]) and str(row.iloc[i]).strip() != '':
                    olcum_verileri[f'olcum_{i-2}'] = str(row.iloc[i]).strip()
            
            # Sadece parametre türü ve adı olan satırları kaydet
            if parametre_turu and parametre_adi and parametre_turu not in ['1.ÖLÇ.', '2.ÖLÇ', '3.ÖLÇ.']:
                print(f"Kayıt ekleniyor: {parametre_turu} - {parametre_adi}")
                # Yeni kayıt oluştur
                new_record = {
                    'id': str(uuid4()),
                    'parametre_turu': parametre_turu,
                    'parametre_adi': parametre_adi,
                    'olcum_verileri': olcum_verileri,
                    'imported_from_excel': True,
                    'created_at': datetime.now().isoformat()
                }
                
                parametre_sahabil_data.append(new_record)
                imported_count += 1
        
        # Verileri kaydet
        if save_parametre_sahabil(parametre_sahabil_data):
            return jsonify({
                'success': True, 
                'message': f'{imported_count} adet parametre ölçümü başarıyla içe aktarıldı'
            })
        else:
            return jsonify({'success': False, 'error': 'Veriler kaydedilemedi'}), 500
            
    except Exception as e:
        print(f"Parametre sahabil import hatası: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/export_all_parametre_sahabil')
def export_all_parametre_sahabil():
    """Tüm parametre sahabil verilerini Excel olarak dışa aktarır."""
    # Session kontrolünü geçici olarak kaldırıyoruz
    # if not session.get('logged_in') or not can_read(session.get('role')):
    #     return jsonify({'success': False, 'error': 'Yetkisiz erişim'}), 403
    
    try:
        parametre_sahabil_data = load_parametre_sahabil()
        
        if not parametre_sahabil_data:
            return jsonify({'success': False, 'error': 'Dışa aktarılacak veri bulunamadı'})
        
        # DataFrame oluştur
        export_data = []
        for record in parametre_sahabil_data:
            row = {
                'Parametre Türü': record.get('parametre_turu', ''),
                'Parametre Adı': record.get('parametre_adi', ''),
                '1.ÖLÇ': record.get('olcum_verileri', {}).get('olcum_1', ''),
                '2.ÖLÇ': record.get('olcum_verileri', {}).get('olcum_2', ''),
                '3.ÖLÇ': record.get('olcum_verileri', {}).get('olcum_3', ''),
                'Oluşturma Tarihi': record.get('created_at', '')
            }
            export_data.append(row)
        
        df = pd.DataFrame(export_data)
        
        # Excel dosyası oluştur
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
            df.to_excel(tmp.name, index=False, sheet_name='Parametre Sahabil', engine='openpyxl')
            tmp.flush()
            
            return send_file(
                tmp.name,
                as_attachment=True,
                download_name=f'parametre_sahabil_export_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx',
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            
    except Exception as e:
        print(f"Parametre sahabil export hatası: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/save_parametre_field', methods=['POST'])
def save_parametre_field():
    """Yeni parametre alanını kaydeder."""
    try:
        data = request.get_json()
        parametre_type = data.get('parametreType')
        field_name = data.get('fieldName')
        
        if not parametre_type or not field_name:
            return jsonify({'success': False, 'message': 'Parametre tipi ve alan adı gerekli!'}), 400
        
        # Mevcut alanları yükle
        fields = load_parametre_fields()
        
        # Parametre tipine göre alanları organize et
        if parametre_type not in fields:
            fields[parametre_type] = []
        
        # Alan zaten var mı kontrol et
        if field_name in fields[parametre_type]:
            return jsonify({'success': False, 'message': f'"{field_name}" alanı zaten mevcut!'}), 400
        
        # Yeni alanı ekle
        fields[parametre_type].append(field_name)
        
        # Kaydet
        if save_parametre_fields(fields):
            return jsonify({
                'success': True, 
                'message': f'"{field_name}" alanı {parametre_type.upper()} parametresine başarıyla eklendi!',
                'fieldName': field_name,
                'parametreType': parametre_type
            })
        else:
            return jsonify({'success': False, 'message': 'Alan kaydedilirken hata oluştu!'}), 500
            
    except Exception as e:
        print(f"Parametre alanı kaydedilirken hata: {e}")
        return jsonify({'success': False, 'message': 'Sunucu hatası!'}), 500

@app.route('/get_parametre_fields/<parametre_type>')
def get_parametre_fields(parametre_type):
    """Belirli bir parametre tipinin alanlarını getirir."""
    try:
        fields = load_parametre_fields()
        parametre_fields = fields.get(parametre_type, [])
        return jsonify({'success': True, 'fields': parametre_fields})
    except Exception as e:
        print(f"Parametre alanları getirilirken hata: {e}")
        return jsonify({'success': False, 'message': 'Sunucu hatası!'}), 500

@app.route('/update_parametre_field', methods=['POST'])
def update_parametre_field():
    """Parametre alanını günceller."""
    try:
        data = request.get_json()
        parametre_type = data.get('parametreType')
        old_field_name = data.get('oldFieldName')
        new_field_name = data.get('newFieldName')
        
        if not parametre_type or not old_field_name or not new_field_name:
            return jsonify({'success': False, 'message': 'Tüm alanlar gerekli!'}), 400
        
        # Mevcut alanları yükle
        fields = load_parametre_fields()
        
        # Parametre tipi var mı kontrol et
        if parametre_type not in fields:
            return jsonify({'success': False, 'message': 'Parametre tipi bulunamadı!'}), 404
        
        # Eski alan var mı kontrol et
        if old_field_name not in fields[parametre_type]:
            return jsonify({'success': False, 'message': f'"{old_field_name}" alanı bulunamadı!'}), 404
        
        # Yeni alan adı zaten var mı kontrol et
        if new_field_name in fields[parametre_type] and new_field_name != old_field_name:
            return jsonify({'success': False, 'message': f'"{new_field_name}" alanı zaten mevcut!'}), 400
        
        # Alanı güncelle
        field_index = fields[parametre_type].index(old_field_name)
        fields[parametre_type][field_index] = new_field_name
        
        # Kaydet
        if save_parametre_fields(fields):
            return jsonify({
                'success': True, 
                'message': f'"{old_field_name}" alanı "{new_field_name}" olarak güncellendi!'
            })
        else:
            return jsonify({'success': False, 'message': 'Alan güncellenirken hata oluştu!'}), 500
            
    except Exception as e:
        print(f"Parametre alanı güncellenirken hata: {e}")
        return jsonify({'success': False, 'message': 'Sunucu hatası!'}), 500

@app.route('/delete_parametre_field', methods=['POST'])
def delete_parametre_field():
    """Parametre alanını siler."""
    try:
        data = request.get_json()
        parametre_type = data.get('parametreType')
        field_name = data.get('fieldName')
        
        if not parametre_type or not field_name:
            return jsonify({'success': False, 'message': 'Parametre tipi ve alan adı gerekli!'}), 400
        
        # Mevcut alanları yükle
        fields = load_parametre_fields()
        
        # Parametre tipi var mı kontrol et
        if parametre_type not in fields:
            return jsonify({'success': False, 'message': 'Parametre tipi bulunamadı!'}), 404
        
        # Alan var mı kontrol et
        if field_name not in fields[parametre_type]:
            return jsonify({'success': False, 'message': f'"{field_name}" alanı bulunamadı!'}), 404
        
        # Alanı sil
        fields[parametre_type].remove(field_name)
        
        # Kaydet
        if save_parametre_fields(fields):
            return jsonify({
                'success': True, 
                'message': f'"{field_name}" alanı başarıyla silindi!'
            })
        else:
            return jsonify({'success': False, 'message': 'Alan silinirken hata oluştu!'}), 500
            
    except Exception as e:
        print(f"Parametre alanı silinirken hata: {e}")
        return jsonify({'success': False, 'message': 'Sunucu hatası!'}), 500

@app.route('/import_parametre_fields', methods=['POST'])
def import_parametre_fields():
    """Tablodaki formatı Excel dosyasından içe aktarır."""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'message': 'Dosya seçilmedi!'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'message': 'Dosya seçilmedi!'}), 400
        
        if not file.filename.endswith(('.xlsx', '.xls')):
            return jsonify({'success': False, 'message': 'Sadece Excel dosyaları (.xlsx, .xls) kabul edilir!'}), 400
        
        # Excel dosyasını header olmadan oku
        df = pd.read_excel(file, header=None)
        
        # Tablodaki formatı kontrol et
        if len(df.columns) < 8:
            return jsonify({'success': False, 'message': 'Dosya formatı uygun değil! En az 8 sütun olmalı.'}), 400
        
        # Başlık satırını kontrol et
        first_row = df.iloc[0].tolist()
        expected_headers = [
            'TOZ,A.MET,SÜLF.A,HF,HCL,AMON.,FORM.,CR6,FOSF.A.,HCN',
            'NEM', 'YG', 'TOC', 'VOC', 'PM10', 'ÇT', 'boş', 'boş'
        ]
        
        # Başlık kontrolü (ilk 6 sütun için)
        for i in range(6):
            if str(first_row[i]).strip() != expected_headers[i]:
                return jsonify({'success': False, 'message': f'Başlık formatı uygun değil! Sütun {i+1}: {first_row[i]}'}), 400
        
        # TARİH satırını kontrol et
        second_row = df.iloc[1].tolist()
        for i in range(6):
            if str(second_row[i]).strip() != 'TARİH':
                return jsonify({'success': False, 'message': f'TARİH satırı uygun değil! Sütun {i+1}: {second_row[i]}'}), 400
        
        # Veri satırlarını al (3. satırdan itibaren)
        data_rows = []
        for i in range(2, len(df)):
            row = df.iloc[i].tolist()
            # Boş satırları atla
            if any(str(cell).strip() for cell in row[:6]):
                data_rows.append([str(cell).strip() if pd.notna(cell) else '' for cell in row])
        
        # Yeni alanlar için dictionary oluştur
        new_fields = {
            'genel': [], 'nem': [], 'yg': [], 'toc': [], 'voc': [], 'pm10': [], 'ct': []
        }
        
        # Her sütun için alanları topla
        for row in data_rows:
            if len(row) >= 8:
                # 1. sütun: genel parametreler
                if row[0]:
                    new_fields['genel'].append(row[0])
                
                # 2. sütun: nem parametreleri
                if row[1]:
                    new_fields['nem'].append(row[1])
                
                # 3. sütun: yg parametreleri
                if row[2]:
                    new_fields['yg'].append(row[2])
                
                # 4. sütun: toc parametreleri
                if row[3]:
                    new_fields['toc'].append(row[3])
                
                # 5. sütun: voc parametreleri
                if row[4]:
                    new_fields['voc'].append(row[4])
                
                # 6. sütun: pm10 parametreleri
                if row[5]:
                    new_fields['pm10'].append(row[5])
                
                # 7. sütun: ct parametreleri
                if row[6]:
                    new_fields['ct'].append(row[6])
        
        # Boş listeleri temizle
        new_fields = {k: v for k, v in new_fields.items() if v}
        
        # Kaydet
        save_parametre_fields(new_fields)
        
        total_fields = sum(len(fields) for fields in new_fields.values())
        return jsonify({
            'success': True, 
            'message': f'{total_fields} alan başarıyla içe aktarıldı!'
        })
            
    except Exception as e:
        print(f"Parametre formatı içe aktarılırken hata: {e}")
        return jsonify({'success': False, 'message': f'Dosya okuma hatası: {str(e)}'}), 500

@app.route('/export_parametre_fields')
def export_parametre_fields():
    """Tablodaki formatı Excel dosyasına dışa aktarır."""
    try:
        # Tabloda gördüğünüz formatı tam olarak oluştur
        excel_data = []
        
        # Başlık satırı
        excel_data.append([
            'TOZ,A.MET,SÜLF.A,HF,HCL,AMON.,FORM.,CR6,FOSF.A.,HCN',
            'NEM', 'YG', 'TOC', 'VOC', 'PM10', 'ÇT', 'boş', 'boş'
        ])
        
        # TARİH satırı
        excel_data.append(['TARİH', 'TARİH', 'TARİH', 'TARİH', 'TARİH', 'TARİH', 'TARİH', '', ''])
        
        # Veri satırları - tablodaki gördüğünüz sırayla
        data_rows = [
            ['METOT', '1İMP-İ', 'B.SIC', 'B.NEM', 'GAZ HAC.', 'METOT', 'T.İÇİ-1', '', ''],
            ['NOZZLE ÇAP', '1-İMP-S', 'O2', 'B.SIC', 'GAZ.SIC.', 'ORT.SIC', 'T.İÇİ-2', '', ''],
            ['TRAVERS', '2-İMP-İ', 'CO', 'TOC(PPM)', 'SEY.GAZ.HAC', 'ORT.NEM', 'T.İÇİ-3', '', ''],
            ['B.HIZ', '2-İMP-S', 'NO', 'KK1-SPAN', 'SEY.GAZ.SIC', 'ORT.RUZ.HIZ', 'T.İÇİ-4', '', ''],
            ['B.SIC', '3-İMP-İ', 'NOX', 'KK1-SPAN', '', 'ÇEK.HACİM', 'T-DIŞ-1', '', ''],
            ['B.BAS(KPA)', '3-İMP-S', 'SO2', 'KK2-SPAN', '', 'SYC.İLK', 'T-DIŞ-2', '', ''],
            ['B.NEM(G/M3)', 'HAC.', 'KK1-O2', 'KK2-SPAN', '', 'SYC.SON', 'T-DIŞ-3', '', ''],
            ['B.NEM(%)', '', 'KK1-CO', '', '', 'ISDL', 'TDİS-4', '', ''],
            ['SYC.HAC.', '', 'KK1-NO', '', '', '', 'İLK KURULUM', '', ''],
            ['SYC.İLK', '', 'KK1-SO2', '', '', '', '2. KURULUM', '', ''],
            ['SYC.SON', '', 'KK2-O2', '', '', '', 'TOPLAMA', '', ''],
            ['SYC.SIC', '', 'KK2-CO', '', '', '', '', '', ''],
            ['DEBİ', '', 'KK2-NO', '', '', '', '', '', ''],
            ['ISDL', '', 'KK2-SO2', '', '', '', '', '', ''],
            ['', '', 'T90', '', '', '', '', '', '']
        ]
        
        excel_data.extend(data_rows)
        
        # DataFrame oluştur
        df = pd.DataFrame(excel_data)
        
        # Excel dosyası oluştur
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
            df.to_excel(tmp.name, index=False, header=False, engine='openpyxl')
            
            return send_file(
                tmp.name,
                as_attachment=True,
                download_name=f'parametre_format_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx',
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            
    except Exception as e:
        print(f"Parametre formatı dışa aktarılırken hata: {e}")
        return jsonify({'success': False, 'message': 'Dışa aktarma hatası!'}), 500


@app.route('/clear_parametre_fields', methods=['POST'])
def clear_parametre_fields():
    """Tüm parametre alanlarını siler."""
    try:
        # Boş dictionary oluştur
        empty_fields = {}
        
        # Kaydet
        save_parametre_fields(empty_fields)
        
        return jsonify({
            'success': True, 
            'message': 'Tüm parametre alanları başarıyla silindi!'
        })
            
    except Exception as e:
        print(f"Parametre alanları silinirken hata: {e}")
        return jsonify({'success': False, 'message': f'Silme hatası: {str(e)}'}), 500

@app.route('/save_parametre_fields_table', methods=['POST'])
def save_parametre_fields_table():
    """Parametre sahabil tablosundaki verileri kaydeder."""
    try:
        data = request.get_json()
        rows = data.get('rows', [])
        
        # Parametre sahabil verilerini kaydet
        if save_parametre_sahabil(rows):
            return jsonify({
                'success': True, 
                'message': f'{len(rows)} satır başarıyla kaydedildi!'
            })
        else:
            return jsonify({'success': False, 'message': 'Veriler kaydedilirken hata oluştu!'}), 500
            
    except Exception as e:
        print(f"Parametre sahabil tablosu kaydedilirken hata: {e}")
        return jsonify({'success': False, 'message': f'Kaydetme hatası: {str(e)}'}), 500

@app.route('/api/save_asama_data', methods=['POST'])
def save_asama_data():
    """Aşama verilerini kaydeder (firma-olcum_kodu-baca kombinasyonu için en son kayıt)."""
    try:
        data = request.get_json()
        
        # Aşama verilerini dosyaya kaydet - trim ile normalize et
        asama_data = {
            'firma': data.get('firma', '').strip(),
            'olcum_kodu': data.get('olcum_kodu', '').strip(),
            'baca': data.get('baca', '').strip(),
            'personel': data.get('personel', '').strip(),
            'asama': data.get('asama', 1),
            'cihazSeri': data.get('cihazSeri', ''),
            'bacaNo': data.get('bacaNo', ''),
            'ortalamaBacaHizi': data.get('ortalamaBacaHizi', ''),
            'bacaSicakligi': data.get('bacaSicakligi', ''),
            'bacaMutlakBasinci': data.get('bacaMutlakBasinci', ''),
            'nem': data.get('nem', ''),
            'sayacGazSicakligi': data.get('sayacGazSicakligi', ''),
            'atmosferBasinci': data.get('atmosferBasinci', ''),
            'kayit_tarihi': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        }
        
        # Mevcut verileri oku - JSONDecodeError yakala
        try:
            with open('asama_verileri.json', 'r', encoding='utf-8') as f:
                content = f.read().strip()
                if not content:
                    mevcut_veriler = []
                else:
                    mevcut_veriler = json.loads(content)
        except (FileNotFoundError, json.JSONDecodeError):
            mevcut_veriler = []
        
        # Aynı firma-olcum_kodu-baca-asama kombinasyonu için eski kayıtları sil
        mevcut_veriler = [
            record for record in mevcut_veriler
            if not (record.get('firma', '').strip() == asama_data['firma'] and
                    record.get('olcum_kodu', '').strip() == asama_data['olcum_kodu'] and
                    record.get('baca', '').strip() == asama_data['baca'] and
                    record.get('asama') == asama_data['asama'])
        ]
        
        # Yeni veriyi ekle
        mevcut_veriler.append(asama_data)
        
        # Dosyaya kaydet
        with open('asama_verileri.json', 'w', encoding='utf-8') as f:
            json.dump(mevcut_veriler, f, ensure_ascii=False, indent=2)
        
        return jsonify({
            'success': True,
            'message': f'{asama_data["firma"]} - Aşama {asama_data["asama"]} verisi kaydedildi!'
        })
        
    except Exception as e:
        print(f"Aşama verisi kaydedilirken hata: {e}")
        return jsonify({'success': False, 'message': f'Kaydetme hatası: {str(e)}'}), 500

@app.route('/api/get_latest_asama_data')
def get_latest_asama_data():
    """En son kaydedilen aşama verisini getirir."""
    try:
        # Dosyayı oku
        try:
            with open('asama_verileri.json', 'r', encoding='utf-8') as f:
                veriler = json.load(f)
            
            if veriler:
                # En son kaydedilen veriyi döndür
                latest_data = veriler[-1]
                return jsonify({
                    'success': True,
                    'data': latest_data
                })
            else:
                return jsonify({
                    'success': False,
                    'message': 'Henüz kayıtlı veri bulunamadı'
                })
                
        except FileNotFoundError:
            return jsonify({
                'success': False,
                'message': 'Veri dosyası bulunamadı'
            })
            
    except Exception as e:
        print(f"Aşama verisi getirilirken hata: {e}")
        return jsonify({'success': False, 'message': f'Veri getirme hatası: {str(e)}'}), 500

@app.route('/api/baca_parameters')
def get_baca_parameters():
    """Belirtilen firma, ölçüm kodu ve bacadaki parametreleri getirir."""
    try:
        firma = request.args.get('firma')
        olcum_kodu = request.args.get('olcum_kodu')
        baca_no = request.args.get('baca_no')
        
        print(f"[API] /api/baca_parameters çağrıldı - Firma={firma}, Ölçüm={olcum_kodu}, Baca={baca_no}")
        
        if not firma or not olcum_kodu or not baca_no:
            print(f"[API] Eksik parametreler - Firma={firma}, Ölçüm={olcum_kodu}, Baca={baca_no}")
            return jsonify({'success': False, 'message': 'Firma, ölçüm kodu ve baca numarası gerekli'}), 400
        
        print(f"[API] Parametre arama: Firma={firma}, Ölçüm Kodu={olcum_kodu}, Baca={baca_no}")
        
        # Baca bilgilerini oku
        try:
            baca_bilgileri_path = data_path('baca_bilgileri.json')
            print(f"[API] Baca bilgileri dosya yolu: {baca_bilgileri_path}")
            with open(baca_bilgileri_path, 'r', encoding='utf-8') as f:
                baca_bilgileri = json.load(f)
            print(f"[API] Baca bilgileri yüklendi: {len(baca_bilgileri)} kayıt")
        except FileNotFoundError as e:
            print(f"[API] Baca bilgileri dosyası bulunamadı: {e}")
            # Dosya yoksa boş liste ile devam et
            baca_bilgileri = []
        
        # Firma, ölçüm kodu ve baca numarasına göre parametreleri bul
        parameters = set()
        for record in baca_bilgileri:
            # Hem eski format (firma) hem yeni format (firma_adi) destekle
            record_firma = record.get('firma_adi', record.get('firma', '')).strip()
            record_olcum = str(record.get('olcum_kodu', '')).strip()
            record_baca = str(record.get('baca_no', '')).strip()
            
            # Karşılaştırma yaparken büyük/küçük harf duyarlılığını kaldır
            if (record_firma.upper() == firma.upper() and 
                record_olcum == str(olcum_kodu).strip() and 
                record_baca.upper() == baca_no.upper()):
                parametre = record.get('parametre', '').strip()
                if parametre:
                    parameters.add(parametre)
                    print(f"Parametre bulundu (baca_bilgileri): {parametre}")
        
        # Eğer baca_bilgileri.json içinde kayıt yoksa, firma_olcum.json'daki
        # baca_parametreleri haritasından geri dönüş (fallback) yap
        if not parameters:
            print(f"Hiç parametre bulunamadı, firma_olcum.json üzerinden aranıyor... Firma={firma}, Ölçüm={olcum_kodu}, Baca={baca_no}")
            try:
                with open('firma_olcum.json', 'r', encoding='utf-8') as f:
                    firma_kayitlari = json.load(f)
            except FileNotFoundError:
                firma_kayitlari = []
            
            for kayit in firma_kayitlari:
                rec_firma = str(kayit.get('firma_adi', '')).strip()
                rec_olcum = str(kayit.get('olcum_kodu', '')).strip()
                if rec_firma.upper() == firma.upper() and rec_olcum == str(olcum_kodu).strip():
                    baca_haritasi = kayit.get('baca_parametreleri') or {}
                    for baca_adi, param_list in baca_haritasi.items():
                        if str(baca_adi).strip().upper() == baca_no.upper():
                            if isinstance(param_list, list):
                                for p in param_list:
                                    p_str = str(p).strip()
                                    if p_str:
                                        parameters.add(p_str)
                                        print(f"Parametre bulundu (firma_olcum fallback): {p_str}")
                            else:
                                p_str = str(param_list).strip()
                                if p_str:
                                    parameters.add(p_str)
                                    print(f"Parametre bulundu (firma_olcum fallback): {p_str}")
        
        if not parameters:
            print(f"[API] ⚠️ Hiç parametre bulunamadı! Firma={firma}, Ölçüm={olcum_kodu}, Baca={baca_no}")
            if baca_bilgileri:
                print(f"[API] İlk 3 baca_bilgileri kaydı:")
                for i, rec in enumerate(baca_bilgileri[:3]):
                    print(f"  [{i}] Firma: {rec.get('firma_adi', rec.get('firma', 'N/A'))}, Ölçüm: {rec.get('olcum_kodu', 'N/A')}, Baca: {rec.get('baca_no', 'N/A')}, Parametre: {rec.get('parametre', 'N/A')}")
            # Varsayılan parametreler döndür (boş liste yerine)
            print(f"[API] Varsayılan parametreler döndürülüyor: ['Toz_EPA5', 'AĞIR MET_EPA 29']")
            parameters = ['Toz_EPA5', 'AĞIR MET_EPA 29']
        else:
            print(f"[API] ✅ {len(parameters)} parametre bulundu: {list(parameters)}")
        
        return jsonify({
            'success': True,
            'parameters': list(parameters)
        })
        
    except Exception as e:
        print(f"[API] ❌ Baca parametreleri getirilirken hata: {e}")
        import traceback
        traceback.print_exc()
        # Hata durumunda bile varsayılan parametreler döndür
        print(f"[API] Hata nedeniyle varsayılan parametreler döndürülüyor")
        return jsonify({
            'success': True,
            'parameters': ['Toz_EPA5', 'AĞIR MET_EPA 29'],
            'warning': f'Veri yükleme hatası, varsayılan parametreler kullanıldı: {str(e)}'
        })
@app.route('/api/save_asama2_data', methods=['POST'])
def save_asama2_data():
    """2. Aşama verilerini kaydeder."""
    try:
        data = request.get_json()
        
        # 2. Aşama verilerini dosyaya kaydet - trim ile normalize et
        # Tarih formatını gg/aa/yy ss:dd formatına çevir
        def format_tarih(tarih_str):
            if not tarih_str:
                return tarih_str
            try:
                # gg.aa.yyyy ss:dd formatını gg/aa/yy ss:dd formatına çevir
                if '.' in tarih_str and len(tarih_str.split('.')[0]) == 2:
                    return tarih_str.replace('.', '/')[:8] + tarih_str[8:]
                # gg/aa/yyyy ss:dd formatını gg/aa/yy ss:dd formatına çevir
                elif '/' in tarih_str:
                    parts = tarih_str.split('/')
                    if len(parts) >= 3:
                        year_part = parts[2].split(' ')[0]
                        if len(year_part) == 4:  # gg/aa/yyyy formatı
                            time_part = tarih_str[tarih_str.find(' '):] if ' ' in tarih_str else ''
                            return f"{parts[0]}/{parts[1]}/{year_part[:2]}{time_part}"
                    return tarih_str
                # yyyy-mm-dd formatını gg/aa/yy formatına çevir
                elif '-' in tarih_str and len(tarih_str.split('-')[0]) == 4:
                    dt = datetime.strptime(tarih_str, '%Y-%m-%d %H:%M:%S')
                    return dt.strftime('%d/%m/%y %H:%M')
                else:
                    return tarih_str
            except:
                return tarih_str
        
        # Parametreler içindeki tarihleri formatla
        parametreler = data.get('parametreler', [])
        for param in parametreler:
            if 'baslangic' in param:
                param['baslangic'] = format_tarih(param['baslangic'])
        
        asama2_data = {
            'firma': data.get('firma', '').strip(),
            'olcum_kodu': data.get('olcumKodu', '').strip(),
            'baca': data.get('baca', '').strip(),
            'personel': data.get('personel', '').strip(),
            'asama': 2,
            'parametreler': parametreler,
            'kayit_tarihi': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        }
        
        # Mevcut verileri oku - JSONDecodeError yakala
        try:
            with open('asama2_verileri.json', 'r', encoding='utf-8') as f:
                content = f.read().strip()
                if not content:
                    mevcut_veriler = []
                else:
                    mevcut_veriler = json.loads(content)
        except (FileNotFoundError, json.JSONDecodeError):
            mevcut_veriler = []
        
        # Aynı firma-olcum_kodu-baca kombinasyonu için mevcut kaydı bul ve güncelle
        updated = False
        for i, record in enumerate(mevcut_veriler):
            if (record.get('firma', '').strip() == asama2_data['firma'] and 
                record.get('olcum_kodu', '').strip() == asama2_data['olcum_kodu'] and
                record.get('baca', '').strip() == asama2_data['baca'] and 
                record.get('asama') == 2):
                mevcut_veriler[i] = asama2_data
                updated = True
                break
        
        # Yeni kayıt ekle
        if not updated:
            mevcut_veriler.append(asama2_data)
        
        # Dosyaya kaydet
        with open('asama2_verileri.json', 'w', encoding='utf-8') as f:
            json.dump(mevcut_veriler, f, ensure_ascii=False, indent=2)
        
        return jsonify({
            'success': True,
            'message': f'{asama2_data["firma"]} - Baca {asama2_data["baca"]} - 2. Aşama verisi kaydedildi!'
        })
        
    except Exception as e:
        print(f"2. Aşama verisi kaydedilirken hata: {e}")
        return jsonify({'success': False, 'message': f'Kaydetme hatası: {str(e)}'}), 500

@app.route('/api/get_asama2_data')
def get_asama2_data():
    """Belirtilen firma-olcum_kodu-baca için 2. Aşama verisini getirir (en son kayıt)."""
    try:
        baca_no = request.args.get('baca_no', '').strip()
        firma = request.args.get('firma', '').strip()
        olcum_kodu = request.args.get('olcum_kodu', '').strip()
        
        if not baca_no:
            return jsonify({'success': False, 'message': 'Baca numarası gerekli'}), 400
        
        # Dosyayı oku - JSONDecodeError yakala
        try:
            with open('asama2_verileri.json', 'r', encoding='utf-8') as f:
                content = f.read().strip()
                if not content:
                    veriler = []
                else:
                    veriler = json.loads(content)
        except (FileNotFoundError, json.JSONDecodeError):
            return jsonify({
                'success': False,
                'message': 'Veri dosyası bulunamadı veya bozuk'
            })
        
        # Belirtilen firma-olcum_kodu-baca kombinasyonu için veriyi bul - en son kaydı döndür (kayit_tarihi'ne göre)
        matching_records = []
        for record in veriler:
            if (record.get('asama') == 2 and
                str(record.get('baca', '')).strip() == baca_no and
                (not firma or str(record.get('firma', '')).strip() == firma) and
                (not olcum_kodu or str(record.get('olcum_kodu', '')).strip() == olcum_kodu)):
                matching_records.append(record)
        
        if matching_records:
            # En son kaydı döndür (kayit_tarihi'ne göre sırala)
            latest_record = max(matching_records, key=lambda r: r.get('kayit_tarihi', ''))
            return jsonify({
                'success': True,
                'data': latest_record
            })
        
        return jsonify({
            'success': False,
            'message': 'Veri bulunamadı'
        })
            
    except Exception as e:
        print(f"2. Aşama verisi getirilirken hata: {e}")
        return jsonify({'success': False, 'message': f'Veri getirme hatası: {str(e)}'}), 500

@app.route('/api/get_existing_asama1_data', methods=['GET'])
def get_existing_asama1_data():
    """Mevcut 1. Aşama verilerini getir (firma, ölçüm kodu, baca bazında - en son kayıt)"""
    try:
        firma = request.args.get('firma', '').strip()
        olcum_kodu = request.args.get('olcum_kodu', '').strip()
        baca = request.args.get('baca', '').strip()
        
        if not firma or not olcum_kodu or not baca:
            return jsonify({'success': False, 'message': 'Firma, ölçüm kodu ve baca gerekli'})
        
        # asama_verileri.json dosyasını oku - JSONDecodeError yakala
        if os.path.exists('asama_verileri.json'):
            try:
                with open('asama_verileri.json', 'r', encoding='utf-8') as f:
                    content = f.read().strip()
                    if not content:
                        all_data = []
                    else:
                        all_data = json.loads(content)
            except json.JSONDecodeError:
                return jsonify({'success': False, 'data': None, 'message': 'Dosya bozuk'})
        else:
            return jsonify({'success': False, 'data': None})
        
        # Aynı firma-olcum_kodu-baca için eşleşen kayıtları bul
        matching_records = []
        for record in all_data:
            if (str(record.get('firma', '')).strip() == firma and 
                str(record.get('olcum_kodu', '')).strip() == olcum_kodu and 
                str(record.get('baca', '')).strip() == baca and
                record.get('asama') == 1):
                matching_records.append(record)
        
        if matching_records:
            # Önce ortalamaBacaHizi dolu olan kayıtları tercih et
            def has_speed(rec):
                val = rec.get('ortalamaBacaHizi', '')
                return str(val).strip() != ''

            records_with_speed = [r for r in matching_records if has_speed(r)]
            source = records_with_speed or matching_records

            # En son kaydı döndür (kayit_tarihi'ne göre sırala)
            latest_record = max(source, key=lambda r: r.get('kayit_tarihi', ''))
            return jsonify({'success': True, 'data': latest_record})
        
        return jsonify({'success': False, 'data': None})
        
    except Exception as e:
        print(f"Mevcut 1. Aşama verileri getirme hatası: {e}")
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/get_existing_asama2_data', methods=['GET'])
def get_existing_asama2_data():
    """Mevcut 2. Aşama verilerini getir (firma, ölçüm kodu, baca bazında - en son kayıt)"""
    try:
        firma = request.args.get('firma', '').strip()
        olcum_kodu = request.args.get('olcum_kodu', '').strip()
        baca = request.args.get('baca', '').strip()
        
        if not firma or not olcum_kodu or not baca:
            return jsonify({'success': False, 'message': 'Firma, ölçüm kodu ve baca gerekli'})
        
        # asama2_verileri.json dosyasını oku - JSONDecodeError yakala
        if os.path.exists('asama2_verileri.json'):
            try:
                with open('asama2_verileri.json', 'r', encoding='utf-8') as f:
                    content = f.read().strip()
                    if not content:
                        all_data = []
                    else:
                        all_data = json.loads(content)
            except json.JSONDecodeError:
                return jsonify({'success': False, 'data': None, 'message': 'Dosya bozuk'})
        else:
            return jsonify({'success': False, 'data': None})
        
        # Aynı firma-olcum_kodu-baca için eşleşen kayıtları bul
        matching_records = []
        for record in all_data:
            if (str(record.get('firma', '')).strip() == firma and 
                str(record.get('olcum_kodu', '')).strip() == olcum_kodu and 
                str(record.get('baca', '')).strip() == baca and
                record.get('asama') == 2):
                matching_records.append(record)
        
        if matching_records:
            # En son kaydı döndür (kayit_tarihi'ne göre sırala)
            latest_record = max(matching_records, key=lambda r: r.get('kayit_tarihi', ''))
            return jsonify({'success': True, 'data': latest_record})
        
        return jsonify({'success': False, 'data': None})
        
    except Exception as e:
        print(f"Mevcut 2. Aşama verileri getirme hatası: {e}")
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/save_parametre_table', methods=['POST'])
def save_parametre_table():
    """Parametre tablosu verilerini kaydet"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'success': False, 'message': 'Veri bulunamadı'})
        
        # Parametre tablosu verilerini parametre_tablosu.json dosyasına kaydet
        with open('parametre_tablosu.json', 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        
        return jsonify({'success': True, 'message': 'Parametre tablosu başarıyla kaydedildi'})
        
    except Exception as e:
        print(f"Parametre tablosu kaydetme hatası: {e}")
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/load_parametre_table', methods=['GET'])
def load_parametre_table():
    """Parametre tablosu verilerini yükle"""
    try:
        # parametre_tablosu.json dosyasını oku
        if os.path.exists('parametre_tablosu.json'):
            with open('parametre_tablosu.json', 'r', encoding='utf-8') as f:
                data = json.load(f)
            return jsonify({'success': True, 'data': data})
        else:
            return jsonify({'success': True, 'data': {'headers': [], 'rows': []}})
            
    except Exception as e:
        print(f"Parametre tablosu yükleme hatası: {e}")
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/save_excel_data', methods=['POST'])
def save_excel_data():
    """Excel benzeri tablo verilerini kaydet"""
    try:
        data = request.get_json()
        
        if not data or 'data' not in data:
            return jsonify({'success': False, 'message': 'Veri bulunamadı'})
        
        # Excel verilerini excel_data.json dosyasına kaydet
        with open('excel_data.json', 'w', encoding='utf-8') as f:
            json.dump(data['data'], f, ensure_ascii=False, indent=2)
        
        return jsonify({'success': True, 'message': 'Excel verileri başarıyla kaydedildi'})
        
    except Exception as e:
        print(f"Excel verileri kaydetme hatası: {e}")
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/load_excel_data', methods=['GET'])
def load_excel_data():
    """Excel benzeri tablo verilerini yükle"""
    try:
        # excel_data.json dosyasını oku
        if os.path.exists('excel_data.json'):
            with open('excel_data.json', 'r', encoding='utf-8') as f:
                data = json.load(f)
            return jsonify({'success': True, 'data': data, 'excelData': data})
        else:
            return jsonify({'success': True, 'data': [], 'excelData': []})
            
    except Exception as e:
        print(f"Excel verileri yükleme hatası: {e}")
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/get_excel_data', methods=['GET'])
def get_excel_data():
    """Excel verilerini sessionStorage için yükle (alias endpoint)"""
    return load_excel_data()

@app.route('/api/save_rapor_table', methods=['POST'])
def save_rapor_table():
    """Rapor tablosu verilerini kaydet"""
    try:
        payload = request.get_json()
        if not payload or 'data' not in payload:
            return jsonify({'success': False, 'message': 'Veri bulunamadı'})

        with open('rapor_table.json', 'w', encoding='utf-8') as f:
            json.dump(payload['data'], f, ensure_ascii=False, indent=2)

        return jsonify({'success': True, 'message': 'Rapor verileri kaydedildi'})
    except Exception as e:
        print(f"Rapor verileri kaydetme hatası: {e}")
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/load_rapor_table', methods=['GET'])
def load_rapor_table():
    """Rapor tablosu verilerini yükle"""
    try:
        if os.path.exists('rapor_table.json'):
            with open('rapor_table.json', 'r', encoding='utf-8') as f:
                data = json.load(f)
            return jsonify({'success': True, 'data': data})
        return jsonify({'success': True, 'data': []})
    except Exception as e:
        print(f"Rapor verileri yükleme hatası: {e}")
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/api/rapor2', methods=['GET', 'POST'])
def rapor2_notes():
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Yetkisiz'}), 401
    if request.method == 'GET':
        content = load_rapor2_content()
        return jsonify({'success': True, 'content': content})
    data = request.get_json() or {}
    content = data.get('content', '')
    if save_rapor2_content(content):
        return jsonify({'success': True, 'message': 'Kaydedildi'})
    return jsonify({'success': False, 'message': 'Kaydedilemedi'}), 500

@app.route('/api/baca_paralar', methods=['GET'])
def get_baca_paralar():
    """Baca parametreleri listesini getir (cihaz seri no için)"""
    try:
        # baca_paralar.json dosyasını oku
        if os.path.exists('baca_paralar.json'):
            with open('baca_paralar.json', 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Cihaz Seri parametresini bul
            for item in data:
                if item.get('baca_par_adi') == 'Cihaz Seri':
                    liste_icerigi = item.get('liste_icerigi', '')
                    if liste_icerigi:
                        # Virgülle ayrılmış değerleri listeye çevir
                        cihaz_listesi = [cihaz.strip() for cihaz in liste_icerigi.split(',') if cihaz.strip()]
                        return jsonify({'success': True, 'data': cihaz_listesi})
            
            return jsonify({'success': False, 'message': 'Cihaz Seri parametresi bulunamadı'})
        
        return jsonify({'success': False, 'message': 'baca_paralar.json dosyası bulunamadı'})
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Hata: {str(e)}'})

@app.route('/export_selected_parametre_sahabil')
def export_selected_parametre_sahabil():
    """Seçilen parametre sahabil verilerini Excel olarak dışa aktarır."""
    # Session kontrolünü geçici olarak kaldırıyoruz
    # if not session.get('logged_in') or not can_read(session.get('role')):
    #     return jsonify({'success': False, 'error': 'Yetkisiz erişim'}), 403
    
    try:
        selected_ids = request.args.get('ids', '').split(',')
        if not selected_ids or selected_ids[0] == '':
            return jsonify({'success': False, 'error': 'Seçilen kayıt bulunamadı'})
        
        parametre_sahabil_data = load_parametre_sahabil()
        
        # Seçilen kayıtları filtrele
        selected_data = [record for record in parametre_sahabil_data if record.get('id') in selected_ids]
        
        if not selected_data:
            return jsonify({'success': False, 'error': 'Seçilen kayıtlar bulunamadı'})
        
        # DataFrame oluştur
        export_data = []
        for record in selected_data:
            row = {
                'Parametre Türü': record.get('parametre_turu', ''),
                'Parametre Adı': record.get('parametre_adi', ''),
                '1.ÖLÇ': record.get('olcum_verileri', {}).get('olcum_1', ''),
                '2.ÖLÇ': record.get('olcum_verileri', {}).get('olcum_2', ''),
                '3.ÖLÇ': record.get('olcum_verileri', {}).get('olcum_3', ''),
                'Oluşturma Tarihi': record.get('created_at', '')
            }
            export_data.append(row)
        
        df = pd.DataFrame(export_data)
        
        # Excel dosyası oluştur
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
            df.to_excel(tmp.name, index=False, sheet_name='Parametre Sahabil', engine='openpyxl')
            tmp.flush()
            
            return send_file(
                tmp.name,
                as_attachment=True,
                download_name=f'parametre_sahabil_selected_export_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx',
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            
    except Exception as e:
        print(f"Seçilen parametre sahabil export hatası: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/baca_bilgileri')
def baca_bilgileri():
    if not session.get('logged_in') or not can_read(session.get('role')):
        return redirect(url_for('login'))
    
    baca_bilgileri_list = load_baca_bilgileri()
    return render_template('baca_bilgileri.html', 
                         username=session.get('username'), 
                         role=session.get('role'),
                         baca_bilgileri=baca_bilgileri_list)

@app.route('/baca_bilgileri/add', methods=['POST'])
def add_baca_bilgisi():
    if not session.get('logged_in') or not can_write(session.get('role')):
        return redirect(url_for('login'))
    
    try:
        baca_bilgileri_list = load_baca_bilgileri()
        
        new_baca = {
            'id': str(uuid4()),
            'firma_adi': request.form.get('firma_adi', ''),
            'olcum_kodu': request.form.get('olcum_kodu', ''),
            'baca_no': request.form.get('baca_no', '').strip(),
            'yakit_turu': request.form.get('yakit_turu', ''),
            'isil_guc': request.form.get('isil_guc', ''),
            'cati_sekli': request.form.get('cati_sekli', ''),
            'kaynak_turu': request.form.get('kaynak_turu', ''),
            'baca_sekli': request.form.get('baca_sekli', ''),
            'baca_olcusu': request.form.get('baca_olcusu', ''),
            'yerden_yuk': request.form.get('yerden_yuk', ''),
            'cati_yuk': request.form.get('cati_yuk', ''),
            'ruzgar_hiz': request.form.get('ruzgar_hiz', ''),
            'ort_sic': request.form.get('ort_sic', ''),
            'ort_nem': request.form.get('ort_nem', ''),
            'ort_bas': request.form.get('ort_bas', ''),
            'a_baca': request.form.get('a_baca', ''),
            'b_baca': request.form.get('b_baca', ''),
            'c_delik': request.form.get('c_delik', ''),
            'foto': ''  # Fotoğraf işlemi daha sonra eklenecek
        }
        
        if not new_baca['baca_no']:
            flash('Baca numarası zorunludur!', 'error')
            return redirect(url_for('baca_bilgileri'))
        
        baca_bilgileri_list.append(new_baca)
        
        if save_baca_bilgileri(baca_bilgileri_list):
            flash('Baca bilgisi başarıyla eklendi!', 'success')
        else:
            flash('Baca bilgisi eklenirken hata oluştu!', 'error')
            
    except Exception as e:
        print(f"Baca bilgisi ekleme hatası: {e}")
        flash(f'Baca bilgisi eklenirken hata oluştu: {str(e)}', 'error')
    
    return redirect(url_for('baca_bilgileri'))

@app.route('/baca_bilgileri/edit/<baca_id>', methods=['POST'])
def edit_baca_bilgisi(baca_id):
    if not session.get('logged_in') or not can_edit(session.get('role')):
        return redirect(url_for('login'))
    
    try:
        baca_bilgileri_list = load_baca_bilgileri()
        baca = next((b for b in baca_bilgileri_list if b['id'] == baca_id), None)
        
        if not baca:
            flash('Baca bilgisi bulunamadı!', 'error')
            return redirect(url_for('baca_bilgileri'))
        
        # Baca bilgilerini güncelle
        baca['firma_adi'] = request.form.get('firma_adi', '')
        baca['olcum_kodu'] = request.form.get('olcum_kodu', '')
        baca['baca_no'] = request.form.get('baca_no', '').strip()
        baca['yakit_turu'] = request.form.get('yakit_turu', '')
        baca['isil_guc'] = request.form.get('isil_guc', '')
        baca['cati_sekli'] = request.form.get('cati_sekli', '')
        baca['kaynak_turu'] = request.form.get('kaynak_turu', '')
        baca['baca_sekli'] = request.form.get('baca_sekli', '')
        baca['baca_olcusu'] = request.form.get('baca_olcusu', '')
        baca['yerden_yuk'] = request.form.get('yerden_yuk', '')
        baca['cati_yuk'] = request.form.get('cati_yuk', '')
        baca['ruzgar_hiz'] = request.form.get('ruzgar_hiz', '')
        baca['ort_sic'] = request.form.get('ort_sic', '')
        baca['ort_nem'] = request.form.get('ort_nem', '')
        baca['ort_bas'] = request.form.get('ort_bas', '')
        baca['a_baca'] = request.form.get('a_baca', '')
        baca['b_baca'] = request.form.get('b_baca', '')
        baca['c_delik'] = request.form.get('c_delik', '')
        
        if not baca['baca_no']:
            flash('Baca numarası zorunludur!', 'error')
            return redirect(url_for('baca_bilgileri'))
        
        if save_baca_bilgileri(baca_bilgileri_list):
            flash('Baca bilgisi başarıyla güncellendi!', 'success')
        else:
            flash('Baca bilgisi güncellenirken hata oluştu!', 'error')
            
    except Exception as e:
        print(f"Baca bilgisi güncelleme hatası: {e}")
        flash(f'Baca bilgisi güncellenirken hata oluştu: {str(e)}', 'error')
    
    return redirect(url_for('baca_bilgileri'))

@app.route('/baca_bilgileri/delete/<baca_id>', methods=['POST'])
def delete_baca_bilgisi(baca_id):
    if not session.get('logged_in') or not can_delete(session.get('role')):
        return redirect(url_for('login'))
    
    try:
        baca_bilgileri_list = load_baca_bilgileri()
        baca_bilgileri_to_keep = [b for b in baca_bilgileri_list if b['id'] != baca_id]
        
        if len(baca_bilgileri_list) == len(baca_bilgileri_to_keep):
            flash('Silinecek baca bilgisi bulunamadı.', 'error')
        else:
            if save_baca_bilgileri(baca_bilgileri_to_keep):
                flash('Baca bilgisi başarıyla silindi.', 'success')
            else:
                flash('Baca bilgisi silinirken hata oluştu.', 'error')
                
    except Exception as e:
        print(f"Baca bilgisi silme hatası: {e}")
        flash(f'Baca bilgisi silinirken hata oluştu: {str(e)}', 'error')
    
    return redirect(url_for('baca_bilgileri'))

# Baca Para Yönetimi Route'ları
@app.route('/add_baca_para', methods=['POST'])
def add_baca_para():
    """Yeni baca parametresi ekler."""
    if not can_write(session.get('role')):
        flash('Bu işlem için yetkiniz yok!', 'error')
        return redirect(url_for('ayarlar', tab='baca-bilgileri'))
    
    try:
        baca_par_adi = request.form.get('baca_par_adi', '').strip()
        liste_icerigi = request.form.get('liste_icerigi', '').strip()
        
        if not baca_par_adi:
            flash('Baca parametre adı zorunludur!', 'error')
            return redirect(url_for('ayarlar', tab='baca-bilgileri'))
        
        # Baca parametrelerini yükle
        baca_paralar = load_baca_paralar()
        
        # Yeni parametre oluştur
        yeni_para = {
            'id': str(uuid4()),
            'baca_par_adi': baca_par_adi,
            'liste_icerigi': liste_icerigi,
            'created_at': datetime.now().isoformat()
        }
        
        baca_paralar.append(yeni_para)
        save_baca_paralar(baca_paralar)
        
        flash(f'Baca parametresi başarıyla eklendi: {baca_par_adi}', 'success')
        
    except Exception as e:
        flash(f'Hata oluştu: {str(e)}', 'error')
    
    # Başarılı ekleme sonrası baca bilgileri sekmesi ile yönlendir
    return redirect(url_for('ayarlar', tab='baca-bilgileri'))

@app.route('/edit_baca_para/<para_id>', methods=['POST'])
def edit_baca_para(para_id):
    """Baca parametresini düzenler."""
    if not can_edit(session.get('role')):
        flash('Bu işlem için yetkiniz yok!', 'error')
        return redirect(url_for('ayarlar', tab='baca-bilgileri'))
    
    try:
        baca_par_adi = request.form.get('baca_par_adi', '').strip()
        liste_icerigi = request.form.get('liste_icerigi', '').strip()
        
        if not baca_par_adi:
            flash('Baca parametre adı zorunludur!', 'error')
            return redirect(url_for('ayarlar', tab='baca-bilgileri'))
        
        # Baca parametrelerini yükle
        baca_paralar = load_baca_paralar()
        
        # Parametreyi bul ve güncelle
        for para in baca_paralar:
            if str(para.get('id')) == str(para_id):
                para['baca_par_adi'] = baca_par_adi
                para['liste_icerigi'] = liste_icerigi
                para['updated_at'] = datetime.now().isoformat()
                break
        
        save_baca_paralar(baca_paralar)
        flash(f'Baca parametresi başarıyla güncellendi: {baca_par_adi}', 'success')
        
    except Exception as e:
        flash(f'Hata oluştu: {str(e)}', 'error')
    
    # Başarılı güncelleme sonrası baca bilgileri sekmesi ile yönlendir
    return redirect(url_for('ayarlar', tab='baca-bilgileri'))

@app.route('/delete_baca_para/<para_id>', methods=['POST'])
def delete_baca_para(para_id):
    """Baca parametresini siler."""
    if not can_delete(session.get('role')):
        return jsonify({'success': False, 'error': 'Yetkiniz yok'})
    
    try:
        baca_paralar = load_baca_paralar()
        
        # Parametreyi bul ve sil
        for i, para in enumerate(baca_paralar):
            if str(para.get('id')) == str(para_id):
                deleted_para = baca_paralar.pop(i)
                save_baca_paralar(baca_paralar)
                return jsonify({'success': True, 'message': f'Baca parametresi silindi: {deleted_para.get("baca_par_adi", "Bilinmeyen")}'})
        
        return jsonify({'success': False, 'error': 'Baca parametresi bulunamadı'})
        
    except Exception as e:
        return jsonify({'success': False, 'error': f'Hata: {str(e)}'})

@app.route('/delete_selected_baca_para', methods=['POST'])
def delete_selected_baca_para():
    """Seçilen baca parametrelerini toplu olarak siler."""
    if not can_delete(session.get('role')):
        return jsonify({'success': False, 'error': 'Yetkiniz yok'})
    
    try:
        data = request.get_json()
        selected_ids = data.get('ids', [])
        
        if not selected_ids:
            return jsonify({'success': False, 'error': 'Silinecek parametre seçilmedi'})
        
        baca_paralar = load_baca_paralar()
        deleted_count = 0
        deleted_names = []
        
        # Seçilen parametreleri sil
        for para_id in selected_ids:
            for i, para in enumerate(baca_paralar):
                if str(para.get('id')) == str(para_id):
                    deleted_para = baca_paralar.pop(i)
                    deleted_count += 1
                    deleted_names.append(deleted_para.get('baca_par_adi', 'Bilinmeyen'))
                    break
        
        if deleted_count > 0:
            save_baca_paralar(baca_paralar)
            return jsonify({
                'success': True, 
                'message': f'{deleted_count} baca parametresi başarıyla silindi: {", ".join(deleted_names)}'
            })
        else:
            return jsonify({'success': False, 'error': 'Silinecek parametre bulunamadı'})
        
    except Exception as e:
        return jsonify({'success': False, 'error': f'Hata: {str(e)}'})

@app.route('/export_baca_bilgileri')
def export_baca_bilgileri():
    """Baca bilgilerini dışa aktarır."""
    if not can_read(session.get('role')):
        flash('Bu işlem için yetkiniz yok!', 'error')
        return redirect(url_for('ayarlar', tab='baca-bilgileri'))
    
    try:
        baca_paralar = load_baca_paralar()
        
        if not baca_paralar:
            flash('Dışa aktarılacak baca parametresi bulunamadı!', 'warning')
            return redirect(url_for('ayarlar', tab='baca-bilgileri'))
        
        # DataFrame oluştur
        df = pd.DataFrame(baca_paralar)
        
        # Geçici dosya oluştur
        with tempfile.NamedTemporaryFile(mode='w', suffix='.xlsx', delete=False, encoding='utf-8') as tmp_file:
            df.to_excel(tmp_file.name, index=False, engine='openpyxl')
            tmp_file_path = tmp_file.name
        
        # Dosyayı gönder
        return send_file(
            tmp_file_path,
            as_attachment=True,
            download_name=f'baca_parametreleri_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx',
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        
    except Exception as e:
        flash(f'Dışa aktarma hatası: {str(e)}', 'error')
        return redirect(url_for('ayarlar', tab='baca-bilgileri'))
def load_baca_paralar():
    """Baca parametrelerini JSON dosyasından yükler."""
    BACA_PARALAR_FILE = 'baca_paralar.json'
    if not os.path.exists(BACA_PARALAR_FILE):
        # Varsayılan parametreler - 26 alan sırasıyla
        default_paralar = [
            {
                'id': str(uuid4()),
                'baca_par_adi': 'BACA NO',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'YAKIT TÜRÜ',
                'liste_icerigi': 'BİOKÜTLE, DOĞAL GAZ, KÖMÜR, SIVI YAKIT',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'ISIL GÜÇ (MW)',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'ÇATI ŞEKLİ',
                'liste_icerigi': 'BAĞIMSIZ, DÜZ, EĞİK',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'KAYNAK TÜRÜ',
                'liste_icerigi': 'YAKMA, PROSES',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'BACA ŞEKLİ',
                'liste_icerigi': 'DÖRTGEN, DAİRESEL',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'BACA ÖLÇÜSÜ',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'YERDEN YÜK.',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'ÇATI YÜK',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'RÜZGAR HIZ (M/S)',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'ORT SIC',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'ORT NEM',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'ORT BAS.',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'A-BACA',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'B-BACA',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'C-DELİK',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'FOTO',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'KK1-O2',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'KK1-CO',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'KK1-NO',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'KK1-SO2',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'KK2-O2',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'KK2-CO',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'KK2-NO',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'KK2-SO2',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            },
            {
                'id': str(uuid4()),
                'baca_par_adi': 'T90',
                'liste_icerigi': '',
                'created_at': datetime.now().isoformat()
            }
        ]
        save_baca_paralar(default_paralar)
        return default_paralar
    
    try:
        with open(BACA_PARALAR_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"Baca parametreleri yüklenirken hata: {e}")
        return []

def save_baca_paralar(baca_paralar_data):
    """Baca parametrelerini JSON dosyasına kaydeder."""
    # Yol başta tanımlandı: BACA_PARALAR_FILE
    try:
        with open(BACA_PARALAR_FILE, 'w', encoding='utf-8') as f:
            json.dump(baca_paralar_data, f, indent=4, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Baca parametreleri kaydedilirken hata: {e}")
        return False

def load_parametre_sahabil():
    """Parametre sahabil verilerini JSON dosyasından yükler."""
    # Yol başta tanımlandı: PARAMETRE_SAHABIL_FILE
    if not os.path.exists(PARAMETRE_SAHABIL_FILE):
        return []
    try:
        with open(PARAMETRE_SAHABIL_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"Parametre sahabil verileri yüklenirken hata: {e}")
        return []

def save_parametre_sahabil(parametre_sahabil_data):
    """Parametre sahabil verilerini JSON dosyasına kaydeder."""
    try:
        with open(PARAMETRE_SAHABIL_FILE, 'w', encoding='utf-8') as f:
            json.dump(parametre_sahabil_data, f, indent=4, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Parametre sahabil verileri kaydedilirken hata: {e}")
        return False

def load_parametre_fields():
    """Parametre alanlarını JSON dosyasından yükler."""
    try:
        if not os.path.exists(PARAMETRE_FIELDS_FILE):
            return {}
        with open(PARAMETRE_FIELDS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"Parametre alanları yüklenirken hata: {e}")
        return {}

def load_asgari_fiyatlar():
    try:
        if not os.path.exists(ASGARI_FIYATLAR_FILE):
            return []
        with open(ASGARI_FIYATLAR_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"Asgari fiyatlar yüklenirken hata: {e}")
        return []

def save_asgari_fiyatlar(data):
    try:
        with open(ASGARI_FIYATLAR_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Asgari fiyatlar kaydedilirken hata: {e}")
        return False

def load_par_saha_headers():
    """PAR_SAHA başlıklarını yükler."""
    try:
        if not os.path.exists(PAR_SAHA_HEADERS_FILE):
            return {"groups": []}
        with open(PAR_SAHA_HEADERS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"PAR SAHA header yüklenirken hata: {e}")
        return {"groups": []}

def save_par_saha_headers(data):
    """PAR_SAHA başlıklarını kaydeder."""
    try:
        with open(PAR_SAHA_HEADERS_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"PAR SAHA header kaydedilirken hata: {e}")
        return False

def save_parametre_fields(parametre_fields_data):
    """Parametre alanlarını JSON dosyasına kaydeder."""
    try:
        with open(PARAMETRE_FIELDS_FILE, 'w', encoding='utf-8') as f:
            json.dump(parametre_fields_data, f, indent=4, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Parametre alanları kaydedilirken hata: {e}")
        return False

def load_forms():
    """Form verilerini JSON dosyasından yükler."""
    if not os.path.exists(FORMS_FILE):
        return []
    try:
        with open(FORMS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"Form verileri yüklenirken hata: {e}")
        return []

def save_forms(forms_data):
    """Form verilerini JSON dosyasına kaydeder."""
    try:
        with open(FORMS_FILE, 'w', encoding='utf-8') as f:
            json.dump(forms_data, f, indent=4, ensure_ascii=False)
        return True
    except Exception as e:
        print(f"Form verileri kaydedilirken hata: {e}")
        return False

@app.route('/api/baca_bilgileri/bulk_delete', methods=['POST'])
def api_bulk_delete_baca_bilgileri():
    """Seçili baca bilgilerini toplu olarak siler."""
    try:
        data = request.get_json()
        if not data or 'ids' not in data:
            return jsonify({'success': False, 'error': 'Silinecek kayıt ID\'leri belirtilmedi'}), 400
        
        record_ids = data['ids']
        if not isinstance(record_ids, list) or len(record_ids) == 0:
            return jsonify({'success': False, 'error': 'Geçersiz kayıt ID listesi'}), 400
        
        saved_baca_bilgileri = load_baca_bilgileri()
        original_count = len(saved_baca_bilgileri)
        
        # Seçili kayıtları filtrele
        filtered_records = []
        deleted_records = []
        
        for record in saved_baca_bilgileri:
            if record.get('id') in record_ids:
                deleted_records.append(record)
            else:
                filtered_records.append(record)
        
        # Silinen kayıt sayısını kontrol et
        if len(deleted_records) == 0:
            return jsonify({'success': False, 'error': 'Silinecek kayıt bulunamadı'}), 404
        
        # Dosyaya kaydet
        if save_baca_bilgileri(filtered_records):
            deleted_info = [f"{r.get('firma_adi')} - {r.get('olcum_kodu')} - {r.get('baca_adi')}" for r in deleted_records]
            print(f"Baca bilgileri toplu silindi: {', '.join(deleted_info)}")
            return jsonify({
                'success': True, 
                'message': f'{len(deleted_records)} kayıt başarıyla silindi',
                'deleted_count': len(deleted_records)
            })
        else:
            return jsonify({'success': False, 'error': 'Veriler kaydedilemedi'}), 500
            
    except Exception as e:
        print(f"Baca bilgileri toplu silinirken hata: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

# PAR_SAHA başlıklarını getir
@app.route('/api/par_saha_headers', methods=['GET'])
def api_get_par_saha_headers():
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'}), 401
    data = load_par_saha_headers()
    return jsonify({'success': True, **data})

# PAR_SAHA başlıklarını kaydet
@app.route('/api/par_saha_headers', methods=['POST'])
def api_save_par_saha_headers():
    if not session.get('logged_in'):
        return jsonify({'success': False, 'message': 'Oturum açmanız gerekiyor'}), 401
    try:
        data = request.get_json() or {}
        if not isinstance(data, dict) or 'groups' not in data or not isinstance(data['groups'], list):
            return jsonify({'success': False, 'message': 'Geçersiz veri'}), 400
        ok = save_par_saha_headers({'groups': data['groups']})
        return jsonify({'success': ok})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@app.route('/api/firma_rapor_export', methods=['POST'])
def api_firma_rapor_export():
    """Bir firmanın tüm bacalarında yapılan tüm parametreleri içeren rapor oluşturur."""
    try:
        print("=== FIRMA RAPOR EXPORT BAŞLADI ===")
        data = request.get_json()
        print(f"Gelen data: {data}")
        
        firma_adi = data.get('firma_adi')
        olcum_kodu = data.get('olcum_kodu')
        print(f"Firma adı: {firma_adi}")
        print(f"Ölçüm kodu: {olcum_kodu}")
        
        if not firma_adi or not olcum_kodu:
            print("Firma adı veya ölçüm kodu eksik!")
            return jsonify({'error': 'Firma adı ve ölçüm kodu gerekli!'}), 400
        
        # Şablon kullanımını kontrol et (varsayılan olarak şablon kullan)
        use_template = data.get('use_template', True)
        print(f"Şablon kullanılacak: {use_template}")
        
        # Word raporu oluştur
        if use_template:
            print("create_firma_raporu_from_template çağrılıyor...")
            result = create_firma_raporu_from_template(firma_adi, olcum_kodu)
            print(f"create_firma_raporu_from_template sonucu: {type(result)}")
            return result
        else:
            print("create_firma_raporu çağrılıyor...")
            return create_firma_raporu(firma_adi, olcum_kodu)
        
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"Firma rapor export hatası: {e}")
        print(f"Hata detayı: {error_details}")
        return jsonify({'error': f'Rapor oluşturma hatası: {str(e)}'}), 500

@app.route('/api/firma_rapor_pdf_export', methods=['POST'])
def api_firma_rapor_pdf_export():
    """Bir firmanın tüm bacalarında yapılan tüm parametreleri içeren PDF rapor oluşturur."""
    try:
        data = request.get_json()
        firma_adi = data.get('firma_adi')
        olcum_kodu = data.get('olcum_kodu')
        
        if not firma_adi or not olcum_kodu:
            return jsonify({'error': 'Firma adı ve ölçüm kodu gerekli!'}), 400
        
        # Word raporu oluştur
        word_response = create_firma_raporu_from_template(firma_adi, olcum_kodu)
        
        # Hata kontrolü
        if hasattr(word_response, 'get_data'):
            # Flask Response objesi
            word_data = word_response.get_data()
        elif isinstance(word_response, tuple):
            # Hata tuple'ı
            return word_response
        else:
            return jsonify({'error': 'Bilinmeyen response tipi'}), 500
        
        # Word dosyasını geçici olarak kaydet
        import tempfile
        import os
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_word:
            tmp_word.write(word_data)
            tmp_word_path = tmp_word.name
        
        # Word dosyasını PDF'e çevir
        try:
            from docx2pdf import convert
            
            # PDF dosya adını oluştur
            def clean_filename(text):
                tr_to_en = {
                    'ç': 'c', 'ğ': 'g', 'ı': 'i', 'ö': 'o', 'ş': 's', 'ü': 'u',
                    'Ç': 'C', 'Ğ': 'G', 'I': 'I', 'İ': 'I', 'Ö': 'O', 'Ş': 'S', 'Ü': 'U'
                }
                for tr_char, en_char in tr_to_en.items():
                    text = text.replace(tr_char, en_char)
                cleaned = "".join(c for c in text if c.isalnum() or c in (' ', '-', '_')).strip()
                cleaned = cleaned.replace(' ', '-')
                while '--' in cleaned:
                    cleaned = cleaned.replace('--', '-')
                return cleaned
            
            firma_adi_clean = clean_filename(firma_adi)
            olcum_kodu_clean = clean_filename(olcum_kodu)
            
            pdf_filename = f"Firma_Raporu_{firma_adi_clean}_{olcum_kodu_clean}.pdf"
            pdf_path = os.path.join(tempfile.gettempdir(), pdf_filename)
            
            # Word dosyasını PDF'e çevir
            convert(tmp_word_path, pdf_path)
            
            # PDF dosyasını oku
            with open(pdf_path, 'rb') as pdf_file:
                pdf_content = pdf_file.read()
            
            # Geçici dosyaları temizle
            os.unlink(tmp_word_path)
            os.unlink(pdf_path)
            
            # PDF dosyasını döndür
            response = make_response(pdf_content)
            response.headers['Content-Type'] = 'application/pdf'
            response.headers['Content-Disposition'] = f'attachment; filename="{pdf_filename}"'
            return response
            
        except Exception as pdf_error:
            # PDF dönüştürme başarısız olursa Word dosyasını döndür
            print(f"PDF dönüştürme hatası: {pdf_error}")
            
            # Geçici Word dosyasını temizle
            if os.path.exists(tmp_word_path):
                os.unlink(tmp_word_path)
            
            # Word dosyasını döndür
            return word_response
        
    except Exception as e:
        print(f"PDF Firma rapor export hatası: {e}")
        return jsonify({'error': f'PDF Rapor oluşturma hatası: {str(e)}'}), 500

@app.route('/api/firma_olcum_detail_word_export', methods=['POST'])
def api_firma_olcum_detail_word_export():
    """Firma ölçüm detay bilgilerini Word formatında dışa aktarır."""
    try:
        data = request.get_json()
        olcum_id = data.get('olcum_id')
        
        if not olcum_id:
            return jsonify({'error': 'Ölçüm ID gerekli!'}), 400
        
        # Firma ölçüm verilerini yükle
        firma_olcumler = load_firma_olcum()
        
        # Seçili ölçümü bul
        olcum = None
        for item in firma_olcumler:
            if str(item.get('id')) == str(olcum_id):
                olcum = item
                break
        
        if not olcum:
            return jsonify({'error': 'Ölçüm bulunamadı'}), 404
        
        # Firma kayıt bilgilerini al
        firma_kayitlar = load_firma_kayit()
        firma_bilgisi = None
        for firma in firma_kayitlar:
            if firma.get('firmaAdi') == olcum.get('firma_adi'):
                firma_bilgisi = firma
                break
        
        # Baca bilgilerini yükle
        baca_bilgileri = load_baca_bilgileri()
        
        # Bu ölçüme ait bacaları bul
        olcum_bacalar = []
        for baca in baca_bilgileri:
            if (baca.get('firma_adi') == olcum.get('firma_adi') and 
                baca.get('olcum_kodu') == olcum.get('olcum_kodu')):
                olcum_bacalar.append(baca)
        
        # Parametre ölçümlerini yükle
        parametre_olcumleri = load_parametre_olcum()
        
        # Bu ölçüme ait parametreleri bul
        olcum_parametreleri = []
        for parametre in parametre_olcumleri:
            if (parametre.get('firma_adi') == olcum.get('firma_adi') and 
                parametre.get('olcum_kodu') == olcum.get('olcum_kodu')):
                olcum_parametreleri.append(parametre)
        
        # Word dokümanı oluştur
        doc = create_firma_olcum_detail_word_document(olcum, firma_bilgisi, olcum_bacalar, olcum_parametreleri)
        return doc
        
    except Exception as e:
        print(f"Firma ölçüm detay Word export hatası: {e}")
        import traceback
        print(f"DEBUG: Hata detayı: {traceback.format_exc()}")
        return jsonify({'error': f'Word export hatası: {str(e)}'}), 500

@app.route('/api/firma_olcum_word_export', methods=['POST'])
def api_firma_olcum_word_export():
    """Firma ölçüm bilgilerini Word formatında dışa aktarır."""
    try:
        data = request.get_json()
        olcum_ids = data.get('olcum_ids', [])
        
        if not olcum_ids:
            return jsonify({'error': 'En az bir ölçüm ID gerekli!'}), 400
        
        # Firma ölçüm verilerini yükle
        firma_olcumler = load_firma_olcum()
        
        # Seçili ölçümleri filtrele
        if olcum_ids:
            filtered_olcumler = [item for item in firma_olcumler if str(item.get('id')) in olcum_ids]
        else:
            filtered_olcumler = firma_olcumler
        
        if not filtered_olcumler:
            return jsonify({'error': 'Dışa aktarılacak veri bulunamadı'}), 404
        
        # Baca bilgilerini yükle
        baca_bilgileri = load_baca_bilgileri()
        
        # Parametre ölçümlerini yükle
        parametre_olcumleri = load_parametre_olcum()
        
        # Her ölçüm için Word dokümanı oluştur
        documents = []
        for olcum in filtered_olcumler:
            # Bu ölçüme ait bacaları bul
            olcum_bacalar = []
            for baca in baca_bilgileri:
                if (baca.get('firma_adi') == olcum.get('firma_adi') and 
                    baca.get('olcum_kodu') == olcum.get('olcum_kodu')):
                    olcum_bacalar.append(baca)
            
            # Bu ölçüme ait parametreleri bul
            olcum_parametreleri = []
            for parametre in parametre_olcumleri:
                if (parametre.get('firma_adi') == olcum.get('firma_adi') and 
                    parametre.get('olcum_kodu') == olcum.get('olcum_kodu')):
                    olcum_parametreleri.append(parametre)
            
            # Word dokümanı oluştur
            doc = create_firma_olcum_word_document(olcum, olcum_bacalar, olcum_parametreleri)
            documents.append(doc)
        
        # Eğer tek doküman varsa onu döndür, birden fazla varsa birleştir
        if len(documents) == 1:
            return documents[0]
        else:
            # Birden fazla dokümanı birleştir
            merged_doc = merge_word_documents(documents)
            return merged_doc
        
    except Exception as e:
        print(f"Firma ölçüm Word export hatası: {e}")
        import traceback
        print(f"DEBUG: Hata detayı: {traceback.format_exc()}")
        return jsonify({'error': f'Word export hatası: {str(e)}'}), 500
@app.route('/api/firma_olcum_pdf_export', methods=['POST'])
def api_firma_olcum_pdf_export():
    """Firma ölçüm bilgilerini PDF formatında dışa aktarır."""
    try:
        data = request.get_json()
        olcum_ids = data.get('olcum_ids', [])
        
        if not olcum_ids:
            return jsonify({'error': 'En az bir ölçüm ID gerekli!'}), 400
        
        # Firma ölçüm verilerini yükle
        firma_olcumler = load_firma_olcum()
        
        # Seçili ölçümleri filtrele
        if olcum_ids:
            filtered_olcumler = [item for item in firma_olcumler if str(item.get('id')) in olcum_ids]
        else:
            filtered_olcumler = firma_olcumler
        
        if not filtered_olcumler:
            return jsonify({'error': 'Dışa aktarılacak veri bulunamadı'}), 404
        
        # Baca bilgilerini yükle
        baca_bilgileri = load_baca_bilgileri()
        
        # Parametre ölçümlerini yükle
        parametre_olcumleri = load_parametre_olcum()
        
        # Her ölçüm için Word dokümanı oluştur
        documents = []
        for olcum in filtered_olcumler:
            # Bu ölçüme ait bacaları bul
            olcum_bacalar = []
            for baca in baca_bilgileri:
                if (baca.get('firma_adi') == olcum.get('firma_adi') and 
                    baca.get('olcum_kodu') == olcum.get('olcum_kodu')):
                    olcum_bacalar.append(baca)
            
            # Bu ölçüme ait parametreleri bul
            olcum_parametreleri = []
            for parametre in parametre_olcumleri:
                if (parametre.get('firma_adi') == olcum.get('firma_adi') and 
                    parametre.get('olcum_kodu') == olcum.get('olcum_kodu')):
                    olcum_parametreleri.append(parametre)
            
            # Word dokümanı oluştur
            doc = create_firma_olcum_word_document(olcum, olcum_bacalar, olcum_parametreleri)
            documents.append(doc)
        
        # Eğer tek doküman varsa onu kullan, birden fazla varsa birleştir
        if len(documents) == 1:
            word_doc = documents[0]
        else:
            # Birden fazla dokümanı birleştir
            word_doc = merge_word_documents(documents)
        
        # Word dosyasını geçici olarak kaydet
        from io import BytesIO
        import tempfile
        import os
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_word:
            tmp_word.write(word_doc.get_data())
            tmp_word_path = tmp_word.name
        
        # Word dosyasını PDF'e çevir
        try:
            from docx2pdf import convert
            
            # PDF dosya adını oluştur
            def clean_filename(text):
                tr_to_en = {
                    'ç': 'c', 'ğ': 'g', 'ı': 'i', 'ö': 'o', 'ş': 's', 'ü': 'u',
                    'Ç': 'C', 'Ğ': 'G', 'I': 'I', 'İ': 'I', 'Ö': 'O', 'Ş': 'S', 'Ü': 'U'
                }
                for tr_char, en_char in tr_to_en.items():
                    text = text.replace(tr_char, en_char)
                cleaned = "".join(c for c in text if c.isalnum() or c in (' ', '-', '_')).strip()
                cleaned = cleaned.replace(' ', '-')
                while '--' in cleaned:
                    cleaned = cleaned.replace('--', '-')
                return cleaned
            
            # Birden fazla ölçüm varsa genel bir isim kullan
            if len(filtered_olcumler) == 1:
                olcum = filtered_olcumler[0]
                firma_adi = olcum.get('firma_adi', '')
                firma_kelimeleri = firma_adi.split()
                firma_ilk_2_kelime = ' '.join(firma_kelimeleri[:2]) if len(firma_kelimeleri) >= 2 else firma_adi
                olcum_kodu_clean = clean_filename(olcum.get('olcum_kodu', ''))
                pdf_filename = f"{clean_filename(firma_ilk_2_kelime)}_{olcum_kodu_clean}_SAHAFORM.pdf"
            else:
                # Birden fazla ölçüm için genel isim
                pdf_filename = f"Firma_Olcum_{len(filtered_olcumler)}_kayit_SAHAFORM.pdf"
            
            pdf_path = os.path.join(tempfile.gettempdir(), pdf_filename)
            
            # Word dosyasını PDF'e çevir
            convert(tmp_word_path, pdf_path)
            
            # PDF dosyasını oku
            with open(pdf_path, 'rb') as pdf_file:
                pdf_content = pdf_file.read()
            
            # Geçici dosyaları temizle
            os.unlink(tmp_word_path)
            os.unlink(pdf_path)
            
            # PDF dosyasını döndür
            response = make_response(pdf_content)
            response.headers['Content-Type'] = 'application/pdf'
            response.headers['Content-Disposition'] = f'attachment; filename="{pdf_filename}"'
            return response
            
        except Exception as pdf_error:
            # PDF dönüştürme başarısız olursa Word dosyasını döndür
            print(f"PDF dönüştürme hatası: {pdf_error}")
            
            # Geçici Word dosyasını temizle
            if os.path.exists(tmp_word_path):
                os.unlink(tmp_word_path)
            
            # Word dosyasını döndür
            return word_doc
        
    except Exception as e:
        print(f"Firma ölçüm PDF export hatası: {e}")
        return jsonify({'error': f'PDF export hatası: {str(e)}'}), 500

@app.route('/api/firma_olcum_detail_excel_export', methods=['POST'])
def api_firma_olcum_detail_excel_export():
    """Firma ölçüm detay bilgilerini Excel formatında dışa aktarır."""
    try:
        data = request.get_json()
        olcum_id = data.get('olcum_id')
        
        if not olcum_id:
            return jsonify({'error': 'Ölçüm ID gerekli!'}), 400
        
        # Firma ölçüm verilerini yükle
        firma_olcumler = load_firma_olcum()
        
        # Seçili ölçümü bul
        olcum = None
        for item in firma_olcumler:
            if str(item.get('id')) == str(olcum_id):
                olcum = item
                break
        
        if not olcum:
            return jsonify({'error': 'Ölçüm bulunamadı'}), 404
        
        # Firma kayıt bilgilerini al
        firma_kayitlar = load_firma_kayit()
        firma_bilgisi = None
        for firma in firma_kayitlar:
            if firma.get('firmaAdi') == olcum.get('firma_adi'):
                firma_bilgisi = firma
                break
        
        # Baca bilgilerini yükle
        baca_bilgileri = load_baca_bilgileri()
        
        # Bu ölçüme ait bacaları bul
        olcum_bacalar = []
        for baca in baca_bilgileri:
            if (baca.get('firma_adi') == olcum.get('firma_adi') and 
                baca.get('olcum_kodu') == olcum.get('olcum_kodu')):
                olcum_bacalar.append(baca)
        
        # Excel için DataFrame oluştur
        df_data = []
        
        # Temel bilgiler
        df_data.append(['TEMEL BİLGİLER', ''])
        df_data.append(['Firma Adı', firma_bilgisi.get('firmaAdi', '') if firma_bilgisi else olcum.get('firma_adi', '')])
        df_data.append(['Ölçüm Kodu', olcum.get('olcum_kodu', '')])
        df_data.append(['Başlangıç Tarihi', olcum.get('baslangic_tarihi', '')])
        df_data.append(['Bitiş Tarihi', olcum.get('bitis_tarihi', '')])
        df_data.append(['Baca Sayısı', olcum.get('baca_sayisi', '')])
        df_data.append(['Durum', olcum.get('durum', '')])
        df_data.append(['', ''])
        
        # İletişim bilgileri
        df_data.append(['İLETİŞİM BİLGİLERİ', ''])
        df_data.append(['İl', firma_bilgisi.get('il', '') if firma_bilgisi else olcum.get('il', '')])
        df_data.append(['İlçe', firma_bilgisi.get('ilce', '') if firma_bilgisi else olcum.get('ilce', '')])
        df_data.append(['Yetkili', firma_bilgisi.get('yetkiliAdi', '') if firma_bilgisi else olcum.get('yetkili', '')])
        df_data.append(['Telefon', str(firma_bilgisi.get('yetkiliTel', '')) if firma_bilgisi else olcum.get('telefon', '')])
        df_data.append(['Personel', ', '.join(olcum.get('personel', []))])
        df_data.append(['', ''])
        
        # Baca-Parametre matrisi
        df_data.append(['BACA-PARAMETRE MATRİSİ', ''])
        
        # Parametre başlıkları
        baca_parametreleri = olcum.get('baca_parametreleri', {})
        all_params = set()
        for baca_params in baca_parametreleri.values():
            for param in baca_params:
                all_params.add(param)
        
        param_headers = ['Baca Adı'] + list(all_params)
        df_data.append(param_headers)
        
        # Baca satırları
        for baca_adi, parametreler in baca_parametreleri.items():
            row = [baca_adi]
            for param in all_params:
                if param in parametreler:
                    row.append('X')
                else:
                    row.append('')
            df_data.append(row)
        
        df = pd.DataFrame(df_data)
        
        # Geçici dosya oluştur
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
            # Excel dosyasını oluştur
            with pd.ExcelWriter(tmp_file.name, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, header=False, sheet_name='Firma Ölçüm Detayı')
                
                # Sütun genişliklerini ayarla
                worksheet = writer.sheets['Firma Ölçüm Detayı']
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    worksheet.column_dimensions[column_letter].width = adjusted_width
            
            # Dosya adını oluştur
            firma_adi = firma_bilgisi.get('firmaAdi', '') if firma_bilgisi else olcum.get('firma_adi', '')
            olcum_kodu = olcum.get('olcum_kodu', '')
            filename = f"{firma_adi}_{olcum_kodu}_DETAY.xlsx"
            
            return send_file(
                tmp_file.name,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
        
    except Exception as e:
        print(f"Firma ölçüm detay Excel export hatası: {e}")
        import traceback
        print(f"DEBUG: Hata detayı: {traceback.format_exc()}")
        return jsonify({'error': f'Excel export hatası: {str(e)}'}), 500

@app.route('/api/firma_olcum_excel_export', methods=['POST'])
def api_firma_olcum_excel_export():
    """Firma ölçüm bilgilerini Excel formatında dışa aktarır - Basit format."""
    try:
        # Pandas kütüphanesini import et
        try:
            import pandas as pd
            print(f"DEBUG: pandas başarıyla import edildi: {pd}")
        except ImportError as e:
            print(f"DEBUG: pandas import hatası: {e}")
            return jsonify({'error': 'pandas kütüphanesi yüklü değil'}), 500
        data = request.get_json()
        olcum_ids = data.get('olcum_ids', [])
        
        if not olcum_ids:
            return jsonify({'error': 'En az bir ölçüm ID gerekli!'}), 400
        
        # Firma ölçümlerini yükle
        firma_olcumler = load_firma_olcum()
        
        # Seçili ölçümleri filtrele
        if olcum_ids:
            filtered_olcumler = [item for item in firma_olcumler if str(item.get('id')) in olcum_ids]
        else:
            filtered_olcumler = firma_olcumler
        
        if not filtered_olcumler:
            return jsonify({'error': 'Dışa aktarılacak veri bulunamadı'}), 404
        
        # Firma kayıt bilgilerini al
        firma_kayitlar = load_firma_kayit()
        
        # Excel için DataFrame oluştur - basit format
        df_data = []
        for index, olcum in enumerate(filtered_olcumler, 1):
            # Firma bilgilerini bul
            firma_bilgisi = None
            for firma in firma_kayitlar:
                if firma.get('firmaAdi') == olcum.get('firma_adi'):
                    firma_bilgisi = firma
                    break
            
            # Tarih formatını düzelt
            baslangic_tarihi = olcum.get('baslangic_tarihi', '')
            bitis_tarihi = olcum.get('bitis_tarihi', '')
            
            # Tarihleri DD.MM.YY formatına çevir
            if baslangic_tarihi:
                try:
                    baslangic_date = datetime.strptime(baslangic_tarihi, '%Y-%m-%d')
                    baslangic_tarihi = baslangic_date.strftime('%d.%m.%y')
                except:
                    pass
            
            if bitis_tarihi:
                try:
                    bitis_date = datetime.strptime(bitis_tarihi, '%Y-%m-%d')
                    bitis_tarihi = bitis_date.strftime('%d.%m.%y')
                except:
                    pass
            
            # Personel listesini string'e çevir
            personel = olcum.get('personel', [])
            if isinstance(personel, list):
                personel_str = ', '.join(personel)
            else:
                personel_str = str(personel) if personel else ''
            
            # Parametreleri al ve sayılarını hesapla
            baca_parametreleri = olcum.get('baca_parametreleri', {})
            parametre_sayilari = {}
            
            for baca_params in baca_parametreleri.values():
                for parametre in baca_params:
                    if parametre in parametre_sayilari:
                        parametre_sayilari[parametre] += 1
                    else:
                        parametre_sayilari[parametre] = 1
            
            # Parametreleri sayılarıyla birlikte formatla
            parametre_listesi = []
            for parametre, sayi in parametre_sayilari.items():
                parametre_listesi.append(f"{parametre} ({sayi})")
            
            parametre_str = ', '.join(parametre_listesi) if parametre_listesi else ''
            
            row = {
                'RA': index,
                'FIRMA': firma_bilgisi.get('firmaAdi', '') if firma_bilgisi else olcum.get('firma_adi', ''),
                'OLC_KOD': olcum.get('olcum_kodu', ''),
                'BAS TRH': baslangic_tarihi,
                'BIT TAR': bitis_tarihi,
                'BACA_SAY': olcum.get('baca_sayisi', ''),
                'PARAMETRE': parametre_str,
                'PER.': personel_str,
                'IL': firma_bilgisi.get('il', '') if firma_bilgisi else olcum.get('il', ''),
                'ILCE': firma_bilgisi.get('ilce', '') if firma_bilgisi else olcum.get('ilce', ''),
                'YETK': firma_bilgisi.get('yetkiliAdi', '') if firma_bilgisi else olcum.get('yetkili', ''),
                'TEL': str(firma_bilgisi.get('yetkiliTel', '')) if firma_bilgisi else olcum.get('telefon', ''),
                'DURU': olcum.get('durum', '')
            }
            
            df_data.append(row)
        
        print(f"DEBUG: df_data hazırlandı: {df_data}")
        df = pd.DataFrame(df_data)
        print(f"DEBUG: DataFrame oluşturuldu: {df}")
        
        # Geçici dosya oluştur
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
            # Excel dosyasını oluştur
            with pd.ExcelWriter(tmp_file.name, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='Firma Ölçümleri')
                
                # Sütun genişliklerini ayarla
                worksheet = writer.sheets['Firma Ölçümleri']
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    worksheet.column_dimensions[column_letter].width = adjusted_width
            
            # Dosya adını oluştur
            if len(filtered_olcumler) == 1:
                olcum = filtered_olcumler[0]
                filename = f"{olcum.get('firma_adi', 'bilinmeyen')}_{olcum.get('olcum_kodu', 'bilinmeyen')}_OLCUM.xlsx"
            else:
                filename = f"TUM_FIRMA_OLCUMLERI_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            
            return send_file(
                tmp_file.name,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
        
    except Exception as e:
        print(f"Firma ölçüm Excel export hatası: {e}")
        return jsonify({'error': f'Excel export hatası: {str(e)}'}), 500

def create_firma_olcum_detail_word_document(olcum, firma_bilgisi, olcum_bacalar, olcum_parametreleri):
    """Firma ölçüm detay bilgilerini Word şablonu kullanarak oluşturur."""
    try:
        print(f"DEBUG: create_firma_olcum_detail_word_document başladı")
        
        # Docx kütüphanesini yükle
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        
        if not Document:
            print(f"DEBUG: python-docx kütüphanesi yüklü değil")
            return jsonify({'error': 'python-docx kütüphanesi yüklü değil'}), 500
        
        # Şablon dosyasını yükle
        template_path = os.path.join(os.path.dirname(__file__), 'static', 'images', 'EMISYON_OLCUM_BILGI_FORMU.docx')
        print(f"DEBUG: Template path = {template_path}")
        print(f"DEBUG: Template exists = {os.path.exists(template_path)}")
        
        if not os.path.exists(template_path):
            print(f"DEBUG: Word şablonu bulunamadı")
            return jsonify({'error': 'Word şablonu bulunamadı'}), 500
        
        # Şablonu yükle
        print(f"DEBUG: Document yükleniyor...")
        doc = Document(template_path)
        print(f"DEBUG: Document başarıyla yüklendi")
        
        # Parametre listesini oluştur (firma_olcum verilerinden)
        baca_parametreleri = olcum.get('baca_parametreleri', {})
        
        # Tüm parametreleri topla
        all_params = set()
        for baca_params in baca_parametreleri.values():
            for param in baca_params:
                all_params.add(param)
        
        # Parametreleri sırala
        sorted_params = sorted(all_params)
        
        # Parametre listesini formatla (parametre adı + sayı)
        parametre_listesi = []
        for param in sorted_params:
            # Bu parametrenin kaç kez kullanıldığını say
            count = sum(1 for baca_params in baca_parametreleri.values() for p in baca_params if p == param)
            parametre_listesi.append(f"{param} ({count})")
        
        # Debug: Parametre listesi oluşturma
        print(f"DEBUG: Baca parametreleri: {baca_parametreleri}")
        print(f"DEBUG: Tüm parametreler: {all_params}")
        print(f"DEBUG: Sıralanmış parametreler: {sorted_params}")
        print(f"DEBUG: Oluşturulan parametre listesi: {parametre_listesi}")
        
        # Şablon verilerini hazırla - Firma kayıt bilgilerini kullan
        data = {
            'FIRMA_ADI': firma_bilgisi.get('firmaAdi', '') if firma_bilgisi else olcum.get('firma_adi', ''),
            'OLCUM_KODU': olcum.get('olcum_kodu', ''),
            'BASLANGIC_TARIHI': olcum.get('baslangic_tarihi', ''),
            'BITIS_TARIHI': olcum.get('bitis_tarihi', ''),
            'BACA_SAYISI': str(olcum.get('baca_sayisi', '')),
            'PARAMETRELER': ', '.join(parametre_listesi),
            'PERSONEL': ', '.join(olcum.get('personel', [])),
            'IL': firma_bilgisi.get('il', '') if firma_bilgisi else olcum.get('il', ''),
            'ILCE': firma_bilgisi.get('ilce', '') if firma_bilgisi else olcum.get('ilce', ''),
            'YETKILI': firma_bilgisi.get('yetkiliAdi', '') if firma_bilgisi else olcum.get('yetkili', ''),
            'TELEFON': str(firma_bilgisi.get('yetkiliTel', '')) if firma_bilgisi else olcum.get('telefon', ''),
            'DURUM': olcum.get('durum', '')
        }
        
        print(f"DEBUG: Data hazırlandı: {data}")
        
        # Şablondaki yer tutucuları değiştir
        print(f"DEBUG: Placeholder'lar değiştiriliyor...")
        replace_placeholders_in_document(doc, data)
        print(f"DEBUG: Placeholder'lar değiştirildi")
        
        # Baca listesi tablosunu ekle
        print(f"DEBUG: Baca listesi tablosu ekleniyor...")
        add_baca_listesi_detail_table(doc, olcum, baca_parametreleri)
        print(f"DEBUG: Baca listesi tablosu eklendi")
        
        # Dosyayı kaydet
        print(f"DEBUG: Dosya kaydediliyor...")
        from io import BytesIO
        docx_io = BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        print(f"DEBUG: Dosya başarıyla kaydedildi")
        
        # Dosya adını oluştur
        print(f"DEBUG: Dosya adı oluşturuluyor...")
        def clean_filename(text):
            tr_to_en = {
                'ç': 'c', 'ğ': 'g', 'ı': 'i', 'ö': 'o', 'ş': 's', 'ü': 'u',
                'Ç': 'C', 'Ğ': 'G', 'I': 'I', 'İ': 'I', 'Ö': 'O', 'Ş': 'S', 'Ü': 'U'
            }
            for tr_char, en_char in tr_to_en.items():
                text = text.replace(tr_char, en_char)
            cleaned = "".join(c for c in text if c.isalnum() or c in (' ', '-', '_')).strip()
            cleaned = cleaned.replace(' ', '-')
            while '--' in cleaned:
                cleaned = cleaned.replace('--', '-')
            return cleaned
        
        # Firma adının ilk 2 kelimesini al
        firma_adi = firma_bilgisi.get('firmaAdi', '') if firma_bilgisi else olcum.get('firma_adi', '')
        firma_kelimeleri = firma_adi.split()
        firma_ilk_2_kelime = ' '.join(firma_kelimeleri[:2]) if len(firma_kelimeleri) >= 2 else firma_adi
        
        # Dosya adını oluştur: "FIRMA ADI İKİ KELİME _OLCUM KODU_SAHAFORM"
        firma_ilk_2_kelime_clean = clean_filename(firma_ilk_2_kelime)
        olcum_kodu_clean = clean_filename(olcum.get('olcum_kodu', ''))
        filename = f"{firma_ilk_2_kelime_clean}_{olcum_kodu_clean}_SAHAFORM.docx"
        
        # Response oluştur
        print(f"DEBUG: Response oluşturuluyor...")
        response = make_response(docx_io.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        print(f"DEBUG: Response başarıyla oluşturuldu, filename: {filename}")
        return response
        
    except Exception as e:
        print(f"Firma ölçüm detay Word doküman oluşturma hatası: {e}")
        import traceback
        print(f"DEBUG: Hata detayı: {traceback.format_exc()}")
        return jsonify({'error': f'Word doküman oluşturma hatası: {str(e)}'}), 500

def create_firma_olcum_word_document(olcum, olcum_bacalar, olcum_parametreleri):
    """Firma ölçüm bilgilerini Word şablonu kullanarak oluşturur."""
    try:
        print(f"DEBUG: create_firma_olcum_word_document başladı")
        print(f"DEBUG: DOCX_AVAILABLE = {DOCX_AVAILABLE}")
        
        # Docx kütüphanesini yükle
        print(f"DEBUG: load_docx çağrılıyor...")
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        print(f"DEBUG: load_docx sonrası DOCX_AVAILABLE = {DOCX_AVAILABLE}")
        print(f"DEBUG: Document = {Document}")
        
        if not Document:
            print(f"DEBUG: python-docx kütüphanesi yüklü değil")
            return jsonify({'error': 'python-docx kütüphanesi yüklü değil'}), 500
        
        # Şablon dosyasını yükle
        template_path = os.path.join(os.path.dirname(__file__), 'static', 'images', 'EMISYON_OLCUM_BILGI_FORMU.docx')
        print(f"DEBUG: Template path = {template_path}")
        print(f"DEBUG: Template exists = {os.path.exists(template_path)}")
        
        if not os.path.exists(template_path):
            print(f"DEBUG: Word şablonu bulunamadı")
            return jsonify({'error': 'Word şablonu bulunamadı'}), 500
        
        # Şablonu yükle
        print(f"DEBUG: Document yükleniyor...")
        doc = Document(template_path)
        print(f"DEBUG: Document başarıyla yüklendi")
        
        # Parametre listesini oluştur (firma_olcum verilerinden)
        priority_params = ['Toz', 'Yg', 'Voc', 'Toc']  # Büyük/küçük harf düzeltildi
        
        # Firma ölçüm verilerinden parametreleri al
        baca_parametreleri = olcum.get('baca_parametreleri', {})
        
        # Tüm parametreleri topla
        all_params = set()
        for baca_params in baca_parametreleri.values():
            for param in baca_params:
                all_params.add(param)
        
        # Parametreleri sırala: önce öncelikli olanlar, sonra diğerleri
        sorted_params = []
        
        # Önce öncelikli parametreleri ekle
        for param in priority_params:
            if param in all_params:
                sorted_params.append(param)
                all_params.remove(param)
        
        # Sonra diğer parametreleri ekle
        sorted_params.extend(sorted(all_params))
        
        # Parametre listesini formatla (parametre adı + sayı)
        parametre_listesi = []
        for param in sorted_params:
            # Bu parametrenin kaç kez kullanıldığını say
            count = sum(1 for baca_params in baca_parametreleri.values() for p in baca_params if p == param)
            parametre_listesi.append(f"{param} ({count})")
        
        # Debug: Parametre listesi oluşturma
        print(f"DEBUG: Baca parametreleri: {baca_parametreleri}")
        print(f"DEBUG: Tüm parametreler: {all_params}")
        print(f"DEBUG: Sıralanmış parametreler: {sorted_params}")
        print(f"DEBUG: Oluşturulan parametre listesi: {parametre_listesi}")
        
        # Şablon verilerini hazırla - Şablondaki gerçek yer tutucuları kullan
        data = {
            'FIRMA_ADI': olcum.get('firma_adi', ''),
            'OLCUM_KODU': olcum.get('olcum_kodu', ''),
            'BASLANGIC_TARIHI': olcum.get('baslangic_tarihi', ''),
            'BITIS_TARIHI': olcum.get('bitis_tarihi', ''),
            'BACA_SAYISI': str(olcum.get('baca_sayisi', '')),
            'PARAMETRELER': ', '.join(parametre_listesi),
            'PERSONEL': ', '.join(olcum.get('personel', [])),
            'IL': olcum.get('il', ''),
            'ILCE': olcum.get('ilce', ''),
            'YETKILI': olcum.get('yetkili', ''),
            'TELEFON': olcum.get('telefon', ''),
            'DURUM': olcum.get('durum', '')
        }
        
        print(f"DEBUG: Data hazırlandı: {data}")
        
        # Şablondaki yer tutucuları değiştir
        print(f"DEBUG: Placeholder'lar değiştiriliyor...")
        replace_placeholders_in_document(doc, data)
        print(f"DEBUG: Placeholder'lar değiştirildi")
        
        # "Buraya baca listesi tablosu eklenecek..." placeholder metnini kaldır
        for paragraph in doc.paragraphs:
            if "Buraya baca listesi tablosu eklenecek" in paragraph.text:
                # Paragrafı temizle
                paragraph.clear()
                break
        
        # Baca listesi tablosunu ekle - firma_olcum verilerinden baca ve parametre bilgilerini geç
        print(f"DEBUG: Baca listesi tablosu ekleniyor...")
        add_baca_listesi_table(doc, olcum, olcum_parametreleri)
        print(f"DEBUG: Baca listesi tablosu eklendi")
        
        # "Buraya baca listesi tablosu eklenecek..." placeholder metnini kaldır
        print(f"DEBUG: Placeholder metni temizleniyor...")
        for paragraph in doc.paragraphs:
            if "Buraya baca listesi tablosu eklenecek" in paragraph.text:
                # Paragrafı temizle
                paragraph.clear()
                break
        print(f"DEBUG: Placeholder metni temizlendi")
        
        # Dosyayı kaydet
        print(f"DEBUG: Dosya kaydediliyor...")
        from io import BytesIO
        docx_io = BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        print(f"DEBUG: Dosya başarıyla kaydedildi")
        
        # Dosya adını oluştur
        print(f"DEBUG: Dosya adı oluşturuluyor...")
        def clean_filename(text):
            tr_to_en = {
                'ç': 'c', 'ğ': 'g', 'ı': 'i', 'ö': 'o', 'ş': 's', 'ü': 'u',
                'Ç': 'C', 'Ğ': 'G', 'I': 'I', 'İ': 'I', 'Ö': 'O', 'Ş': 'S', 'Ü': 'U'
            }
            for tr_char, en_char in tr_to_en.items():
                text = text.replace(tr_char, en_char)
            cleaned = "".join(c for c in text if c.isalnum() or c in (' ', '-', '_')).strip()
            cleaned = cleaned.replace(' ', '-')
            while '--' in cleaned:
                cleaned = cleaned.replace('--', '-')
            return cleaned
        
        # Firma adının ilk 2 kelimesini al
        firma_adi = olcum.get('firma_adi', '')
        firma_kelimeleri = firma_adi.split()
        firma_ilk_2_kelime = ' '.join(firma_kelimeleri[:2]) if len(firma_kelimeleri) >= 2 else firma_adi
        
        # Dosya adını oluştur: "FIRMA ADI İKİ KELİME _OLCUM KODU_SAHAFORM"
        firma_ilk_2_kelime_clean = clean_filename(firma_ilk_2_kelime)
        olcum_kodu_clean = clean_filename(olcum.get('olcum_kodu', ''))
        filename = f"{firma_ilk_2_kelime_clean}_{olcum_kodu_clean}_SAHAFORM.docx"
        
        # Response oluştur
        print(f"DEBUG: Response oluşturuluyor...")
        response = make_response(docx_io.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        print(f"DEBUG: Response başarıyla oluşturuldu, filename: {filename}")
        return response
        
    except Exception as e:
        print(f"Firma ölçüm Word doküman oluşturma hatası: {e}")
        import traceback
        print(f"DEBUG: Hata detayı: {traceback.format_exc()}")
        return jsonify({'error': f'Word doküman oluşturma hatası: {str(e)}'}), 500

def add_baca_listesi_detail_table(doc, olcum, baca_parametreleri):
    """Baca listesi tablosunu ekler (detay sayfası için)."""
    try:
        # Docx kütüphanesini yükle
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        
        if not Document:
            print("Word rapor oluşturma için python-docx kütüphanesi gerekli")
            return
        
        # Tüm parametreleri topla
        all_params = set()
        for baca_params in baca_parametreleri.values():
            for param in baca_params:
                all_params.add(param)
        
        # Parametreleri sırala
        sorted_params = sorted(all_params)
        
        # Tablo oluştur
        if baca_parametreleri:
            # Başlık satırı
            doc.add_paragraph()
            title = doc.add_paragraph('BACA-PARAMETRE MATRİSİ')
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title.runs:
                run.font.bold = True
                run.font.size = Pt(14)
            
            # Tablo oluştur
            table = doc.add_table(rows=1, cols=len(sorted_params) + 1)
            table.style = 'Table Grid'
            
            # Başlık satırı
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Baca Adı'
            header_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_cells[0].paragraphs[0].runs[0].font.bold = True
            
            for i, param in enumerate(sorted_params):
                header_cells[i + 1].text = param
                header_cells[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                header_cells[i + 1].paragraphs[0].runs[0].font.bold = True
            
            # Veri satırları
            for baca_adi, parametreler in baca_parametreleri.items():
                row_cells = table.add_row().cells
                row_cells[0].text = baca_adi
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[0].paragraphs[0].runs[0].font.bold = True
                
                for i, param in enumerate(sorted_params):
                    if param in parametreler:
                        row_cells[i + 1].text = 'X'
                        row_cells[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        row_cells[i + 1].paragraphs[0].runs[0].font.bold = True
                        row_cells[i + 1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Yeşil
                    else:
                        row_cells[i + 1].text = ''
                        row_cells[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Tablo genişliğini ayarla
            table.autofit = True
            table.allow_autofit = True
            
    except Exception as e:
        print(f"Baca listesi detay tablosu ekleme hatası: {e}")
def add_baca_listesi_table(doc, olcum, olcum_parametreleri):
    """Baca listesi tablosunu ekler (özel parametre sıralaması ile)."""
    try:
        # Docx kütüphanesini yükle
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        
        if not Document:
            print("Word rapor oluşturma için python-docx kütüphanesi gerekli")
            return
        # Özel parametre sıralaması: Toz-Yg-Voc-Toc (ilk 4 sıra) - büyük/küçük harf düzeltildi
        priority_params = ['Toz', 'Yg', 'Voc', 'Toc']
        
        # Firma ölçüm verilerinden baca ve parametre bilgilerini al
        baca_parametreleri = olcum.get('baca_parametreleri', {})
        
        # Baca isimlerini topla
        baca_isimleri = list(baca_parametreleri.keys())
        
        # Tüm parametreleri topla
        all_params = set()
        for parametreler in baca_parametreleri.values():
            for parametre in parametreler:
                all_params.add(parametre)
        
        # Parametreleri sırala: önce öncelikli olanlar, sonra diğerleri
        sorted_params = []
        
        # Önce öncelikli parametreleri ekle
        for param in priority_params:
            if param in all_params:
                sorted_params.append(param)
                all_params.remove(param)
        
        # Sonra diğer parametreleri ekle
        sorted_params.extend(sorted(all_params))
        
        if not baca_isimleri or not sorted_params:
            return
        
        # Şablonda zaten başlık olabilir, bu yüzden başlık eklemeyelim
        # Sadece tablo ekleyelim - boşluk azaltıldı
        
        # Tablo oluştur (sıra numarası + baca + parametreler)
        table = doc.add_table(rows=len(baca_isimleri) + 1, cols=len(sorted_params) + 2)
        table.style = 'Table Grid'
        
        # Başlık satırı
        header_row = table.rows[0]
        header_row.cells[0].text = "SIRA"
        header_row.cells[0].paragraphs[0].runs[0].font.bold = True
        header_row.cells[1].text = "BACA"
        header_row.cells[1].paragraphs[0].runs[0].font.bold = True
        
        for i, param in enumerate(sorted_params):
            # Parametre sayısını hesapla
            param_count = 0
            for baca_params in baca_parametreleri.values():
                if param in baca_params:
                    param_count += 1
            
            # Başlığa sayıyı ekle
            header_text = f"{param} ({param_count})"
            header_row.cells[i + 2].text = header_text
            header_row.cells[i + 2].paragraphs[0].runs[0].font.bold = True
        
        # Veri satırları
        for i, baca_adi in enumerate(baca_isimleri):
            row = table.rows[i + 1]
            
            # Sıra numarası
            row.cells[0].text = str(i + 1)
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Baca adı
            row.cells[1].text = baca_adi
            row.cells[1].paragraphs[0].runs[0].font.bold = True
            
            # Her parametre için X veya boş kontrol et
            for j, param in enumerate(sorted_params):
                # Bu bacaya ait bu parametre seçili mi kontrol et
                baca_parametreleri = olcum.get('baca_parametreleri', {}).get(baca_adi, [])
                
                if param in baca_parametreleri:
                    # Parametre seçili - X yaz
                    row.cells[j + 2].text = "X"
                    row.cells[j + 2].paragraphs[0].runs[0].font.bold = True
                else:
                    # Parametre seçili değil - boş bırak
                    row.cells[j + 2].text = ""
        
        # Sütun genişliklerini ayarla
        # SIRA sütunu - 0.5 cm
        table.columns[0].width = Inches(0.5 / 2.54)
        
        # BACA sütunu - 3 cm
        table.columns[1].width = Inches(3.0 / 2.54)
        
        # Parametre sütunları - 1.5 cm
        for i in range(2, len(table.columns)):
            table.columns[i].width = Inches(1.5 / 2.54)
        
        # Tablo formatını ayarla
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    except Exception as e:
        print(f"Baca listesi tablosu ekleme hatası: {e}")

@app.route('/api/birlesik_veriler/<firma_adi>/<olcum_kodu>')
def api_birlesik_veriler(firma_adi, olcum_kodu):
    """Bir firmanın baca bilgileri ve parametre ölçümlerini birleştirerek döndürür."""
    try:
        # Verileri yükle
        baca_bilgileri = load_baca_bilgileri()
        parametre_olcumleri = load_parametre_olcum()
        
        # Firma verilerini filtrele
        firma_bacalar = [b for b in baca_bilgileri if b['firma_adi'] == firma_adi and b['olcum_kodu'] == olcum_kodu]
        firma_parametreleri = [p for p in parametre_olcumleri if p['firma_adi'] == firma_adi and p['olcum_kodu'] == olcum_kodu]
        
        # Birleştirilmiş veri yapısı oluştur
        birlesik_veriler = []
        
        for baca in firma_bacalar:
            baca_verisi = {
                'baca_id': baca['id'],
                'baca_adi': baca['baca_adi'],
                'baca_bilgileri': baca['baca_bilgileri'],
                'created_at': baca['created_at'],
                'updated_at': baca['updated_at'],
                'parametreler': []
            }
            
            # Bu bacaya ait parametreleri bul
            baca_parametreleri = [p for p in firma_parametreleri if p['baca_adi'] == baca['baca_adi']]
            
            for parametre in baca_parametreleri:
                baca_verisi['parametreler'].append({
                    'parametre_id': parametre['id'],
                    'parametre_adi': parametre['parametre_adi'],
                    'parametre_verileri': parametre['parametre_verileri'],
                    'created_at': parametre['created_at'],
                    'updated_at': parametre['updated_at']
                })
            
            birlesik_veriler.append(baca_verisi)
        
        return jsonify({
            'firma_adi': firma_adi,
            'olcum_kodu': olcum_kodu,
            'baca_sayisi': len(birlesik_veriler),
            'toplam_parametre_sayisi': len(firma_parametreleri),
            'veriler': birlesik_veriler
        })
        
    except Exception as e:
        print(f"Birleştirilmiş veri hatası: {e}")
        return jsonify({'error': f'Veri birleştirme hatası: {str(e)}'}), 500

@app.route('/api/baca_word_export', methods=['POST'])
def api_baca_word_export():
    """Seçili baca bilgilerini ve ölçüm parametrelerini Word formatında dışa aktarır."""
    try:
        print("=== Word Export Başladı ===")
        data = request.get_json()
        print(f"Gelen veri: {data}")
        
        firma = data.get('firma')
        olcum_kodu = data.get('olcum_kodu')
        baca_adi = data.get('baca_adi')
        
        print(f"Firma: {firma}")
        print(f"Ölçüm Kodu: {olcum_kodu}")
        print(f"Baca Adı: {baca_adi}")
        
        if not all([firma, olcum_kodu, baca_adi]):
            print("Eksik parametreler!")
            return jsonify({'error': 'Eksik parametreler'}), 400
        
        # Baca bilgilerini al
        print("Baca bilgileri yükleniyor...")
        baca_bilgileri = load_baca_bilgileri()
        print(f"Toplam {len(baca_bilgileri)} baca bilgisi bulundu")
        
        baca_bilgisi = None
        
        for baca in baca_bilgileri:
            print(f"Kontrol edilen baca: {baca.get('firma_adi')} - {baca.get('olcum_kodu')} - {baca.get('baca_adi')}")
            if (baca.get('firma_adi') == firma and 
                baca.get('olcum_kodu') == olcum_kodu and 
                baca.get('baca_adi') == baca_adi):
                baca_bilgisi = baca
                print("Baca bilgisi bulundu!")
                break
        
        if not baca_bilgisi:
            print("Baca bilgisi bulunamadı!")
            return jsonify({'error': 'Baca bilgisi bulunamadı'}), 404
        
        # Parametre ölçümlerini al
        parametre_olcumleri = load_parametre_olcum()
        baca_parametreleri = []
        
        for parametre in parametre_olcumleri:
            if (parametre.get('firma_adi') == firma and 
                parametre.get('olcum_kodu') == olcum_kodu and 
                parametre.get('baca_adi') == baca_adi):
                baca_parametreleri.append(parametre)
        
        # Form bilgilerini al (logo ve üst bilgiler için)
        forms = load_forms()
        form_bilgisi = None
        if forms:
            # İlk formu kullan (veya EMİSYON FORMU olanı)
            for form in forms:
                if 'EMİSYON' in form.get('form_adi', '').upper():
                    form_bilgisi = form
                    break
            if not form_bilgisi:
                form_bilgisi = forms[0]
        
        # Word dosyası oluştur (sadece şablon kullan)
        print("Word dosyası oluşturuluyor...")
        doc = create_baca_word_document_from_template(baca_bilgisi, baca_parametreleri, form_bilgisi)
        
        # Eğer doc None ise, hata döndür
        if doc is None:
            return jsonify({'error': 'Word dosyası oluşturulurken hata oluştu'}), 500
        
        # Eğer doc bir Document objesi ise, Flask response'a çevir
        if hasattr(doc, 'save'):
            # Geçici dosya oluştur
            import tempfile
            import os
            from io import BytesIO
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                doc.save(tmp_file.name)
                tmp_file_path = tmp_file.name
            
            # Dosyayı oku
            with open(tmp_file_path, 'rb') as f:
                doc_content = f.read()
            
            # Geçici dosyayı sil
            os.unlink(tmp_file_path)
            
            # Dosya adını oluştur
            firma_adi = baca_bilgisi.get('firma_adi', '')
            olcum_kodu = baca_bilgisi.get('olcum_kodu', '')
            baca_adi = baca_bilgisi.get('baca_adi', '')
            
            def clean_filename(text):
                tr_to_en = {
                    'ç': 'c', 'ğ': 'g', 'ı': 'i', 'ö': 'o', 'ş': 's', 'ü': 'u',
                    'Ç': 'C', 'Ğ': 'G', 'I': 'I', 'İ': 'I', 'Ö': 'O', 'Ş': 'S', 'Ü': 'U'
                }
                for tr_char, en_char in tr_to_en.items():
                    text = text.replace(tr_char, en_char)
                cleaned = "".join(c for c in text if c.isalnum() or c in (' ', '-', '_')).strip()
                cleaned = cleaned.replace(' ', '-')
                while '--' in cleaned:
                    cleaned = cleaned.replace('--', '-')
                return cleaned
            
            firma_adi_clean = clean_filename(firma_adi)
            olcum_kodu_clean = clean_filename(olcum_kodu)
            baca_adi_clean = clean_filename(baca_adi)
            
            filename = f"{firma_adi_clean}-{baca_adi_clean}-{olcum_kodu_clean}.docx"
            
            # Flask response oluştur
            response = make_response(doc_content)
            response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
        else:
            # Eğer doc bir hata response'u ise, direkt döndür
            return doc
        
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"Baca Word export hatası: {e}")
        print(f"Hata detayı: {error_details}")
        return jsonify({'error': f'Word dosyası oluşturulurken hata oluştu: {str(e)}', 'details': error_details}), 500

def export_forms_excel(forms):
    """Formları Excel formatında dışa aktarır."""
    try:
        # DataFrame oluştur
        df_data = []
        for form in forms:
            df_data.append({
                'FORM ADI': form.get('formAdi', ''),
                'FORM KODU': form.get('formKodu', ''),
                'YAYIN TARİHİ': form.get('yayinTarihi', ''),
                'REVİZYON TARİHİ': form.get('revizyonTarihi', ''),
                'REVİZYON NO': form.get('revizyonNo', ''),
                'OLUŞTURMA TARİHİ': form.get('created_at', ''),
                'GÜNCELLEME TARİHİ': form.get('updated_at', '')
            })
        
        df = pd.DataFrame(df_data)
        
        # Excel dosyası oluştur
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Formlar', index=False)
            
            # Sütun genişliklerini ayarla
            worksheet = writer.sheets['Formlar']
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                worksheet.column_dimensions[column_letter].width = adjusted_width
        
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'Formlar_{datetime.now().strftime("%Y%m%d_%H%M")}.xlsx'
        )
        
    except Exception as e:
        print(f"Excel dışa aktarma hatası: {e}")
        raise e

def export_forms_word(forms):
    """Formları Word formatında dışa aktarır."""
    try:
        # Docx kütüphanesini yükle
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        
        if not Document:
            return jsonify({'error': 'python-docx kütüphanesi yüklü değil'}), 500
        
        doc = Document()
        
        # Başlık
        title = doc.add_heading('FORM LİSTESİ', 0)
        title.alignment = 1  # Ortalı
        
        # Tarih
        doc.add_paragraph(f'Oluşturma Tarihi: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
        doc.add_paragraph('')
        
        # Tablo oluştur
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Başlık satırı
        header_cells = table.rows[0].cells
        headers = ['SIRA', 'FORM ADI', 'FORM KODU', 'YAYIN TARİHİ', 'REVİZYON TARİHİ', 'REVİZYON NO']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        # Veri satırları
        for i, form in enumerate(forms, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = form.get('formAdi', '')
            row_cells[2].text = form.get('formKodu', '')
            row_cells[3].text = form.get('yayinTarihi', '')
            row_cells[4].text = form.get('revizyonTarihi', '')
            row_cells[5].text = form.get('revizyonNo', '')
        
        # Dosyayı kaydet
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'Formlar_{datetime.now().strftime("%Y%m%d_%H%M")}.docx'
        )
        
    except Exception as e:
        print(f"Word dışa aktarma hatası: {e}")
        raise e

def export_forms_pdf(forms):
    return 'PDF özelliği bu ortamda devre dışı (WeasyPrint eksik veya desteklenmiyor).', 501

def create_firma_raporu(firma_adi, olcum_kodu):
    """
    Bir firmanın tüm bacalarında yapılan tüm parametreleri içeren rapor oluşturur.
    JSON şablonunu kullanarak esnek rapor formatı sağlar.
    """
    try:
        # Şablonu yükle
        with open('rapor_sablonu.json', 'r', encoding='utf-8') as f:
            sablon = json.load(f)
        
        # Verileri yükle
        baca_bilgileri = load_baca_bilgileri()
        parametre_olcumleri = load_parametre_olcum()
        
        # Firma verilerini filtrele
        firma_bacalar = [b for b in baca_bilgileri if b['firma_adi'] == firma_adi and b['olcum_kodu'] == olcum_kodu]
        firma_parametreleri = [p for p in parametre_olcumleri if p['firma_adi'] == firma_adi and p['olcum_kodu'] == olcum_kodu]
        
        if not firma_bacalar:
            return "Firma verisi bulunamadı!", 404
        
        # Toplam sayfa sayısını hesapla (sadece header var)
        toplam_sayfa = 1  # Sadece 1 sayfa (header ile)
        
        # Debug için sayfa hesaplamasını yazdır
        print(f"Sayfa hesaplaması: Sadece header, Toplam={toplam_sayfa} sayfa")
        
        # Emisyon formu verilerini yükle
        emisyon_formu = None
        try:
            with open('forms.json', 'r', encoding='utf-8') as f:
                forms_data = json.load(f)
                emisyon_formu = next((form for form in forms_data if form['formAdi'] == 'EMİSYON ÖLÇÜM FORMU'), None)
        except Exception as e:
            print(f"Form verileri yüklenirken hata: {e}")
        
        # Tarih formatını GG.AA.YY şeklinde düzenle
        def format_tarih(tarih_str):
            try:
                if not tarih_str:
                    return ""
                # YYYY-MM-DD formatını GG.AA.YY'ye çevir
                dt = datetime.strptime(tarih_str, '%Y-%m-%d')
                return dt.strftime('%d.%m.%y')
            except:
                return tarih_str
        
        # Form bilgilerini hazırla
        if emisyon_formu:
            yayin_tarihi = format_tarih(emisyon_formu['yayinTarihi'])
            revizyon_tarihi = format_tarih(emisyon_formu['revizyonTarihi'])
            form_kodu = emisyon_formu['formKodu']
            revizyon_no = emisyon_formu['revizyonNo']
        else:
            yayin_tarihi = "01.08.15"
            revizyon_tarihi = "29.02.24"
            form_kodu = "AÇ.F.52"
            revizyon_no = "03"
        
        # Docx kütüphanesini yükle
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        
        if not Document:
            return "Word rapor oluşturma için python-docx kütüphanesi gerekli", 500
        
        # Word dokümanı oluştur
        doc = Document()
        
        # Sayfa ayarları
        section = doc.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
        # Header oluştur - Sadece 1 satır
        header = section.header
        header_table = header.add_table(rows=1, cols=3, width=Inches(8))
        header_table.style = 'Table Grid'
        
        # Logo hücresi - 3cm genişlik
        logo_cell = header_table.cell(0, 0)
        logo_cell.width = Inches(1.2)  # 3cm
        
        # Logo ekle - FORM54.png
        try:
            logo_path = "static/images/FORM54.png"
            if os.path.exists(logo_path):
                logo_cell.text = ""
                logo_paragraph = logo_cell.paragraphs[0]
                logo_run = logo_paragraph.add_run()
                logo_run.add_picture(logo_path, width=Inches(0.8), height=Inches(0.6))
                logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"Logo eklenirken hata: {e}")
        
        # Başlık hücresi - 9cm genişlik
        baslik_cell = header_table.cell(0, 1)
        baslik_cell.width = Inches(3.5)  # 9cm
        baslik_cell.text = "EMİSYON ÖLÇÜM FORMU"
        baslik_paragraph = baslik_cell.paragraphs[0]
        baslik_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        baslik_run = baslik_paragraph.runs[0]
        baslik_run.font.size = Pt(16)
        baslik_run.font.bold = True
        baslik_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Form bilgileri hücresi (sağ sütun) - Dinamik sayfa numarası ile
        form_cell = header_table.cell(0, 2)
        form_cell.text = f"Form Kodu: {form_kodu}\n"
        form_cell.text += f"Yayın Tarihi: {yayin_tarihi}\n"
        form_cell.text += f"Revizyon No: {revizyon_no}\n"
        form_cell.text += f"Revizyon Tarihi: {revizyon_tarihi}\n"
        form_cell.text += f"Sayfa No: 1/{toplam_sayfa}"  # Dinamik sayfa numarası
        
        # Header satır yüksekliğini ayarla - 1 satır 0.4cm
        for row in header_table.rows:
            row.height = Inches(0.16)  # 0.4cm
        
        # Header ile sayfa üstü arasında 3mm boşluk
        section.top_margin = Inches(0.12)  # 3mm
        section.header_distance = Inches(0.12)  # 3mm
        
        # Ana içerik - Baca bilgileri tablosu
        doc.add_paragraph()  # Boşluk
        
        # Baca bilgileri başlığı
        baca_baslik = doc.add_heading("BACA BİLGİLERİ", level=1)
        baca_baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if firma_bacalar:
            # Baca bilgileri tablosu - 4 sütun x 4 satır
            baca_table = doc.add_table(rows=4, cols=4)
            baca_table.style = 'Table Grid'
            
            # Tablo genişliğini ayarla
            baca_table.autofit = False
            for row in baca_table.rows:
                row.height = Inches(0.3)  # Satır yüksekliği
            
            # Sütun genişliklerini ayarla
            for i, col in enumerate(baca_table.columns):
                if i == 0:  # İlk sütun (etiketler)
                    col.width = Inches(2.5)
                else:  # Diğer sütunlar (değerler)
                    col.width = Inches(1.5)
            
            # 1. Satır
            baca_table.cell(0, 0).text = "FİRMA ADI:"
            baca_table.cell(0, 1).text = firma_adi
            baca_table.cell(0, 2).text = "ÖLÇÜM KODU:"
            baca_table.cell(0, 3).text = olcum_kodu
            
            # 2. Satır
            baca_table.cell(1, 0).text = "BACA ADI:"
            baca_table.cell(1, 1).text = firma_bacalar[0]['baca_adi'] if firma_bacalar else ""
            baca_table.cell(1, 2).text = "BACA NO:"
            baca_table.cell(1, 3).text = firma_bacalar[0]['baca_bilgileri'].get('6774fc7e-c124-4272-b826-482d064f3215', '') if firma_bacalar else ""
            
            # 3. Satır
            baca_table.cell(2, 0).text = "ISIL GÜÇ:"
            baca_table.cell(2, 1).text = firma_bacalar[0]['baca_bilgileri'].get('2360aa46-81dd-41db-a6b6-5c23e62900fe', '') if firma_bacalar else ""
            baca_table.cell(2, 2).text = "KAYNAK TÜRÜ:"
            baca_table.cell(2, 3).text = firma_bacalar[0]['baca_bilgileri'].get('7647078a-4f69-4908-99f7-3fff5651cc9b', '') if firma_bacalar else ""
            
            # 4. Satır
            baca_table.cell(3, 0).text = "YAKIT TÜRÜ:"
            baca_table.cell(3, 1).text = firma_bacalar[0]['baca_bilgileri'].get('2517064c-c286-4210-8f60-fa0a6b9e22a9', '') if firma_bacalar else ""
            baca_table.cell(3, 2).text = "ÇATI ŞEKLİ:"
            baca_table.cell(3, 3).text = firma_bacalar[0]['baca_bilgileri'].get('f1e7b875-ddf4-4628-87a3-eda6afa775b7', '') if firma_bacalar else ""
            
            # Tüm hücrelerde metin hizalaması
            for row in baca_table.rows:
                for cell in row.cells:
                    # Paragraf hizalaması
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        # İlk karakteri kalın yap (etiketler için)
                        if paragraph.text and paragraph.text.endswith(':'):
                            if paragraph.runs:
                                paragraph.runs[0].font.bold = True
                    
                    # Hücre dikey hizalaması
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Dosyayı kaydet
        from io import BytesIO
        docx_io = BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        
        # Response oluştur
        response = make_response(docx_io.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename={firma_adi}_{olcum_kodu}_Raporu.docx'
        
        return response
        
    except Exception as e:
        print(f"Rapor oluşturma hatası: {e}")
        return f"Rapor oluşturma hatası: {str(e)}", 500



def create_baca_word_document_from_template(baca_bilgisi, baca_parametreleri, form_bilgisi):
    """Word şablonunu kullanarak baca raporu oluşturur."""
    try:
        print("=== Şablon Word Dosyası Oluşturma Başladı ===")
        print(f"Baca bilgisi: {baca_bilgisi}")
        print(f"Parametre sayısı: {len(baca_parametreleri) if baca_parametreleri else 0}")
        
        # Docx kütüphanesini yükle
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        
        if not Document:
            print("python-docx kütüphanesi yüklü değil!")
            return jsonify({'error': 'python-docx kütüphanesi yüklü değil'}), 500
        
        # Şablon dosyasını yükle
        template_path = os.path.join(os.path.dirname(__file__), 'static', 'images', 'F52 EMISYON.docx')
        print(f"Şablon yolu: {template_path}")
        print(f"Şablon dosyası var mı: {os.path.exists(template_path)}")
        
        if not os.path.exists(template_path):
            print("Word şablonu bulunamadı!")
            return jsonify({'error': 'Word şablonu bulunamadı'}), 500
        
        # Şablonu kopyala
        doc = Document(template_path)
        
        # Sayfa kenar boşluklarını ayarla - Sol ve sağ 5 mm
        sections = doc.sections
        for section in sections:
            section.left_margin = Inches(5/25.4)  # 5mm = 5/25.4 inç
            section.right_margin = Inches(5/25.4)  # 5mm = 5/25.4 inç
        
        # Boş değerleri "-" ile değiştiren yardımcı fonksiyon
        def get_value_or_dash(value):
            if value is None or value == '' or str(value).strip() == '':
                return '-'
            return str(value).strip()
        
        # Veri hazırla - Doğru ID'leri kullan ve boş değerleri "*" ile değiştir
        data = {
            'FIRMA_ADI': baca_bilgisi.get('firma_adi', '').title(),  # İlk karakterler büyük
            'OLCUM_KODU': baca_bilgisi.get('olcum_kodu', ''),
            'BACA_ADI': baca_bilgisi.get('baca_adi', '').title(),  # İlk karakterler büyük
            'BACA_NO': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('597fad80-d28f-40ea-bd28-a76c61c5203d', '')),
            'ISIL_GUC': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('22867c9a-ca3c-4d80-b017-b73dafdd7fef', '')),
            'KAYNAK_TURU': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('98399625-5bbc-465e-8e09-de454f231ae4', '')),
            'YAKIT_TURU': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('ddca398d-0e55-4662-b661-3731e0975bd2', '')),
            'CATI_SEKLI': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('6b3546e0-184c-49de-82e4-e2835e81923b', '')),
            'BACA_SEKLI': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('d9958774-43f7-4bc3-8e12-436614a6193a', '')),
            'BACA_OLCUSU': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('b1b6fc38-98c0-4048-8b8e-795cf7d44c48', '')),
            'YERDEN_YUKSEKLIK': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('6a301d72-f21b-485b-b8fb-116ad5cb223f', '')),
            'CATI_YUK': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('8ec9ecc9-ecda-4bf2-9802-02fa2e3fda4c', '')),
            'RUZGAR_HIZ': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('ab2b67dd-16a5-4bee-9b6e-b60b8cfc2d0c', '')),
            'ORT_SIC': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('eca60e54-ec39-4412-8884-caa17faed0be', '')),
            'ORT_NEM': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('64238e0a-6387-4c31-9bf4-d7f800ef17e1', '')),
            'ORT_BAS': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('b09ad69a-e4d4-4219-b055-2cf923ffd499', '')),
            'A_BACA': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('9c8c8bcf-c98e-4109-8b10-63b08b26460e', '')),
            'B_BACA': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('af55c55f-f83b-4b90-a655-ee76bf6bb2ac', '')),
            'C_DELIK': get_value_or_dash(baca_bilgisi.get('baca_bilgileri', {}).get('20881447-f7c8-4a6b-8583-76c7246082ef', '')),
            'PERSONEL': get_value_or_dash(baca_bilgisi.get('personel_adi', '')),
            'TARIH': datetime.now().strftime('%d.%m.%Y')
        }
        
        # Şablondaki yer tutucuları değiştir
        replace_placeholders_in_document(doc, data)
        
        # Parametre ölçümlerini ekle
        if baca_parametreleri:
            add_parametre_measurements_to_document(doc, baca_parametreleri)
        
        # Dosya adını temizle
        def clean_filename(text):
            # Türkçe karakterleri İngilizce karşılıklarıyla değiştir
            tr_to_en = {
                'ç': 'c', 'ğ': 'g', 'ı': 'i', 'ö': 'o', 'ş': 's', 'ü': 'u',
                'Ç': 'C', 'Ğ': 'G', 'I': 'I', 'İ': 'I', 'Ö': 'O', 'Ş': 'S', 'Ü': 'U'
            }
            for tr_char, en_char in tr_to_en.items():
                text = text.replace(tr_char, en_char)
            # Sadece alfanumerik karakterler ve bazı özel karakterler
            cleaned = "".join(c for c in text if c.isalnum() or c in (' ', '-', '_')).strip()
            # Boşlukları tire ile değiştir
            cleaned = cleaned.replace(' ', '-')
            # Birden fazla tireyi tek tireye çevir
            while '--' in cleaned:
                cleaned = cleaned.replace('--', '-')
            return cleaned
        
        # Dosya adını oluştur
        firma_adi = baca_bilgisi.get('firma_adi', '')
        olcum_kodu = baca_bilgisi.get('olcum_kodu', '')
        baca_adi = baca_bilgisi.get('baca_adi', '')
        
        firma_adi_clean = clean_filename(firma_adi)
        olcum_kodu_clean = clean_filename(olcum_kodu)
        baca_adi_clean = clean_filename(baca_adi)
        
        filename = f"{firma_adi_clean}-{baca_adi_clean}-{olcum_kodu_clean}.docx"
        
        # Document objesini döndür (sayfa sonu ekleme işlemi merge_word_documents'da yapılacak)
        return doc
        
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"Şablon rapor oluşturma hatası: {e}")
        print(f"Hata detayı: {error_details}")
        # Hata durumunda jsonify döndür
        return jsonify({'error': f'Şablon rapor oluşturma hatası: {str(e)}'}), 500

def add_parametre_measurements_to_document(doc, baca_parametreleri):
    """Word dokümanına parametre ölçümlerini ekler."""
    try:
        if baca_parametreleri:
            # Boşluk ekle
            doc.add_paragraph()
            
            # Parametre tablolarını ekle
            baca_adi = baca_parametreleri[0].get('baca_adi', '') if baca_parametreleri else ''
            for parametre in baca_parametreleri:
                add_parametre_details_to_document(doc, parametre, baca_adi)
            
    except Exception as e:
        print(f"Parametre ölçümleri ekleme hatası: {e}")
def add_parametre_details_to_document(doc, parametre, baca_adi=""):
    """Her parametre için detay bilgileri ekler."""
    try:
        # Docx kütüphanesini yükle
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        
        if not Document:
            print("Word rapor oluşturma için python-docx kütüphanesi gerekli")
            return
        parametre_adi = parametre.get('parametre_adi', '')
        parametre_verileri = parametre.get('parametre_verileri', {})
        
        if not parametre_verileri:
            return
        
        # Parametre başlığı - 2-3 satır sola kaydır
        if baca_adi:
            subtitle_para = doc.add_paragraph(f"    {parametre_adi.upper()} PARAMETRE ÖLÇÜM DETAYLARI ({baca_adi})")
        else:
            subtitle_para = doc.add_paragraph(f"    {parametre_adi.upper()} PARAMETRE ÖLÇÜM DETAYLARI")
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        subtitle_run = subtitle_para.runs[0]
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.bold = True
        subtitle_run.font.name = 'Arial'
        try:
            subtitle_run.font.color.rgb = RGBColor(0, 0, 0)
        except:
            pass
        
                    # Detay bilgileri tablosu
        if parametre_verileri:
            # Boş değerleri "-" ile değiştiren yardımcı fonksiyon
            def get_value_or_dash(value):
                if value is None or value == '' or str(value).strip() == '':
                    return '-'
                return str(value).strip()
            
            # Parametre verilerini al ve sırala
            details = []
            for key, value in parametre_verileri.items():
                key_name = key
                details.append([key_name, get_value_or_dash(value)])
            
            # Detayları mantıklı sıraya koy
            priority_order = [
                'TARİH', 'METOT', 'NOZZLE ÇAP', 'TRAVERS', 'B.HIZ', 'B.SIC', 'B.BAS(KPA)', 
                'B.NEM(G/M3)', 'B.NEM(%)', 'SYC.HAC.', 'SYC.İLK', 'SYC.SON', 'SYC.SIC', 
                'DEBİ', 'ISDL', 'AA', 'SS', 'DD'
            ]
            
            # Öncelikli sıraya göre sırala
            sorted_details = []
            for priority_key in priority_order:
                for detail in details:
                    if detail[0] == priority_key:
                        sorted_details.append(detail)
                        break
            
            # Sıralanmamış olanları da ekle
            for detail in details:
                if detail not in sorted_details:
                    sorted_details.append(detail)
            
            if sorted_details:
                # 6 sütunlu detay tablosu oluştur (daha dar tablo)
                detail_table = doc.add_table(rows=1, cols=6)
                detail_table.style = 'Table Grid'
                detail_table.autofit = True
                
                # Tablo genişliğini manuel olarak ayarla
                detail_table.allow_autofit = False
                detail_table.autofit = False
                
                # Tabloyu 5 mm sola kaydır (Word şablonu gibi)
                detail_table.alignment = WD_ALIGN_PARAGRAPH.LEFT
                detail_table.left_indent = Inches(0.5/2.54)  # 5 mm sola kaydır
                
                # Her sütun için genişlik ayarla - 3,5 cm = 3.5/2.54 inç (daha geniş)
                for i, column in enumerate(detail_table.columns):
                    column.width = Inches(3.5/2.54)  # Her sütun 3,5 cm genişlik
                
                # Başlık satırı
                header_row = detail_table.rows[0]
                header_row.cells[0].text = "Ölç.Param."
                header_row.cells[1].text = "Değer"
                header_row.cells[2].text = "Ölç.Param."
                header_row.cells[3].text = "Değer"
                header_row.cells[4].text = "Ölç.Param."
                header_row.cells[5].text = "Değer"
                
                # Başlık formatla
                for cell in header_row.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].runs[0].font.size = Pt(11)  # 11 punto
                        cell.paragraphs[0].runs[0].font.name = 'Arial'
                    
                    # Başlık hücre genişliğini ayarla - 3,5 cm
                    cell.width = Inches(3.5/2.54)
                
                # Detay verilerini 6 sütunlu olarak ekle (3 parametre-değer çifti)
                for i in range(0, len(sorted_details), 3):
                    row = detail_table.add_row()
                    
                    # İlk parametre
                    row.cells[0].text = sorted_details[i][0]
                    row.cells[1].text = sorted_details[i][1]
                    
                    # İkinci parametre (varsa)
                    if i + 1 < len(sorted_details):
                        row.cells[2].text = sorted_details[i + 1][0]
                        row.cells[3].text = sorted_details[i + 1][1]
                    else:
                        row.cells[2].text = ""
                        row.cells[3].text = ""
                    
                    # Üçüncü parametre (varsa)
                    if i + 2 < len(sorted_details):
                        row.cells[4].text = sorted_details[i + 2][0]
                        row.cells[5].text = sorted_details[i + 2][1]
                    else:
                        row.cells[4].text = ""
                        row.cells[5].text = ""
                    
                    # Hücre formatla
                    for j in range(6):
                        if j % 2 == 0:  # Parametre adı sütunları
                            row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        else:  # Değer sütunları
                            row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Run kontrolü yap
                        if row.cells[j].paragraphs[0].runs:
                            row.cells[j].paragraphs[0].runs[0].font.size = Pt(11)  # 11 punto
                            row.cells[j].paragraphs[0].runs[0].font.name = 'Arial'
                        
                        # Hücre genişliğini ayarla - 3,5 cm
                        row.cells[j].width = Inches(3.5/2.54)
 
        # Tablo sonrası 1 boşluk
        doc.add_paragraph()
          
        
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"Parametre detayları ekleme hatası: {e}")
        print(f"Hata detayı: {error_details}")
        # Hatayı yukarı fırlatma, sadece logla
        raise e

def replace_placeholders_in_document(doc, data):
    """Word dokümanındaki yer tutucuları gerçek verilerle değiştirir."""
    # Docx kütüphanesini yükle
    Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
    
    if not Document:
        print("Word rapor oluşturma için python-docx kütüphanesi gerekli")
        return
    
    # Paragraflardaki yer tutucuları değiştir
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        
        for key, value in data.items():
            # Hem {{KEY}} hem de {KEY} formatlarını kontrol et
            placeholder1 = f"{{{{{key}}}}}"
            placeholder2 = f"{{{key}}}"
            if placeholder1 in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder1, str(value))
            elif placeholder2 in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder2, str(value))
        
        # Eğer metin değiştiyse ve başlık satırıysa formatla
        if paragraph.text != original_text:
            # Başlık satırlarını kontrol et
            if any(key in original_text for key in ['FIRMA_ADI', 'OLCUM_KODU', 'BACA_ADI']):
                # Başlık formatlaması
                for run in paragraph.runs:
                    run.font.size = Pt(14)  # 14 punto
                    run.font.bold = True
                    run.font.name = 'Arial'
                    # RGBColor kullanımını güvenli hale getir
                    try:
                        run.font.color.rgb = RGBColor(0, 0, 0)  # Siyah renk
                    except:
                        # RGBColor kullanılamıyorsa varsayılan rengi kullan
                        pass
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Sola dayalı
    
    # Tablolardaki yer tutucuları değiştir
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        # Hem {{KEY}} hem de {KEY} formatlarını kontrol et
                        placeholder1 = f"{{{{{key}}}}}"
                        placeholder2 = f"{{{key}}}"
                        if placeholder1 in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder1, str(value))
                        elif placeholder2 in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder2, str(value))
    
    # Header ve footer'daki yer tutucuları değiştir
    for section in doc.sections:
        # Header
        if section.header:
            for paragraph in section.header.paragraphs:
                for key, value in data.items():
                    # Hem {{KEY}} hem de {KEY} formatlarını kontrol et
                    placeholder1 = f"{{{{{key}}}}}"
                    placeholder2 = f"{{{key}}}"
                    if placeholder1 in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder1, str(value))
                    elif placeholder2 in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder2, str(value))
        
        # Footer
        if section.footer:
            for paragraph in section.footer.paragraphs:
                for key, value in data.items():
                    # Hem {{KEY}} hem de {KEY} formatlarını kontrol et
                    placeholder1 = f"{{{{{key}}}}}"
                    placeholder2 = f"{{{key}}}"
                    if placeholder1 in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder1, str(value))
                    elif placeholder2 in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder2, str(value))

def merge_word_documents(documents):
    """Birden fazla Word dokümanını tek bir dokümanda birleştirir."""
    if not documents:
        return None
    
    # İlk dokümanı ana doküman olarak kullan
    merged_doc = documents[0]
    
    # Diğer dokümanları ekle
    for i in range(1, len(documents)):
        doc = documents[i]
        # Dokümanın tüm içeriğini ana dokümana ekle
        for element in doc.element.body:
            merged_doc.element.body.append(element)
        
        # Sayfa sonu ekle (son doküman hariç)
        if i < len(documents) - 1:
            merged_doc.add_page_break()
    
    return merged_doc

def create_firma_raporu_from_template(firma_adi, olcum_kodu):
    """Her baca için ayrı Word dosyası oluşturup birleştirir."""
    try:
        print("=== CREATE_FIRMA_RAPORU_FROM_TEMPLATE BAŞLADI ===")
        
        # Docx kütüphanesini yükle
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        print(f"load_docx sonucu - Document: {Document is not None}")
        
        if not Document:
            print("python-docx kütüphanesi yüklü değil!")
            return jsonify({'error': 'python-docx kütüphanesi yüklü değil'}), 500
        
        # Firma verilerini al
        print("Baca bilgileri yükleniyor...")
        baca_bilgileri = load_baca_bilgileri()
        print(f"Toplam baca bilgisi sayısı: {len(baca_bilgileri)}")
        
        firma_bacalar = [baca for baca in baca_bilgileri if baca.get('firma_adi') == firma_adi and baca.get('olcum_kodu') == olcum_kodu]
        print(f"Firma için bulunan baca sayısı: {len(firma_bacalar)}")
        
        if not firma_bacalar:
            print("Firma için baca bilgisi bulunamadı!")
            return jsonify({'error': 'Firma için baca bilgisi bulunamadı'}), 404
        
        # Parametre ölçümlerini al
        parametre_olcumleri = load_parametre_olcum()
        
        # Form bilgilerini al
        forms = load_forms()
        form_bilgisi = None
        for form in forms:
            if form.get('formKodu') == 'AÇ.F.52':
                form_bilgisi = form
                break
        if not form_bilgisi:
            form_bilgisi = forms[0] if forms else {}
        
        # Her baca için ayrı Word dosyası oluştur
        baca_documents = []
        
        for baca_bilgisi in firma_bacalar:
            # Baca parametrelerini al
            baca_adi = baca_bilgisi.get('baca_adi', '')
            baca_parametreleri = [param for param in parametre_olcumleri 
                                 if param.get('firma_adi') == firma_adi 
                                 and param.get('olcum_kodu') == olcum_kodu 
                                 and param.get('baca_adi') == baca_adi]
            
            print(f"Baca: {baca_adi}, Parametre sayısı: {len(baca_parametreleri)}")
            for param in baca_parametreleri:
                print(f"  - {param.get('parametre_adi', '')}: {len(param.get('parametre_verileri', {}))} veri")
            
            # Her baca için ayrı Word dosyası oluştur (mevcut fonksiyonu kullan)
            baca_doc = create_baca_word_document_from_template(baca_bilgisi, baca_parametreleri, form_bilgisi)
            
            # Hata kontrolü
            if baca_doc is None:
                print(f"Baca {baca_adi} için doküman oluşturulamadı!")
                continue
            
            # Dokümanı listeye ekle
            baca_documents.append(baca_doc)
        
        # Tüm dokümanları birleştir
        if baca_documents:
            # İlk dokümanı ana doküman olarak kullan
            merged_doc = baca_documents[0]
            
            # Diğer dokümanları ekle (her biri yeni sayfadan başlasın)
            for i in range(1, len(baca_documents)):
                # Sayfa sonu ekle
                merged_doc.add_page_break()
                
                # Dokümanın tüm içeriğini ana dokümana ekle
                doc = baca_documents[i]
                for element in doc.element.body:
                    merged_doc.element.body.append(element)
        else:
            return jsonify({'error': 'Hiç baca dokümanı oluşturulamadı'}), 500
        
        # Dosya adını temizle (Türkçe karakterleri ve geçersiz karakterleri kaldır)
        def clean_filename(text):
            # Türkçe karakterleri İngilizce karşılıklarıyla değiştir
            tr_to_en = {
                'ç': 'c', 'ğ': 'g', 'ı': 'i', 'ö': 'o', 'ş': 's', 'ü': 'u',
                'Ç': 'C', 'Ğ': 'G', 'I': 'I', 'İ': 'I', 'Ö': 'O', 'Ş': 'S', 'Ü': 'U'
            }
            for tr_char, en_char in tr_to_en.items():
                text = text.replace(tr_char, en_char)
            # Sadece alfanumerik karakterler ve bazı özel karakterler
            cleaned = "".join(c for c in text if c.isalnum() or c in (' ', '-', '_')).strip()
            # Boşlukları tire ile değiştir
            cleaned = cleaned.replace(' ', '-')
            # Birden fazla tireyi tek tireye çevir
            while '--' in cleaned:
                cleaned = cleaned.replace('--', '-')
            return cleaned
        
        # Dosya adını oluştur
        firma_adi_clean = clean_filename(firma_adi)
        olcum_kodu_clean = clean_filename(olcum_kodu)
        filename = f"F54-{firma_adi_clean}-{olcum_kodu_clean}-FirmaRaporu.docx"
        
        # Dosyayı kaydet
        from io import BytesIO
        docx_io = BytesIO()
        merged_doc.save(docx_io)
        docx_io.seek(0)
        
        # Response oluştur
        response = make_response(docx_io.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except Exception as e:
        print(f"Firma şablon rapor oluşturma hatası: {e}")
        return jsonify({'error': f'Firma şablon rapor oluşturma hatası: {str(e)}'}), 500

# Parametre ölçümleri export API'leri
@app.route('/api/parametre_olcumleri_excel_export', methods=['POST'])
def api_parametre_olcumleri_excel_export():
    """Parametre ölçümlerini Excel formatında dışa aktarır."""
    try:
        print("Excel export başlatıldı...")
        data = request.get_json()
        olcum_ids = data.get('olcum_ids', [])
        
        print(f"Gelen olcum_ids: {olcum_ids}")
        
        if not olcum_ids:
            return jsonify({'error': 'En az bir ölçüm ID gerekli!'}), 400
        
        # Parametre ölçümlerini yükle
        parametre_olcumleri = load_parametre_olcum()
        print(f"Toplam parametre ölçümü sayısı: {len(parametre_olcumleri)}")
        
        # Seçili ölçümleri filtrele
        if olcum_ids:
            filtered_olcumler = [item for item in parametre_olcumleri if str(item.get('id')) in olcum_ids]
        else:
            filtered_olcumler = parametre_olcumleri
        
        print(f"Filtrelenmiş ölçüm sayısı: {len(filtered_olcumler)}")
        
        if not filtered_olcumler:
            return jsonify({'error': 'Dışa aktarılacak veri bulunamadı'}), 404
        
        # Tüm parametre alanlarını tanımla (frontend ile aynı sırada)
        allParametreFields = [
            'TARİH', 'METOT', 'NOZZLE ÇAP', 'TRAVERS', 'B.HIZ', 'B.SIC', 'B.BAS(KPA)', 
            'B.NEM(G/M3)', 'B.NEM(%)', 'SYC.HAC.', 'SYC.İLK', 'SYC.SON', 'SYC.SIC', 'DEBİ', 'ISDL',
            '1İMP-İ', '1-İMP-S', '2-İMP-İ', '2-İMP-S', '3-İMP-İ', '3-İMP-S', 'HAC.',
            'O2', 'CO', 'NO', 'NOX', 'SO2', 'KK1-O2', 'KK1-CO', 'KK1-NO', 'KK1-SO2', 
            'KK2-O2', 'KK2-CO', 'KK2-NO', 'KK2-SO2', 'T90',
                            'TOC(PPM)', 'KK1-SPAN', 'KK1-SPAN', 'KK2-O', 'KK2-SPAN',
            'GAZ HAC.', 'GAZ.SIC.', 'SEY.GAZ.HAC', 'SEY.GAZ.SIC',
            'ORT.SIC', 'ORT.NEM', 'ORT.RUZ.HIZ', 'ÇEK.HACİM',
            'T.İÇİ-1', 'T.İÇİ-2', 'T.İÇİ-3', 'T.İÇİ-4', 'T-DIŞ-1', 'T.DIŞ-2', 
            'T.DIŞ-3', 'T.DIŞ-4', 'İLK KURULUM', '2. KURULUM'
        ]
        
        print("DataFrame oluşturuluyor...")
        
        # Excel için DataFrame oluştur - ekranda gördüğünüz format
        df_data = []
        for index, olcum in enumerate(filtered_olcumler, 1):
            # Temel bilgiler (ekranda gördüğünüz sütunlar)
            row = {
                '#': index,
                'Firma': olcum.get('firma_adi', ''),
                'Ölçüm Kodu': olcum.get('olcum_kodu', ''),
                'Baca': olcum.get('baca_adi', ''),
                'Parametre': olcum.get('parametre_adi', ''),
                'Kayıt Tarihi': olcum.get('created_at', '')[:19].replace('T', ' ') if olcum.get('created_at') else ''
            }
            
            # Parametre verilerini ekle (ekranda gördüğünüz parametre sütunları)
            parametre_verileri = olcum.get('parametre_verileri', {})
            for alan in allParametreFields:
                row[alan] = parametre_verileri.get(alan, '*')
            
            df_data.append(row)
        
        print(f"DataFrame verisi hazırlandı, satır sayısı: {len(df_data)}")
        
        df = pd.DataFrame(df_data)
        print(f"DataFrame oluşturuldu, boyut: {df.shape}")
        
        # Geçici dosya oluştur
        print("Geçici dosya oluşturuluyor...")
        tmp_file_path = None
        
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
                tmp_file_path = tmp_file.name
                print(f"Geçici dosya yolu: {tmp_file_path}")
            
            # Excel dosyasını oluştur
            with pd.ExcelWriter(tmp_file_path, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='Parametre Ölçümleri')
                
                # Sütun genişliklerini ayarla
                worksheet = writer.sheets['Parametre Ölçümleri']
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    worksheet.column_dimensions[column_letter].width = adjusted_width
            
            print("Excel dosyası oluşturuldu, okunuyor...")
            
            # Dosyayı oku ve response oluştur
            with open(tmp_file_path, 'rb') as f:
                file_content = f.read()
            
            print(f"Dosya okundu, boyut: {len(file_content)} bytes")
            
        finally:
            # Geçici dosyayı sil (hata olsa bile)
            if tmp_file_path and os.path.exists(tmp_file_path):
                try:
                    os.unlink(tmp_file_path)
                    print("Geçici dosya silindi")
                except Exception as delete_error:
                    print(f"Geçici dosya silinirken hata: {delete_error}")
                    # Dosya silinemezse sorun değil, sistem temizleyecek
        
        # Response oluştur
        response = make_response(file_content)
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        
        # Dosya adını oluştur
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"Parametre_Olcumleri_{len(filtered_olcumler)}_kayit_{timestamp}.xlsx"
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        print(f"Response hazırlandı, dosya adı: {filename}")
        return response
        
    except Exception as e:
        print(f"Parametre ölçümleri Excel export hatası: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Excel export hatası: {str(e)}'}), 500

@app.route('/api/parametre_olcumleri_word_export', methods=['POST'])
def api_parametre_olcumleri_word_export():
    """Parametre ölçümlerini Word formatında dışa aktarır."""
    try:
        data = request.get_json()
        olcum_ids = data.get('olcum_ids', [])
        
        if not olcum_ids:
            return jsonify({'error': 'En az bir ölçüm ID gerekli!'}), 400
        
        # Parametre ölçümlerini yükle
        parametre_olcumleri = load_parametre_olcum()
        
        # Seçili ölçümleri filtrele
        if olcum_ids:
            filtered_olcumler = [item for item in parametre_olcumleri if str(item.get('id')) in olcum_ids]
        else:
            filtered_olcumler = parametre_olcumleri
        
        if not filtered_olcumler:
            return jsonify({'error': 'Dışa aktarılacak veri bulunamadı'}), 404
        
        # Docx kütüphanesini yükle
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        
        if not Document:
            return jsonify({'error': 'Word rapor oluşturma için python-docx kütüphanesi gerekli'}), 500
        
        # Word dokümanı oluştur
        doc = Document()
        
        # Başlık
        title = doc.add_heading('Parametre Ölçümleri Raporu', 0)
        title.alignment = 1  # Ortalı
        
        # Tarih
        doc.add_paragraph(f'Oluşturma Tarihi: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
        doc.add_paragraph('')
        
        # Her ölçüm için tablo oluştur
        for i, olcum in enumerate(filtered_olcumler, 1):
            # Ölçüm başlığı
            doc.add_heading(f'{i}. Ölçüm Kaydı', level=1)
            
            # Temel bilgiler tablosu
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Alan'
            hdr_cells[1].text = 'Değer'
            
            # Temel bilgileri ekle
            basic_info = [
                ('Firma Adı', olcum.get('firma_adi', '')),
                ('Ölçüm Kodu', olcum.get('olcum_kodu', '')),
                ('Baca Adı', olcum.get('baca_adi', '')),
                ('Parametre Adı', olcum.get('parametre_adi', '')),
                ('Kayıt Tarihi', olcum.get('created_at', '')[:19].replace('T', ' ') if olcum.get('created_at') else '')
            ]
            
            for field, value in basic_info:
                row_cells = table.add_row().cells
                row_cells[0].text = field
                row_cells[1].text = str(value)
            
            # Parametre verileri tablosu
            parametre_verileri = olcum.get('parametre_verileri', {})
            if parametre_verileri:
                doc.add_paragraph('')
                doc.add_heading('Parametre Verileri', level=2)
                
                param_table = doc.add_table(rows=1, cols=2)
                param_table.style = 'Table Grid'
                param_hdr_cells = param_table.rows[0].cells
                param_hdr_cells[0].text = 'Parametre'
                param_hdr_cells[1].text = 'Değer'
                
                for key, value in parametre_verileri.items():
                    param_row_cells = param_table.add_row().cells
                    param_row_cells[0].text = str(key)
                    param_row_cells[1].text = str(value)
            
            # Sayfa sonu ekle (son ölçüm hariç)
            if i < len(filtered_olcumler):
                doc.add_page_break()
        
        # Dosyayı kaydet
        from io import BytesIO
        docx_io = BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        
        # Response oluştur
        response = make_response(docx_io.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        
        # Dosya adını oluştur
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"Parametre_Olcumleri_{len(filtered_olcumler)}_kayit_{timestamp}.docx"
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except Exception as e:
        print(f"Parametre ölçümleri Word export hatası: {e}")
        return jsonify({'error': f'Word export hatası: {str(e)}'}), 500

@app.route('/api/parametre_olcumleri_pdf_export', methods=['POST'])
def api_parametre_olcumleri_pdf_export():
    """Parametre ölçümlerini PDF formatında dışa aktarır."""
    try:
        data = request.get_json()
        olcum_ids = data.get('olcum_ids', [])
        
        if not olcum_ids:
            return jsonify({'error': 'En az bir ölçüm ID gerekli!'}), 400
        
        # Parametre ölçümlerini yükle
        parametre_olcumleri = load_parametre_olcum()
        
        # Seçili ölçümleri filtrele
        if olcum_ids:
            filtered_olcumler = [item for item in parametre_olcumleri if str(item.get('id')) in olcum_ids]
        else:
            filtered_olcumler = parametre_olcumleri
        
        if not filtered_olcumler:
            return jsonify({'error': 'Dışa aktarılacak veri bulunamadı'}), 404
        
        # Docx kütüphanesini yükle
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        
        if not Document:
            return jsonify({'error': 'Word rapor oluşturma için python-docx kütüphanesi gerekli'}), 500
        
        # Word dokümanı oluştur (Word export fonksiyonunu kullan)
        doc = Document()
        
        # Başlık
        title = doc.add_heading('Parametre Ölçümleri Raporu', 0)
        title.alignment = 1  # Ortalı
        
        # Tarih
        doc.add_paragraph(f'Oluşturma Tarihi: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
        doc.add_paragraph('')
        
        # Her ölçüm için tablo oluştur
        for i, olcum in enumerate(filtered_olcumler, 1):
            # Ölçüm başlığı
            doc.add_heading(f'{i}. Ölçüm Kaydı', level=1)
            
            # Temel bilgiler tablosu
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Alan'
            hdr_cells[1].text = 'Değer'
            
            # Temel bilgileri ekle
            basic_info = [
                ('Firma Adı', olcum.get('firma_adi', '')),
                ('Ölçüm Kodu', olcum.get('olcum_kodu', '')),
                ('Baca Adı', olcum.get('baca_adi', '')),
                ('Parametre Adı', olcum.get('parametre_adi', '')),
                ('Kayıt Tarihi', olcum.get('created_at', '')[:19].replace('T', ' ') if olcum.get('created_at') else '')
            ]
            
            for field, value in basic_info:
                row_cells = table.add_row().cells
                row_cells[0].text = field
                row_cells[1].text = str(value)
            
            # Parametre verileri tablosu
            parametre_verileri = olcum.get('parametre_verileri', {})
            if parametre_verileri:
                doc.add_paragraph('')
                doc.add_heading('Parametre Verileri', level=2)
                
                param_table = doc.add_table(rows=1, cols=2)
                param_table.style = 'Table Grid'
                param_hdr_cells = param_table.rows[0].cells
                param_hdr_cells[0].text = 'Parametre'
                param_hdr_cells[1].text = 'Değer'
                
                for key, value in parametre_verileri.items():
                    param_row_cells = param_table.add_row().cells
                    param_row_cells[0].text = str(key)
                    param_row_cells[1].text = str(value)
            
            # Sayfa sonu ekle (son ölçüm hariç)
            if i < len(filtered_olcumler):
                doc.add_page_break()
        
        # Word dosyasını geçici olarak kaydet
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_word:
            doc.save(tmp_word.name)
            tmp_word_path = tmp_word.name
        
        # Word dosyasını PDF'e çevir
        try:
            from docx2pdf import convert
            
            # PDF dosya adını oluştur
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            pdf_filename = f"Parametre_Olcumleri_{len(filtered_olcumler)}_kayit_{timestamp}.pdf"
            pdf_path = os.path.join(tempfile.gettempdir(), pdf_filename)
            
            # Word'ü PDF'e çevir
            convert(tmp_word_path, pdf_path)
            
            # PDF dosyasını oku
            with open(pdf_path, 'rb') as f:
                pdf_content = f.read()
            
            # Geçici dosyaları sil
            os.unlink(tmp_word_path)
            os.unlink(pdf_path)
            
            # Response oluştur
            response = make_response(pdf_content)
            response.headers['Content-Type'] = 'application/pdf'
            response.headers['Content-Disposition'] = f'attachment; filename="{pdf_filename}"'
            
            return response
            
        except Exception as pdf_error:
            print(f"PDF dönüştürme hatası: {pdf_error}")
            # PDF dönüştürme başarısız olursa Word dosyasını döndür
            with open(tmp_word_path, 'rb') as f:
                word_content = f.read()
            
            os.unlink(tmp_word_path)
            
            response = make_response(word_content)
            response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            response.headers['Content-Disposition'] = f'attachment; filename="Parametre_Olcumleri_{len(filtered_olcumler)}_kayit_{timestamp}.docx"'
            
            return response
        
    except Exception as e:
        print(f"Parametre ölçümleri PDF export hatası: {e}")
        return jsonify({'error': f'PDF export hatası: {str(e)}'}), 500
# create_word_teklif fonksiyonu silindi
    try:
        Document, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL = load_docx()
        
        if not DOCX_AVAILABLE:
            return jsonify({'success': False, 'message': 'Word dosyası oluşturma için gerekli kütüphane yüklü değil'})
        
        # Yeni Word dokümanı oluştur - daha basit yaklaşım
        doc = Document()
        
        # Sayfa kenar boşluklarını ayarla
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Header (üst bilgi) ekle
        try:
            header_img_path = 'static/images/tek_ust1.jpg'
            if os.path.exists(header_img_path):
                # Header section'a resim ekle
                section = doc.sections[0]
                header = section.header
                header_paragraph = header.paragraphs[0]
                header_run = header_paragraph.add_run()
                header_run.add_picture(header_img_path, width=Inches(7.5))
                header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"Header resmi eklenirken hata: {e}")
        
        # Footer (alt bilgi) - Sayfa numarası ekle
        try:
            section = doc.sections[0]
            footer = section.footer
            footer_paragraph = footer.paragraphs[0]
            footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Sayfa numarası ekle - basit yöntem
            footer_run = footer_paragraph.add_run()
            footer_run.add_text("Sayfa ")
            # Word'ün otomatik sayfa numarası field'ını ekle
            from docx.oxml.shared import qn
            from docx.oxml import OxmlElement
            
            # PAGE field ekle
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            fldChar1.set(qn('w:dirty'), 'true')
            footer_run._r.append(fldChar1)
            
            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"
            footer_run._r.append(instrText)
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            footer_run._r.append(fldChar2)
            
            footer_run.add_text(" / ")
            
            # NUMPAGES field ekle
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'begin')
            fldChar3.set(qn('w:dirty'), 'true')
            footer_run._r.append(fldChar3)
            
            instrText2 = OxmlElement('w:instrText')
            instrText2.text = "NUMPAGES"
            footer_run._r.append(instrText2)
            
            fldChar4 = OxmlElement('w:fldChar')
            fldChar4.set(qn('w:fldCharType'), 'end')
            footer_run._r.append(fldChar4)
            
        except Exception as e:
            print(f"Footer eklenirken hata: {e}")
            # Fallback: basit metin
            try:
                footer_paragraph = doc.sections[0].footer.paragraphs[0]
                footer_paragraph.text = "Sayfa 1 / 5"
                footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                pass
        
        # 1. SAYFA - TEKLİF FORMU
        
        # TEKLİF FORMU başlığı kaldırıldı
        
        # 3 satır boşluk ekle
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Firma bilgileri tablosu (3 sütunlu - 4 cm, 0.5 cm, 13 cm) - Kenarlıksız
        firma_tablo = doc.add_table(rows=8, cols=3)
        # Kenarlıkları kaldır - basit yöntem
        firma_tablo.style = None  # Varsayılan stil kaldır
        
        # Sütun genişliklerini ve satır yüksekliklerini ayarla
        for row in firma_tablo.rows:
            row.cells[0].width = Inches(1.57)  # 4 cm
            row.cells[1].width = Inches(0.2)   # 0.5 cm
            row.cells[2].width = Inches(5.12)  # 13 cm
            row.height = Inches(0.37)  # Satır yüksekliği 9.5 mm
        
        # Sol sütun - Etiketler
        etiketler = [
            "Firma Adı",
            "Firma Yetkili",
            "Firma Adresi",
            "Tel / E-Posta",
            "Talep",
            "Teklif Kodu",
            "Sayfa Adedi",
            "Tarih"
        ]
        
        # Değerler - Sayfa adedi şimdilik boş, sonra doldurulacak
        tel_email = f"{firma.get('telefon', '') if firma else ''} / {firma.get('email', '') if firma else ''}"
        degerler = [
            firma.get('firmaAdi', '') if firma else '',
            firma.get('yetkiliAdi', '') if firma else '',
            firma.get('adres', '') if firma else '',
            tel_email,
            teklif.get('teklif_tipi', ''),
            teklif.get('teklif_no', ''),
            "",  # Sayfa adedi sonra doldurulacak
            teklif.get('teklif_tarihi', '')
        ]
        
        # Tabloyu doldur
        for i in range(8):
            # Sol sütun - etiketler (kalın ve altı çizili)
            left_cell = firma_tablo.rows[i].cells[0]
            left_paragraph = left_cell.paragraphs[0]
            left_run = left_paragraph.add_run(etiketler[i])
            left_run.bold = True
            left_run.underline = True
            
            # Orta sütun - ":" işareti (0.5 cm)
            middle_cell = firma_tablo.rows[i].cells[1]
            middle_cell.text = ":"
            
            # Sağ sütun - değerler
            right_cell = firma_tablo.rows[i].cells[2]
            right_cell.text = str(degerler[i]) if degerler[i] else ''
        
        # Tablo ile metin arası 2 satır boşluk
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Teklif giriş metni (localStorage'dan gelen veri) - 1. sayfada
        # localStorage_data'dan giriş metnini al
        request_data = request.get_json() or {}
        localStorage_data = request_data.get('localStorage', {})
        giris_metni = localStorage_data.get('teklifGirisText', '')
        
        # Eğer localStorage'dan gelmediyse, teklif verisinden al
        if not giris_metni:
            giris_metni = teklif.get('teklif_giris_metni', '')
        
        # Eğer 3. aşamadan metin varsa onu kullan, yoksa varsayılan
        if not giris_metni or giris_metni.strip() == '':
            giris_metni = '''Sayın Yetkili;

Talebiniz doğrultusunda hazırlanan fiyat teklifimiz bilginize sunulmuştur.

Laboratuvarımız çalışmalarını "TÜRKAK Akreditasyon Belgesi" ve "Çevre Analizleri Yeterlilik Belgesi" kapsamında gerçekleştirmektedir.

Teklifimizi uygun bulacağınızı umar, iyi çalışmalar dilerim.

Saygılarımızla

Teklifi Hazırlayan
Hafize Demet Fazli'''
        
        # HTML içeriğini bold formatlamayı koruyarak ekle
        print(f"DEBUG - Giriş metni verisi: {giris_metni[:200]}...")
        if giris_metni:
            # HTML içeriğini parse et ve bold formatlamayı koru
            from bs4 import BeautifulSoup
            import re
            
            try:
                # Word field tag'lerini temizle - Daha agresif
                clean_html = giris_metni
                clean_html = re.sub(r'\[if\s+!supportLists\][\s\S]*?\[endif\]', '', clean_html)
                clean_html = re.sub(r'\[if\s+supportLists\][\s\S]*?\[endif\]', '', clean_html)
                clean_html = re.sub(r'\[if\s+!mso\][\s\S]*?\[endif\]', '', clean_html)
                clean_html = re.sub(r'\[if\s+mso\][\s\S]*?\[endif\]', '', clean_html)
                clean_html = re.sub(r'<!--\[if[^>]*>.*?<!\[endif\]-->', '', clean_html, flags=re.DOTALL)
                clean_html = re.sub(r'<!--.*?-->', '', clean_html, flags=re.DOTALL)
                clean_html = re.sub(r'style="[^"]*"', '', clean_html)
                clean_html = re.sub(r'class="[^"]*"', '', clean_html)
                clean_html = re.sub(r'lang="[^"]*"', '', clean_html)
                
                soup = BeautifulSoup(clean_html, 'html.parser')
                
                # Paragraf oluştur
                p = doc.add_paragraph()
                
                # HTML içeriğini işle - Daha akıllı parsing
                def process_element(element):
                    if element.name is None:  # Text node
                        text = element.strip()
                        if text:
                            # Özel karakterleri temizle
                            text = text.replace('&nbsp;', ' ')
                            text = text.replace('&amp;', '&')
                            text = text.replace('&lt;', '<')
                            text = text.replace('&gt;', '>')
                            text = text.replace('&quot;', '"')
                            text = text.replace('&apos;', "'")
                            text = re.sub(r'\s+', ' ', text)
                            
                            if text.strip():
                                run = p.add_run(text.strip() + ' ')
                                # Font ve stil ayarları - Normal metin
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(11)
                                run.font.bold = False
                    elif element.name in ['strong', 'b']:  # Bold tag
                        text = element.get_text().strip()
                        if text:
                            # Özel karakterleri temizle
                            text = text.replace('&nbsp;', ' ')
                            text = text.replace('&amp;', '&')
                            text = text.replace('&lt;', '<')
                            text = text.replace('&gt;', '>')
                            text = text.replace('&quot;', '"')
                            text = text.replace('&apos;', "'")
                            text = re.sub(r'\s+', ' ', text)
                            
                            if text.strip():
                                run = p.add_run(text.strip() + ' ')
                                # Font ve stil ayarları - Bold metin
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(11)
                                run.font.bold = True
                    else:
                        # Diğer elementler için recursive işleme
                        for child in element.children:
                            process_element(child)
                
                # Tüm elementleri işle
                for element in soup.children:
                    process_element(element)
                
                # Eğer hiç içerik yoksa, fallback olarak basit temizleme yap
                if not p.runs:
                    clean_text = re.sub(r'<[^>]+>', '', clean_html)
                    clean_text = clean_text.replace('&nbsp;', ' ')
                    clean_text = clean_text.replace('<br>', '\n')
                    clean_text = clean_text.replace('<br/>', '\n')
                    clean_text = clean_text.replace('<br />', '\n')
                    
                    # Satırları temizle ve birleştir
                    lines = [line.strip() for line in clean_text.split('\n') if line.strip()]
                    if lines:
                        paragraph_text = ' '.join(lines)
                        p = doc.add_paragraph(paragraph_text)
                        # Font ve stil ayarları - Fallback
                        for run in p.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)
                            run.font.bold = False
                        
            except Exception as e:
                print(f"Giriş metni formatlaması hatası: {e}")
                # Fallback: Basit temizleme
                clean_text = re.sub(r'<[^>]+>', '', giris_metni)
                clean_text = re.sub(r'\[if\s+!supportLists\][\s\S]*?\[endif\]', '', clean_text)
                clean_text = re.sub(r'\[if\s+supportLists\][\s\S]*?\[endif\]', '', clean_text)
                clean_text = re.sub(r'<!--\[if[^>]*>.*?<!\[endif\]-->', '', clean_text, flags=re.DOTALL)
                clean_text = re.sub(r'<!--.*?-->', '', clean_text, flags=re.DOTALL)
                clean_text = clean_text.replace('&nbsp;', ' ')
                clean_text = clean_text.replace('<br>', '\n')
                clean_text = clean_text.replace('<br/>', '\n')
                clean_text = clean_text.replace('<br />', '\n')
                
                lines = [line.strip() for line in clean_text.split('\n') if line.strip()]
                if lines:
                    paragraph_text = ' '.join(lines)
                    p = doc.add_paragraph(paragraph_text)
                    # Font ve stil ayarları - Fallback
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(11)
                        run.font.bold = False
        
        # Sayfa sonu ekle
        doc.add_page_break()
        
        # 2. sayfa başlangıcında 1 satır boşluk
        doc.add_paragraph()
        
        # 2. SAYFA - ÖLÇÜM METOTLARI VE GENEL HÜKÜMLER
        
        # ÖLÇÜM METOTLARI VE ÜCRETLENDİRME başlığı
        olcum_baslik = doc.add_heading('ÖLÇÜM METOTLARI VE ÜCRETLENDİRME', level=1)
        olcum_baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Parametre tablosu
        parametreler = teklif.get('parametreler', [])
        if parametreler:
            parametre_tablo = doc.add_table(rows=1, cols=5)
            parametre_tablo.style = 'Table Grid'
            
            # Tablo başlıkları
            hdr_cells = parametre_tablo.rows[0].cells
            hdr_cells[0].text = 'Ölçülecek\nParametre'
            hdr_cells[1].text = 'Ölçüm\nMetodu'
            hdr_cells[2].text = 'Adet'
            hdr_cells[3].text = 'Birim\nFiyatı'
            hdr_cells[4].text = 'Toplam\nFiyat'
            
            # Başlık hücrelerini kalın yap ve hizalama ayarla
            for i, cell in enumerate(hdr_cells):
                cell.paragraphs[0].runs[0].bold = True
                # Birim Fiyat ve Toplam Fiyat sütunlarını sağdan hizala
                if i == 3 or i == 4:  # Birim Fiyat ve Toplam Fiyat sütunları
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Sütun genişliklerini ayarla
            # Ölçülecek Parametre: 4.5cm, Ölçüm Metodu: 6.5cm, Adet: 1cm, Birim Fiyat: 3.5cm, Toplam Fiyat: 3.5cm
            for row in parametre_tablo.rows:
                row.cells[0].width = Inches(1.77)  # 4.5 cm
                row.cells[1].width = Inches(2.56)  # 6.5 cm  
                row.cells[2].width = Inches(0.39)  # 1 cm (aynı)
                row.cells[3].width = Inches(1.38)  # 3.5 cm
                row.cells[4].width = Inches(1.38)  # 3.5 cm
                row.height = Inches(0.37)  # 9.5 mm
            
            # Parametre verilerini ekle
            for parametre in parametreler:
                if parametre.get('adet', 0) > 0:  # Sadece adet > 0 olanları ekle
                    new_row = parametre_tablo.add_row()
                    new_row.height = Inches(0.37)  # Yeni satırın yüksekliği 9.5 mm
                    
                    # Yeni satır için de sütun genişliklerini ayarla
                    new_row.cells[0].width = Inches(1.77)  # 4.5 cm
                    new_row.cells[1].width = Inches(2.56)  # 6.5 cm  
                    new_row.cells[2].width = Inches(0.39)  # 1 cm (aynı)
                    new_row.cells[3].width = Inches(1.38)  # 3.5 cm
                    new_row.cells[4].width = Inches(1.38)  # 3.5 cm
                    
                    row_cells = new_row.cells
                    row_cells[0].text = parametre.get('parametre', '')
                    row_cells[1].text = parametre.get('metot', '')
                    row_cells[2].text = str(parametre.get('adet', 0))
                    row_cells[3].text = f"{parametre.get('birimFiyat', 0):.2f}"
                    row_cells[4].text = f"{parametre.get('topFiyat', 0):.2f}"
                    
                    # Birim Fiyat ve Toplam Fiyat sütunlarını sağdan hizala
                    row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Toplam bilgileri için 3 satır ekle (tablo içinde)
            # Toplam satırı
            toplam_row = parametre_tablo.add_row()
            toplam_row.height = Inches(0.37)
            toplam_row.cells[0].merge(toplam_row.cells[1])
            toplam_row.cells[0].merge(toplam_row.cells[2])
            toplam_row.cells[0].merge(toplam_row.cells[3])  # Birim Fiyatı da dahil et
            toplam_row.cells[0].text = "Toplam:"
            toplam_row.cells[0].paragraphs[0].runs[0].bold = True
            toplam_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            toplam_row.cells[4].text = f"{teklif.get('toplam', 0):.2f} TL"  # Toplam Fiyat sütunu
            toplam_row.cells[4].paragraphs[0].runs[0].bold = True
            toplam_row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # İndirim Tutarı satırı
            indirim_row = parametre_tablo.add_row()
            indirim_row.height = Inches(0.37)
            indirim_row.cells[0].merge(indirim_row.cells[1])
            indirim_row.cells[0].merge(indirim_row.cells[2])
            indirim_row.cells[0].merge(indirim_row.cells[3])  # Birim Fiyatı da dahil et
            indirim_row.cells[0].text = "İndirim Tutarı:"
            indirim_row.cells[0].paragraphs[0].runs[0].bold = True
            indirim_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            indirim_row.cells[4].text = f"{teklif.get('indirim', 0):.2f} TL"  # Toplam Fiyat sütunu
            indirim_row.cells[4].paragraphs[0].runs[0].bold = True
            indirim_row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Toplam Tutar satırı
            net_toplam_row = parametre_tablo.add_row()
            net_toplam_row.height = Inches(0.37)
            net_toplam_row.cells[0].merge(net_toplam_row.cells[1])
            net_toplam_row.cells[0].merge(net_toplam_row.cells[2])
            net_toplam_row.cells[0].merge(net_toplam_row.cells[3])  # Birim Fiyatı da dahil et
            net_toplam_row.cells[0].text = "Toplam Tutar:"
            net_toplam_row.cells[0].paragraphs[0].runs[0].bold = True
            net_toplam_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            net_toplam_row.cells[4].text = f"{teklif.get('netToplam', 0):.2f} TL"  # Toplam Fiyat sütunu
            net_toplam_row.cells[4].paragraphs[0].runs[0].bold = True
            net_toplam_row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # 2. sayfa bitiminde 1 satır boşluk bırak ve sayfa sonu ekle
        doc.add_paragraph()
        doc.add_page_break()
        
        # 3. SAYFA - GENEL HÜKÜMLER
        
        # 3. sayfa başlangıcında 1 satır boşluk
        doc.add_paragraph()
        
        # Başlık ekleme - içerikte zaten var
        

        
        # Genel hükümler metni (localStorage'dan gelen veri) - İyileştirilmiş
        # localStorage_data'dan genel hükümleri al
        request_data = request.get_json() or {}
        localStorage_data = request_data.get('localStorage', {})
        genel_hukumler = localStorage_data.get('genelHukumlerText', '')
        
        # Eğer localStorage'dan gelmediyse, teklif verisinden al
        if not genel_hukumler:
            genel_hukumler = teklif.get('genel_hukumler', '')
        
        # HTML'den temiz metin çıkar - Akıllı temizleme (paragraf yapısını koru)
        def clean_html_to_text(html_content):
            if not html_content:
                return ""
            
            import re
            from bs4 import BeautifulSoup
            
            try:
                # BeautifulSoup ile HTML'i parse et
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Paragraf yapısını koruyarak temizle
                paragraphs = []
                
                # <p> tag'lerini bul ve işle
                for p in soup.find_all('p'):
                    text = p.get_text().strip()
                    if text and len(text) > 3:
                        paragraphs.append(text)
                
                # Eğer <p> tag'i yoksa, genel metni al
                if not paragraphs:
                    text = soup.get_text()
                    # Satırları temizle
                    lines = []
                    for line in text.split('\n'):
                        line = line.strip()
                        if line and len(line) > 3:
                            lines.append(line)
                    paragraphs = lines
                
                # Word'den gelen özel karakterleri temizle
                clean_paragraphs = []
                for para in paragraphs:
                    para = para.replace('&nbsp;', ' ')
                    para = para.replace('&amp;', '&')
                    para = para.replace('&lt;', '<')
                    para = para.replace('&gt;', '>')
                    para = para.replace('&quot;', '"')
                    para = para.replace('&apos;', "'")
                    
                    # Fazla boşlukları temizle
                    para = re.sub(r'\s+', ' ', para)
                    
                    if para.strip():
                        clean_paragraphs.append(para.strip())
                
                # Paragrafları birleştir (her paragraf ayrı satır)
                return '\n\n'.join(clean_paragraphs)
                
            except Exception as e:
                print(f"HTML temizleme hatası: {e}")
                # Fallback: basit regex temizleme
                clean_text = re.sub(r'<[^>]+>', '', html_content)
                clean_text = clean_text.replace('&nbsp;', ' ')
                clean_text = clean_text.replace('<br>', '\n')
                clean_text = clean_text.replace('<br/>', '\n')
                clean_text = clean_text.replace('<br />', '\n')
                clean_text = re.sub(r'\s+', ' ', clean_text)
                return clean_text.strip()
        
        # Genel hükümler - GENEL_HUKUM.docx dosyasından al
        try:
            # GENEL_HUKUM.docx dosyasını oku
            genel_hukum_doc_path = 'static/images/GENEL_HUKUM.docx'
            if os.path.exists(genel_hukum_doc_path):
                print(f"DEBUG - GENEL_HUKUM.docx dosyası bulundu, içerik alınıyor...")
                
                # Word belgesini aç ve içeriği al
                from docx import Document as DocxDocument
                genel_hukum_doc = DocxDocument(genel_hukum_doc_path)
                
                # Tüm paragrafları al
                for paragraph in genel_hukum_doc.paragraphs:
                    if paragraph.text.strip():
                        p = doc.add_paragraph()
                        
                        # Paragraf metnini al
                        paragraph_text = paragraph.text.strip()
                        
                        # İlk paragraf "1. GENEL HÜKÜMLER" ise başlık olarak işle
                        if 'GENEL HÜKÜMLER' in paragraph_text and paragraph_text.startswith('1.'):
                            # Başlık olarak ekle (madde işareti olmadan)
                            title_run = p.add_run(paragraph_text)
                            title_run.font.name = 'Times New Roman'
                            title_run.font.size = Pt(11)
                            title_run.font.bold = True
                            title_run.font.color.rgb = RGBColor(0, 0, 0)
                        else:
                            # Normal madde olarak ekle
                            # Yuvarlak madde işareti ekle (Bold)
                            bullet_run = p.add_run('• ')
                            bullet_run.font.name = 'Times New Roman'
                            bullet_run.font.size = Pt(11)
                            bullet_run.font.bold = True
                            bullet_run.font.color.rgb = RGBColor(0, 0, 0)
                            
                            # Metin ekle
                            text_run = p.add_run(paragraph_text)
                            text_run.font.name = 'Times New Roman'
                            text_run.font.size = Pt(11)
                            text_run.font.bold = False
                            text_run.font.color.rgb = RGBColor(0, 0, 0)
                        
                        # Paragraf formatını ayarla
                        p.paragraph_format.space_after = Pt(12)
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.line_spacing = 1.15
                        p.paragraph_format.first_line_indent = Inches(0)
                        p.paragraph_format.left_indent = Inches(0)
                        p.paragraph_format.right_indent = Inches(0)
                
                print(f"DEBUG - GENEL_HUKUM.docx içeriği başarıyla eklendi")
                
            else:
                print(f"DEBUG - GENEL_HUKUM.docx dosyası bulunamadı, localStorage'dan alınıyor...")
                # Fallback: localStorage'dan al
                localStorage_data = request_data.get('localStorage', {})
                genel_hukumler = localStorage_data.get('genelHukumlerText', '')
                
                # Eğer localStorage'dan gelmediyse, teklif verisinden al
                if not genel_hukumler:
                    genel_hukumler = teklif.get('genel_hukumler', '')
                
                print(f"DEBUG - Genel hükümler verisi: {genel_hukumler[:200]}...")
                
                if genel_hukumler and genel_hukumler.strip():
                    try:
                        from bs4 import BeautifulSoup
                        import re
                        
                        # HTML'i temizle ve parse et
                        soup = BeautifulSoup(genel_hukumler, 'html.parser')
                        
                        # Önce <ul> ve <li> tag'lerini kontrol et
                        ul_tags = soup.find_all('ul')
                        
                        if ul_tags:
                            # <ul> tag'i varsa, <li> elementlerini işle
                            for ul in ul_tags:
                                li_tags = ul.find_all('li')
                                for i, li_tag in enumerate(li_tags):
                                    text = li_tag.get_text().strip()
                                    if not text or len(text) < 3:
                                        continue
                                    
                                    # Paragraf oluştur
                                    p = doc.add_paragraph()
                                    
                                    # Yuvarlak madde işareti ekle (Bold)
                                    bullet_run = p.add_run('• ')
                                    bullet_run.font.name = 'Times New Roman'
                                    bullet_run.font.size = Pt(11)
                                    bullet_run.font.bold = True
                                    bullet_run.font.color.rgb = RGBColor(0, 0, 0)
                                    
                                    # Metin ekle
                                    text_run = p.add_run(text)
                                    text_run.font.name = 'Times New Roman'
                                    text_run.font.size = Pt(11)
                                    text_run.font.bold = False
                                    text_run.font.color.rgb = RGBColor(0, 0, 0)
                                    
                                    # Paragraf formatını ayarla
                                    p.paragraph_format.space_after = Pt(12)
                                    p.paragraph_format.space_before = Pt(0)
                                    p.paragraph_format.line_spacing = 1.15
                                    p.paragraph_format.first_line_indent = Inches(0)
                                    p.paragraph_format.left_indent = Inches(0)
                                    p.paragraph_format.right_indent = Inches(0)
                            
                            # <strong> tag'lerini de kontrol et (başlık için)
                            strong_tags = soup.find_all('strong')
                            for strong_tag in strong_tags:
                                text = strong_tag.get_text().strip()
                                if 'GENEL HÜKÜMLER' in text:
                                    # Başlık olarak ekle (madde işareti olmadan)
                                    p = doc.add_paragraph()
                                    title_run = p.add_run(text)
                                    title_run.font.name = 'Times New Roman'
                                    title_run.font.size = Pt(11)
                                    title_run.font.bold = True
                                    title_run.font.color.rgb = RGBColor(0, 0, 0)
                                    
                                    # Paragraf formatını ayarla
                                    p.paragraph_format.space_after = Pt(12)
                                    p.paragraph_format.space_before = Pt(0)
                                    p.paragraph_format.line_spacing = 1.15
                                    p.paragraph_format.first_line_indent = Inches(0)
                                    p.paragraph_format.left_indent = Inches(0)
                                    p.paragraph_format.right_indent = Inches(0)
                                    break
                        
                        # Tüm <p> tag'lerini bul
                        paragraphs = soup.find_all('p')
                        
                        if paragraphs:
                            # Her paragrafı işle
                            for i, p_tag in enumerate(paragraphs):
                                text = p_tag.get_text().strip()
                                if not text or len(text) < 3:
                                    continue
                                
                                # Paragraf oluştur
                                p = doc.add_paragraph()
                                
                                # İlk paragraf "1. GENEL HÜKÜMLER" ise başlık olarak işle
                                if i == 0 and 'GENEL HÜKÜMLER' in text:
                                    # Başlık olarak ekle (madde işareti olmadan)
                                    title_run = p.add_run(text)
                                    title_run.font.name = 'Times New Roman'
                                    title_run.font.size = Pt(11)
                                    title_run.font.bold = True
                                    title_run.font.color.rgb = RGBColor(0, 0, 0)
                                else:
                                    # Normal madde olarak ekle
                                    # Yuvarlak madde işareti ekle (Bold)
                                    bullet_run = p.add_run('• ')
                                    bullet_run.font.name = 'Times New Roman'
                                    bullet_run.font.size = Pt(11)
                                    bullet_run.font.bold = True
                                    bullet_run.font.color.rgb = RGBColor(0, 0, 0)
                                    
                                    # Metin ekle
                                    text_run = p.add_run(text)
                                    text_run.font.name = 'Times New Roman'
                                    text_run.font.size = Pt(11)
                                    text_run.font.bold = False
                                    text_run.font.color.rgb = RGBColor(0, 0, 0)
                                
                                # Paragraf formatını ayarla
                                p.paragraph_format.space_after = Pt(12)
                                p.paragraph_format.space_before = Pt(0)
                                p.paragraph_format.line_spacing = 1.15
                                p.paragraph_format.first_line_indent = Inches(0)
                                p.paragraph_format.left_indent = Inches(0)
                                p.paragraph_format.right_indent = Inches(0)
                        
                        else:
                            # <p> tag'i yoksa, düz metni satırlara böl
                            clean_text = soup.get_text()
                            lines = [line.strip() for line in clean_text.split('\n') if line.strip() and len(line.strip()) > 3]
                            
                            for i, line in enumerate(lines):
                                p = doc.add_paragraph()
                                
                                # İlk satır "1. GENEL HÜKÜMLER" ise başlık olarak işle
                                if i == 0 and 'GENEL HÜKÜMLER' in line:
                                    # Başlık olarak ekle (madde işareti olmadan)
                                    title_run = p.add_run(line)
                                    title_run.font.name = 'Times New Roman'
                                    title_run.font.size = Pt(11)
                                    title_run.font.bold = True
                                    title_run.font.color.rgb = RGBColor(0, 0, 0)
                                else:
                                    # Normal madde olarak ekle
                                    # Yuvarlak madde işareti ekle (Bold)
                                    bullet_run = p.add_run('• ')
                                    bullet_run.font.name = 'Times New Roman'
                                    bullet_run.font.size = Pt(11)
                                    bullet_run.font.bold = True
                                    bullet_run.font.color.rgb = RGBColor(0, 0, 0)
                                    
                                    # Metin ekle
                                    text_run = p.add_run(line)
                                    text_run.font.name = 'Times New Roman'
                                    text_run.font.size = Pt(11)
                                    text_run.font.bold = False
                                    text_run.font.color.rgb = RGBColor(0, 0, 0)
                                
                                # Paragraf formatını ayarla
                                p.paragraph_format.space_after = Pt(12)
                                p.paragraph_format.space_before = Pt(0)
                                p.paragraph_format.line_spacing = 1.15
                                p.paragraph_format.first_line_indent = Inches(0)
                                p.paragraph_format.left_indent = Inches(0)
                                p.paragraph_format.right_indent = Inches(0)
                        
                    except Exception as e:
                        print(f"localStorage genel hükümler formatlaması hatası: {e}")
                        # Fallback: Basit madde işaretli liste
                        clean_text = re.sub(r'<[^>]+>', '', genel_hukumler)
                        lines = [line.strip() for line in clean_text.split('\n') if line.strip() and len(line.strip()) > 3]
                        
                        for i, line in enumerate(lines):
                            p = doc.add_paragraph()
                            
                            # İlk satır "1. GENEL HÜKÜMLER" ise başlık olarak işle
                            if i == 0 and 'GENEL HÜKÜMLER' in line:
                                # Başlık olarak ekle (madde işareti olmadan)
                                title_run = p.add_run(line)
                                title_run.font.name = 'Times New Roman'
                                title_run.font.size = Pt(11)
                                title_run.font.bold = True
                                title_run.font.color.rgb = RGBColor(0, 0, 0)
                            else:
                                # Normal madde olarak ekle
                                # Yuvarlak madde işareti ekle (Bold)
                                bullet_run = p.add_run('• ')
                                bullet_run.font.name = 'Times New Roman'
                                bullet_run.font.size = Pt(11)
                                bullet_run.font.bold = True
                                bullet_run.font.color.rgb = RGBColor(0, 0, 0)
                                
                                # Metin ekle
                                text_run = p.add_run(line)
                                text_run.font.name = 'Times New Roman'
                                text_run.font.size = Pt(11)
                                text_run.font.bold = False
                                text_run.font.color.rgb = RGBColor(0, 0, 0)
                            
                            # Paragraf formatını ayarla
                            p.paragraph_format.space_after = Pt(12)
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.line_spacing = 1.15
                            p.paragraph_format.first_line_indent = Inches(0)
                            p.paragraph_format.left_indent = Inches(0)
                            p.paragraph_format.right_indent = Inches(0)
        except Exception as e:
            print(f"GENEL_HUKUM.docx okuma hatası: {e}")
            # Fallback: localStorage'dan al
            localStorage_data = request_data.get('localStorage', {})
            genel_hukumler = localStorage_data.get('genelHukumlerText', '')
            
            # Eğer localStorage'dan gelmediyse, teklif verisinden al
            if not genel_hukumler:
                genel_hukumler = teklif.get('genel_hukumler', '')
            
            print(f"DEBUG - Genel hükümler verisi: {genel_hukumler[:200]}...")
            
            if genel_hukumler and genel_hukumler.strip():
                try:
                    from bs4 import BeautifulSoup
                    import re
                    
                    # HTML'i temizle ve parse et
                    soup = BeautifulSoup(genel_hukumler, 'html.parser')
                    
                    # Tüm <p> tag'lerini bul
                    paragraphs = soup.find_all('p')
                    
                    if paragraphs:
                        # Her paragrafı işle
                        for i, p_tag in enumerate(paragraphs):
                            text = p_tag.get_text().strip()
                            if not text or len(text) < 3:
                                continue
                            
                            # Paragraf oluştur
                            p = doc.add_paragraph()
                            
                            # İlk paragraf "1. GENEL HÜKÜMLER" ise başlık olarak işle
                            if i == 0 and 'GENEL HÜKÜMLER' in text:
                                # Başlık olarak ekle (madde işareti olmadan)
                                title_run = p.add_run(text)
                                title_run.font.name = 'Times New Roman'
                                title_run.font.size = Pt(11)
                                title_run.font.bold = True
                                title_run.font.color.rgb = RGBColor(0, 0, 0)
                            else:
                                # Normal madde olarak ekle
                                # Yuvarlak madde işareti ekle (Bold)
                                bullet_run = p.add_run('• ')
                                bullet_run.font.name = 'Times New Roman'
                                bullet_run.font.size = Pt(11)
                                bullet_run.font.bold = True
                                bullet_run.font.color.rgb = RGBColor(0, 0, 0)
                                
                                # Metin ekle
                                text_run = p.add_run(text)
                                text_run.font.name = 'Times New Roman'
                                text_run.font.size = Pt(11)
                                text_run.font.bold = False
                                text_run.font.color.rgb = RGBColor(0, 0, 0)
                            
                            # Paragraf formatını ayarla
                            p.paragraph_format.space_after = Pt(12)
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.line_spacing = 1.15
                            p.paragraph_format.first_line_indent = Inches(0)
                            p.paragraph_format.left_indent = Inches(0)
                            p.paragraph_format.right_indent = Inches(0)
                    
                    else:
                        # <p> tag'i yoksa, düz metni satırlara böl
                        clean_text = soup.get_text()
                        lines = [line.strip() for line in clean_text.split('\n') if line.strip() and len(line.strip()) > 3]
                        
                        for i, line in enumerate(lines):
                            p = doc.add_paragraph()
                            
                            # İlk satır "1. GENEL HÜKÜMLER" ise başlık olarak işle
                            if i == 0 and 'GENEL HÜKÜMLER' in line:
                                # Başlık olarak ekle (madde işareti olmadan)
                                title_run = p.add_run(line)
                                title_run.font.name = 'Times New Roman'
                                title_run.font.size = Pt(11)
                                title_run.font.bold = True
                                title_run.font.color.rgb = RGBColor(0, 0, 0)
                            else:
                                # Normal madde olarak ekle
                                # Yuvarlak madde işareti ekle (Bold)
                                bullet_run = p.add_run('• ')
                                bullet_run.font.name = 'Times New Roman'
                                bullet_run.font.size = Pt(11)
                                bullet_run.font.bold = True
                                bullet_run.font.color.rgb = RGBColor(0, 0, 0)
                                
                                # Metin ekle
                                text_run = p.add_run(line)
                                text_run.font.name = 'Times New Roman'
                                text_run.font.size = Pt(11)
                                text_run.font.bold = False
                                text_run.font.color.rgb = RGBColor(0, 0, 0)
                            
                            # Paragraf formatını ayarla
                            p.paragraph_format.space_after = Pt(12)
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.line_spacing = 1.15
                            p.paragraph_format.first_line_indent = Inches(0)
                            p.paragraph_format.left_indent = Inches(0)
                            p.paragraph_format.right_indent = Inches(0)
                            
                except Exception as e:
                    print(f"Genel hükümler formatlaması hatası: {e}")
                    # Fallback: Basit madde işaretli liste
                    clean_text = re.sub(r'<[^>]+>', '', genel_hukumler)
                    lines = [line.strip() for line in clean_text.split('\n') if line.strip() and len(line.strip()) > 3]
                    
                    for i, line in enumerate(lines):
                        p = doc.add_paragraph()
                        
                        # İlk satır "1. GENEL HÜKÜMLER" ise başlık olarak işle
                        if i == 0 and 'GENEL HÜKÜMLER' in line:
                            # Başlık olarak ekle (madde işareti olmadan)
                            title_run = p.add_run(line)
                            title_run.font.name = 'Times New Roman'
                            title_run.font.size = Pt(11)
                            title_run.font.bold = True
                            title_run.font.color.rgb = RGBColor(0, 0, 0)
                        else:
                            # Normal madde olarak ekle
                            # Yuvarlak madde işareti ekle (Bold)
                            bullet_run = p.add_run('• ')
                            bullet_run.font.name = 'Times New Roman'
                            bullet_run.font.size = Pt(11)
                            bullet_run.font.bold = True
                            bullet_run.font.color.rgb = RGBColor(0, 0, 0)
                            
                            # Metin ekle
                            text_run = p.add_run(line)
                            text_run.font.name = 'Times New Roman'
                            text_run.font.size = Pt(11)
                            text_run.font.bold = False
                            text_run.font.color.rgb = RGBColor(0, 0, 0)
                        
                        # Paragraf formatını ayarla
                        p.paragraph_format.space_after = Pt(12)
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.line_spacing = 1.15
                        p.paragraph_format.first_line_indent = Inches(0)
                        p.paragraph_format.left_indent = Inches(0)
                        p.paragraph_format.right_indent = Inches(0)
        
        
        # Geçici dosya oluştur ve kaydet
        temp_file_path = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file_path = temp_file.name
            
            # Belgeyi kaydet
            doc.save(temp_file_path)
            
            # Sayfa adedi güncelle - daha doğru hesaplama
            try:
                # Word belgesini açıp sayfa sayısını hesapla
                from docx import Document as DocxDocument
                temp_doc = DocxDocument(temp_file_path)
                
                # Daha doğru sayfa hesaplama - basit ve doğru
                total_paragraphs = len(temp_doc.paragraphs)
                
                # Basit hesaplama: her 15 paragraf = 1 sayfa
                estimated_pages = max(3, (total_paragraphs // 15) + 1)
                
                # Genel hükümler varsa sayfa sayısını artır
                if genel_hukumler and len(genel_hukumler) > 500:
                    estimated_pages += 1
                
                # Maksimum 8 sayfa ile sınırla
                estimated_pages = min(estimated_pages, 8)
                
                firma_tablo.rows[6].cells[2].text = str(estimated_pages)
                
                # Güncellenmiş belgeyi tekrar kaydet
                doc.save(temp_file_path)
            except Exception as e:
                print(f"Sayfa sayısı hesaplama hatası: {e}")
                # Fallback: varsayılan sayfa sayısı
                firma_tablo.rows[6].cells[2].text = "5"
                doc.save(temp_file_path)
            
            # Dosya adını oluştur (yeni format)
            firma_adi = firma.get('firmaAdi', '') if firma else ''
            if not firma_adi:
                # Eğer firma objesinden alamazsak, teklif verisinden al
                firma_adi = teklif.get('firma_adi', '')
            
            print(f"DEBUG - Firma adı: {firma_adi}")
            print(f"DEBUG - Firma objesi: {firma}")
            
            # Firma adını temizle ve ilk 2 kelimeyi al
            firma_adi_temiz = firma_adi.strip() if firma_adi else ''
            print(f"DEBUG - Firma adı temiz: '{firma_adi_temiz}'")
            
            firma_kelimeler = firma_adi_temiz.split()[:2] if firma_adi_temiz else []
            print(f"DEBUG - Firma kelimeler: {firma_kelimeler}")
            
            firma_kisa = '_'.join(firma_kelimeler) if firma_kelimeler else 'Firma'
            print(f"DEBUG - Firma kısa: '{firma_kisa}'")
            
            teklif_tarihi = teklif.get('teklif_tarihi', '')
            if teklif_tarihi:
                try:
                    # Tarihi parse et ve ggaayy formatına çevir
                    tarih_obj = datetime.strptime(teklif_tarihi, '%Y-%m-%d')
                    tarih_format = tarih_obj.strftime('%d%m%y')
                except:
                    tarih_format = datetime.now().strftime('%d%m%y')
            else:
                tarih_format = datetime.now().strftime('%d%m%y')
            
            # Tutarı KDV hariç olarak ekle
            tutar_kdv_haric = teklif.get('netToplam', 0)
            
            dosya_adi = f"{firma_kisa}_{teklif.get('teklif_no', '')}_{tarih_format}_{tutar_kdv_haric:.0f}TL.docx"
            print(f"DEBUG - Dosya adı: {dosya_adi}")
            
            # Dosyayı gönder
            return send_file(
                temp_file_path,
                as_attachment=True,
                download_name=dosya_adi,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        
        except Exception as save_error:
            # Geçici dosyayı temizle
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    os.unlink(temp_file_path)
                except:
                    pass
            raise save_error
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Word dosyası oluşturma hatası: {str(e)}'})

def create_pdf_teklif(teklif, firma):
    """PDF formatında teklif oluşturur (HTML olarak)"""
    try:
        # HTML içeriği oluştur
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Teklif - {teklif.get('teklif_no', '')}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                .header {{ text-align: center; margin-bottom: 30px; }}
                .title {{ font-size: 24px; font-weight: bold; margin-bottom: 20px; }}
                table {{ width: 100%; border-collapse: collapse; margin-bottom: 20px; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background-color: #f2f2f2; font-weight: bold; }}
                .total {{ font-weight: bold; margin-top: 20px; }}
            </style>
        </head>
        <body>
            <div class="header">
                <h1>FİYAT TEKLİFİ</h1>
                <p>Firma: {firma.get('firmaAdi', '') if firma else ''}</p>
                <p>Teklif No: {teklif.get('teklif_no', '')}</p>
                <p>Tarih: {format_tarih_gg_aa_yyyy(teklif.get('teklif_tarihi', ''))}</p>
            </div>
            
            <h2>Parametreler</h2>
            <table>
                <tr>
                    <th>Parametre</th>
                    <th>Metot</th>
                    <th>Adet</th>
                    <th>Birim Fiyat</th>
                    <th>Toplam</th>
                </tr>
        """
        
        # Parametreleri ekle
        for parametre in teklif.get('parametreler', []):
            if parametre.get('adet', 0) > 0:
                html_content += f"""
                <tr>
                    <td>{parametre.get('parametre', '')}</td>
                    <td>{parametre.get('metot', '')}</td>
                    <td>{parametre.get('adet', 0)}</td>
                    <td>{parametre.get('birimFiyat', 0):.2f} TL</td>
                    <td>{parametre.get('topFiyat', 0):.2f} TL</td>
                </tr>
                """
        
        html_content += f"""
            </table>
            
            <div class="total">
                <p>Toplam: {teklif.get('toplam', 0):.2f} TL</p>
                <p>İndirim: {teklif.get('indirim', 0):.2f} TL</p>
                <p>Net Toplam: {teklif.get('netToplam', 0):.2f} TL</p>
            </div>
            
            <h2>Genel Hükümler</h2>
            <p>{teklif.get('genel_hukumler', 'Genel hükümler buraya gelecek.')}</p>
        </body>
        </html>
        """
        
        # Geçici HTML dosyası oluştur
        temp_html = tempfile.NamedTemporaryFile(delete=False, suffix='.html')
        temp_html.write(html_content.encode('utf-8'))
        temp_html.close()
        
        # Dosya adını oluştur
        firma_adi = firma.get('firmaAdi', '') if firma else ''
        firma_kelimeler = firma_adi.split()[:2]  # İlk 2 kelime
        firma_kisa = '_'.join(firma_kelimeler) if firma_kelimeler else 'Firma'
        
        teklif_tarihi = teklif.get('teklif_tarihi', '')
        if teklif_tarihi:
            try:
                tarih_obj = datetime.strptime(teklif_tarihi, '%Y-%m-%d')
                tarih_format = tarih_obj.strftime('%d%m%y')
            except:
                tarih_format = datetime.now().strftime('%d%m%y')
        else:
            tarih_format = datetime.now().strftime('%d%m%y')
        
        tutar_kdv_haric = teklif.get('netToplam', 0)
        dosya_adi = f"{firma_kisa}_{teklif.get('teklif_no', '')}_{tarih_format}_{tutar_kdv_haric:.0f}.html"
        
        # HTML dosyasını gönder
        return send_file(
            temp_html.name,
            as_attachment=True,
            download_name=dosya_adi,
            mimetype='text/html'
        )
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'PDF dosyası oluşturma hatası: {str(e)}'})

# @app.route('/export_graph_and_data', methods=['POST']) - KK Eğri sistemi silindi
# def export_graph_and_data(): - KK Eğri sistemi silindi
    try:
        # Grafiği oluştur ve PNG olarak kaydet
        plt, np, mdates, Rectangle = load_matplotlib()
        plt.plot([1, 2, 3, 4], [10, 20, 25, 30])
        plt.title("Örnek Grafik")
        graph_path = tempfile.NamedTemporaryFile(delete=False, suffix='.png').name
        plt.savefig(graph_path)
        plt.close()

        # Veri tablosunu oluştur
        data = [
            ["Sıra", "Tarih", "Firma", "Kod", "Baca", "Personel", "Cihaz", "Değer"],
            [1, "22.08.25", "AKARE ÇE", "E-250281-02", "SSS", "Admin", "515"],
            [2, "11.08.25", "ATALAY T", "EE-250811-02", "DDDD", "Admin", "490"],
        ]

        # Excel dosyasını oluştur
        output = BytesIO()
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Grafik ve Veri"

        # PNG'yi ekle
        from openpyxl.drawing.image import Image
        img = Image(graph_path)
        ws.add_image(img, "A1")

        # PNG'nin altına veri tablosunu ekle (örneğin A20'den başlat)
        start_row = 20
        for i, row in enumerate(data):
            for j, value in enumerate(row):
                ws.cell(row=start_row + i, column=1 + j, value=value)

        wb.save(output)
        output.seek(0)

        return send_file(
            output,
            as_attachment=True,
            download_name="grafik_ve_veri.xlsx",
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    except Exception as e:
        print(f"Grafik ve veri dışa aktarılırken hata: {e}")
        return jsonify({"success": False, "error": str(e)}), 500

# @app.route('/api/kk_excel_with_graph', methods=['POST']) - KK Eğri sistemi silindi
# def kk_excel_with_graph(): - KK Eğri sistemi silindi
    try:
        data = request.get_json()
        parametre = data.get('parametre')
        tarih_baslangic = data.get('tarih_baslangic')
        tarih_bitis = data.get('tarih_bitis')
        table_data = data.get('localStorage_verileri', [])
        graph_png = data.get('graph_png')

        # PNG'yi base64'ten dosyaya çevir
        import base64
        from openpyxl import Workbook
        from openpyxl.drawing.image import Image
        import tempfile
        import re
        img_data = re.sub('^data:image/.+;base64,', '', graph_png)
        img_bytes = base64.b64decode(img_data)
        img_path = tempfile.NamedTemporaryFile(delete=False, suffix='.png').name
        with open(img_path, 'wb') as f:
            f.write(img_bytes)

        # Excel oluştur
        wb = Workbook()
        ws = wb.active
        ws.title = "Grafik ve Veri"
        # PNG'yi ekle
        img = Image(img_path)
        ws.add_image(img, "A1")
        # Tabloyu ekle (10. satırdan itibaren)
        headers = ["Sıra", "Tarih", "Firma", "Kod", "Baca", "Personel", "Cihaz", "Değer"]
        for col, val in enumerate(headers, start=1):
            ws.cell(row=10, column=col, value=val)
        for i, row in enumerate(table_data, start=11):
            for col, key in enumerate(["sira", "tarih", "firma", "kod", "baca", "personel", "cihaz", "deger"], start=1):
                ws.cell(row=i, column=col, value=row.get(key, ''))
        # Dosyayı kaydet
        output = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx')
        wb.save(output.name)
        return send_file(
            output.name,
            as_attachment=True,
            download_name=f'KK_Rapor_{parametre}_{tarih_baslangic}_{tarih_bitis}.xlsx',
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
    except Exception as e:
        print(f"Excel export with graph error: {e}")
        return jsonify({"success": False, "error": str(e)}), 500

# @app.route('/api/pivot/compare') - Pivot sistemi silindi
# def api_pivot_compare(): - Pivot sistemi silindi
    """PIVOT karşılaştırma API endpoint'i - yıllara göre baca sayısı, parametre sayısı, fiyat ve personel performansını döndürür."""
    try:
        years = request.args.get('years', '')
        if not years:
            return jsonify({'success': False, 'error': 'Yıl parametresi gerekli'}), 400
        
        year_list = [int(y.strip()) for y in years.split(',') if y.strip()]
        if not year_list:
            return jsonify({'success': False, 'error': 'Geçerli yıl listesi gerekli'}), 400
        
        # Verileri yükle
        baca_bilgileri = load_baca_bilgileri()
        parametre_olcumleri = load_parametre_olcum()
        # teklifler = load_teklif() - teklif sistemi silindi
        
        result_data = {}
        
        for year in year_list:
            # O yıl için teklifleri filtrele - teklif sistemi silindi
            year_teklifler = []  # teklif sistemi silindi
            
            # Toplam teklif sayısı ve tutarı - teklif sistemi silindi
            toplam_teklif_adedi = 0  # teklif sistemi silindi
            toplam_teklif_tutari = 0  # teklif sistemi silindi
            
            # Kabul olan teklifleri filtrele
            def _norm(sval: str) -> str:
                m = (sval or '').strip().lower()
                tr_map = str.maketrans({'ı':'i','İ':'i','ş':'s','Ş':'s','ç':'c','Ç':'c','ğ':'g','Ğ':'g','ü':'u','Ü':'u','ö':'o','Ö':'o'})
                return m.translate(tr_map)
            
            kabul_teklifler = []  # teklif sistemi silindi
            # for t in year_teklifler: - teklif sistemi silindi
            
            kabul_adet = 0  # teklif sistemi silindi
            kabul_tutari = 0  # teklif sistemi silindi
            
            # Kapsam içi/dışı teklifleri hesapla - teklif sistemi silindi
            kapsam_ici_teklifler = []  # teklif sistemi silindi
            kapsam_disi_teklifler = []  # teklif sistemi silindi
            is_birligi_teklifler = []  # teklif sistemi silindi
            
            # Kapsam içi/dışı tutarları - teklif sistemi silindi
            kapsam_ici_kabul_teklifler = []  # teklif sistemi silindi
            kapsam_disi_kabul_teklifler = []  # teklif sistemi silindi
            is_birligi_kabul_teklifler = []  # teklif sistemi silindi
            
            # Adet sayıları - teklif sistemi silindi
            kapsam_ici_adet = 0  # teklif sistemi silindi
            kapsam_disi_adet = 0  # teklif sistemi silindi
            is_birligi_adet = 0  # teklif sistemi silindi
            
            kapsam_ici_tutar = 0  # teklif sistemi silindi
            kapsam_disi_tutar = 0  # teklif sistemi silindi
            is_birligi_tutar = 0  # teklif sistemi silindi
            
            # Baca sayısını hesapla - Firma ölçümlerindeki BACA SAY değerlerinin toplamı
            firma_olcumler = load_firma_olcum()
            toplam_baca_adedi = 0
            for olcum in firma_olcumler:
                olcum_date = olcum.get('baslangic_tarihi', '')
                if olcum_date and olcum_date[:4] == str(year):
                    baca_sayisi = olcum.get('baca_sayisi', '0')
                    try:
                        toplam_baca_adedi += int(baca_sayisi)
                    except (ValueError, TypeError):
                        # Baca sayısı sayı değilse 0 olarak kabul et
                        pass
            
            # Parametre isim eşleştirme tablosu - önce tanımla
            parametre_eslesme = {
                'YG': 'YANMA GAZI',
                'YANMA GAZI': 'YANMA GAZI',  # Eksik olan eşleştirme eklendi
                'AMET': 'AĞIR METAL', 
                'SÜLF.A': 'SÜLFÜRİK ASİT',
                'SÜLFÜRİK ASİT': 'SÜLFÜRİK ASİT',  # Eksik olan eşleştirme eklendi
                'TOZ': 'TOZ',
                'VOC': 'VOC',
                'TOC': 'TOC',
                'NEM': 'NEM',
                'PM10': 'PM10',
                'HIZ': 'HIZ',
                'HF': 'HF',
                'HCL': 'HCL',
                'AMONYAK': 'AMONYAK',
                'FORMALDEHİT': 'FORMALDEHİT',
                'CR+6': 'CR+6',
                'CR6': 'CR+6',  # Alternatif yazım
                'FOSFORİK ASİT': 'FOSFORİK ASİT',
                'HCN': 'HCN',
                'DİOKSİN FURAN': 'DİOKSİN FURAN',
                'PAH': 'PAH',
                'ÇÖKEN TOZ': 'ÇÖKEN TOZ',
                'ÇT': 'ÇÖKEN TOZ',  # Alternatif yazım
                'MODELLEME': 'MODELLEME'
            }
            
            # Firma ölçümlerinden parametre sayılarını hesapla
            from collections import defaultdict
            parametre_sayilari = defaultdict(int)
            toplam_parametre_adedi = 0
            
            # Firma ölçümlerindeki parametreleri say
            for olcum in firma_olcumler:
                olcum_date = olcum.get('baslangic_tarihi', '')
                if olcum_date and olcum_date[:4] == str(year):
                    baca_parametreleri = olcum.get('baca_parametreleri', {})
                    for baca_adi, parametre_listesi in baca_parametreleri.items():
                        for parametre in parametre_listesi:
                            if parametre:  # Boş parametre değilse
                                norm_param = parametre.upper().strip()
                                # Parametre eşleştirmesi yap
                                eslesen_param = parametre_eslesme.get(norm_param, norm_param)
                                parametre_sayilari[eslesen_param] += 1
                                toplam_parametre_adedi += 1
            
            # Parametre tutarlarını hesapla: Parametre Adedi × Asgari Fiyat
            parametre_tl = defaultdict(float)
            asgari_fiyatlar = load_asgari_fiyatlar()
            
            # Asgari fiyatları yıla göre filtrele
            for param_adi, adet in parametre_sayilari.items():
                if adet > 0:  # Sadece adedi olan parametreler için hesapla
                    # Parametre ismini eşleştir (zaten eşleştirilmiş halde geliyor)
                    eslesen_param = param_adi  # Artık zaten eşleştirilmiş halde geliyor
                    
                    # Asgari fiyat tablosundan o yılın fiyatını bul
                    fiyat = 0
                    for asg_fiyat in asgari_fiyatlar:
                        if asg_fiyat.get('parametre', '').upper().strip() == eslesen_param:
                            yillik_fiyatlar = asg_fiyat.get('yillik', {})
                            fiyat = float(yillik_fiyatlar.get(str(year), 0))
                            break
                    
                    # Tutar = Adet × Fiyat
                    parametre_tl[param_adi] = adet * fiyat
                    print(f"DEBUG PIVOT: {param_adi} -> {eslesen_param} -> {adet} adet × {fiyat} TL = {adet * fiyat} TL")
            
            # Personel performansını hesapla (baca bilgileri kayıtlarından)
            personel_performans = defaultdict(int)
            personel_tutarlar = defaultdict(float)  # Personel başına toplam tutar
            personel_parametre_performans = defaultdict(int)  # Personel-Parametre kombinasyonu
            
            for baca_kayit in baca_bilgileri:
                created_at = baca_kayit.get('created_at', '')
                if created_at and created_at[:4] == str(year):
                    personel_adi = baca_kayit.get('personel_adi', '')
                    if personel_adi:
                        # Sadece ad kısmını al (boşluktan önceki kısım)
                        personel_adi_kisa = personel_adi.split(' ')[0] if personel_adi else ''
                        # Baca sayısını artır - kısa ad ile
                        personel_performans[personel_adi_kisa] += 1
                        
                        # Bu personelin girdiği baca için parametreleri bul ve fiyat hesapla
                        firma_adi = baca_kayit.get('firma_adi', '')
                        olcum_kodu = baca_kayit.get('olcum_kodu', '')
                        baca_adi = baca_kayit.get('baca_adi', '')
                        
                        # Dinamik olarak firma_olcum.json'dan parametreleri oku
                        baca_parametreleri = []
                        
                        # Firma ölçümlerinden bu baca için parametreleri bul
                        for olcum in firma_olcumler:
                            if (olcum.get('firma_adi') == firma_adi and 
                                olcum.get('olcum_kodu') == olcum_kodu and
                                olcum.get('baca_parametreleri')):
                                
                                baca_parametreleri_dict = olcum['baca_parametreleri']
                                if isinstance(baca_parametreleri_dict, dict) and baca_adi in baca_parametreleri_dict:
                                    baca_parametreleri_list = baca_parametreleri_dict[baca_adi]
                                    if isinstance(baca_parametreleri_list, list):
                                        # Parametreleri normalize et (büyük harf, trim)
                                        baca_parametreleri = [p.upper().strip() for p in baca_parametreleri_list if p and p.strip()]
                                        break
                        
                        # Eğer parametre bulunamadıysa, eski sabit kodlanmış listeyi kullan (fallback)
                        if not baca_parametreleri:
                            # HAFİZE A bacası için: TOZ, YG, VOC, AĞIR METAL
                            if firma_adi == "HAFİZE" and baca_adi == "a":
                                baca_parametreleri = ["TOZ", "YANMA GAZI", "VOC", "AĞIR METAL"]
                            # HAFİZE B bacası için: TOZ, YG, SÜLFÜRİK ASİT, PM10  
                            elif firma_adi == "HAFİZE" and baca_adi == "b":
                                baca_parametreleri = ["TOZ", "YANMA GAZI", "SÜLFÜRİK ASİT", "PM10"]
                            # AKARE ÇEVRE FÜZYON BACASI için: TOZ, NEM, YG (sadece bu 3 parametre!)
                            elif "AKARE ÇEVRE" in firma_adi and "FÜZYON" in baca_adi:
                                baca_parametreleri = ["TOZ", "NEM", "YANMA GAZI"]
                            # ATALAY TOKER adem bacası için: TOZ, YANMA GAZI, VOC, AĞIR METAL
                            elif firma_adi == "ATALAY TOKER" and baca_adi == "adem":
                                baca_parametreleri = ["TOZ", "YANMA GAZI", "VOC", "AĞIR METAL"]
                        
                        # Bu parametreler için o yılın asgari fiyatlarını hesapla
                        toplam_fiyat = 0
                        for parametre in baca_parametreleri:
                            fiyat = get_parametre_fiyati(parametre, year)
                            toplam_fiyat += fiyat
                        
                        # Debug log
                        print(f"Personel: {personel_adi} ({personel_adi_kisa}), Firma: {firma_adi}, Baca: {baca_adi}")
                        print(f"Bulunan parametreler: {baca_parametreleri}")
                        print(f"Toplam Fiyat: {toplam_fiyat}")
                        
                        # Kısa ad ile grupla
                        personel_tutarlar[personel_adi_kisa] += toplam_fiyat
                        
                        # Personel-Parametre kombinasyonu hesapla
                        for parametre in baca_parametreleri:
                            kombinasyon_adi = f"{parametre}-{personel_adi_kisa}"
                            personel_parametre_performans[kombinasyon_adi] += 1
            
            summary = {
                'toplam_baca_adedi': toplam_baca_adedi,
                'toplam_parametre_adedi': toplam_parametre_adedi,
                'toplam_teklif_adedi': toplam_teklif_adedi,
                'kabul_adet': kabul_adet,
                'kapsam_ici_adet': kapsam_ici_adet,
                'kapsam_disi_adet': kapsam_disi_adet,
                'is_birligi_adet': is_birligi_adet,
                'toplam_teklif_tutari': round(toplam_teklif_tutari, 2),
                'kabul_tutari': round(kabul_tutari, 2),
                'kapsam_ici_tutar': round(kapsam_ici_tutar, 2),
                'kapsam_disi_tutar': round(kapsam_disi_tutar, 2),
                'is_birligi_tutar': round(is_birligi_tutar, 2),
                'parametre_sayilari': dict(parametre_sayilari),
                'personel_performans': dict(personel_performans),
                'personel_tutarlar': dict(personel_tutarlar),
                'personel_parametre_performans': dict(personel_parametre_performans)
            }
            
            result_data[str(year)] = {
                'summary': summary,
                'parametre_tl': parametre_tl
            }
        
        return jsonify({
            'success': True,
            'data': result_data
        })
        
    except Exception as e:
        print(f"Pivot compare API hatası: {e}")
        import traceback
        print(f"Hata detayı: {traceback.format_exc()}")
        return jsonify({'success': False, 'error': str(e)}), 500
# @app.route('/api/pivot/export-xlsx') - Pivot sistemi silindi
# def api_pivot_export_xlsx(): - Pivot sistemi silindi
        
        wb = Workbook()
        ws = wb.active
        ws.title = "PIVOT Karşılaştırma"
        
        # Stil tanımlamaları
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        section_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Başlık satırı
        headers = ['METRİK']
        for year in year_list:
            headers.extend([f'{year} Adet', f'{year} TL'])
        
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = border
        
        row_num = 2
        
        # TEKLİF TAKİP bölümü
        section_cell = ws.cell(row=row_num, column=1, value='TEKLİF TAKİP')
        section_cell.font = Font(bold=True)
        section_cell.fill = section_fill
        section_cell.border = border
        ws.merge_cells(start_row=row_num, start_column=1, end_row=row_num, end_column=len(headers))
        row_num += 1
        
        # Teklif satırları
        teklif_rows = [
            ('TOP. TEKLİF', lambda d: d.get('summary', {}).get('toplam_teklif_adedi', 0), lambda d: d.get('summary', {}).get('toplam_teklif_tutari', 0)),
            ('TOP KABUL OLAN TEK.', lambda d: d.get('summary', {}).get('kabul_adet', 0), lambda d: d.get('summary', {}).get('kabul_tutari', 0)),
            ('KAPSAM İÇİ TEKLİF', lambda d: d.get('summary', {}).get('kapsam_ici_adet', 0), lambda d: d.get('summary', {}).get('kapsam_ici_tutar', 0)),
            ('KAPSAM DIŞI TEKLİF', lambda d: d.get('summary', {}).get('kapsam_disi_adet', 0), lambda d: d.get('summary', {}).get('kapsam_disi_tutar', 0)),
            ('İŞ BİRLİĞİ TEKLİF', lambda d: d.get('summary', {}).get('is_birligi_adet', 0), lambda d: d.get('summary', {}).get('is_birligi_tutar', 0))
        ]
        
        for title, adet_func, tutar_func in teklif_rows:
            ws.cell(row=row_num, column=1, value=title).border = border
            
            for col, year in enumerate(year_list, 2):
                year_data = data['data'].get(str(year), {})
                
                # Adet sütunu
                adet_value = adet_func(year_data)
                adet_cell = ws.cell(row=row_num, column=col, value=adet_value)
                adet_cell.number_format = '0'  # Sayı formatı
                adet_cell.alignment = Alignment(horizontal='right')
                adet_cell.border = border
                
                # TL sütunu
                tutar_value = tutar_func(year_data)
                tutar_cell = ws.cell(row=row_num, column=col + len(year_list), value=tutar_value)
                tutar_cell.number_format = '#,##0.00'  # Para formatı
                tutar_cell.alignment = Alignment(horizontal='right')
                tutar_cell.border = border
            
            row_num += 1
        
        # PARAMETRE bölümü
        section_cell = ws.cell(row=row_num, column=1, value='PARAMETRE')
        section_cell.font = Font(bold=True)
        section_cell.fill = section_fill
        section_cell.border = border
        ws.merge_cells(start_row=row_num, start_column=1, end_row=row_num, end_column=len(headers))
        row_num += 1
        
        # Parametre satırları
        parametre_rows = [
            ('TOP. BACA SAYISI', lambda d: d.get('summary', {}).get('toplam_baca_adedi', 0), lambda d: d.get('summary', {}).get('kabul_tutari', 0)),
            ('TOP. PARAMETRE SAYISI', lambda d: d.get('summary', {}).get('toplam_parametre_adedi', 0), lambda d: 0)
        ]
        
        for title, adet_func, tutar_func in parametre_rows:
            ws.cell(row=row_num, column=1, value=title).border = border
            
            for col, year in enumerate(year_list, 2):
                year_data = data['data'].get(str(year), {})
                
                # Adet sütunu
                adet_value = adet_func(year_data)
                adet_cell = ws.cell(row=row_num, column=col, value=adet_value)
                adet_cell.number_format = '0'  # Sayı formatı
                adet_cell.alignment = Alignment(horizontal='right')
                adet_cell.border = border
                
                # TL sütunu
                tutar_value = tutar_func(year_data)
                tutar_cell = ws.cell(row=row_num, column=col + len(year_list), value=tutar_value)
                tutar_cell.number_format = '#,##0.00'  # Para formatı
                tutar_cell.alignment = Alignment(horizontal='right')
                tutar_cell.border = border
            
            row_num += 1
        
        # Dinamik parametre satırları
        all_params = set()
        for year in year_list:
            year_data = data['data'].get(str(year), {})
            params = year_data.get('summary', {}).get('parametre_sayilari', {})
            all_params.update(params.keys())
        
        # Asgari fiyat fallback
        asgari_fiyatlar_fallback = {
            'TOZ': 6655, 'YG': 3855, 'YANMA GAZI': 3855, 'AMET': 19000, 'AĞIR METAL': 19000,
            'SÜLF.A': 6290, 'SÜLFÜRİK ASİT': 6290, 'VOC': 13170, 'TOC': 13170, 'HIZ': 1430,
            'NEM': 1430, 'HF': 7505, 'HCL': 7505, 'AMONYAK': 6645, 'FORMALDEHİT': 6685,
            'CR+6': 6290, 'FOSFORİK ASİT': 6290, 'HCN': 6810, 'DİOKSİN FURAN': 83950,
            'PAH': 42220, 'PM10': 7370, 'ÇÖKEN TOZ': 6640, 'MODELLEME': 10560
        }
        
        for param in sorted(all_params):
            title = f'TOP. {param} SAYISI'
            ws.cell(row=row_num, column=1, value=title).border = border
            
            for col, year in enumerate(year_list, 2):
                year_data = data['data'].get(str(year), {})
                
                # Adet sütunu
                adet_value = year_data.get('summary', {}).get('parametre_sayilari', {}).get(param, 0)
                adet_cell = ws.cell(row=row_num, column=col, value=adet_value)
                adet_cell.number_format = '0'  # Sayı formatı
                adet_cell.alignment = Alignment(horizontal='right')
                adet_cell.border = border
                
                # TL sütunu (parametre sayısı × asgari fiyat)
                fiyat = asgari_fiyatlar_fallback.get(param, 0)
                tutar_value = adet_value * fiyat
                tutar_cell = ws.cell(row=row_num, column=col + len(year_list), value=tutar_value)
                tutar_cell.number_format = '#,##0.00'  # Para formatı
                tutar_cell.alignment = Alignment(horizontal='right')
                tutar_cell.border = border
            
            row_num += 1
        
        # PERSONEL PERFORMANS bölümü
        section_cell = ws.cell(row=row_num, column=1, value='PERSONEL PERFORMANS')
        section_cell.font = Font(bold=True)
        section_cell.fill = section_fill
        section_cell.border = border
        ws.merge_cells(start_row=row_num, start_column=1, end_row=row_num, end_column=len(headers))
        row_num += 1
        
        # Personel performans satırları
        all_personel = set()
        for year in year_list:
            year_data = data['data'].get(str(year), {})
            personel = year_data.get('summary', {}).get('personel_performans', {})
            all_personel.update(personel.keys())
        
        for personel in sorted(all_personel):
            title = f'{personel} BACA SAYISI'
            ws.cell(row=row_num, column=1, value=title).border = border
            
            for col, year in enumerate(year_list, 2):
                year_data = data['data'].get(str(year), {})
                
                # Adet sütunu
                adet_value = year_data.get('summary', {}).get('personel_performans', {}).get(personel, 0)
                adet_cell = ws.cell(row=row_num, column=col, value=adet_value)
                adet_cell.number_format = '0'  # Sayı formatı
                adet_cell.alignment = Alignment(horizontal='right')
                adet_cell.border = border
                
                # TL sütunu
                tutar_value = year_data.get('summary', {}).get('personel_tutarlar', {}).get(personel, 0)
                tutar_cell = ws.cell(row=row_num, column=col + len(year_list), value=tutar_value)
                tutar_cell.number_format = '#,##0.00'  # Para formatı
                tutar_cell.alignment = Alignment(horizontal='right')
                tutar_cell.border = border
            
            row_num += 1
        
        # Sütun genişliklerini ayarla
        ws.column_dimensions['A'].width = 30
        for col in range(2, len(headers) + 1):
            ws.column_dimensions[get_column_letter(col)].width = 15
        
        # Dosyayı memory'de oluştur
        from io import BytesIO
        output = BytesIO()
        wb.save(output)
        output.seek(0)
        
        # Dosya adı oluştur
        filename = f'PIVOT_Karsilastirma_{"_".join(map(str, year_list))}.xlsx'
        
        return send_file(
            output,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        
    except Exception as e:
        print(f"Pivot XLSX export hatası: {e}")
        import traceback
        print(f"Hata detayı: {traceback.format_exc()}")
        return jsonify({'success': False, 'error': str(e)}), 500

# Admin Backup/Restore endpoints
# @app.route('/api/admin/backup', methods=['POST']) - Admin sistemi silindi
# def api_admin_backup(): - Admin sistemi silindi
    """Admin için veri yedekleme"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'error': 'Oturum açmanız gerekiyor'}), 401
    
    # Sadece admin kullanıcısı erişebilir
    if session.get('username') != 'admin':
        return jsonify({'success': False, 'error': 'Bu işlem için admin yetkisi gerekiyor'}), 403
    
    try:
        from datetime import datetime
        import shutil
        
        # Yedekleme dizini oluştur
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        backup_dir = f'backups/backup_{timestamp}'
        os.makedirs(backup_dir, exist_ok=True)
        
        # Yedeklenecek dosyalar
        data_files = [
            'firma_kayit.json',
            # 'teklif.json' silindi
            'firma_olcum.json',
            'saha_olc.json',
            'parameters.json',
            'baca_bilgileri.json',
            'parametre_olcum.json',
            'parametre_sahabil.json',
            'asgari_fiyatlar.json',
            'forms.json',
            'users.json',
            # 'used_teklif_numbers.json' silindi
            'par_saha_header_groups.json'
        ]
        
        backed_up_files = []
        for file in data_files:
            if os.path.exists(file):
                shutil.copy2(file, backup_dir)
                backed_up_files.append(file)
        
        # Eski yedekleri temizle (7 günden eski)
        cleanup_old_backups()
        
        return jsonify({
            'success': True, 
            'message': f'{len(backed_up_files)} dosya yedeklendi. Yedek klasörü: {backup_dir}',
            'backup_dir': backup_dir,
            'files': backed_up_files
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': f'Yedekleme hatası: {str(e)}'}), 500

@app.route('/api/save_asama3_data', methods=['POST'])
def save_asama3_data():
    """3. Aşama verilerini kaydeder (firma-olcum_kodu-baca kombinasyonu için en son kayıt)."""
    try:
        data = request.get_json()
        
        # 3. Aşama verilerini dosyaya kaydet - trim ile normalize et
        # Tarih formatını gg/aa/yy ss:dd formatına çevir
        def format_tarih(tarih_str):
            if not tarih_str:
                return tarih_str
            try:
                # gg.aa.yyyy ss:dd formatını gg/aa/yy ss:dd formatına çevir
                if '.' in tarih_str and len(tarih_str.split('.')[0]) == 2:
                    return tarih_str.replace('.', '/')[:8] + tarih_str[8:]
                # gg/aa/yyyy ss:dd formatını gg/aa/yy ss:dd formatına çevir
                elif '/' in tarih_str:
                    parts = tarih_str.split('/')
                    if len(parts) >= 3:
                        year_part = parts[2].split(' ')[0]
                        if len(year_part) == 4:  # gg/aa/yyyy formatı
                            time_part = tarih_str[tarih_str.find(' '):] if ' ' in tarih_str else ''
                            return f"{parts[0]}/{parts[1]}/{year_part[:2]}{time_part}"
                    return tarih_str
                # yyyy-mm-dd formatını gg/aa/yy formatına çevir
                elif '-' in tarih_str and len(tarih_str.split('-')[0]) == 4:
                    dt = datetime.strptime(tarih_str, '%Y-%m-%d %H:%M:%S')
                    return dt.strftime('%d/%m/%y %H:%M')
                else:
                    return tarih_str
            except:
                return tarih_str
        
        # Parametreler içindeki tarihleri formatla
        parametreler = data.get('parametreler', [])
        for param in parametreler:
            if 'hesaplamalar' in param:
                for hesaplama in param['hesaplamalar']:
                    if 'degerler' in hesaplama:
                        degerler = hesaplama['degerler']
                        if 'deger_2' in degerler:
                            degerler['deger_2'] = format_tarih(degerler['deger_2'])
        
        # Kilit bilgisini al ve normalize et (0/1)
        locked_raw = data.get('locked', 0)
        locked = 0
        if isinstance(locked_raw, bool):
            locked = 1 if locked_raw else 0
        else:
            try:
                locked = 1 if str(locked_raw).strip().lower() in ('1', 'true', 'evet', 'yes', 'on') else 0
            except Exception:
                locked = 0

        asama3_data = {
            'firma': data.get('firma', '').strip(),
            'olcum_kodu': data.get('olcumKodu', '').strip(),
            'baca': data.get('baca', '').strip(),
            'personel': data.get('personel', '').strip(),
            'asama': 3,
            'parametreler': parametreler,
            'kayit_tarihi': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'locked': locked,
        }
        
        # Mevcut verileri oku - JSONDecodeError yakala
        try:
            with open('asama3_verileri.json', 'r', encoding='utf-8') as f:
                content = f.read().strip()
                if not content:
                    mevcut_veriler = []
                else:
                    mevcut_veriler = json.loads(content)
        except (FileNotFoundError, json.JSONDecodeError):
            mevcut_veriler = []
        
        # Aynı firma-olcum_kodu-baca kombinasyonu için eski kayıtları sil
        mevcut_veriler = [
            record for record in mevcut_veriler
            if not (record.get('firma', '').strip() == asama3_data['firma'] and
                    record.get('olcum_kodu', '').strip() == asama3_data['olcum_kodu'] and
                    record.get('baca', '').strip() == asama3_data['baca'] and
                    record.get('asama') == 3)
        ]
        
        # Yeni kayıt ekle
        mevcut_veriler.append(asama3_data)
        
        # Dosyaya kaydet
        with open('asama3_verileri.json', 'w', encoding='utf-8') as f:
            json.dump(mevcut_veriler, f, ensure_ascii=False, indent=2)
        
        return jsonify({
            'success': True,
            'message': f'{asama3_data["firma"]} - Baca {asama3_data["baca"]} - 3. Aşama verisi kaydedildi!'
        })
        
    except Exception as e:
        print(f"3. Aşama verisi kaydedilirken hata: {e}")
        return jsonify({'success': False, 'message': f'Kaydetme hatası: {str(e)}'}), 500

@app.route('/api/get_asama3_data')
def get_asama3_data():
    """Belirtilen firma-olcum_kodu-baca için 3. Aşama verisini getirir (en son kayıt)."""
    try:
        baca_no = request.args.get('baca_no', '').strip()
        firma = request.args.get('firma', '').strip()
        olcum_kodu = request.args.get('olcum_kodu', '').strip()
        
        if not baca_no:
            return jsonify({'success': False, 'message': 'Baca numarası gerekli'}), 400
        
        # Dosyayı oku - JSONDecodeError yakala
        try:
            with open('asama3_verileri.json', 'r', encoding='utf-8') as f:
                content = f.read().strip()
                if not content:
                    veriler = []
                else:
                    veriler = json.loads(content)
        except (FileNotFoundError, json.JSONDecodeError):
            return jsonify({'success': False, 'message': '3. Aşama verisi bulunamadı veya bozuk'}), 404
        
        # Belirtilen firma-olcum_kodu-baca kombinasyonu için veriyi bul - en son kayıt
        matching_records = []
        for veri in veriler:
            if (veri.get('asama') == 3 and
                str(veri.get('baca', '')).strip() == baca_no and
                (not firma or str(veri.get('firma', '')).strip() == firma) and
                (not olcum_kodu or str(veri.get('olcum_kodu', '')).strip() == olcum_kodu)):
                matching_records.append(veri)
        
        if matching_records:
            # En son kaydı döndür (kayit_tarihi'ne göre sırala)
            latest_record = max(matching_records, key=lambda r: r.get('kayit_tarihi', ''))
            return jsonify({'success': True, 'data': latest_record})
        
        return jsonify({'success': False, 'message': '3. Aşama verisi bulunamadı'}), 404
        
    except Exception as e:
        print(f"3. Aşama verisi getirilirken hata: {e}")
        return jsonify({'success': False, 'message': f'Veri getirme hatası: {str(e)}'}), 500

@app.route('/api/export_eurometric', methods=['GET'])
def export_eurometric():
    """3. Aşama verilerini eurometric.txt formatına dönüştürür.

    Varsayılan olarak dosya indirme başlatır. Eğer `inline=1` (veya true) parametresi
    gönderilirse, içeriği doğrudan text/plain olarak döner (önizleme için)."""
    try:
        firma = request.args.get('firma', '').strip()
        olcum_kodu = request.args.get('olcum_kodu', '').strip()
        baca = request.args.get('baca', '').strip()
        olcum_no = request.args.get('olcum_no', '1')  # Varsayılan 1. ölçüm
        param_index = request.args.get('param_index', '0')  # Varsayılan ilk parametre
        inline = request.args.get('inline', '').strip().lower() in ('1', 'true', 'yes')
        
        if not all([firma, olcum_kodu, baca]):
            return jsonify({'success': False, 'message': 'Firma, ölçüm kodu ve baca bilgisi gerekli'}), 400
        
        # 3. Aşama verilerini yükle
        try:
            with open('asama3_verileri.json', 'r', encoding='utf-8') as f:
                content = f.read().strip()
                if not content:
                    asama3_veriler = []
                else:
                    asama3_veriler = json.loads(content)
        except (FileNotFoundError, json.JSONDecodeError):
            return jsonify({'success': False, 'message': '3. Aşama verisi bulunamadı'}), 404
        
        # İlgili kaydı bul
        asama3_data = None
        for veri in asama3_veriler:
            if (veri.get('asama') == 3 and
                str(veri.get('firma', '')).strip() == firma and
                str(veri.get('olcum_kodu', '')).strip() == olcum_kodu and
                str(veri.get('baca', '')).strip() == baca):
                asama3_data = veri
                break
        
        if not asama3_data:
            return jsonify({'success': False, 'message': '3. Aşama verisi bulunamadı'}), 404
        
        # 1. ve 2. Aşama verilerini de yükle
        asama1_data = None
        asama2_data = None
        
        try:
            with open('asama_verileri.json', 'r', encoding='utf-8') as f:
                content = f.read().strip()
                if content:
                    asama1_veriler = json.loads(content)
                    for veri in asama1_veriler:
                        if (veri.get('asama') == 1 and
                            str(veri.get('firma', '')).strip() == firma and
                            str(veri.get('olcum_kodu', '')).strip() == olcum_kodu and
                            str(veri.get('baca', '')).strip() == baca):
                            asama1_data = veri
                            break
        except:
            pass
        
        try:
            with open('asama2_verileri.json', 'r', encoding='utf-8') as f:
                content = f.read().strip()
                if content:
                    asama2_veriler = json.loads(content)
                    for veri in asama2_veriler:
                        if (veri.get('asama') == 2 and
                            str(veri.get('firma', '')).strip() == firma and
                            str(veri.get('olcum_kodu', '')).strip() == olcum_kodu and
                            str(veri.get('baca', '')).strip() == baca):
                            asama2_data = veri
                            break
        except:
            pass
        
        # Eurometric şablonunu oku
        try:
            with open('formlar/eurometric.txt', 'r', encoding='utf-8') as f:
                template = f.read()
        except FileNotFoundError:
            return jsonify({'success': False, 'message': 'Eurometric şablonu bulunamadı'}), 404
        
        # Şablonu doldur
        olcum_no_int = int(olcum_no) if olcum_no.isdigit() else 1
        param_idx = int(param_index) if param_index.isdigit() else 0
        output = generate_eurometric_report(template, asama1_data, asama2_data, asama3_data, olcum_no_int, param_idx)
        
        # Eğer inline isteniyorsa, içeriği doğrudan döndür (önizleme)
        if inline:
            return output, 200, {'Content-Type': 'text/plain; charset=utf-8'}

        # Dosyayı oluştur ve gönder (indirme)
        output_bytes = BytesIO()
        output_bytes.write(output.encode('utf-8'))
        output_bytes.seek(0)

        # Kullanıcı isteğine göre dosya adı:
        # firma adı (ilk kelime)_baca adı(ölçüm kodu)_parametre_olcNo
        # Örn: wwww_proses(5)_Toz-1olc

        # Firma ilk kelime
        firma_first = (firma.split()[0] if firma else '').strip() or 'Firma'

        # Baca adı (gelen baca parametresi)
        baca_adi = (baca or '').strip() or 'Baca'

        # Parametre adı: öncelik sırası 2. aşama, yoksa 3. aşama (param_index'e göre)
        parametre_adi = ''
        try:
            if asama2_data and asama2_data.get('parametreler'):
                params = asama2_data['parametreler']
                if param_idx < len(params):
                    p = params[param_idx]
                    parametre_adi = str(p.get('parametre', '')).strip()
            if (not parametre_adi) and asama3_data and asama3_data.get('parametreler'):
                params3 = asama3_data['parametreler']
                if param_idx < len(params3):
                    p3 = params3[param_idx]
                    parametre_adi = str(p3.get('parametre', '') or p3.get('parametre_adi', '')).strip()
        except Exception:
            pass

        if not parametre_adi:
            parametre_adi = 'Parametre'

        # Parametre adını sadeleştir (ilk satır / ilk kelime, boşlukları tireye çevir)
        parametre_clean = str(parametre_adi).split('\n')[0].split('|')[0].strip()
        if ' ' in parametre_clean:
            parametre_clean = parametre_clean.split(' ')[0]
        parametre_clean = parametre_clean.replace(' ', '_')

        # Ölçüm numarası etiketi
        try:
            olc_no_int = int(olcum_no_int)
        except Exception:
            olc_no_int = 1
        olc_suffix = f"{olc_no_int}olc"

        filename = f"{firma_first}_{baca_adi}({olcum_kodu})_{parametre_clean}-{olc_suffix}.txt"
        
        return send_file(
            output_bytes,
            mimetype='text/plain',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        print(f"Eurometric export hatası: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'message': f'Export hatası: {str(e)}'}), 500

def format_date_for_eurometric(date_str):
    """Tarihi gg/aa/yy ss:dd formatına çevirir."""
    if not date_str:
        return ''
    
    from datetime import datetime
    
    try:
        # Eğer zaten gg/aa/yy ss:dd formatındaysa
        if '/' in date_str and ':' in date_str:
            parts = date_str.split()
            if len(parts) == 2:
                date_part = parts[0]
                time_part = parts[1]
                date_parts = date_part.split('/')
                if len(date_parts) == 3:
                    day = date_parts[0].zfill(2)
                    month = date_parts[1].zfill(2)
                    year = date_parts[2]
                    # Yıl 4 haneli ise 2 haneye düşür
                    if len(year) == 4:
                        year = year[-2:]
                    return f"{day}/{month}/{year} {time_part}"
            return date_str
        
        # ISO format veya diğer formatlar için
        if 'T' in date_str:
            dt = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
        else:
            # Çeşitli formatları dene
            for fmt in ['%Y-%m-%d %H:%M:%S', '%d/%m/%Y %H:%M', '%d.%m.%Y %H:%M', '%d-%m-%Y %H:%M']:
                try:
                    dt = datetime.strptime(date_str, fmt)
                    break
                except:
                    continue
            else:
                return date_str
        
        day = dt.strftime('%d')
        month = dt.strftime('%m')
        year = dt.strftime('%y')  # 2 haneli yıl
        hour = dt.strftime('%H')
        minute = dt.strftime('%M')
        return f"{day}/{month}/{year} {hour}:{minute}"
    except:
        return date_str

def generate_eurometric_report(template, asama1_data, asama2_data, asama3_data, olcum_no=1, param_index=0):
    """Eurometric şablonunu verilerle doldurur."""
    import re
    # Yardımcı formatlayıcı: genişlik ve ondalık sayısına göre sağa hizalı metin üretir
    def fmt(value, width=None, precision=None):
        try:
            if value is None or str(value).strip() == '':
                s = ''
            else:
                v = str(value).replace(',', '.')
                num = float(v)
                if precision is not None:
                    s = f"{num:.{precision}f}"
                else:
                    s = str(num)
        except Exception:
            s = str(value) if value is not None else ''
        if width is not None and width > 0:
            s = s.rjust(width)[:width]
        return s
    
    # Belirtilen parametreyi al (param_index'e göre)
    if not asama3_data or not asama3_data.get('parametreler'):
        return template
    
    parametreler = asama3_data['parametreler']
    if param_index >= len(parametreler):
        param_index = 0  # Geçersiz index ise ilk parametreyi kullan
    
    first_parametre = parametreler[param_index]
    hesaplamalar = first_parametre.get('hesaplamalar', [])
    
    # Ortalama satırlarını filtrele
    all_travers_rows = [h for h in hesaplamalar if not h.get('ortalama', False)]
    
    # Ölçüm için travers sayısını belirle (2. aşamadan veya ilk 6 travers)
    travers_count = 6
    if asama2_data and asama2_data.get('parametreler'):
        params_a2 = asama2_data['parametreler']
        if param_index < len(params_a2):
            param_a2 = params_a2[param_index]
            travers_count = int(param_a2.get('traversSayisi', 6)) if param_a2.get('traversSayisi') else 6
    
    # Belirtilen ölçüm numarasına göre travers'leri al
    # 1. ölçüm: 0-5, 2. ölçüm: 6-11, 3. ölçüm: 12-17, vs.
    start_idx = (olcum_no - 1) * travers_count
    end_idx = start_idx + travers_count
    travers_rows = all_travers_rows[start_idx:end_idx]
    
    if not travers_rows:
        return template
    
    # Genel bilgileri hazırla
    first_row = travers_rows[0].get('degerler', {})
    
    # Cihaz seri
    cihaz_seri = first_row.get('deger_19', '') or (asama1_data.get('cihazSeri', '') if asama1_data else '')
    
    # Başlangıç: öncelik 2. aşama verisi (ilgili parametre + ölçüm no), yoksa 3. aşama ilk travers zamanı
    baslangic_raw = first_row.get('deger_2', '')
    if asama2_data and asama2_data.get('parametreler'):
        try:
            params_a2 = asama2_data['parametreler']
            param_a2 = params_a2[param_index] if param_index < len(params_a2) else params_a2[0]
            if olcum_no == 1 and param_a2.get('baslangic'):
                baslangic_raw = param_a2.get('baslangic')
            elif olcum_no == 2 and param_a2.get('olcum2Baslangic'):
                baslangic_raw = param_a2.get('olcum2Baslangic')
            elif olcum_no == 3 and param_a2.get('olcum3Baslangic'):
                baslangic_raw = param_a2.get('olcum3Baslangic')
        except Exception:
            # Herhangi bir hata olursa 3. aşama travers zamanına geri düş
            pass
    baslangic = format_date_for_eurometric(baslangic_raw)
    
    # Ayarlanan süre (2. aşamadan veya 3. aşamadan) - sadece dakika cinsinden (dd)
    ayarlanan_sure_dk = ''
    if asama2_data and asama2_data.get('parametreler'):
        try:
            params_a2 = asama2_data['parametreler']
            first_param = params_a2[param_index] if param_index < len(params_a2) else params_a2[0]
            ayarlanan_sure_raw = first_param.get('ayarlananSure') or first_param.get('sure', '')
            if ayarlanan_sure_raw:
                # Saniyeyi dakikaya çevir (720 sn = 12 dakika)
                try:
                    sure_value = float(str(ayarlanan_sure_raw).replace(',', '.'))
                    if sure_value >= 60:
                        # Saniye cinsinden (720 sn = 12 dakika)
                        minutes = int(sure_value // 60)
                        ayarlanan_sure_dk = str(minutes)
                    else:
                        # Zaten dakika cinsinden
                        ayarlanan_sure_dk = str(int(sure_value))
                except:
                    ayarlanan_sure_dk = str(ayarlanan_sure_raw)
        except Exception:
            pass

    if not ayarlanan_sure_dk and travers_rows:
        try:
            last_sure_raw = (travers_rows[-1].get('degerler', {}).get('deger_5', '') or '')[:5]
            if last_sure_raw and ':' in last_sure_raw:
                mm_part = last_sure_raw.split(':')[0].strip()
                if mm_part:
                    ayarlanan_sure_dk = str(int(mm_part))
        except Exception:
            pass
    
    # Baca No - önce 1. aşamadan bacaNo değerini al, yoksa 1. aşamadan baca adını al, yoksa 2. aşamadan baca adını al
    baca_no = ''
    if asama1_data:
        # Önce bacaNo değerini al (1. aşamadan)
        baca_no = asama1_data.get('bacaNo', '') or asama1_data.get('baca', '')
    elif asama2_data and asama2_data.get('baca'):
        # 1. aşama yoksa 2. aşamadan baca adını al
        baca_no = asama2_data.get('baca', '')

    # "3.0" gibi değerleri raporda "3" olarak göstermek için sadeleştir
    if baca_no is not None and str(baca_no).strip() != '':
        s = str(baca_no).strip()
        try:
            v = float(s.replace(',', '.'))
            if v.is_integer():
                baca_no = str(int(v))
            else:
                baca_no = s
        except Exception:
            baca_no = s
    else:
        baca_no = ''
    
    # Travers sayısı
    travers_sayisi = len(travers_rows)
    
    # Nozzle çapı (ilk travers'ten)
    nozzle_capi_mm = first_row.get('deger_4', '')
    
    # Atmosfer basıncı (1. aşamadan veya 3. aşamadan)
    atmosfer_kpa = first_row.get('deger_18', '') or (asama1_data.get('atmosferBasinci', '') if asama1_data else '')
    # kPa -> mmHg dönüşümü (1 kPa ≈ 7.50062 mmHg)
    atmosfer_mmhg = ''
    try:
        if str(atmosfer_kpa).strip() != '':
            kpa_val = float(str(atmosfer_kpa).replace(',', '.'))
            mmhg_val = kpa_val * 7.50062
            atmosfer_mmhg = f"{mmhg_val:.2f}"
    except Exception:
        atmosfer_mmhg = ''
    
    # Pitot katsayı
    pitot_k = first_row.get('deger_22', '0.84')
    
    # Nem
    nem = first_row.get('deger_23', '') or (asama1_data.get('nem', '') if asama1_data else '')
    
    # Şablonu doldur - genel bilgiler
    template = template.replace('{cihaz_seri|12}', (cihaz_seri or '')[:12])
    template = template.replace('{baslangic|16}', (baslangic or '')[:16])
    # Ayarlanan süre ve travers sayısı tam sayı olarak (örn: 32 dk., 8)
    template = template.replace('{ayarlanan_sure_dk|3}', fmt(ayarlanan_sure_dk, 3, 0))
    template = template.replace('{baca_no|3}', fmt(baca_no, 3, 0))
    template = template.replace('{travers_sayisi|3}', fmt(travers_sayisi, 3, 0))
    # Nozzle çapı: önce 3. aşamadan, yoksa 2. aşamadaki nozzleCapi'den al
    try:
        if (nozzle_capi_mm is None or str(nozzle_capi_mm).strip() == '') and asama2_data and asama2_data.get('parametreler'):
            params_a2_nz = asama2_data['parametreler']
            p2_nz = params_a2_nz[param_index] if param_index < len(params_a2_nz) else params_a2_nz[0]
            nozzle_capi_mm = p2_nz.get('nozzleCapi', '')
    except Exception:
        pass
    template = template.replace('{nozzle_capi_mm|3,1}', fmt(nozzle_capi_mm, 3, 1))
    template = template.replace('{atmosfer_kpa|6,2}', fmt(atmosfer_kpa, 6, 2))
    template = template.replace('{atmosfer_mmhg|6,2}', fmt(atmosfer_mmhg, 6, 2))
    template = template.replace('{pitot_k|5,2}', fmt(pitot_k, 5, 2))
    template = template.replace('{nem|5,2}', fmt(nem, 5, 2))
    
    # Travers verilerini oluştur
    travers_section = ''
    total_travers = len(travers_rows)
    for idx, row in enumerate(travers_rows, 1):
        degerler = row.get('degerler', {})
        
        t_no = str(idx)
        t_sure_mmss = (degerler.get('deger_5', '00:00') or '00:00')[:5]
        t_ort_baca_hizi_ms = degerler.get('deger_7', '')
        t_sayac_hacmi_l = degerler.get('deger_17', '') or degerler.get('deger_12', '')
        t_baca_sicak_c = degerler.get('deger_8', '')
        # Sayaç gaz sicaklığı: 3. aşama tablosunda 16. sütun (deger_16)
        t_sayac_gaz_c = degerler.get('deger_16', '')
        t_baca_mutlak_kpa = degerler.get('deger_9', '')
        # kPa -> mmHg
        t_baca_mutlak_mmhg = ''
        try:
            if str(t_baca_mutlak_kpa).strip() != '':
                _k = float(str(t_baca_mutlak_kpa).replace(',', '.'))
                t_baca_mutlak_mmhg = f"{_k * 7.50062:.2f}"
        except Exception:
            t_baca_mutlak_mmhg = ''
        t_pitot_dp_pa = degerler.get('deger_20', '')
        # Pa -> mmHg
        t_pitot_dp_mmhg = ''
        try:
            if str(t_pitot_dp_pa).strip() != '':
                _p = float(str(t_pitot_dp_pa).replace(',', '.'))
                t_pitot_dp_mmhg = f"{_p * 0.00750062:.4f}"
        except Exception:
            t_pitot_dp_mmhg = ''
        t_izok_verim = degerler.get('deger_6', '')

        # Son travers için altına çizgi koyma, diğerlerinde koy
        is_last = (idx == total_travers)
        travers_block = f"""BACA NO       : {fmt(baca_no, 3, 0)}
TRAVERS NO    : {fmt(t_no, 3, 0)}
Sure          : {t_sure_mmss[:5]}
Ort.baca hizi :{fmt(t_ort_baca_hizi_ms, 6, 1)} m/s
Sayac hacmi   :{fmt(t_sayac_hacmi_l, 7, 1)} l
Baca sicakligi: {fmt(t_baca_sicak_c, 6, 1)} C
Sayac gaz sic.: {fmt(t_sayac_gaz_c, 6, 1)} C
Baca mutlk bas:{fmt(t_baca_mutlak_kpa, 7, 2)} kPa
Baca mutlk bas:{fmt(t_baca_mutlak_mmhg, 7, 2)} mmHg
Pitot dP      :{fmt(t_pitot_dp_pa, 7, 2)} Pa
Pitot dP      :{fmt(t_pitot_dp_mmhg, 7, 4)} mmHg
izokin.verimi : {fmt(t_izok_verim, 5, 2)}
"""
        if not is_last:
            # Ara traversler: bir ayırıcı çizgi ve ardından bir boş satır
            travers_block += " --------------------------- \n\n"
        travers_section += travers_block
    
    # Travers bölümünü değiştir
    template = re.sub(
        r'\[\[TRAVERS_START\]\].*?\[\[TRAVERS_END\]\]',
        travers_section,
        template,
        flags=re.DOTALL
    )
    
    # Özet rapor bölümünü doldur (ortalama değerler)
    avg_rows = [h for h in hesaplamalar if h.get('ortalama', False)]
    if avg_rows:
        # Her ölçüm (1., 2., 3. ölçüm vb.) için ayrı ORTALAMA satırı varsa,
        # istenen ölçüme karşılık gelen ortalama satırını kullan.
        try:
            avg_index = max(0, min(len(avg_rows) - 1, olcum_no - 1))
        except Exception:
            avg_index = 0

        avg_row = avg_rows[avg_index].get('degerler', {})

        o_baslangic = baslangic

        # Bitiş: Öncelik 2. aşamadaki ilgili ölçüm bitiş zamanı (olcum1Bitis/olcum2Bitis/olcum3Bitis)
        o_bitis = ''
        o_bitis_raw = ''
        if asama2_data and asama2_data.get('parametreler'):
            try:
                params_a2_bitis = asama2_data['parametreler']
                param_a2_bitis = params_a2_bitis[param_index] if param_index < len(params_a2_bitis) else params_a2_bitis[0]
                bitis_key = f"olcum{olcum_no}_Bitis"
                # Eski kayıtlarda alt çizgisiz olabilir
                alt_keys = [bitis_key, bitis_key.replace('_', ''), f"olcum{olcum_no}Bitis"]
                for k in alt_keys:
                    if k in param_a2_bitis and param_a2_bitis.get(k):
                        o_bitis_raw = param_a2_bitis.get(k)
                        break
            except Exception:
                o_bitis_raw = ''

        # Eğer 2. aşamadan bitiş alınamadıysa, son travers'in başlangıcını kullan (eski davranış)
        if not o_bitis_raw and travers_rows:
            o_bitis_raw = travers_rows[-1].get('degerler', {}).get('deger_2', '')

        o_bitis = format_date_for_eurometric(o_bitis_raw) if o_bitis_raw else ''

        # Örnekleme süresi: son travers'in süresi (hh:mm)
        last_travers_sure = travers_rows[-1].get('degerler', {}).get('deger_5', '00:00')[:5] if travers_rows else ''
        o_ornekleme_sure_hhmm = last_travers_sure

        # Ortalama baca hızı (7. sütun) - ölçüme özel ortalama
        o_ort_baca_hizi_ms = avg_row.get('deger_7', '')

        # Ortalama vakum debisi: 3. aşama ortalama satırındaki 16. ve 15. sütunlar
        # 15. sütun: Sayaç L/D, 16. sütun: Sayaç NL/D
        o_ort_vakum_debi_l_d = avg_row.get('deger_15', '')
        o_ort_vakum_debi_nl_d = avg_row.get('deger_16', '')

        o_kurugaz_nl = avg_row.get('deger_11', '')
        o_sayac_hacmi_l = avg_row.get('deger_12', '')
        o_baca_sicak_c = avg_row.get('deger_8', '')
        o_sayac_gaz_c = avg_row.get('deger_16', '')  # Sayaç gaz için 16. sütun değerini de kullan
        o_ref_sic_c = ''  # Referans sıcaklık

        # Baca mutlak basıncı (kPa) ve mmHg dönüşümü
        o_baca_mutlak_kpa = avg_row.get('deger_9', '')
        o_baca_mutlak_mmhg = ''
        try:
            if str(o_baca_mutlak_kpa).strip() != '':
                _k = float(str(o_baca_mutlak_kpa).replace(',', '.'))
                o_baca_mutlak_mmhg = f"{_k * 7.50062:.2f}"
        except Exception:
            o_baca_mutlak_mmhg = ''

        # Sayaç basıncı: Atmosfer basıncı ile aynı olacak
        o_sayac_basinci_kpa = atmosfer_kpa
        o_sayac_basinci_mmhg = atmosfer_mmhg

        # Referans basınç (şimdilik boş bırakıyoruz)
        o_ref_basinc_kpa = ''
        o_ref_basinc_mmhg = ''

        # Pitot dP: 1. ölçüm ORTALAMA 20. sütun
        o_pitot_dp_pa = avg_row.get('deger_20', '')
        o_pitot_dp_mmhg = ''
        try:
            if str(o_pitot_dp_pa).strip() != '':
                _p = float(str(o_pitot_dp_pa).replace(',', '.'))
                o_pitot_dp_mmhg = f"{_p * 0.00750062:.4f}"
        except Exception:
            o_pitot_dp_mmhg = ''
        o_izok_verim = avg_row.get('deger_6', '')
        
        # Özet rapor placeholder'larını değiştir
        template = template.replace('{o_baslangic|16}', (o_baslangic or '')[:16])
        template = template.replace('{o_bitis|16}', (o_bitis or '')[:16])
        template = template.replace('{o_ornekleme_sure_hhmm|6}', (o_ornekleme_sure_hhmm or '')[:6])
        template = template.replace('{o_ort_baca_hizi_ms|6,1}', fmt(o_ort_baca_hizi_ms, 6, 1))
        template = template.replace('{o_ort_vakum_debi_nl_d|5,1}', fmt(o_ort_vakum_debi_nl_d, 5, 1))
        template = template.replace('{o_ort_vakum_debi_l_d|5,1}', fmt(o_ort_vakum_debi_l_d, 5, 1))
        template = template.replace('{o_kurugaz_nl|7,1}', fmt(o_kurugaz_nl, 7, 1))
        template = template.replace('{o_sayac_hacmi_l|7,1}', fmt(o_sayac_hacmi_l, 7, 1))
        template = template.replace('{o_baca_sicak_c|6,1}', fmt(o_baca_sicak_c, 6, 1))
        template = template.replace('{o_sayac_gaz_c|6,1}', fmt(o_sayac_gaz_c, 6, 1))
        template = template.replace('{o_ref_sic_c|4,1}', fmt(o_ref_sic_c, 4, 1))
        template = template.replace('{o_baca_mutlak_kpa|7,2}', fmt(o_baca_mutlak_kpa, 7, 2))
        template = template.replace('{o_baca_mutlak_mmhg|7,2}', fmt(o_baca_mutlak_mmhg, 7, 2))
        template = template.replace('{o_sayac_basinci_kpa|7,2}', fmt(o_sayac_basinci_kpa, 7, 2))
        template = template.replace('{o_sayac_basinci_mmhg|7,2}', fmt(o_sayac_basinci_mmhg, 7, 2))
        template = template.replace('{o_ref_basinc_kpa|7,2}', fmt(o_ref_basinc_kpa, 7, 2))
        template = template.replace('{o_ref_basinc_mmhg|7,2}', fmt(o_ref_basinc_mmhg, 7, 2))
        template = template.replace('{o_pitot_dp_pa|7,2}', fmt(o_pitot_dp_pa, 7, 2))
        template = template.replace('{o_pitot_dp_mmhg|7,4}', fmt(o_pitot_dp_mmhg, 7, 4))
        template = template.replace('{o_izok_verim|5,2}', fmt(o_izok_verim, 5, 2))
    
    return template

# @app.route('/api/admin/restore', methods=['POST']) - Admin sistemi silindi
# def api_admin_restore(): - Admin sistemi silindi
    """Admin için veri geri yükleme"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'error': 'Oturum açmanız gerekiyor'}), 401
    
    # Sadece admin kullanıcısı erişebilir
    if session.get('username') != 'admin':
        return jsonify({'success': False, 'error': 'Bu işlem için admin yetkisi gerekiyor'}), 403
    
    try:
        import shutil
        from datetime import datetime
        
        # En son yedek klasörünü bul
        backup_base_dir = 'backups'
        if not os.path.exists(backup_base_dir):
            return jsonify({'success': False, 'error': 'Yedek klasörü bulunamadı'}), 404
        
        # Yedek klasörlerini listele
        backup_dirs = [d for d in os.listdir(backup_base_dir) if d.startswith('backup_')]
        if not backup_dirs:
            return jsonify({'success': False, 'error': 'Yedek dosyası bulunamadı'}), 404
        
        # En son yedek klasörünü seç
        latest_backup = sorted(backup_dirs)[-1]
        latest_backup_path = os.path.join(backup_base_dir, latest_backup)
        
        # Mevcut dosyaları yedekle (güvenlik için)
        current_backup_time = datetime.now().strftime("%Y%m%d_%H%M%S")
        current_backup_dir = f'backups/current_backup_{current_backup_time}'
        os.makedirs(current_backup_dir, exist_ok=True)
        
        data_files = [
            'firma_kayit.json',
            # 'teklif.json' silindi
            'firma_olcum.json',
            'saha_olc.json',
            'parameters.json',
            'baca_bilgileri.json',
            'parametre_olcum.json',
            'parametre_sahabil.json',
            'asgari_fiyatlar.json',
            'forms.json',
            'users.json',
            # 'used_teklif_numbers.json' silindi
            'par_saha_header_groups.json'
        ]
        
        # Mevcut dosyaları yedekle
        for file in data_files:
            if os.path.exists(file):
                shutil.copy2(file, current_backup_dir)
        
        # Yedek dosyalarını geri yükle
        restored_files = []
        for file in data_files:
            backup_file = os.path.join(latest_backup_path, file)
            if os.path.exists(backup_file):
                shutil.copy2(backup_file, file)
                restored_files.append(file)
        
        return jsonify({
            'success': True, 
            'message': f'{len(restored_files)} dosya geri yüklendi. Yedek: {latest_backup}',
            'restored_files': restored_files,
            'backup_used': latest_backup
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': f'Geri yükleme hatası: {str(e)}'}), 500

def cleanup_old_backups():
    """7 günden eski yedekleri temizler"""
    import shutil
    backup_dir = 'backups'
    if not os.path.exists(backup_dir):
        return
    
    from datetime import datetime
    current_time = datetime.now()
    for item in os.listdir(backup_dir):
        item_path = os.path.join(backup_dir, item)
        if os.path.isdir(item_path):
            # Dizin oluşturma zamanını kontrol et
            creation_time = datetime.fromtimestamp(os.path.getctime(item_path))
            if (current_time - creation_time).days > 7:
                shutil.rmtree(item_path)
                print(f"🗑️ Eski yedek silindi: {item}")

# @app.route('/api/admin/backup-download', methods=['POST']) - Admin sistemi silindi
# def api_admin_backup_download(): - Admin sistemi silindi
    """Admin için veri yedekleme - ZIP dosyası olarak indir"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'error': 'Oturum açmanız gerekiyor'}), 401
    
    # Sadece admin kullanıcısı erişebilir
    if session.get('username') != 'admin':
        return jsonify({'success': False, 'error': 'Bu işlem için admin yetkisi gerekiyor'}), 403
    
    try:
        import zipfile
        import io
        from datetime import datetime
        
        # Yedeklenecek dosyalar
        data_files = [
            'firma_kayit.json',
            # 'teklif.json' silindi
            'firma_olcum.json',
            'saha_olc.json',
            'parameters.json',
            'baca_bilgileri.json',
            'parametre_olcum.json',
            'parametre_sahabil.json',
            'asgari_fiyatlar.json',
            'forms.json',
            'users.json',
            # 'used_teklif_numbers.json' silindi
            'par_saha_header_groups.json'
        ]
        
        # ZIP dosyası oluştur
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for file in data_files:
                if os.path.exists(file):
                    zip_file.write(file, file)
        
        zip_buffer.seek(0)
        
        # Dosya adı oluştur
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'emisyon_backup_{timestamp}.zip'
        
        return zip_buffer.getvalue(), 200, {
            'Content-Type': 'application/zip',
            'Content-Disposition': f'attachment; filename="{filename}"'
        }
        
    except Exception as e:
        return jsonify({'success': False, 'error': f'Yedekleme hatası: {str(e)}'}), 500

        return jsonify({'success': False, 'error': f'Dosya okuma hatası: {str(e)}'}), 500

# @app.route('/api/admin/restore-upload', methods=['POST']) - Admin sistemi silindi
# def api_admin_restore_upload(): - Admin sistemi silindi
    """Admin için veri geri yükleme - Yüklenen dosyalardan"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'error': 'Oturum açmanız gerekiyor'}), 401
    
    # Sadece admin kullanıcısı erişebilir
    if session.get('username') != 'admin':
        return jsonify({'success': False, 'error': 'Bu işlem için admin yetkisi gerekiyor'}), 403
    
    try:
        import zipfile
        import io
        import shutil
        from datetime import datetime
        
        if 'files' not in request.files:
            return jsonify({'success': False, 'error': 'Dosya bulunamadı'}), 400
        
        files = request.files.getlist('files')
        if not files or files[0].filename == '':
            return jsonify({'success': False, 'error': 'Dosya seçilmedi'}), 400
        
        # Mevcut dosyaları yedekle (güvenlik için)
        current_backup_time = datetime.now().strftime("%Y%m%d_%H%M%S")
        current_backup_dir = f'backups/current_backup_{current_backup_time}'
        os.makedirs(current_backup_dir, exist_ok=True)
        
        data_files = [
            'firma_kayit.json',
            # 'teklif.json' silindi
            'firma_olcum.json',
            'saha_olc.json',
            'parameters.json',
            'baca_bilgileri.json',
            'parametre_olcum.json',
            'parametre_sahabil.json',
            'asgari_fiyatlar.json',
            'forms.json',
            'users.json',
            # 'used_teklif_numbers.json' silindi
            'par_saha_header_groups.json'
        ]
        
        # Mevcut dosyaları yedekle
        for file in data_files:
            if os.path.exists(file):
                shutil.copy2(file, current_backup_dir)
        
        restored_files = []
        
        # ZIP dosyası kontrolü
        for file in files:
            if file.filename.endswith('.zip'):
                # ZIP dosyasını çıkar
                zip_data = file.read()
                with zipfile.ZipFile(io.BytesIO(zip_data), 'r') as zip_file:
                    for zip_info in zip_file.infolist():
                        if zip_info.filename in data_files:
                            # Dosyayı çıkar ve kaydet
                            file_data = zip_file.read(zip_info.filename)
                            with open(zip_info.filename, 'wb') as f:
                                f.write(file_data)
                            restored_files.append(zip_info.filename)
            else:
                # Tekil JSON dosyası
                filename = file.filename
                if filename in data_files:
                    file.save(filename)
                    restored_files.append(filename)
        
        return jsonify({
            'success': True, 
            'message': f'{len(restored_files)} dosya geri yüklendi.',
            'restored_files': restored_files
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': f'Geri yükleme hatası: {str(e)}'}), 500


@app.route('/api/backups/run-now', methods=['POST'])
def api_backups_run_now():
    if not session.get('logged_in'):
        return jsonify({'success': False, 'error': 'Oturum açmanız gerekiyor'}), 401
    if session.get('username') != 'admin':
        return jsonify({'success': False, 'error': 'Bu işlem için admin yetkisi gerekiyor'}), 403

    try:
        from datetime import datetime as _dt
        import zipfile

        base_dir = os.path.dirname(os.path.abspath(__file__))
        backup_base_dir = os.path.join(base_dir, 'backups')
        os.makedirs(backup_base_dir, exist_ok=True)

        timestamp = _dt.now().strftime('%Y%m%d_%H%M%S')
        backup_name = f'backup_{timestamp}'
        backup_dir = os.path.join(backup_base_dir, backup_name)
        os.makedirs(backup_dir, exist_ok=True)

        excluded_dirs = {'backups', 'yedek', '__pycache__', '.git', '.idea', 'venv', '.venv', 'env'}

        copied = []
        for root_dir, dirs, files in os.walk(base_dir):
            dirs[:] = [d for d in dirs if d not in excluded_dirs]

            for file_name in files:
                src = os.path.join(root_dir, file_name)
                rel_path = os.path.relpath(src, base_dir)
                dest_path = os.path.join(backup_dir, rel_path)
                dest_dir = os.path.dirname(dest_path)
                if dest_dir and not os.path.exists(dest_dir):
                    os.makedirs(dest_dir, exist_ok=True)
                shutil.copy2(src, dest_path)
                copied.append(rel_path)

        cleanup_old_backups()

        zip_root = os.path.join(base_dir, 'yedek')
        os.makedirs(zip_root, exist_ok=True)
        zip_path = os.path.join(zip_root, f'emisyon_backup_{timestamp}.zip')

        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for root_dir, _dirs, files in os.walk(backup_dir):
                for file_name in files:
                    full_path = os.path.join(root_dir, file_name)
                    rel_path = os.path.relpath(full_path, backup_dir)
                    zip_file.write(full_path, arcname=rel_path)

        return jsonify({
            'success': True,
            'message': f'{len(copied)} dosya yedeklendi.',
            'backup_name': backup_name,
            'backup_dir': backup_dir,
            'zip_file': zip_path,
            'files': copied
        })

    except Exception as e:
        return jsonify({'success': False, 'error': f'Yedekleme hatası: {str(e)}'}), 500


@app.route('/api/backups/list', methods=['GET'])
def api_backups_list():
    if not session.get('logged_in'):
        return jsonify({'success': False, 'error': 'Oturum açmanız gerekiyor'}), 401
    if session.get('username') != 'admin':
        return jsonify({'success': False, 'error': 'Bu işlem için admin yetkisi gerekiyor'}), 403

    backup_base_dir = 'backups'
    if not os.path.exists(backup_base_dir):
        return jsonify({'success': True, 'backups': []})

    backups = []
    for item in os.listdir(backup_base_dir):
        if not item.startswith('backup_'):
            continue
        item_path = os.path.join(backup_base_dir, item)
        if not os.path.isdir(item_path):
            continue

        try:
            created_ts = os.path.getctime(item_path)
        except OSError:
            created_ts = 0

        total_size = 0
        for root_dir, _dirs, files in os.walk(item_path):
            for fname in files:
                fpath = os.path.join(root_dir, fname)
                try:
                    total_size += os.path.getsize(fpath)
                except OSError:
                    pass

        created_dt = datetime.fromtimestamp(created_ts) if created_ts else None
        backups.append({
            'name': item,
            'created_at': created_dt.isoformat(sep=' ') if created_dt else '',
            'display_date': created_dt.strftime('%d.%m.%Y %H:%M') if created_dt else '',
            'size_bytes': total_size
        })

    backups.sort(key=lambda b: b.get('created_at', ''), reverse=True)

    return jsonify({'success': True, 'backups': backups})


@app.route('/api/backups/delete', methods=['POST'])
def api_backups_delete():
    if not session.get('logged_in'):
        return jsonify({'success': False, 'error': 'Oturum açmanız gerekiyor'}), 401
    if session.get('username') != 'admin':
        return jsonify({'success': False, 'error': 'Bu işlem için admin yetkisi gerekiyor'}), 403

    data = request.get_json(silent=True) or {}
    names = data.get('names') or []
    if not isinstance(names, list) or not names:
        return jsonify({'success': False, 'error': 'Silinecek yedek bulunamadı'}), 400

    backup_base_dir = 'backups'
    deleted = []

    for name in names:
        if not isinstance(name, str):
            continue
        if not name.startswith('backup_'):
            continue
        if '/' in name or '\\' in name or '..' in name:
            continue

        item_path = os.path.join(backup_base_dir, name)
        if os.path.isdir(item_path):
            try:
                shutil.rmtree(item_path)
                deleted.append(name)
            except Exception:
                continue

    return jsonify({'success': True, 'deleted': deleted, 'count': len(deleted)})


@app.route('/api/backups/restore', methods=['POST'])
def api_backups_restore():
    if not session.get('logged_in'):
        return jsonify({'success': False, 'error': 'Oturum açmanız gerekiyor'}), 401
    if session.get('username') != 'admin':
        return jsonify({'success': False, 'error': 'Bu işlem için admin yetkisi gerekiyor'}), 403

    data = request.get_json(silent=True) or {}
    name = data.get('name') or ''
    if not isinstance(name, str) or not name.startswith('backup_') or '/' in name or '\\' in name or '..' in name:
        return jsonify({'success': False, 'error': 'Geçersiz yedek adı'}), 400

    backup_base_dir = 'backups'
    backup_path = os.path.join(backup_base_dir, name)
    if not os.path.isdir(backup_path):
        return jsonify({'success': False, 'error': 'Yedek klasörü bulunamadı'}), 404

    try:
        current_backup_time = datetime.now().strftime('%Y%m%d_%H%M%S')
        current_backup_dir = f'backups/current_backup_{current_backup_time}'
        os.makedirs(current_backup_dir, exist_ok=True)

        data_files = [
            'firma_kayit.json',
            'firma_olcum.json',
            'saha_olc.json',
            'parameters.json',
            'baca_bilgileri.json',
            'parametre_olcum.json',
            'parametre_sahabil.json',
            'asgari_fiyatlar.json',
            'forms.json',
            'users.json',
            'par_saha_header_groups.json'
        ]

        for fname in data_files:
            if os.path.exists(fname):
                shutil.copy2(fname, os.path.join(current_backup_dir, fname))

        restored = []
        for fname in data_files:
            src = os.path.join(backup_path, fname)
            if os.path.exists(src):
                shutil.copy2(src, fname)
                restored.append(fname)

        return jsonify({'success': True, 'restored_files': restored, 'backup_used': name})

    except Exception as e:
        return jsonify({'success': False, 'error': f'Geri yükleme hatası: {str(e)}'}), 500


# HESAPLAMALAR TABLOSU API ENDPOINT'LERİ
@app.route('/api/save_hesaplama_data', methods=['POST'])
def save_hesaplama_data():
    try:
        data = request.get_json()
        
        with open('hesaplama_data.json', 'w', encoding='utf-8') as f:
            json.dump(data['data'], f, ensure_ascii=False, indent=2)
        
        return jsonify({'success': True, 'message': 'Hesaplama verileri kaydedildi'})
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/load_hesaplama_data', methods=['GET'])
def load_hesaplama_data():
    try:
        if os.path.exists('hesaplama_data.json'):
            with open('hesaplama_data.json', 'r', encoding='utf-8') as f:
                data = json.load(f)
            return jsonify({'success': True, 'data': data})
        else:
            return jsonify({'success': False, 'message': 'Dosya bulunamadı'})
            
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/get_rapor2_data', methods=['GET'])
def get_rapor2_data():
    """RAPOR2 verilerini yükle"""
    try:
        if os.path.exists(RAPOR2_FILE):
            with open(RAPOR2_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return jsonify({'success': True, 'content': data.get('content', '')})
        else:
            return jsonify({'success': True, 'content': ''})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/save_rapor2_data', methods=['POST'])
def save_rapor2_data():
    """RAPOR2 verilerini kaydet"""
    if not session.get('logged_in'):
        return jsonify({'success': False, 'error': 'Oturum açmanız gerekiyor'}), 401
    
    try:
        data = request.get_json()
        content = data.get('content', '')
        
        # rapor2_sablonu.json dosyasına kaydet
        with open(RAPOR2_FILE, 'w', encoding='utf-8') as f:
            json.dump({'content': content}, f, ensure_ascii=False, indent=2)
        
        return jsonify({'success': True, 'message': 'Rapor şablonu başarıyla kaydedildi'})
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/save_cikti_raporu', methods=['POST'])
def save_cikti_raporu():
    """Çıktı raporunu kaydet"""
    try:
        data = request.get_json()
        content = data.get('content', '')
        
        # cikti_raporu.txt dosyasına kaydet
        with open('cikti_raporu.txt', 'w', encoding='utf-8') as f:
            f.write(content)
        
        return jsonify({'success': True, 'message': 'Çıktı raporu kaydedildi'})
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

if __name__ == '__main__':
    # Teklif numarası benzersizliği için migration çalıştır
    # migrate_existing_teklif_numbers() - teklif sistemi silindi
    
    # Mevcut teklif numaralarını 3 haneli formata dönüştür - teklif sistemi silindi
    # convert_teklif_numbers_to_3_digit() - teklif sistemi silindi
    
    # Mevcut teklifleri 001'den başlayarak yeniden sırala - teklif sistemi silindi
    # resequence_teklif_numbers() - teklif sistemi silindi
    
    # Render için port ayarı (Render'ın verdiği PORT değişkenini kullan, yoksa 5001 kullan)
    port = int(os.environ.get('PORT', 7000))
    
    # Geliştirme modu kontrolü - her zaman debug=True
    debug = True
    
    # Uygulamayı başlat
    if __name__ == '__main__':
        app.run(
            host='0.0.0.0',
            port=port,
            debug=debug,
            threaded=True,
            use_reloader=True,
            use_debugger=True
        )